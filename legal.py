import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF for PDF processing
from docx import Document
from PIL import Image
import io
import pathlib
import tempfile
import os
from datetime import datetime

# --- Configuration ---
st.set_page_config(
    page_title="LegalVet AI - Legal Document Analysis",
    page_icon="⚖️",
    layout="wide"
)

# Get Google API key from Streamlit secrets
try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except KeyError:
    st.error("""
    🔑 **Google API Key Not Found!**
    
    Please add your Google API key to Streamlit secrets:
    
    **For local development:**
    Create a `.streamlit/secrets.toml` file in your project directory with:
    ```toml
    GOOGLE_API_KEY = "your_google_api_key_here"
    ```
    
    **For Streamlit Cloud deployment:**
    1. Go to your app's settings in Streamlit Cloud
    2. Click on "Secrets" 
    3. Add the following:
    ```toml
    GOOGLE_API_KEY = "your_google_api_key_here"
    ```
    
    **How to get a Google API Key:**
    1. Go to [Google AI Studio](https://aistudio.google.com/app/apikey)
    2. Create a new API key
    3. Copy the key and add it to your secrets
    """)
    st.stop()

# Initialize Google Gemini Client
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel(model_name="gemini-2.0-flash-exp")
except Exception as e:
    st.error(f"Error configuring Google Gemini client: {e}")
    st.stop()

# --- Helper Functions ---

def extract_text_from_pdf(pdf_bytes):
    """Extract text from PDF bytes."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_word(word_bytes):
    """Extract text from Word document bytes."""
    try:
        # Save bytes to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(word_bytes)
            tmp_file.flush()
            
            # Read with python-docx
            doc = Document(tmp_file.name)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Clean up temp file
            os.unlink(tmp_file.name)
            return text
    except Exception as e:
        st.error(f"Error extracting text from Word document: {e}")
        return None

def analyze_legal_document(file_bytes, filename, document_type, analysis_type, jurisdiction="Delhi, India"):
    """Analyze legal document using Gemini API."""
    
    try:
        # Save the uploaded file bytes to a temporary file
        file_extension = pathlib.Path(filename).suffix.lower()
        temp_file_path = pathlib.Path(f"temp_legal_{filename}")
        
        with open(temp_file_path, "wb") as f:
            f.write(file_bytes)

        st.info(f"Uploading '{filename}' to LegalVet AI...")
        
        # Determine MIME type
        if file_extension == '.pdf':
            mime_type = "application/pdf"
        elif file_extension in ['.docx', '.doc']:
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            mime_type = "application/octet-stream"
        
        # Upload the file using the File API
        uploaded_file = genai.upload_file(
            path=temp_file_path, 
            display_name=filename, 
            mime_type=mime_type
        )
        st.success(f"'{filename}' uploaded successfully to LegalVet AI")

        # Create legal analysis prompt based on analysis type
        prompt = create_legal_analysis_prompt(document_type, analysis_type, jurisdiction, filename)
        
        st.info(f"Analyzing {document_type} for {analysis_type}...")
        
        # Generate analysis
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=131072,
            temperature=0.2,  # Lower temperature for more factual legal analysis
        )
        
        response = model.generate_content(
            [uploaded_file, prompt],
            generation_config=generation_config
        )

        # Check for issues
        if response.candidates and response.candidates[0].finish_reason == 4:
            st.error("Content policy issue detected. Please ensure the document contains appropriate legal content.")
            return None, "Content Policy Error"

        # Extract response text
        if hasattr(response, 'text') and response.text:
            analysis_result = response.text
        else:
            if response.candidates and hasattr(response.candidates[0], 'content') and response.candidates[0].content:
                parts = response.candidates[0].content.parts
                if parts:
                    analysis_result = ''.join([part.text for part in parts if hasattr(part, 'text')])
                else:
                    raise Exception("No content parts found in response")
            else:
                raise Exception("No valid text content found in response")

        # Clean up resources
        try:
            genai.delete_file(uploaded_file.name)
            st.info(f"Temporary file '{uploaded_file.display_name}' deleted from Gemini.")
        except Exception as e_del:
            st.warning(f"Could not delete temporary file: {e_del}")

        if temp_file_path.exists():
            temp_file_path.unlink()

        return analysis_result, "Legal analysis completed successfully."

    except Exception as e:
        st.error(f"Error during legal document analysis: {e}")
        # Clean up on error
        if 'uploaded_file' in locals():
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
        if 'temp_file_path' in locals() and temp_file_path.exists():
            temp_file_path.unlink()
        
        return None, f"Error: {str(e)}"

def create_legal_analysis_prompt(document_type, analysis_type, jurisdiction, filename):
    """Create specialized legal analysis prompts."""
    
    base_prompt = f"""You are LegalVet AI, an expert legal assistant specializing in Indian law, particularly for the jurisdiction of {jurisdiction}.
You are assisting a practicing lawyer with document analysis and legal vetting.

**Document:** {filename}
**Document Type:** {document_type}
**Analysis Type:** {analysis_type}
**Jurisdiction:** {jurisdiction}

**IMPORTANT LEGAL DISCLAIMERS:**
- This analysis is for professional legal assistance only
- This is NOT a substitute for professional legal judgment
- The practicing lawyer retains full responsibility for legal decisions
- This analysis is based on Indian legal framework and precedents
"""

    if analysis_type == "Document Vetting":
        return base_prompt + """
**COMPREHENSIVE DOCUMENT VETTING ANALYSIS:**

Please provide a thorough legal vetting of this document including:

## 1. EXECUTIVE SUMMARY
- Brief overview of the document
- Key legal issues identified
- Overall risk assessment (High/Medium/Low)
- Immediate action items

## 2. LEGAL COMPLIANCE ANALYSIS
- Compliance with applicable Indian laws and regulations
- Statutory requirements check
- Regulatory compliance status
- Missing legal formalities

## 3. CONTRACTUAL ANALYSIS (if applicable)
- Key terms and conditions review
- Identification of ambiguous clauses
- Enforceability assessment
- Standard vs. non-standard provisions
- Termination and dispute resolution clauses

## 4. RISK ASSESSMENT
- Legal risks identified
- Commercial risks
- Operational risks
- Compliance risks
- Mitigation strategies

## 5. RED FLAGS AND CONCERNS
- Critical issues requiring immediate attention
- Potentially problematic clauses
- Legal loopholes or vulnerabilities
- Areas of non-compliance

## 6. RECOMMENDATIONS
- Specific amendments suggested
- Additional clauses to be included
- Legal protections to be added
- Negotiation points

## 7. PRECEDENT AND CASE LAW
- Relevant Indian court decisions
- Applicable statutory provisions
- Recent legal developments affecting this type of document

## 8. ACTION ITEMS
- Immediate steps required
- Documents/approvals needed
- Timeline for corrections
- Priority levels

**Please be thorough, practical, and provide actionable legal advice suitable for a practicing lawyer in Delhi.**
"""

    elif analysis_type == "Contract Review":
        return base_prompt + """
**DETAILED CONTRACT REVIEW:**

Please provide comprehensive contract analysis including:

## 1. CONTRACT OVERVIEW
- Type of contract and its purpose
- Parties involved and their legal status
- Contract value and duration
- Key commercial terms

## 2. LEGAL FRAMEWORK ANALYSIS
- Applicable laws and regulations
- Governing law and jurisdiction clauses
- Statutory compliance requirements
- Regulatory approvals needed

## 3. TERMS AND CONDITIONS REVIEW
- Payment terms and conditions
- Delivery/performance obligations
- Warranties and representations
- Indemnity and liability clauses
- Force majeure provisions

## 4. RISK MATRIX
- Performance risks
- Financial risks
- Legal and regulatory risks
- Termination risks
- Dispute resolution risks

## 5. CLAUSE-BY-CLAUSE ANALYSIS
- Problematic or unclear clauses
- Missing essential clauses
- One-sided or unfair terms
- Enforceability concerns

## 6. NEGOTIATION STRATEGY
- Key points to negotiate
- Alternative clause suggestions
- Deal-breaker issues
- Compromise positions

## 7. COMPLIANCE CHECKLIST
- Statutory requirements
- Documentation needed
- Approvals and registrations
- Tax implications

## 8. LEGAL PRECEDENTS
- Relevant case law
- Industry-specific considerations
- Recent judicial trends

**Focus on practical, actionable advice for contract negotiation and finalization.**
"""

    elif analysis_type == "Legal Opinion":
        return base_prompt + """
**COMPREHENSIVE LEGAL OPINION:**

Please provide a detailed legal opinion including:

## 1. FACTUAL MATRIX
- Summary of facts presented
- Key legal questions raised
- Scope of opinion

## 2. LEGAL ANALYSIS
- Applicable laws and regulations
- Statutory provisions analysis
- Constitutional considerations (if any)
- Procedural requirements

## 3. CASE LAW RESEARCH
- Relevant Supreme Court decisions
- High Court precedents
- Recent judgments on similar issues
- Distinguishing factors

## 4. LEGAL POSITION
- Clear legal opinion on the matter
- Strength of legal position
- Alternative interpretations
- Potential challenges

## 5. PRACTICAL IMPLICATIONS
- Real-world consequences
- Business impact assessment
- Operational considerations
- Strategic recommendations

## 6. RISK ANALYSIS
- Legal risks involved
- Probability of success
- Potential liabilities
- Mitigation strategies

## 7. RECOMMENDATIONS
- Suggested course of action
- Alternative approaches
- Precautionary measures
- Documentation requirements

## 8. CONCLUSION
- Summary of opinion
- Key takeaways
- Next steps

**Provide a well-reasoned, citable legal opinion suitable for professional use.**
"""

    elif analysis_type == "Due Diligence":
        return base_prompt + """
**LEGAL DUE DILIGENCE REPORT:**

Please conduct thorough legal due diligence including:

## 1. DOCUMENT AUTHENTICITY
- Verification of document validity
- Signature and authorization check
- Stamp duty and registration compliance
- Chain of documents analysis

## 2. LEGAL TITLE AND OWNERSHIP
- Title verification (if applicable)
- Ownership rights analysis
- Encumbrance assessment
- Third-party rights

## 3. REGULATORY COMPLIANCE
- Licenses and approvals status
- Compliance with sectoral regulations
- Environmental clearances (if applicable)
- Tax compliance verification

## 4. LITIGATION ANALYSIS
- Pending litigation assessment
- Potential legal disputes
- Regulatory actions
- Compliance defaults

## 5. FINANCIAL AND LEGAL OBLIGATIONS
- Outstanding liabilities
- Contingent liabilities
- Guarantee obligations
- Security interests

## 6. CORPORATE GOVERNANCE
- Board resolutions and approvals
- Shareholder consents
- Regulatory filings
- Compliance certificates

## 7. MATERIAL AGREEMENTS
- Key contract analysis
- Performance obligations
- Termination rights
- Assignment restrictions

## 8. RED FLAGS AND RECOMMENDATIONS
- Critical issues identified
- Risk mitigation strategies
- Further investigation needed
- Deal structure recommendations

**Provide comprehensive due diligence suitable for M&A, investments, or major transactions.**
"""

    else:  # General Analysis
        return base_prompt + """
**GENERAL LEGAL DOCUMENT ANALYSIS:**

Please provide comprehensive analysis including:

## 1. DOCUMENT SUMMARY
- Type and nature of document
- Key provisions and terms
- Legal significance

## 2. LEGAL VALIDITY
- Formal requirements compliance
- Legal capacity of parties
- Consideration and mutuality
- Statutory compliance

## 3. ENFORCEABILITY
- Legal enforceability assessment
- Practical enforcement challenges
- Jurisdiction and governing law
- Dispute resolution mechanisms

## 4. LEGAL RISKS
- Identified legal risks
- Compliance gaps
- Potential liabilities
- Enforcement challenges

## 5. RECOMMENDATIONS
- Suggested improvements
- Additional legal protections
- Compliance requirements
- Risk mitigation strategies

## 6. LEGAL PRECEDENTS
- Relevant case law
- Statutory provisions
- Regulatory guidelines

**Provide practical, actionable legal analysis suitable for professional legal practice.**
"""

def generate_legal_chat_response(user_prompt, chat_history, uploaded_files, jurisdiction="Delhi, India"):
    """Generate legal consultation response."""
    try:
        # Build context from chat history
        conversation_context = ""
        if len(chat_history) > 1:
            recent_messages = chat_history[-8:]  # Keep last 8 messages for context
            for msg in recent_messages[:-1]:  # Exclude the current message
                conversation_context += f"{msg['role'].title()}: {msg['content']}\n"

        # Build system prompt for legal consultation
        system_prompt = f"""You are LegalVet AI, an expert legal assistant and support lawyer specializing in Indian law, particularly for the jurisdiction of {jurisdiction}.

You are assisting a practicing lawyer with legal consultation, research, and advisory services.

**Your Expertise Includes:**
- Indian Constitutional Law, Civil Law, Criminal Law, Corporate Law
- Contract drafting and review
- Legal research and case law analysis
- Statutory interpretation and compliance
- Court procedures and legal strategy
- Legal documentation and drafting
- Risk assessment and legal opinions

**Jurisdiction Specific Knowledge:**
- Delhi High Court procedures and precedents
- Delhi-specific regulations and local laws
- National Capital Territory of Delhi laws
- Central government regulations applicable in Delhi

**Guidelines for Legal Assistance:**
- Provide accurate, well-researched legal information
- Cite relevant case law and statutory provisions
- Offer practical, actionable legal advice
- Identify potential legal risks and mitigation strategies
- Suggest appropriate legal procedures and documentation
- Maintain professional legal standards
- Always remind that final legal decisions rest with the practicing lawyer

**IMPORTANT DISCLAIMERS:**
- This is professional legal assistance for a practicing lawyer
- This is not a substitute for independent legal judgment
- The practicing lawyer retains full professional responsibility
- Advice is based on Indian legal framework and {jurisdiction} jurisdiction

Previous conversation:
{conversation_context}

Current legal query: {user_prompt}

Please provide comprehensive legal assistance with proper legal reasoning, relevant precedents, and practical recommendations.
"""

        # Prepare content for generation
        content_parts = [system_prompt]
        
        # Add uploaded files if any
        uploaded_file_objects = []
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # Save uploaded file temporarily
                    file_extension = pathlib.Path(uploaded_file.name).suffix.lower()
                    temp_path = pathlib.Path(f"temp_chat_{uploaded_file.name}")
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    
                    # Determine MIME type
                    if file_extension == '.pdf':
                        mime_type = "application/pdf"
                    elif file_extension in ['.docx', '.doc']:
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        mime_type = "application/octet-stream"
                    
                    # Upload to Gemini
                    gemini_file = genai.upload_file(
                        path=temp_path, 
                        display_name=uploaded_file.name, 
                        mime_type=mime_type
                    )
                    uploaded_file_objects.append(gemini_file)
                    content_parts.append(gemini_file)
                    
                    # Clean up temp file
                    temp_path.unlink()
                    
                except Exception as e:
                    st.warning(f"Could not process {uploaded_file.name}: {e}")

        # Generate response
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=131072,
            temperature=0.3,
        )

        response = model.generate_content(
            content_parts,
            generation_config=generation_config
        )

        # Clean up uploaded files from Gemini
        for gemini_file in uploaded_file_objects:
            try:
                genai.delete_file(gemini_file.name)
            except Exception as e:
                pass

        # Extract response text
        if hasattr(response, 'text') and response.text:
            return response.text
        elif response.candidates and response.candidates[0].content:
            parts = response.candidates[0].content.parts
            if parts:
                return ''.join([part.text for part in parts if hasattr(part, 'text')])
            else:
                return "I apologize, but I couldn't generate a proper legal response. Could you please rephrase your legal query?"
        else:
            return "I'm having trouble processing your legal request right now. Please try again."

    except Exception as e:
        return f"I encountered an error while processing your legal query: {str(e)}. Please try asking your question again or check if your uploaded files are valid legal documents."

def create_legal_word_document(analysis_text, document_title="Legal Analysis Report"):
    """Creates a professional legal Word document from analysis text."""
    doc = Document()
    
    # Add title
    title = doc.add_heading(document_title, 0)
    title.alignment = 1  # Center alignment
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph("Generated by: LegalVet AI - Legal Document Analysis Tool")
    doc.add_paragraph().add_run("CONFIDENTIAL LEGAL ANALYSIS").bold = True
    
    doc.add_page_break()
    
    # Process the analysis text
    lines = analysis_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle different markdown formats
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Add footer disclaimer
    doc.add_page_break()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("LEGAL DISCLAIMER: ").bold = True
    disclaimer.add_run("This analysis is generated by LegalVet AI for professional legal assistance. "
                      "It is not a substitute for independent legal judgment. The practicing lawyer "
                      "retains full professional responsibility for all legal decisions and advice. "
                      "This analysis is based on Indian legal framework and should be verified with "
                      "current laws and regulations.")
    
    return doc

# --- Streamlit App ---

def main():
    # App Header
    st.title("⚖️ LegalVet AI - Legal Document Analysis & Support")
    st.markdown("""
    **Professional Legal Assistant for Indian Lawyers**
    
    Comprehensive document vetting, contract review, legal analysis, and professional support for legal practice in India.
    """)

    # Sidebar for settings
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # Jurisdiction selector
        jurisdiction = st.selectbox(
            "Jurisdiction:",
            ["Delhi, India", "Mumbai, India", "Bangalore, India", "Chennai, India", "Kolkata, India", "Pune, India", "Other"],
            index=0
        )
        
        if jurisdiction == "Other":
            jurisdiction = st.text_input("Specify Jurisdiction:", "India")
        
        st.markdown("---")
        st.markdown("**📞 Emergency Legal Contacts**")
        st.markdown("- Delhi High Court: 011-2389-1419")
        st.markdown("- Bar Council of Delhi: 011-2338-9371")
        st.markdown("- Legal Aid Services: 15100")
        
        st.markdown("---")
        st.markdown("**📚 Quick Legal Resources**")
        st.markdown("- [Supreme Court of India](https://main.sci.gov.in/)")
        st.markdown("- [Delhi High Court](http://delhihighcourt.nic.in/)")
        st.markdown("- [Bar Council of India](https://www.barcouncilofindia.org/)")

    # Create tabs for different functionalities
    tab1, tab2 = st.tabs(["📋 Document Analysis", "💬 Legal Consultation"])

    with tab1:
        st.header("📋 Legal Document Analysis & Vetting")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Document upload
            uploaded_file = st.file_uploader(
                "Upload Legal Document:",
                type=['pdf', 'docx', 'doc'],
                help="Upload PDF or Word documents for legal analysis"
            )
        
        with col2:
            if uploaded_file:
                st.success(f"📄 **{uploaded_file.name}** uploaded successfully!")
                
                # File info
                file_size = len(uploaded_file.getvalue())
                st.info(f"File size: {file_size / 1024:.1f} KB")
        
        if uploaded_file is not None:
            # Analysis options
            col1, col2 = st.columns(2)
            
            with col1:
                document_type = st.selectbox(
                    "Document Type:",
                    [
                        "Contract/Agreement",
                        "Legal Notice",
                        "Petition/Application", 
                        "Property Document",
                        "Corporate Document",
                        "Compliance Document",
                        "Court Filing",
                        "Legal Opinion",
                        "MOU/LOI",
                        "Power of Attorney",
                        "Will/Testament",
                        "Other Legal Document"
                    ]
                )
            
            with col2:
                analysis_type = st.selectbox(
                    "Analysis Type:",
                    [
                        "Document Vetting",
                        "Contract Review", 
                        "Legal Opinion",
                        "Due Diligence",
                        "General Analysis"
                    ]
                )
            
            # Analysis button
            if st.button("🔍 Analyze Document", key="analyze_doc", type="primary"):
                with st.spinner(f"Analyzing {uploaded_file.name} for {analysis_type}..."):
                    file_bytes = uploaded_file.getvalue()
                    
                    analysis_result, status_message = analyze_legal_document(
                        file_bytes, 
                        uploaded_file.name, 
                        document_type, 
                        analysis_type, 
                        jurisdiction
                    )
                    
                    if analysis_result:
                        st.success(f"✅ {status_message}")
                        
                        # Display analysis
                        st.subheader("📊 Legal Analysis Report")
                        st.markdown(analysis_result)
                        
                        # Download options
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Download as Word document
                            doc = create_legal_word_document(
                                analysis_result, 
                                f"Legal Analysis - {uploaded_file.name}"
                            )
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            
                            st.download_button(
                                label="📥 Download Analysis Report (.docx)",
                                data=doc_io,
                                file_name=f"legal_analysis_{uploaded_file.name.split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        with col2:
                            # Download as text file
                            st.download_button(
                                label="📥 Download as Text (.txt)",
                                data=analysis_result,
                                file_name=f"legal_analysis_{uploaded_file.name.split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                mime="text/plain"
                            )
                    else:
                        st.error(f"❌ Analysis failed: {status_message}")

    with tab2:
        st.header("💬 Legal Consultation with LegalVet AI")
        st.markdown("""
        Chat with LegalVet AI for legal guidance, research assistance, and professional support.
        Upload documents for context or ask direct legal questions.
        """)
        
        # Document upload for context
        consultation_files = st.file_uploader(
            "Upload documents for consultation context:",
            type=['pdf', 'docx', 'doc'],
            accept_multiple_files=True,
            key="consultation_files"
        )
        
        if consultation_files:
            st.success(f"📄 {len(consultation_files)} document(s) uploaded for consultation context")
            for file in consultation_files:
                st.text(f"• {file.name}")
        
        st.divider()
        
        # Initialize chat session state
        if 'legal_chat_messages' not in st.session_state:
            st.session_state.legal_chat_messages = [
                {
                    "role": "assistant", 
                    "content": f"Hello! I'm LegalVet AI, your legal support assistant. I'm here to help with legal document analysis, case research, legal opinions, and general legal consultation for practice in {jurisdiction}.\n\nHow can I assist you with your legal work today?"
                }
            ]
        
        # Chat interface
        chat_container = st.container()
        
        with chat_container:
            # Display chat messages
            for message in st.session_state.legal_chat_messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask your legal question..."):
            # Add user message
            st.session_state.legal_chat_messages.append({"role": "user", "content": prompt})
            
            # Display user message
            with chat_container:
                with st.chat_message("user"):
                    st.markdown(prompt)
            
            # Generate and display response
            with chat_container:
                with st.chat_message("assistant"):
                    with st.spinner("Researching legal response..."):
                        response = generate_legal_chat_response(
                            prompt,
                            st.session_state.legal_chat_messages,
                            consultation_files if consultation_files else [],
                            jurisdiction
                        )
                        st.markdown(response)
                        
                        # Add response to chat history
                        st.session_state.legal_chat_messages.append({"role": "assistant", "content": response})
        
        # Clear chat button
        if st.button("🗑️ Clear Consultation History"):
            st.session_state.legal_chat_messages = [
                {
                    "role": "assistant", 
                    "content": f"Hello! I'm LegalVet AI, your legal support assistant. I'm here to help with legal document analysis, case research, legal opinions, and general legal consultation for practice in {jurisdiction}.\n\nHow can I assist you with your legal work today?"
                }
            ]
            st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666;'>
    <p><strong>LegalVet AI</strong> - Professional Legal Document Analysis & Support</p>
    <p>🔒 Confidential • ⚖️ Professional • 🇮🇳 India-Focused</p>
    <p><em>Disclaimer: This tool provides legal assistance to practicing lawyers. It does not replace professional legal judgment or constitute legal advice to clients.</em></p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
