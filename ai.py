import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import json
import pathlib

# Get Google API key from Streamlit secrets
try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except KeyError:
    st.error("""
    🔑 **Google API Key Not Found!**
    
    Please add your Google API key to Streamlit secrets.
    See instructions in app.py for details.
    """)
    st.stop()

# Initialize Google Gemini Client
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-05-06")
except Exception as e:
    st.error(f"Error configuring Google Gemini client: {e}")
    st.stop()

# UAE Curriculum Structure
UAE_CURRICULUM_STRUCTURE = {
    "Kindergarten": {
        "age": "4-5",
        "focus": "Introduction to AI through stories and play",
        "themes": ["What is AI?", "People and Machines", "AI in Our World", "Talking to Machines", "Learning and Helping", "Be Kind with AI"]
    },
    "Cycle 1": {
        "grades": "1-4",
        "focus": "Understanding differences between machines and humans, digital thinking skills, AI applications",
        "core_areas": ["Foundational concepts", "Digital thinking", "AI recognition", "Basic ethics"]
    },
    "Cycle 2": {
        "grades": "5-8",
        "focus": "Designing and evaluating AI systems, learning about bias and algorithms, ethical AI use",
        "core_areas": ["AI design", "Algorithms", "Bias awareness", "Ethical considerations", "Data understanding"]
    },
    "Cycle 3": {
        "grades": "9-12",
        "focus": "Preparing for higher education and careers, command engineering, real-world scenarios",
        "core_areas": ["Advanced AI concepts", "Command engineering", "Real-world applications", "Career preparation", "Innovation"]
    }
}

# --- Helper Functions ---

def extract_units_from_curriculum(curriculum_text):
    """Extract unit information from generated curriculum"""
    units = []
    lines = curriculum_text.split('\n')
    
    # Look for the Units/Modules section
    in_units_section = False
    current_unit = None
    
    for i, line in enumerate(lines):
        # Check if we're in the units section
        if "Units/Modules" in line or "Units and Modules" in line:
            in_units_section = True
            continue
            
        # Check if we've left the units section
        if in_units_section and line.startswith("##") and not line.startswith("###"):
            break
            
        # Extract unit information
        if in_units_section:
            # Look for unit titles (usually marked with ### or **Unit X:**)
            if line.startswith("###") or ("Unit" in line and (":" in line or "-" in line)):
                # Clean up the unit title
                unit_title = line.replace("###", "").replace("**", "").strip()
                if unit_title:
                    # Extract unit details from following lines
                    unit_details = []
                    j = i + 1
                    while j < len(lines) and not (lines[j].startswith("###") or "Unit" in lines[j]):
                        if lines[j].strip() and not lines[j].startswith("##"):
                            unit_details.append(lines[j].strip())
                        j += 1
                        if len(unit_details) >= 5:  # Limit details to prevent too much content
                            break
                    
                    units.append({
                        "title": unit_title,
                        "details": "\n".join(unit_details[:5])  # Keep first 5 detail lines
                    })
    
    # If no units found with the above method, try alternative parsing
    if not units:
        for i, line in enumerate(lines):
            if "Unit" in line and any(char in line for char in [":", "-", "–"]):
                unit_title = line.strip()
                # Get some context
                unit_details = []
                j = i + 1
                while j < len(lines) and j < i + 6:
                    if lines[j].strip():
                        unit_details.append(lines[j].strip())
                    j += 1
                
                units.append({
                    "title": unit_title,
                    "details": "\n".join(unit_details[:4])
                })
    
    return units

def generate_uae_curriculum(level, cycle_info, specific_grade=None, language="English"):
    """Generates UAE-aligned AI curriculum for a specific level"""
    
    # Prepare grade-specific context
    grade_context = f"for Grade {specific_grade}" if specific_grade else f"for {level}"
    
    # Language-specific instructions
    language_instruction = f"""
**IMPORTANT LANGUAGE REQUIREMENT:**
Generate this entire curriculum in {language}. 
{"Use Arabic language throughout, including headings, content, and examples. Provide Arabic equivalents for technical AI terms." if language == "Arabic" else "Use clear, simple English suitable for UAE context."}
All content must be culturally appropriate for UAE schools.
"""
    
    prompt = f"""You are an expert in educational curriculum development, specifically for AI education in UAE schools.
You need to create a comprehensive AI curriculum {grade_context} following UAE's AI curriculum framework.

{language_instruction}

**UAE AI CURRICULUM FRAMEWORK:**
The UAE AI curriculum encompasses seven core areas:
1. Foundational concepts
2. Data and algorithms
3. Software usage
4. Ethical awareness
5. Real-world applications
6. Innovation and project design
7. Policies and community engagement

**SPECIFIC LEVEL INFORMATION:**
{json.dumps(cycle_info, indent=2)}
{f"**Target Grade: {specific_grade}**" if specific_grade else ""}

**CURRICULUM REQUIREMENTS:**

## 1. **Course Overview**
- {f"Grade {specific_grade} specific focus" if specific_grade else "Age group/grades covered"}
- Integration with Computing, Creative Design, and Innovation subject
- Total hours per week (within existing subject hours)
- Assessment approach aligned with UAE standards

## 2. **Learning Objectives**
- 8-10 specific, measurable objectives {grade_context}
- Aligned with UAE's seven core AI areas
- Age-appropriate and culturally relevant
- Focus on practical skills and ethical awareness

## 3. **Units/Modules** (Generate 6-8 units)
For each unit:
- Unit title {"(in Arabic)" if language == "Arabic" else "(English and Arabic consideration)"}
- Duration (in weeks)
- Key concepts aligned with UAE framework
- Learning outcomes
- Integration with existing subjects
- Cultural relevance and local examples

## 4. **Teaching Methodology**
Based on the level:
{"- Kindergarten: Stories, play-based learning, interactive activities" if level == "Kindergarten" else ""}
{"- Cycle 1: Hands-on activities, digital thinking games, simple projects" if "Cycle 1" in level else ""}
{"- Cycle 2: Project-based learning, AI system design, ethical discussions" if "Cycle 2" in level else ""}
{"- Cycle 3: Real-world scenarios, career exploration, advanced projects" if "Cycle 3" in level else ""}

## 5. **Assessment Framework**
- Formative assessment strategies
- Project-based assessments
- Skills demonstration
- Portfolio development
- Ethical reasoning evaluation

## 6. **Resources and Materials**
- Digital tools and platforms
- Unplugged activities
- Local AI examples (UAE context)
- {"Arabic language resources" if language == "Arabic" else "Bilingual resources consideration"}

## 7. **Ethical AI Focus**
- Age-appropriate discussions on AI ethics
- Bias awareness activities
- Responsible AI usage
- Digital citizenship in UAE context

## 8. **Real-World Connections**
- UAE's AI initiatives and vision
- Local AI applications
- Career pathways in UAE
- Connection to UAE's digital transformation

## 9. **Innovation Projects**
- Age-appropriate innovation challenges
- Connection to UAE's innovation goals
- Cross-curricular projects
- Community engagement opportunities

## 10. **Teacher Support**
- Professional development needs
- Lesson plan templates
- Activity guides
- Assessment rubrics

Generate a comprehensive curriculum that aligns with UAE's AI education vision, is culturally appropriate, and builds progressive AI literacy {grade_context}.

Format the output in Markdown with clear headings and structure.
{"Remember: Generate all content in Arabic language." if language == "Arabic" else ""}
"""
    
    try:
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=32768,
            temperature=0.7,
        )
        
        response = model.generate_content(
            [prompt],
            generation_config=generation_config
        )
        
        if hasattr(response, 'text') and response.text:
            return response.text
        else:
            return "Error: Could not generate curriculum"
            
    except Exception as e:
        st.error(f"Error generating curriculum: {e}")
        return None

def generate_uae_textbook_unit(level, unit_info, curriculum_context, specific_grade=None, language="English"):
    """Generates a textbook unit for UAE AI curriculum"""
    
    # Get example structure from Kindergarten
    kg_example = """
Example from UAE Kindergarten AI Textbook:
- Story-based learning (e.g., "Robo the Helper")
- Interactive activities (coloring, matching, crafts)
- Songs and movement
- Visual aids and characters
- Simple vocabulary introduction
- Assessment through observation
"""
    
    # Prepare grade-specific context
    grade_context = f"Grade {specific_grade}" if specific_grade else level
    
    # Language-specific instructions
    language_instruction = f"""
**IMPORTANT LANGUAGE REQUIREMENT:**
Generate this entire textbook unit in {language}. 
{"Use Arabic language throughout, including stories, activities, instructions, and assessments. Provide Arabic names for characters and examples." if language == "Arabic" else "Use clear, simple English suitable for UAE students."}
Ensure all content is culturally appropriate for UAE schools.
"""
    
    prompt = f"""You are an expert in creating educational content for AI education in UAE schools.
Create a complete textbook unit for **{grade_context}** following UAE's AI curriculum approach.

{language_instruction}

**UNIT INFORMATION:**
{unit_info}

**CURRICULUM CONTEXT:**
{curriculum_context}

{"**KINDERGARTEN EXAMPLE STRUCTURE:**" + kg_example if level == "Kindergarten" else ""}

**TEXTBOOK UNIT REQUIREMENTS:**

## 1. **Unit Opening**
- Unit title {"(in Arabic)" if language == "Arabic" else "(consider bilingual needs)"}
- Learning objectives (aligned with UAE curriculum)
- Unit overview (age-appropriate introduction)
- Connection to daily life in UAE
- Visual mascot/character introduction {"(with Arabic name)" if language == "Arabic" else ""}

## 2. **Core Content Structure**

{"### For Kindergarten (Ages 4-5):" if level == "Kindergarten" else ""}
{"- Story-based lesson (like 'Robo the Helper')" if level == "Kindergarten" else ""}
{"- Interactive elements (colors, shapes, sounds)" if level == "Kindergarten" else ""}
{"- Simple vocabulary with visuals" if level == "Kindergarten" else ""}
{"- Play-based activities" if level == "Kindergarten" else ""}

{"### For Cycle 1 (Grades 1-4):" if "Cycle 1" in level else ""}
{"- Engaging narratives with AI concepts" if "Cycle 1" in level else ""}
{"- Hands-on activities and experiments" if "Cycle 1" in level else ""}
{"- Digital thinking challenges" if "Cycle 1" in level else ""}
{"- Simple coding concepts (unplugged)" if "Cycle 1" in level else ""}

{"### For Cycle 2 (Grades 5-8):" if "Cycle 2" in level else ""}
{"- Conceptual explanations with examples" if "Cycle 2" in level else ""}
{"- AI system design activities" if "Cycle 2" in level else ""}
{"- Bias and ethics discussions" if "Cycle 2" in level else ""}
{"- Algorithm exploration" if "Cycle 2" in level else ""}

{"### For Cycle 3 (Grades 9-12):" if "Cycle 3" in level else ""}
{"- Advanced concepts and theories" if "Cycle 3" in level else ""}
{"- Real-world case studies" if "Cycle 3" in level else ""}
{"- Command engineering practice" if "Cycle 3" in level else ""}
{"- Career exploration activities" if "Cycle 3" in level else ""}

## 3. **Activities Section**
Include 3-4 activities appropriate for the level:
- Activity objectives
- Materials needed
- Step-by-step instructions
- Expected outcomes
- Extension ideas
- Group vs individual work

## 4. **Technology Integration**
- Digital tools usage (age-appropriate)
- Unplugged alternatives
- Screen time considerations
- Safe online practices

## 5. **Assessment Strategies**
{"- Observation checklists" if level == "Kindergarten" else ""}
{"- Simple skill demonstrations" if "Cycle 1" in level else ""}
{"- Project-based assessments" if "Cycle 2" in level else ""}
{"- Portfolio development" if "Cycle 3" in level else ""}
- Formative assessment ideas
- Self-assessment tools (age-appropriate)

## 6. **Cultural Integration**
- UAE context and examples
- Local AI applications
- Arabic language considerations
- Cultural values integration

## 7. **Home-School Connection**
- Parent/guardian guidance
- Home activities
- Safety guidelines
- Family engagement ideas

## 8. **Additional Resources**
- Vocabulary list {"(Arabic AI terms)" if language == "Arabic" else "(English/Arabic)"}
- Visual aids descriptions
- Songs/rhymes (if applicable) {"in Arabic" if language == "Arabic" else ""}
- Digital resources links
- Teacher notes {"(in Arabic)" if language == "Arabic" else ""}

## 9. **Unit Summary**
- Key concepts recap
- Skills developed
- Connection to next unit
- Celebration of learning

**IMPORTANT GUIDELINES:**
1. Use age-appropriate language
2. Include visual learning elements
3. Ensure cultural sensitivity
4. Balance digital and unplugged activities
5. Focus on ethical AI use
6. Make content engaging and interactive
7. {"Use Arabic throughout the unit" if language == "Arabic" else "Consider bilingual needs"}

Generate a complete textbook unit that aligns with UAE's AI curriculum vision and engages students effectively.

Format in Markdown with clear sections and visual descriptions.
{"Remember: All content must be in Arabic." if language == "Arabic" else ""}
"""
    
    try:
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=65536,
            temperature=0.7,
        )
        
        response = model.generate_content(
            [prompt],
            generation_config=generation_config
        )
        
        if hasattr(response, 'text') and response.text:
            return response.text
        else:
            return "Error: Could not generate textbook unit"
            
    except Exception as e:
        st.error(f"Error generating textbook unit: {e}")
        return None

def create_word_document(content):
    """Creates a Word document from the markdown content"""
    doc = Document()
    
    # Basic markdown to docx conversion
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    return doc

def create_combined_units_document(level, units_dict):
    """Creates a combined Word document with all units for a level"""
    doc = Document()
    
    # Add title page
    doc.add_heading(f"UAE AI Curriculum - {level}", level=1)
    doc.add_paragraph(f"Complete Textbook Units")
    
    # Add language info if present
    if "Arabic" in level:
        doc.add_paragraph("النسخة العربية")
    elif "English" in level:
        doc.add_paragraph("English Version")
        
    doc.add_page_break()
    
    # Add table of contents
    doc.add_heading("Table of Contents" if "English" in level or "English" not in level and "Arabic" not in level else "جدول المحتويات", level=1)
    for i, (unit_name, _) in enumerate(units_dict.items(), 1):
        doc.add_paragraph(f"{i}. {unit_name}", style='List Number')
    doc.add_page_break()
    
    # Add each unit
    for i, (unit_name, unit_content) in enumerate(units_dict.items(), 1):
        # Add unit separator
        if i > 1:
            doc.add_page_break()
        
        # Convert markdown content to docx
        lines = unit_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    return doc

# --- Streamlit App ---
st.set_page_config(page_title="UAE AI Curriculum Generator", layout="wide")

st.title("🇦🇪 UAE AI Curriculum & Textbook Generator 🤖")

st.markdown("""
This tool generates AI curriculum and textbook content aligned with the UAE's AI education framework.
The UAE AI curriculum integrates into existing Computing, Creative Design, and Innovation subjects.
""")

# UAE Curriculum Overview
with st.expander("📋 UAE AI Curriculum Framework Overview"):
    st.markdown("""
    ### 🎯 Seven Core Areas
    1. **Foundational concepts** - Basic AI understanding
    2. **Data and algorithms** - How AI processes information
    3. **Software usage** - Practical AI tools
    4. **Ethical awareness** - Responsible AI use
    5. **Real-world applications** - AI in daily life
    6. **Innovation and project design** - Creating with AI
    7. **Policies and community engagement** - AI in society
    
    ### 📚 Learning Cycles
    
    **🧸 Kindergarten (Ages 4-5)**
    - Introduction through stories and play
    - Simple concepts like "smart machines"
    - Interactive, game-based learning
    
    **📖 Cycle 1 (Grades 1-4)**
    - Understanding machines vs humans
    - Developing digital thinking skills
    - Exploring AI applications
    
    **💻 Cycle 2 (Grades 5-8)**
    - Designing and evaluating AI systems
    - Learning about bias and algorithms
    - Focusing on ethical AI use
    
    **🚀 Cycle 3 (Grades 9-12)**
    - Command engineering
    - Real-world scenarios
    - Career preparation
    
    ### ✨ Key Features
    - Integrated into existing subjects (no extra hours)
    - Focus on ethical and practical skills
    - Comprehensive teacher support materials
    - Age-appropriate progression
    """)

# Create tabs
tab1, tab2, tab3, tab4 = st.tabs(["📋 Generate Curriculum", "📖 Generate Textbook Units", "📥 Download Materials", "💡 Examples"])

with tab1:
    st.header("Generate UAE AI Curriculum")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Language selection
        language = st.radio(
            "Select Language / اختر اللغة:",
            ["English", "Arabic"],
            horizontal=True,
            key="curriculum_language"
        )
        
        selected_level = st.selectbox(
            "Select Education Level:",
            ["Kindergarten", "Cycle 1 (Grades 1-4)", "Cycle 2 (Grades 5-8)", "Cycle 3 (Grades 9-12)"]
        )
        
        # Add grade selection for cycles
        specific_grade = None
        if selected_level != "Kindergarten":
            if "Cycle 1" in selected_level:
                specific_grade = st.selectbox(
                    "Select Specific Grade:",
                    [1, 2, 3, 4],
                    format_func=lambda x: f"Grade {x}",
                    key="grade_selector_cycle1"
                )
            elif "Cycle 2" in selected_level:
                specific_grade = st.selectbox(
                    "Select Specific Grade:",
                    [5, 6, 7, 8],
                    format_func=lambda x: f"Grade {x}",
                    key="grade_selector_cycle2"
                )
            elif "Cycle 3" in selected_level:
                specific_grade = st.selectbox(
                    "Select Specific Grade:",
                    [9, 10, 11, 12],
                    format_func=lambda x: f"Grade {x}",
                    key="grade_selector_cycle3"
                )
        
        # Get cycle info
        if selected_level == "Kindergarten":
            cycle_info = UAE_CURRICULUM_STRUCTURE["Kindergarten"]
        else:
            cycle_key = selected_level.split(" ")[0] + " " + selected_level.split(" ")[1]
            cycle_info = UAE_CURRICULUM_STRUCTURE[cycle_key]
        
        generate_curriculum_btn = st.button("🎯 Generate Curriculum", key="gen_curr_uae")
        
        if generate_curriculum_btn:
            # Create display name for the level
            display_level = f"Grade {specific_grade}" if specific_grade else selected_level
            display_level_with_lang = f"{display_level} ({language})"
            
            with st.spinner(f"Generating UAE AI curriculum for {display_level} in {language}..."):
                curriculum = generate_uae_curriculum(selected_level, cycle_info, specific_grade, language)
                
                if curriculum:
                    # Save to session state with grade-specific and language-specific key
                    if 'uae_curricula' not in st.session_state:
                        st.session_state.uae_curricula = {}
                    st.session_state.uae_curricula[display_level_with_lang] = curriculum
                    
                    # Extract and save units
                    units = extract_units_from_curriculum(curriculum)
                    if 'uae_curriculum_units' not in st.session_state:
                        st.session_state.uae_curriculum_units = {}
                    st.session_state.uae_curriculum_units[display_level_with_lang] = units
                    
                    st.success(f"✅ Curriculum generated successfully for {display_level} in {language}!")
                    if units:
                        st.info(f"📚 Found {len(units)} units in the curriculum")
                    
                    # Display curriculum
                    st.markdown("---")
                    st.markdown(curriculum)
                    
                    # Download button
                    doc = create_word_document(curriculum)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.download_button(
                        label="📥 Download Curriculum as Word Document",
                        data=doc_io,
                        file_name=f"UAE_AI_Curriculum_{display_level.replace(' ', '_').replace('(', '').replace(')', '')}_{language}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    
    with col2:
        st.info("""
        **UAE Curriculum Features:**
        - Seven core AI areas
        - Age-appropriate content
        - Integrated approach
        - Ethical focus
        - Local context
        - Bilingual support (English/Arabic)
        """)
        
        if specific_grade:
            st.success(f"🎯 Selected: Grade {specific_grade}")
            st.caption("Perfect for CBSE, IB, American, and other boards")
        
        if language == "Arabic":
            st.info("📝 سيتم إنشاء المنهج باللغة العربية")

with tab2:
    st.header("Generate AI Textbook Units")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Language selection
        unit_language = st.radio(
            "Select Language / اختر اللغة:",
            ["English", "Arabic"],
            horizontal=True,
            key="unit_language"
        )
        
        textbook_level = st.selectbox(
            "Select Education Level:",
            ["Kindergarten", "Cycle 1 (Grades 1-4)", "Cycle 2 (Grades 5-8)", "Cycle 3 (Grades 9-12)"],
            key="textbook_level"
        )
        
        # Add grade selection for cycles
        textbook_grade = None
        if textbook_level != "Kindergarten":
            if "Cycle 1" in textbook_level:
                textbook_grade = st.selectbox(
                    "Select Specific Grade:",
                    [1, 2, 3, 4],
                    format_func=lambda x: f"Grade {x}",
                    key="textbook_grade_cycle1"
                )
            elif "Cycle 2" in textbook_level:
                textbook_grade = st.selectbox(
                    "Select Specific Grade:",
                    [5, 6, 7, 8],
                    format_func=lambda x: f"Grade {x}",
                    key="textbook_grade_cycle2"
                )
            elif "Cycle 3" in textbook_level:
                textbook_grade = st.selectbox(
                    "Select Specific Grade:",
                    [9, 10, 11, 12],
                    format_func=lambda x: f"Grade {x}",
                    key="textbook_grade_cycle3"
                )
        
        # Create display level name
        display_level = f"Grade {textbook_grade}" if textbook_grade else textbook_level
        display_level_with_lang = f"{display_level} ({unit_language})"
        
        # Check if curriculum and units exist
        curriculum_exists = 'uae_curricula' in st.session_state and display_level_with_lang in st.session_state.uae_curricula
        units_exist = 'uae_curriculum_units' in st.session_state and display_level_with_lang in st.session_state.uae_curriculum_units
        
        if curriculum_exists and units_exist:
            st.success(f"✓ Curriculum found for {display_level} in {unit_language}")
            curriculum_context = st.session_state.uae_curricula[display_level_with_lang]
            available_units = st.session_state.uae_curriculum_units[display_level_with_lang]
            
            if available_units:
                # Create unit selection dropdown
                unit_options = [f"{i+1}. {unit['title']}" for i, unit in enumerate(available_units)]
                selected_unit_index = st.selectbox(
                    "Select Unit to Generate:",
                    range(len(unit_options)),
                    format_func=lambda x: unit_options[x],
                    key="unit_selector"
                )
                
                selected_unit = available_units[selected_unit_index]
                
                # Show unit details
                st.info(f"**Unit Details:**\n{selected_unit['details']}")
                
                # Check if unit already generated
                unit_already_generated = False
                if 'uae_units' in st.session_state and display_level_with_lang in st.session_state.uae_units:
                    if selected_unit['title'] in st.session_state.uae_units[display_level_with_lang]:
                        unit_already_generated = True
                        st.warning("⚠️ This unit has already been generated. Generating again will replace the existing content.")
                
                # Generate button
                generate_unit_btn = st.button("📝 Generate Textbook Unit", key="gen_unit")
                
                if generate_unit_btn:
                    # Prepare unit info
                    unit_info = f"{selected_unit['title']}\n{selected_unit['details']}"
                    
                    with st.spinner(f"Generating textbook unit: {selected_unit['title']} in {unit_language}..."):
                        unit_content = generate_uae_textbook_unit(textbook_level, unit_info, curriculum_context, textbook_grade, unit_language)
                        
                        if unit_content:
                            # Save to session state
                            if 'uae_units' not in st.session_state:
                                st.session_state.uae_units = {}
                            if display_level_with_lang not in st.session_state.uae_units:
                                st.session_state.uae_units[display_level_with_lang] = {}
                            
                            st.session_state.uae_units[display_level_with_lang][selected_unit['title']] = unit_content
                            
                            st.success(f"✅ Textbook unit generated successfully in {unit_language}!")
                            
                            # Display unit
                            st.markdown("---")
                            st.markdown(unit_content)
                            
                            # Download button
                            doc = create_word_document(unit_content)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            
                            st.download_button(
                                label="📥 Download Unit as Word Document",
                                data=doc_io,
                                file_name=f"UAE_AI_{display_level.replace(' ', '_')}_{selected_unit['title'].replace(' ', '_').replace(':', '')}_{unit_language}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                
                # Show generation progress
                st.markdown("---")
                st.subheader("📊 Generation Progress")
                if 'uae_units' in st.session_state and display_level_with_lang in st.session_state.uae_units:
                    generated_units = st.session_state.uae_units[display_level_with_lang]
                    progress = len(generated_units) / len(available_units)
                    st.progress(progress)
                    st.write(f"Generated {len(generated_units)} out of {len(available_units)} units")
                    
                    # List generated units
                    st.write("**Generated Units:**")
                    for unit in available_units:
                        if unit['title'] in generated_units:
                            st.write(f"✅ {unit['title']}")
                        else:
                            st.write(f"⬜ {unit['title']}")
                else:
                    st.progress(0.0)
                    st.write("No units generated yet")
            else:
                st.warning("No units found in the curriculum. Please regenerate the curriculum.")
        else:
            st.warning(f"⚠️ No curriculum found for {display_level} in {unit_language}. Please generate curriculum first in the 'Generate Curriculum' tab.")
    
    with col2:
        st.info("""
        **Unit Features:**
        - Story-based learning
        - Interactive activities
        - Cultural integration
        - Ethical discussions
        - Assessment strategies
        - Parent engagement
        """)
        
        if textbook_grade:
            st.success(f"🎯 Selected: Grade {textbook_grade}")
            st.caption("Customized for specific grade level")
        
        if unit_language == "Arabic":
            st.info("📝 سيتم إنشاء الوحدة باللغة العربية")
        
        # Add manual entry option
        with st.expander("📝 Manual Unit Entry (Optional)"):
            st.markdown("If you want to create a custom unit not in the curriculum:")
            custom_unit_title = st.text_input("Unit Title:")
            custom_unit_details = st.text_area(
                "Unit Details:",
                placeholder="Enter unit concepts and objectives...",
                height=100
            )
            
            if st.button("Generate Custom Unit", key="gen_custom_unit"):
                if custom_unit_title and custom_unit_details:
                    unit_info = f"{custom_unit_title}\n{custom_unit_details}"
                    curriculum_context = st.session_state.uae_curricula.get(display_level_with_lang, "General UAE AI curriculum")
                    
                    with st.spinner(f"Generating custom unit: {custom_unit_title} in {unit_language}..."):
                        unit_content = generate_uae_textbook_unit(textbook_level, unit_info, curriculum_context, textbook_grade, unit_language)
                        
                        if unit_content:
                            # Save to session state
                            if 'uae_units' not in st.session_state:
                                st.session_state.uae_units = {}
                            if display_level_with_lang not in st.session_state.uae_units:
                                st.session_state.uae_units[display_level_with_lang] = {}
                            
                            st.session_state.uae_units[display_level_with_lang][custom_unit_title] = unit_content
                            st.success(f"✅ Custom unit generated in {unit_language}!")
                            st.rerun()

with tab3:
    st.header("Download Generated Materials")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📋 Available Curricula")
        if 'uae_curricula' in st.session_state and st.session_state.uae_curricula:
            for level, curriculum in st.session_state.uae_curricula.items():
                st.write(f"✓ {level}")
                
                doc = create_word_document(curriculum)
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.download_button(
                    label=f"Download {level} Curriculum",
                    data=doc_io,
                    file_name=f"UAE_AI_Curriculum_{level.replace(' ', '_').replace('(', '').replace(')', '')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_curr_{level}"
                )
        else:
            st.info("No curricula generated yet.")
    
    with col2:
        st.subheader("📖 Available Textbook Units")
        if 'uae_units' in st.session_state and st.session_state.uae_units:
            for level, units in st.session_state.uae_units.items():
                st.write(f"**{level}:**")
                
                # Show individual units with download buttons
                with st.expander(f"Individual Units ({len(units)} units)"):
                    for unit_name, content in units.items():
                        col_a, col_b = st.columns([3, 1])
                        with col_a:
                            st.write(f"✓ {unit_name}")
                        with col_b:
                            doc = create_word_document(content)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            
                            st.download_button(
                                label="📥",
                                data=doc_io,
                                file_name=f"UAE_AI_{level.replace(' ', '_')}_{unit_name.replace(' ', '_').replace(':', '')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_unit_{level}_{unit_name}"
                            )
                
                # Combined download button
                if len(units) > 1:
                    st.write("**Download All Units Combined:**")
                    combined_doc = create_combined_units_document(level, units)
                    combined_io = io.BytesIO()
                    combined_doc.save(combined_io)
                    combined_io.seek(0)
                    
                    st.download_button(
                        label=f"📚 Download All {level} Units (Combined)",
                        data=combined_io,
                        file_name=f"UAE_AI_{level.replace(' ', '_').replace('(', '').replace(')', '')}_All_Units_Combined.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_combined_{level}"
                    )
                
                st.markdown("---")
        else:
            st.info("No textbook units generated yet.")
    
    # Add summary section
    st.markdown("---")
    st.subheader("📊 Generation Summary")
    
    summary_col1, summary_col2, summary_col3 = st.columns(3)
    
    with summary_col1:
        curriculum_count = len(st.session_state.get('uae_curricula', {}))
        st.metric("Curricula Generated", curriculum_count)
    
    with summary_col2:
        total_units = sum(len(units) for units in st.session_state.get('uae_units', {}).values())
        st.metric("Total Units Generated", total_units)
    
    with summary_col3:
        levels_with_units = len(st.session_state.get('uae_units', {}))
        st.metric("Levels with Units", levels_with_units)

with tab4:
    st.header("📚 Examples from UAE AI Curriculum")
    
    # Language toggle for examples
    example_lang = st.radio(
        "View Examples in:",
        ["English", "Arabic عربي"],
        horizontal=True,
        key="example_language"
    )
    
    if example_lang == "English":
        st.subheader("🧸 Kindergarten Example: 'Robo the Helper'")
        st.markdown("""
        ### Story Summary
        Robo is a friendly robot that helps Leo and his family at home. He turns on lights, 
        reminds Leo to brush his teeth, plays music, and even helps find Leo's lost shoe!
        
        ### Learning Activities
        1. **Color the AI Helpers** - Children color smart devices
        2. **AI vs Not AI Game** - Identify smart helpers
        3. **"I Am a Smart Machine" Song** - Music and movement
        4. **Make Your Own Robo** - Craft activity
        
        ### Assessment
        - Observation checklist
        - "Junior AI Explorer" badges
        - Portfolio of creative work
        """)
    else:
        st.subheader("🧸 مثال رياض الأطفال: 'روبو المساعد'")
        st.markdown("""
        ### ملخص القصة
        روبو هو روبوت ودود يساعد ليو وعائلته في المنزل. يقوم بتشغيل الأضواء،
        ويذكّر ليو بتنظيف أسنانه، ويشغّل الموسيقى، وحتى يساعد في العثور على حذاء ليو المفقود!
        
        ### الأنشطة التعليمية
        1. **تلوين المساعدين الأذكياء** - يلون الأطفال الأجهزة الذكية
        2. **لعبة الذكاء الاصطناعي أم لا** - تحديد المساعدين الأذكياء
        3. **أغنية "أنا آلة ذكية"** - موسيقى وحركة
        4. **اصنع روبو الخاص بك** - نشاط حرفي
        
        ### التقييم
        - قائمة مراقبة الملاحظة
        - شارات "مستكشف الذكاء الاصطناعي الصغير"
        - ملف الأعمال الإبداعية
        """)
    
    st.subheader("💡 Teaching Strategies by Level / استراتيجيات التدريس حسب المستوى")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if example_lang == "English":
            st.markdown("""
            **Kindergarten & Cycle 1:**
            - Story-based learning
            - Interactive play
            - Visual aids
            - Songs and rhymes
            - Hands-on crafts
            - Simple vocabulary
            """)
        else:
            st.markdown("""
            **رياض الأطفال والحلقة الأولى:**
            - التعلم القائم على القصص
            - اللعب التفاعلي
            - الوسائل البصرية
            - الأغاني والقوافي
            - الحرف اليدوية
            - المفردات البسيطة
            """)
    
    with col2:
        if example_lang == "English":
            st.markdown("""
            **Cycle 2 & 3:**
            - Project-based learning
            - Ethical discussions
            - Real-world applications
            - Digital tools exploration
            - Career connections
            - Innovation challenges
            """)
        else:
            st.markdown("""
            **الحلقة الثانية والثالثة:**
            - التعلم القائم على المشاريع
            - المناقشات الأخلاقية
            - التطبيقات الواقعية
            - استكشاف الأدوات الرقمية
            - الروابط المهنية
            - تحديات الابتكار
            """)
    
    # Add sample AI terms glossary
    st.subheader("📚 Sample AI Terms / مصطلحات الذكاء الاصطناعي")
    glossary_col1, glossary_col2 = st.columns(2)
    
    with glossary_col1:
        st.markdown("""
        **English Terms:**
        - Artificial Intelligence
        - Machine Learning
        - Algorithm
        - Data
        - Robot
        - Smart Device
        - Pattern Recognition
        - Digital Assistant
        """)
    
    with glossary_col2:
        st.markdown("""
        **المصطلحات العربية:**
        - الذكاء الاصطناعي
        - التعلم الآلي
        - الخوارزمية
        - البيانات
        - الروبوت
        - الجهاز الذكي
        - التعرف على الأنماط
        - المساعد الرقمي
        """)

# Sidebar
st.sidebar.header("About UAE AI Curriculum")
st.sidebar.markdown("""
The UAE AI curriculum represents a pioneering educational initiative to prepare students for an AI-driven future.

**Key Principles:**
- 🌍 Culturally relevant
- 🤝 Ethically focused
- 💡 Innovation-oriented
- 📚 Integrated learning
- 🎯 Practical skills
- 🌐 **Bilingual Generation** (English/Arabic)

**Language Support:**
- 🇬🇧 **English**: Clear, simple language suitable for international schools
- 🇦🇪 **Arabic**: Full Arabic content with appropriate terminology
- 📝 All materials can be generated in either language
- 🔄 Create parallel versions for bilingual programs

**Grade-Specific Content:**
- 📖 Generate for individual grades (1-12)
- 🎓 Suitable for different boards:
  - CBSE
  - IB (International Baccalaureate)
  - American Curriculum
  - British Curriculum
  - Other international boards

**Integration:**
AI education is seamlessly integrated into existing subjects:
- Computing
- Creative Design
- Innovation

No additional teaching hours required!
""")

st.sidebar.markdown("---")
st.sidebar.header("🎯 Quick Start Guide")
st.sidebar.markdown("""
1. **Select Language** - Choose English or Arabic
2. **Select Grade Level** - Choose cycle and specific grade
3. **Generate Curriculum** first for your chosen grade
4. **Create Units** using the curriculum as context
5. **Download Materials** for classroom use
6. **Check Examples** for inspiration

**Language Tips:**
- Generate curriculum in your preferred language first
- Units must match the curriculum language
- Create parallel versions by generating in both languages

**Recommended Workflow:**
- Start with Kindergarten for foundational concepts
- Progress through grades systematically
- Customize units for your board's requirements
- Use grade-specific content for differentiation
""")

st.sidebar.markdown("---")
st.sidebar.info("Powered by Google Gemini AI | دعم باللغة العربية")
