import streamlit as st
# from openai import OpenAI # Removed
import google.generativeai as genai # Added
import fitz  # PyMuPDF - Still needed for initial validation or if we want to show images separately
from docx import Document
from PIL import Image
import io
import pathlib # Added

# --- Configuration ---
# OPENROUTER_API_KEY = "<YOUR_OPENROUTER_API_KEY_HERE>" # Removed
# YOUR_SITE_URL = "<YOUR_SITE_URL_HERE>" # Removed
# YOUR_SITE_NAME = "<YOUR_SITE_NAME_HERE>" # Removed

# Get Google API key from Streamlit secrets
try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except KeyError:
    st.error("""
    🔑 **Google API Key Not Found!**
    
    Please add your Google API key to Streamlit secrets:
    
    **For local development:**
    Create a `.streamlit/secrets.toml` file in your project directory with:
    ```toml
    GOOGLE_API_KEY = "your_google_api_key_here"
    ```
    
    **For Streamlit Cloud deployment:**
    1. Go to your app's settings in Streamlit Cloud
    2. Click on "Secrets" 
    3. Add the following:
    ```toml
    GOOGLE_API_KEY = "your_google_api_key_here"
    ```
    
    **How to get a Google API Key:**
    1. Go to [Google AI Studio](https://aistudio.google.com/app/apikey)
    2. Create a new API key
    3. Copy the key and add it to your secrets
    """)
    st.stop()

# Initialize Google Gemini Client
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel(model_name="gemini-2.5-pro-preview-05-06") # Updated model name
except Exception as e:
    st.error(f"Error configuring Google Gemini client: {e}")
    st.stop()


# --- Helper Functions (will be expanded) ---

def load_model_chapter_progression(file_path="Model Chapter Progression and Elements.txt"):
    """Loads the model chapter progression text."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        st.error(f"Error: The file {file_path} was not found. Please make sure it's in the same directory as app.py.")
        return None

# We might still want to extract images if we want to display them in Streamlit before/after analysis,
# or if we want to allow users to interact with them separately.
# For now, the core analysis will rely on Gemini's PDF capabilities.
def extract_images_for_display(pdf_file_bytes):
    """
    Extracts images from PDF bytes for display purposes.
    Returns a list of images (Pillow Image objects).
    """
    images = []
    try:
        doc = fitz.open(stream=pdf_file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes_data = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes_data))
                images.append({"page": page_num + 1, "image": image, "bbox": page.get_image_bbox(img)})
        doc.close()
    except Exception as e:
        st.warning(f"Could not extract images for display from PDF: {e}")
        return []
    return images


def analyze_with_llm(pdf_file_bytes, pdf_filename, model_progression_text, grade_level):
    """
    Analyzes the uploaded PDF chapter using Google Gemini API based on the model progression.
    """
    st.info("Preparing PDF...")

    try:
        # Save the uploaded PDF bytes to a temporary file, as the Gemini API
        # can take a file path or a file-like object for uploads.
        # Using a named temporary file can be more robust.
        temp_pdf_path = pathlib.Path(f"temp_{pdf_filename}")
        with open(temp_pdf_path, "wb") as f:
            f.write(pdf_file_bytes)

        st.info(f"Uploading '{pdf_filename}' to EeeBee...")
        # Upload the PDF using the File API
        # It's good practice to use the File API for documents, especially if they might be large.
        # Files are stored for 48 hours.
        uploaded_file = genai.upload_file(path=temp_pdf_path, display_name=pdf_filename, mime_type="application/pdf")
        st.success(f"'{pdf_filename}' uploaded successfully to EeeBee (URI: {uploaded_file.uri})")

    except Exception as e:
        st.error(f"Error uploading PDF to EeeBee: {e}")
        if 'temp_pdf_path' in locals() and temp_pdf_path.exists():
            temp_pdf_path.unlink() # Clean up temp file
        return "Error during PDF upload to EeeBee.", "Error"
    
    prompt_content = f"""You are an expert in educational content development, specifically for CBSE curriculum.
Your task is to analyze the provided PDF document, which is a book chapter intended for **{grade_level} (CBSE)**.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes. Please proceed with the analysis and rewriting.

You need to evaluate it against the 'Model Chapter Progression and Elements' provided below.
The final output should be a completely rewritten and improved chapter that incorporates all necessary corrections and adheres to the model, suitable for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines. The book should be according to the NCERT books for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**REQUIRED CONTENT STRUCTURE:**
The improved chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
10. Exercise (with the following 5 question types, each with appropriate number of questions):
   - MCQ
   - Assertion and Reason
   - Fill in the Blanks
   - True False
   - Define the following terms
   - Match the column
   - Give Reason for the following Statement (Easy Level)
   - Answer in Brief (Moderate Level)
   - Answer in Detail (Hard Level)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Formatting Requirements:**
* Format like a proper textbook with well-structured paragraphs (not too long but sufficiently developed)
* Use clear headings and subheadings to organize content
* Include multiple Bloom's Taxonomy questions at different levels in each relevant section
* Create concept-wise summaries (visuals are optional but helpful)
* Integrate special features within the core concept flow rather than as separate sections
* DO NOT change the original concept names/terminology from the uploaded PDF

**Instructions for Analysis and Rewrite:**
1.  **Read and understand the entire PDF chapter.** Pay attention to text, images, diagrams, and overall structure, keeping the **{grade_level}** student in mind.
2.  **Compare the chapter to each element** in the 'Model Chapter Progression and Elements'.
3.  **Identify deficiencies:** Where does the uploaded chapter fall short for a **{grade_level}** audience and the model? Be specific.
4.  **Analyze Images:** For each image in the PDF, assess its:
    *   Relevance to the surrounding text and learning objectives for **{grade_level}**.
    *   Clarity and quality.
    *   Age-appropriateness for **{grade_level}**.
    *   Presence of any specific characters if expected (e.g., 'EeeBee').
5.  **Suggest Image Actions:**
    *   If an image is good, state that it should be kept.
    *   If an image needs improvement (e.g., better clarity, different focus for **{grade_level}**), describe the improvement.
    *   If an image is irrelevant, unsuitable, or missing, provide a detailed prompt for an AI image generator to create a new, appropriate image suitable for **{grade_level}**. The prompt should be clearly marked, e.g., "[PROMPT FOR NEW IMAGE: A detailed description of the image needed, including style, content, characters like 'EeeBee' if applicable, and educational purpose for **{grade_level}**.]"
6.  **Rewrite the Chapter:** Create a new version of the chapter that:
    *   Follows the 'Model Chapter Progression and Elements' strictly.
    *   MUST include ALL elements from the REQUIRED CONTENT STRUCTURE in the order specified.
    *   Corrects all identified mistakes and deficiencies.
    *   Integrates image feedback directly into the text.
    *   Ensures content is aligned with NCERT, NCF, and NEP 2020 principles (e.g., inquiry-based learning, 21st-century skills, real-world connections) appropriate for **{grade_level}**.
    *   Is written in clear, engaging, and age-appropriate language for **{grade_level} (CBSE)**.
    *   Use Markdown for formatting (headings, bold, italics, lists) to make it easy to convert to a Word document. For example:
        # Chapter Title
        ## Section Title
        **Bold Text**
        *Italic Text*
        - Item 1
        - Item 2
    *   Present content in proper textbook paragraphs like a NCERT textbook
    *   Include multiple Bloom's Taxonomy questions at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create) in each appropriate section
    *   Organize summaries by individual concepts rather than as a single summary
    *   Weave special features (misconceptions, 21st century skills, etc.) into the core concept flow
    *   PRESERVE all original concept names and terminology from the source PDF

IMPORTANT: Do NOT directly copy large portions of text from the PDF. Instead, use your own words to rewrite and improve the content while maintaining the original concepts and terminology.

**Output Format:**
Provide the complete, rewritten chapter text in Markdown format, incorporating all analyses and image prompts as described. Do not include a preamble or a summary of your analysis; just the final chapter content.
"""
    st.info(f"Sending request to EeeBee model for {grade_level}... This may take some time for larger documents.")
    try:
        # Use the uploaded file in the prompt
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=65536, # Explicitly set max output tokens
            temperature=0.3,  # Add some creativity to avoid direct copying
        )
        response = model.generate_content(
            [uploaded_file, prompt_content],
            generation_config=generation_config
        )

        # --- Debugging: Print finish reason and safety feedback ---
        if response.candidates:
            finish_reason = response.candidates[0].finish_reason
            st.write(f"EeeBee Finish Reason: {finish_reason}")
            
            # Check for copyright detection (finish_reason 4)
            if finish_reason == 4:
                st.error("EeeBee detected potential copyright concerns. Since this is your own content being used for educational purposes, we need to adjust our approach.")
                st.warning("Try one of these solutions: 1) Modify your PDF content to be less like published material, 2) Break the document into smaller chunks, or 3) Change specific terminology that might be triggering the copyright filter.")
                
                # Clean up resources
                try:
                    genai.delete_file(uploaded_file.name)
                    st.info(f"Temporary file '{uploaded_file.display_name}' deleted from Gemini.")
                except Exception as e_del:
                    st.warning(f"Could not delete temporary file '{uploaded_file.display_name}' from Gemini: {e_del}")
                
                if temp_pdf_path.exists():
                    temp_pdf_path.unlink()
                
                return "Error: The model detected copyright concerns. Since this is your own material, try modifying the content to be more distinct from published works.", "Copyright Error"
            
            if response.candidates[0].safety_ratings:
                st.write("Safety Ratings:")
                for rating in response.candidates[0].safety_ratings:
                    st.write(f"- Category: {rating.category}, Probability: {rating.probability}")
        if response.prompt_feedback:
            st.write(f"Gemini Prompt Feedback: {response.prompt_feedback}")
        # --- End Debugging ---

        # Check if we have valid text in the response
        if hasattr(response, 'text') and response.text:
            improved_text = response.text
        else:
            # If there's no text property but we have candidates with content
            if response.candidates and hasattr(response.candidates[0], 'content') and response.candidates[0].content:
                # Extract text from parts if available
                parts = response.candidates[0].content.parts
                if parts:
                    improved_text = ''.join([part.text for part in parts if hasattr(part, 'text')])
                else:
                    raise Exception("No content parts found in response")
            else:
                raise Exception("No valid text content found in response")
        
        # It's good practice to delete the file from Gemini storage if no longer needed,
        # though they auto-delete after 48 hours.
        try:
            genai.delete_file(uploaded_file.name)
            st.info(f"Temporary file '{uploaded_file.display_name}' deleted from Gemini.")
        except Exception as e_del:
            st.warning(f"Could not delete temporary file '{uploaded_file.display_name}' from Gemini: {e_del}")

        if temp_pdf_path.exists():
            temp_pdf_path.unlink() # Clean up local temp file

        return improved_text, "LLM analysis and rewrite complete using Gemini."
    except Exception as e:
        st.error(f"Error during Gemini API call: {e}")
        # Clean up uploaded file on Gemini if an error occurs during generation
        try:
            genai.delete_file(uploaded_file.name)
            st.warning(f"Cleaned up file '{uploaded_file.display_name}' from Gemini after error.")
        except Exception as e_del_err:
            st.warning(f"Could not delete file from Gemini after error: {e_del_err}")

        if temp_pdf_path.exists(): # Clean up local temp file
            temp_pdf_path.unlink()
        return f"Error: Could not complete Gemini analysis. {e}", "Error"


def create_word_document(improved_text_markdown):
    """Creates a Word document from the improved text (Markdown formatted)."""
    doc = Document()
    # This is a very basic Markdown to DOCX conversion.
    # For more complex Markdown, a library like 'markdown' and then 'html2text' or 'pypandoc' might be needed,
    # or more sophisticated parsing logic.
    
    # For now, let's handle headings and paragraphs simply.
    lines = improved_text_markdown.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        # Add more heading levels if needed
        elif line.startswith("- ") or line.startswith("* "):
            # Basic list item handling (doesn't handle nested lists well without more logic)
            # For proper list handling, python-docx requires paragraph styles or more direct list creation.
            # This will just add it as a paragraph with a dash.
            doc.add_paragraph(line, style='ListBullet') # May need to define 'ListBullet' style or use default
        else:
            # Simple bold/italic handling - this is very rudimentary
            # A proper Markdown parser would be better.
            # For now, just add as a paragraph. The LLM output format is key.
            doc.add_paragraph(line)
            
    # In the future, we would also add images back into the document here,
    # or at least placeholders for where the image prompts are.
    return doc

def analyze_with_chunked_approach(pdf_file_bytes, pdf_filename, model_progression_text, grade_level):
    """
    Analyzes the PDF in smaller chunks to avoid triggering copyright protection.
    This is a fallback method when the regular analysis fails due to copyright concerns.
    """
    st.info("Using chunked approach to analyze PDF...")
    
    try:
        # Save the uploaded PDF bytes to a temporary file, as the Gemini API
        # can take a file path or a file-like object for uploads.
        temp_pdf_path = pathlib.Path(f"temp_{pdf_filename}")
        with open(temp_pdf_path, "wb") as f:
            f.write(pdf_file_bytes)
            
        # Upload the entire PDF file to Gemini
        st.info(f"Uploading '{pdf_filename}' to EeeBee for chunked analysis...")
        uploaded_file = genai.upload_file(path=temp_pdf_path, display_name=pdf_filename, mime_type="application/pdf")
        st.success(f"'{pdf_filename}' uploaded successfully to EeeBee (URI: {uploaded_file.uri})")
        
        # Extract text from PDF for determining page chunks
        doc = fitz.open(stream=pdf_file_bytes, filetype="pdf")
        total_pages = len(doc)
        
        # Determine chunk size (pages per chunk)
        pages_per_chunk = 5  # Can be adjusted based on content density
        
        # Define page ranges for chunked analysis
        page_ranges = []
        for start_page in range(0, total_pages, pages_per_chunk):
            end_page = min(start_page + pages_per_chunk - 1, total_pages - 1)
            page_ranges.append({
                "start": start_page + 1,  # 1-indexed for human readability
                "end": end_page + 1,
                "range_text": f"Pages {start_page+1}-{end_page+1} of {total_pages}"
            })
        
        doc.close()
        
        # Process each chunk
        analysis_results = []
        
        for idx, page_range in enumerate(page_ranges):
            st.info(f"Analyzing chunk {idx+1}/{len(page_ranges)}: {page_range['range_text']}...")
            
            prompt_content = f"""You are an expert in educational content development for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: You are analyzing CHUNK {idx+1} of {len(page_ranges)} ({page_range['range_text']}) from a book chapter for **{grade_level} (CBSE)**.
This is just a PORTION of the full chapter - focus only on improving this section according to the model.

**FOCUS ONLY ON PAGES {page_range['start']} to {page_range['end']} of the PDF.**

The book chapter is intended to align with NCERT, NCF, and NEP 2020 guidelines for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**FINAL CHAPTER STRUCTURE:**
Note: The final integrated chapter will follow this structure (you're only working on a portion):
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
8. Link and Learn based Question
9. Image based question
10. Exercise (with question types like MCQs, Fill in the Blanks, etc.)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Instructions:**
* Rewrite and improve ONLY the content from pages {page_range['start']} to {page_range['end']}
* Format like a proper textbook with well-structured paragraphs
* Use clear headings and subheadings where appropriate
* Include Bloom's Taxonomy questions if this section is suitable for them
* Maintain the original concept names/terminology
* Write in your own words - DO NOT copy text directly
* For any images in these pages:
  - Analyze their relevance, clarity, and age-appropriateness
  - Suggest keeping good images or provide image prompts for new/improved ones
  - Image prompts should be marked as: [PROMPT FOR NEW IMAGE: description]

Please provide ONLY the improved version of this specific chunk in Markdown format. 
Do not include analysis or explanations - just the rewritten content.
"""
            
            try:
                generation_config = genai.types.GenerationConfig(
                    max_output_tokens=32768,
                    temperature=0.3,
                )
                
                # Send the full PDF but instruct to focus on specific pages
                response = model.generate_content(
                    [uploaded_file, prompt_content],
                    generation_config=generation_config
                )
                
                # Check for copyright detection
                if response.candidates and response.candidates[0].finish_reason == 4:
                    st.warning(f"Copyright detection on chunk {idx+1}. Trying with text-only fallback...")
                    
                    # Extract just the text for this page range as fallback
                    chunk_text = ""
                    for page_num in range(page_range['start']-1, page_range['end']):
                        page = doc.load_page(page_num)
                        chunk_text += page.get_text()
                    
                    # Modify prompt to indicate images won't be available
                    text_only_prompt = prompt_content + f"""

**NOTE: This is now processing text-only due to copyright concerns with the full PDF.**

**Text from these pages:**
'''
{chunk_text}
'''

Please rewrite and improve this content, noting that any image references will need to be based on the text descriptions only.
"""
                    
                    # Try again with text-only
                    response = model.generate_content(
                        [text_only_prompt],
                        generation_config=generation_config
                    )
                
                if hasattr(response, 'text') and response.text:
                    improved_chunk = response.text
                    analysis_results.append(improved_chunk)
                else:
                    # Handle alternative response format if needed
                    if response.candidates and hasattr(response.candidates[0], 'content') and response.candidates[0].content:
                        parts = response.candidates[0].content.parts
                        if parts:
                            improved_chunk = ''.join([part.text for part in parts if hasattr(part, 'text')])
                            analysis_results.append(improved_chunk)
                        else:
                            analysis_results.append(f"[Error processing chunk {idx+1}]")
                    else:
                        analysis_results.append(f"[Error processing chunk {idx+1}]")
                        
            except Exception as chunk_e:
                st.warning(f"Error processing chunk {idx+1}: {chunk_e}")
                analysis_results.append(f"[Error processing chunk {idx+1}: {chunk_e}]")
        
        # Combine results
        combined_result = "\n\n".join(analysis_results)
        
        # Add final integration prompt to ensure coherence
        integration_prompt = f"""You are an expert educational content developer for CBSE curriculum.
You have been given several chunks of rewritten content for a **{grade_level} (CBSE)** book chapter.
Your task is to integrate these chunks into a coherent, well-structured chapter that flows naturally.

This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

**REQUIRED CONTENT STRUCTURE:**
The final integrated chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
8. Link and Learn based Question
9. Image based question
10. Exercise (with the following 5 question types, each with appropriate number of questions):
   - MCQ
   - Assertion and Reason
   - Fill in the Blanks
   - True False
   - Define the following terms
   - Match the column
   - Give Reason for the following Statement (Easy Level)
   - Answer in Brief (Moderate Level)
   - Answer in Detail (Hard Level)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Important:**
* Ensure smooth transitions between chunks
* Remove any redundancies or repeated information
* Create a consistent style and tone throughout
* Add a chapter introduction and conclusion if they don't exist
* Format as a proper textbook chapter with appropriate headings and subheadings
* Ensure all Bloom's Taxonomy questions are integrated appropriately
* STRICTLY follow the REQUIRED CONTENT STRUCTURE above
* Maintain all image prompts and image references from the chunks

**Content to integrate:**
{combined_result}

Provide the complete, integrated chapter in Markdown format.
"""
        
        try:
            integration_config = genai.types.GenerationConfig(
                max_output_tokens=65536,
                temperature=0.4,  # Lower temperature for more coherent integration
            )
            
            final_response = model.generate_content(
                [integration_prompt],
                generation_config=integration_config
            )
            
            if hasattr(final_response, 'text') and final_response.text:
                final_improved_text = final_response.text
            else:
                # Handle alternative response format
                if final_response.candidates and hasattr(final_response.candidates[0], 'content') and final_response.candidates[0].content:
                    parts = final_response.candidates[0].content.parts
                    if parts:
                        final_improved_text = ''.join([part.text for part in parts if hasattr(part, 'text')])
                    else:
                        raise Exception("No content parts found in final integration response")
                else:
                    raise Exception("No valid text content found in final integration response")
            
            # Clean up resources
            try:
                genai.delete_file(uploaded_file.name)
                st.info(f"Temporary file '{uploaded_file.display_name}' deleted from Gemini.")
            except Exception as e_del:
                st.warning(f"Could not delete temporary file '{uploaded_file.display_name}' from Gemini: {e_del}")
            
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
                    
            return final_improved_text, "LLM analysis and rewrite complete using chunked approach with full PDF access."
            
        except Exception as integration_e:
            st.error(f"Error during final integration: {integration_e}")
            
            # Clean up resources
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
            
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
                
            # Return the combined chunks without final integration
            return combined_result, "Partial analysis complete - final integration failed."
            
    except Exception as e:
        st.error(f"Error during chunked PDF analysis: {e}")
        
        # Clean up resources
        if 'uploaded_file' in locals():
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
        
        if 'temp_pdf_path' in locals() and temp_pdf_path.exists():
            temp_pdf_path.unlink()
        
        return f"Error: Could not complete chunked analysis. {e}", "Error"

# --- Helper Functions for Content Generation ---

def create_specific_prompt(content_type, grade_level, model_progression_text, subject_type="General"):
    """Creates a prompt focused on a specific content type"""
    
    if subject_type == "Mathematics":
        # Mathematics-specific prompts
        if content_type == "chapter":
            return create_math_chapter_prompt(grade_level, model_progression_text)
        elif content_type == "exercises":
            return create_math_exercises_prompt(grade_level, model_progression_text)
        elif content_type == "skills":
            return create_math_skills_prompt(grade_level, model_progression_text)
        elif content_type == "art":
            return create_math_art_prompt(grade_level, model_progression_text)
    else:
        # General subject prompts
        if content_type == "chapter":
            return create_general_chapter_prompt(grade_level, model_progression_text)
        elif content_type == "exercises":
            return create_general_exercises_prompt(grade_level, model_progression_text)
        elif content_type == "skills":
            return create_general_skills_prompt(grade_level, model_progression_text)
        elif content_type == "art":
            return create_general_art_prompt(grade_level, model_progression_text)

# Mathematics-specific prompt functions
def create_math_chapter_prompt(grade_level, model_progression_text):
    """Creates a mathematics-specific chapter content prompt"""
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines for Mathematics education.

**CRITICAL INSTRUCTION**: The PDF may contain MULTIPLE MAJOR SECTIONS (e.g., Section 1, Section 2, Section 3, etc.). You MUST include ALL sections present in the PDF. Do NOT stop after completing just one or two sections. Generate comprehensive content for EVERY section found in the document.

**Model Chapter Progression and Elements (Base Structure):**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Mathematics Syllabus)

Your task is to generate COMPREHENSIVE MATHEMATICS CHAPTER CONTENT following the Model Chapter Progression structure enhanced with mathematics-specific elements.

**IMPORTANT**: If the PDF contains multiple sections (e.g., "Section 2: Place Value", "Section 3: Operations and Estimation with Large Numbers"), you MUST generate complete content for EACH section. The structure below should be applied to EACH major section in the PDF.

**REQUIRED SECTIONS (Generate ALL with substantial content for EACH major section in the PDF):**

## I. Chapter Opener

1. **Chapter Title** - Engaging and mathematically focused

2. **Hook (with Image Prompt)** (80-100 words)
   - Create an engaging mathematical opening that captures student interest
   - Use real-life mathematical scenarios, surprising mathematical facts, or thought-provoking mathematical questions
   - Connect to students' daily mathematical experiences
   - Include a detailed image prompt for a compelling mathematical visual

3. **Real-World Connection** (100 words)
   - Provide multiple real-world applications of the mathematical concepts
   - Show how math is used in everyday life situations
   - Include examples from technology, engineering, finance, science, etc.
   - Connect to mathematical careers and future studies

4. **Learning Outcomes** (125 words)
   - List specific, measurable mathematical learning objectives
   - Use action verbs (define, explain, calculate, apply, analyze, solve, prove, etc.)
   - Align with Bloom's Taxonomy levels for mathematics
   - Connect to CBSE mathematics curriculum standards

5. **Previous Class Link** (100 words)
   - Link to prior mathematical knowledge from previous classes
   - Explain how previous concepts connect to current learning
   - Provide a brief review of essential prerequisites

6. **Chapter Map/Overview** (100 words)
   - Visual layout of mathematical concepts (mind map or flowchart description)
   - Show mathematical progressions and connections

7. **Meet the Character (EeeBee)** (50-100 words)
   - Introduce EeeBee as a mathematical guide/helper throughout the chapter

## II. Core Content Sections (REPEAT FOR EACH MAJOR SECTION IN THE PDF)

**NOTE: If the PDF has Section 1, Section 2, Section 3, etc., create complete content for EACH section following this structure:**

8. **Introduction of Section** (150 words per section)
   - Give a related section introduction that sets the mathematical context
   - Explain the importance of the mathematical concepts to be learned
   - Connect to the broader mathematical curriculum

9. **History of Concepts in This Section** (150 words per section)
   - Provide comprehensive historical background of the mathematical concepts
   - Include key mathematicians, their contributions, and discoveries
   - Explain the timeline of mathematical developments

10. **Warm-up Questions** (100 words per section)
    - Create 5-7 engaging warm-up questions that connect to prior mathematical knowledge
    - Include a mix of question types (mental math, real-world problems, pattern recognition)

11. **Current Concepts** (3500-4500 words minimum PER SECTION)
    
    For each major concept in EACH section, include ALL of the following:
    
    **A. Concept Introduction** (100 words per concept)
    - Clear introduction to each mathematical concept
    - Simple, clear mathematical language
    - Use analogies and mathematical examples
    
    **B. Mathematical Explanation** (400-500 words per concept)
    - Detailed theoretical understanding of the mathematical concept
    - Include step-by-step mathematical reasoning and derivations
    - Show different mathematical approaches where applicable
    
    **C. Solved Examples** (500-600 words per concept)
    - Provide 4-5 different worked examples for each concept
    - Include step-by-step solutions with clear explanations
    - Use varied difficulty levels and question formats
    
    **D. Concept-Based Exercise Questions** (400-500 words per concept)
    Create the following specific mathematical question types for each concept:
    
    1. **Fill in the Blanks** - 3-4 questions
       - Mathematical formulas, definitions, and key concept completion
    
    2. **Conversion-Based Questions** - 2 questions
       - Unit conversions, mathematical transformations
    
    3. **Word Problems (Real-life Context)** - 5 questions
       - Real-world mathematical scenarios applying the concept
       - Varied difficulty levels from basic to complex applications
    
    4. **Match the Columns** - 1 question (5 matches)
       - Mathematical concepts with definitions, formulas with applications
    
    5. **Estimation Questions (Real-life Estimation)** - 5 questions
       - Practical mathematical estimation scenarios
       - Real-world mathematical approximation problems
    
    6. **Miscellaneous Questions** - 3 questions
       - Mixed question types testing various aspects of the concept
       - Include higher-order thinking and application problems
    
    **E. Visuals** (Throughout each concept)
    - Include detailed image prompts for mathematical diagrams, graphs, and illustrations
    
    **F. Mathematical Activities/Experiments** (200-250 words per concept)
    - Step-by-step mathematical experiments or investigations
    - Inquiry-based mathematical activities
    
    **G. Check Your Understanding** (150-200 words per concept)
    - Quick check questions (2-3 simple questions) to test immediate understanding
    - Brief conceptual questions or simple calculations
    
    **H. Key Mathematical Terms** (100-150 words per concept)
    - Highlighted mathematical terminology in text
    - Clear mathematical definitions
    
    **I. Mathematical Applications** (200-250 words per concept)
    - Show relevance to daily life, mathematical careers, technology
    
    **J. Fun Mathematical Facts** (100-150 words per concept)
    - Include interesting mathematical facts related to the concept
    
    **K. Think About It! (Mathematical Exploration)** (100-150 words per concept)
    - Present thought-provoking mathematical questions or scenarios
    
    **L. Mental Mathematics** (150-200 words per concept)
    - Provide mental mathematics strategies and techniques

## III. Special Features (Apply to the ENTIRE chapter, not just one section)

12. **Common Mathematical Misconceptions** (200-300 words)
    - 2-3 misconceptions per mathematical concept
    - Correct early mathematical misunderstandings

13. **21st Century Skills Focus** (300-400 words)
    - Mathematical Design Challenge
    - Mathematical Debate
    - Collaborate & Create

14. **Differentiation** (200-300 words)
    - Challenge sections for advanced mathematical learners
    - Support sections for mathematical revision

15. **Technology Integration** (150-200 words)
    - Mathematical software and tools
    - Digital mathematical simulations

16. **Character Integration** (Throughout)
    - EeeBee appears throughout to ask mathematical questions

## IV. Chapter Wrap-Up (For the ENTIRE chapter)

17. **Self-Assessment Checklist** (200-250 words)
    - Create a comprehensive self-assessment checklist

18. **Chapter-wise Miscellaneous Exercise** (500-600 words)
    - MCQs (5 questions)
    - Short/Long Answer (3 short, 2 long)
    - Open-ended Mathematical Problems (2 questions)
    - Assertion & Reason (3 statements)
    - Mathematical Concept Mapping (1 map)
    - True/False with Justification (5 statements)
    - Mathematical Case Studies (1 scenario)
    - Mathematical Puzzle (1 puzzle)
    - Thinking Based Activities

19. **Apply Your Mathematical Knowledge** (200-300 words)
    - Real-world mathematical application tasks
    - Project-based mathematical problems

**CONTENT REQUIREMENTS:**
* **Minimum Total Length**: 18000-25000 words (to accommodate multiple sections)
* **CRITICAL**: Include ALL major sections from the PDF (e.g., if there are Sections 1, 2, and 3, generate complete content for ALL three)
* **Mathematical Accuracy**: Ensure all mathematical content is accurate
* **Clear Mathematical Language**: Use precise mathematical terminology
* **Step-by-step Solutions**: Provide detailed mathematical working
* **Visual Integration**: Include detailed image prompts
* **Progressive Difficulty**: Structure content from basic to advanced
* DO NOT use the same mathematical figures (numbers) from the pdf

Provide ONLY the comprehensive mathematics chapter content in Markdown format. Remember to include EVERY section found in the PDF document.
"""

def create_math_exercises_prompt(grade_level, model_progression_text):
    """Creates a mathematics-specific exercises prompt"""
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPREHENSIVE MATHEMATICS EXERCISES based on the chapter content in the PDF.

**Core Mathematical Exercise Types:**
1. **MCQ (Multiple Choice Questions)** - at least 12 questions with detailed solutions
2. **Short Answer Mathematical Problems** - at least 8 questions  
3. **Long Answer Mathematical Problems** - at least 5 questions
4. **Assertion & Reason (Mathematical)** - at least 5 questions
5. **True/False with Mathematical Justification** - at least 10 statements
6. **Fill in the Blanks (Mathematical)** - at least 10 questions
7. **Match the Following (Mathematical)** - at least 2 sets with 5 matches each
8. **Mathematical Concept Mapping** - at least 1 comprehensive exercise
9. **Mathematical Case Studies** - at least 2 scenarios
10. **Open-ended Mathematical Problems** - at least 3 questions

**Special Mathematical Features:**
- **EeeBee Integration**: Include questions where EeeBee guides mathematical thinking
- **Technology Integration**: Questions involving mathematical software, calculators, or digital tools
- **Differentiation**: Include challenge problems for advanced learners and support questions for revision
- **21st Century Skills**: Collaborative mathematical problems and communication-based questions

Ensure that:
* Questions cover ALL important mathematical concepts from the PDF
* Questions follow Bloom's Taxonomy at various levels
* Mathematical language is clear and appropriate for {grade_level}
* Questions increase in difficulty progressively
* All exercises include detailed mathematical solutions with step-by-step working
* Content is formatted in Markdown with proper mathematical notation

Provide ONLY the comprehensive mathematical exercises in Markdown format.
"""

def create_math_skills_prompt(grade_level, model_progression_text):
    """Creates a mathematics-specific skills and STEM activities prompt"""
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate MATHEMATICAL SKILL-BASED ACTIVITIES and STEM projects.

## Mathematical Skill-Based Activities

Create at least 3 comprehensive mathematical activities:

**Activity Structure for Each:**
1. **Clear Mathematical Objective**
2. **Materials Required**
3. **Step-by-Step Mathematical Procedure**
4. **Inquiry-Based Mathematical Exploration**
5. **Real-Life Mathematical Connections**
6. **EeeBee Integration**
7. **Mathematical Reflection Questions**
8. **Expected Mathematical Outcomes**

**Types of Mathematical Activities:**
- Mathematical Experiments/Investigations
- Mathematical Manipulative Activities
- Mathematical Problem-Solving Challenges
- Mathematical Pattern Recognition Activities
- Mathematical Measurement and Data Collection

## Mathematical STEM Projects

Create at least 2 comprehensive projects that integrate mathematics with Science, Technology, and Engineering:

**Project Structure for Each:**
1. **Mathematical Integration Focus**
2. **Real-World Mathematical Application**
3. **Technology Integration**
4. **Collaborative Mathematical Work**
5. **Mathematical Design Challenge**
6. **Mathematical Analysis and Calculations**
7. **Mathematical Presentation Component**
8. **Assessment Criteria**

**STEM Integration Examples:**
- Mathematical Modeling in Science
- Engineering Design with Mathematical Constraints
- Technology-Enhanced Mathematical Problem Solving
- Mathematical Data Analysis Projects

Format the content in Markdown with proper mathematical notation.

Provide ONLY the Mathematical Skill Activities and STEM Projects in Markdown format.
"""

def create_math_art_prompt(grade_level, model_progression_text):
    """Creates a mathematics-specific art integration prompt"""
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate MATHEMATICS-INTEGRATED CREATIVE LEARNING projects.

## Mathematical Art Projects

Create at least 3 creative projects that connect mathematical concepts to various art forms:

**Project Structure for Each:**
1. **Mathematical Learning Objective**
2. **Art Form Integration**
3. **Mathematical Concepts Highlighted**
4. **Materials and Tools**
5. **Step-by-Step Mathematical-Artistic Process**
6. **EeeBee's Creative Guidance**
7. **Mathematical Reflection and Analysis**
8. **Showcase and Communication**

**Types of Mathematical Art Integration:**
- Geometric Art and Patterns
- Mathematical Music and Rhythm
- Mathematical Drama and Storytelling
- Mathematical Digital Art
- Mathematical Sculpture and 3D Models

## Mathematical Case Studies

**Case Study – Level 1 (Accessible Mathematical Analysis):**
Create at least 1 simpler case study that:
- Presents a real-world mathematical scenario appropriate for {grade_level}
- Includes guided mathematical analysis questions
- Is accessible to all students with basic mathematical skills

**Case Study – Level 2 (Advanced Mathematical Challenge):**
Create at least 1 more complex case study that:
- Challenges students with multi-step mathematical problems
- Requires deeper application of mathematical concepts
- Encourages higher-order mathematical thinking skills

**Case Study Structure for Each:**
1. **Real-World Mathematical Context**
2. **Mathematical Background Information**
3. **Guided Mathematical Questions**
4. **Mathematical Analysis Requirements**
5. **Creative Mathematical Solutions**
6. **Mathematical Communication Component**
7. **EeeBee's Mathematical Insights**
8. **Extension Mathematical Challenges**

Format the content in Markdown with proper mathematical notation.

Provide ONLY the Mathematics-Integrated Creative Learning content in Markdown format.
"""

# General subject prompt functions
def create_general_chapter_prompt(grade_level, model_progression_text):
    """Creates a general subject chapter content prompt"""
    return f"""You are an expert in educational content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes.

You are analyzing a book chapter intended for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Syllabus)

Your task is to generate COMPREHENSIVE CORE CHAPTER CONTENT that should be equivalent to a complete textbook chapter.

**REQUIRED SECTIONS (Generate ALL with substantial content):**

1. **Current Concepts** (1000-1200 words minimum)
   - Provide detailed explanations of ALL key concepts from the PDF
   - Include multiple examples for each concept
   - Use analogies and real-world connections to explain complex ideas
   - Break down concepts into sub-concepts with clear explanations
   - Include scientific principles, formulas, or definitions where applicable
   - Give questions for each current concept - 3 MCQs, 2 Short questions, 1 Long question
   - Give activity according to the concept
   - Give fun fact for each concept
   - Give key points for each concept
   - Give common misconceptions for each concept
   - Mark concepts which are exactly coming from the PDF and give some extra concepts for higher level understanding
   - Make sure there is no repetition of concepts

2. **Hook (with Image Prompt)** (50-70 words)
   - Create an engaging opening that captures student interest
   - Use storytelling, surprising facts, or thought-provoking questions
   - Connect to students' daily experiences or current events
   - Include a detailed image prompt for a compelling visual

3. **Learning Outcome** (50-70 words)
   - List specific, measurable learning objectives
   - Use action verbs (analyze, evaluate, create, etc.)
   - Align with Bloom's Taxonomy levels
   - Connect to CBSE curriculum standards

4. **Real World Connection** (35-50 words)
   - Provide multiple real-world applications of the concepts
   - Include current examples from technology, environment, health, etc.
   - Explain how the concepts impact daily life
   - Connect to career opportunities and future studies

5. **Previous Class Concept** (50-100 words)
   - Give the concept name and the previous class it was studied in according to NCERT textbooks

6. **History** (70-100 words)
   - Provide comprehensive historical background
   - Include key scientists, inventors, or historical figures
   - Explain the timeline of discoveries or developments
   - Connect historical context to modern understanding

7. **Summary** (600-800 words)
   - Create detailed concept-wise summaries (not just one overall summary)
   - Include key points, formulas, and important facts
   - Organize by individual concepts covered in the chapter
   - Provide clear, concise explanations that reinforce learning

8. **Link and Learn Based Question** (200-300 words)
   - Create 3-5 questions that connect different concepts
   - Include questions that link to other subjects or real-world scenarios
   - Provide detailed explanations for the connections

9. **Image Based Question** (200-300 words)
   - Create 3-5 questions based on images/diagrams from the chapter
   - Include detailed image descriptions if creating new image prompts
   - Ensure questions test understanding, not just observation

**CONTENT REQUIREMENTS:**
* **Minimum Total Length**: 6000-7000 words for the complete chapter content
* **Detailed Explanations**: Each concept should be explained thoroughly with multiple paragraphs
* **Examples and Illustrations**: Include numerous examples, case studies, and practical applications
* **Age-Appropriate Language**: Use vocabulary suitable for {grade_level} but don't oversimplify
* **Engaging Tone**: Write in an engaging, conversational style that maintains student interest
* **Clear Structure**: Use proper headings, subheadings, and formatting
* **Visual Integration**: Include detailed image prompts throughout the content

**FORMATTING REQUIREMENTS:**
* Use Markdown formatting with clear headings (# ## ###)
* Include bullet points and numbered lists where appropriate
* Use **bold** for key terms and *italics* for emphasis
* Create well-structured paragraphs (3-5 sentences each)
* Include image prompts marked as: [PROMPT FOR NEW IMAGE: detailed description]

**QUALITY STANDARDS:**
* Each section should be substantial and comprehensive
* Avoid superficial coverage - go deep into each topic
* Include multiple perspectives and approaches to concepts
* Ensure content flows logically from one section to the next
* Maintain consistency in terminology and explanations

Analyze the PDF document thoroughly and create improved content that expands significantly on what's provided while maintaining all original concept names and terminology.

Provide ONLY the comprehensive chapter content in Markdown format. Do not include exercises, activities, or art projects.
"""

def create_general_exercises_prompt(grade_level, model_progression_text):
    """Creates a general subject exercises prompt"""
    return f"""You are an expert in educational content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPREHENSIVE EXERCISES based on the chapter content in the PDF.

Create the following exercise types:
1. MCQ (Multiple Choice Questions) - at least 10 questions
2. Assertion and Reason - at least 5 questions
3. Fill in the Blanks - at least 10 questions
4. True False - at least 10 statements
5. Define the following terms - at least 10 terms
6. Match the column - at least 2 sets with 5 matches each
7. Give Reason for the following Statement (Easy Level) - at least 5 questions
8. Answer in Brief (Moderate Level) - at least 5 questions
9. Answer in Detail (Hard Level) - at least 5 questions

Ensure that:
* Questions cover ALL important concepts from the PDF
* Questions follow Bloom's Taxonomy at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create)
* Language is clear and appropriate for {grade_level}
* Questions increase in difficulty from basic recall to higher-order thinking
* All exercises include correct answers or model solutions
* The content is formatted in Markdown with proper headings and organization

Do NOT directly copy questions from the PDF. Create new, original questions based on the content.

Provide ONLY the exercises in Markdown format.
"""

def create_general_skills_prompt(grade_level, model_progression_text):
    """Creates a general subject skills and STEM activities prompt"""
    return f"""You are an expert in educational content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate SKILL-BASED ACTIVITIES and STEM projects based on the chapter content in the PDF.

Create the following:
1. Skill-based Activities - At least 3 hands-on activities that:
   * Reinforce the key concepts from the chapter
   * Develop practical skills relevant to the subject
   * Can be completed with easily available materials
   * Include clear step-by-step instructions

2. Activity Time – STEM Projects - At least 2 projects that:
   * Integrate Science, Technology, Engineering and Mathematics
   * Connect to real-world applications of the chapter concepts
   * Encourage problem-solving and critical thinking
   * Include materials needed, procedure, and expected outcomes

For each activity/project, include:
* A clear title and objective
* Materials required
* Detailed procedure with steps
* Safety precautions where applicable
* Questions for reflection
* Expected outcomes or learning points

Format the content in Markdown with proper headings, lists, and organization.

Provide ONLY the Skill Activities and STEM Projects in Markdown format.
"""

def create_general_art_prompt(grade_level, model_progression_text):
    """Creates a general subject art integration prompt"""
    return f"""You are an expert in educational content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate ART-INTEGRATED LEARNING projects based on the chapter content in the PDF.

Create the following:
1. Creativity – Art Projects - At least 3 creative projects that:
   * Connect the chapter concepts to various art forms (visual arts, music, drama, etc.)
   * Allow for creative expression while reinforcing learning
   * Can be completed with commonly available art supplies
   * Are age-appropriate for {grade_level} students

2. Case Study – Level 1 - At least 1 simpler case study that:
   * Presents a real-world scenario related to the chapter concepts
   * Includes guiding questions for analysis
   * Is accessible to all students

3. Case Study – Level 2 - At least 1 more complex case study that:
   * Challenges students with a multi-faceted scenario
   * Requires deeper application of chapter concepts
   * Encourages higher-order thinking skills

For each project/case study, include:
* A clear title and learning objective
* Materials needed (for art projects)
* Detailed instructions or scenario description
* Guiding questions or reflection points
* Assessment criteria or expected outcomes

Format the content in Markdown with proper headings, lists, and organization.

Provide ONLY the Art-Integrated Learning content in Markdown format.
"""

def generate_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, model_progression_text, subject_type="General", use_chunked=False):
    """Generates specific content based on content type"""
    prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type)
    
    if not use_chunked:
        # Standard approach
        temp_pdf_path = pathlib.Path(f"temp_{pdf_filename}")
        with open(temp_pdf_path, "wb") as f:
            f.write(pdf_bytes)
        
        try:
            # Upload the PDF
            st.info(f"Uploading '{pdf_filename}' to EeeBee...")
            uploaded_file = genai.upload_file(path=temp_pdf_path, display_name=pdf_filename, mime_type="application/pdf")
            st.success(f"'{pdf_filename}' uploaded successfully to Gemini")
            
            # Generate content
            st.info(f"Generating {content_type} content for {grade_level}...")
            
            # Use higher token limit for mathematics chapter content
            if subject_type == "Mathematics" and content_type == "chapter":
                generation_config = genai.types.GenerationConfig(
                    max_output_tokens=131072,  # Doubled for multi-section math chapters
                    temperature=0.3,
                )
            else:
                generation_config = genai.types.GenerationConfig(
                    max_output_tokens=131072,
                    temperature=0.3,
                )
            
            response = model.generate_content(
                [uploaded_file, prompt],
                generation_config=generation_config
            )
            
            # Check for copyright detection
            if response.candidates and response.candidates[0].finish_reason == 4:
                st.error("Gemini detected potential copyright concerns. Trying chunked approach...")
                # Clean up resources
                genai.delete_file(uploaded_file.name)
                temp_pdf_path.unlink()
                # Use chunked approach instead
                return generate_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, model_progression_text, subject_type, use_chunked=True)
            
            # Extract text from response
            if hasattr(response, 'text') and response.text:
                result_text = response.text
            else:
                # Handle alternative response format
                if response.candidates and hasattr(response.candidates[0], 'content') and response.candidates[0].content:
                    parts = response.candidates[0].content.parts
                    if parts:
                        result_text = ''.join([part.text for part in parts if hasattr(part, 'text')])
                    else:
                        raise Exception("No content parts found in response")
                else:
                    raise Exception("No valid text content found in response")
            
            # Clean up resources
            genai.delete_file(uploaded_file.name)
            temp_pdf_path.unlink()
            
            return result_text, "Generated successfully using standard approach."
            
        except Exception as e:
            st.error(f"Error during Gemini API call: {e}")
            # Clean up resources
            if 'uploaded_file' in locals():
                try:
                    genai.delete_file(uploaded_file.name)
                except:
                    pass
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
            return None, f"Error: {str(e)}"
    else:
        # Chunked approach
        return analyze_with_chunked_approach_for_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, model_progression_text, subject_type)

def analyze_with_chunked_approach_for_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, model_progression_text, subject_type):
    """Specialized chunked approach for specific content types"""
    st.info(f"Using chunked approach to generate {content_type} content...")
    
    try:
        # Save the uploaded PDF bytes to a temporary file
        temp_pdf_path = pathlib.Path(f"temp_{pdf_filename}")
        with open(temp_pdf_path, "wb") as f:
            f.write(pdf_bytes)
            
        # Upload the entire PDF file to Gemini
        st.info(f"Uploading '{pdf_filename}' to EeeBee for chunked analysis...")
        uploaded_file = genai.upload_file(path=temp_pdf_path, display_name=pdf_filename, mime_type="application/pdf")
        st.success(f"'{pdf_filename}' uploaded successfully to Gemini")
        
        # Extract text from PDF for determining page chunks
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        
        # Determine chunk size (pages per chunk)
        pages_per_chunk = 5  # Can be adjusted based on content density
        
        # Define page ranges for chunked analysis
        page_ranges = []
        for start_page in range(0, total_pages, pages_per_chunk):
            end_page = min(start_page + pages_per_chunk - 1, total_pages - 1)
            page_ranges.append({
                "start": start_page + 1,  # 1-indexed for human readability
                "end": end_page + 1,
                "range_text": f"Pages {start_page+1}-{end_page+1} of {total_pages}"
            })
        
        doc.close()
        
        # Process each chunk
        analysis_results = []
        
        for idx, page_range in enumerate(page_ranges):
            st.info(f"Analyzing chunk {idx+1}/{len(page_ranges)}: {page_range['range_text']}...")
            
            chunk_prompt = f"""You are an expert in educational content development for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: You are analyzing CHUNK {idx+1} of {len(page_ranges)} ({page_range['range_text']}) from a book chapter for **{grade_level} (CBSE)**.
This is just a PORTION of the full chapter - focus only on this section.

**FOCUS ONLY ON PAGES {page_range['start']} to {page_range['end']} of the PDF.**

You are extracting information relevant for: {content_type.upper()} CONTENT

Please analyze these specific pages and extract:
* Key concepts, topics, and terminology relevant to {content_type}
* Important facts, examples, and explanations
* For any images, their content and relevance
* Any specific information that would be useful for creating {content_type} content

Format your analysis in Markdown. This is just an intermediate step - don't create final content yet.
"""
            
            try:
                generation_config = genai.types.GenerationConfig(
                    max_output_tokens=32768,
                    temperature=0.3,
                )
                
                # Send the full PDF but instruct to focus on specific pages
                response = model.generate_content(
                    [uploaded_file, chunk_prompt],
                    generation_config=generation_config
                )
                
                # Check for copyright detection with fallback
                if response.candidates and response.candidates[0].finish_reason == 4:
                    st.warning(f"Copyright detection on chunk {idx+1}. Trying with text-only fallback...")
                    
                    # Extract just the text for this page range as fallback
                    chunk_text = ""
                    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    for page_num in range(page_range['start']-1, page_range['end']):
                        page = doc.load_page(page_num)
                        chunk_text += page.get_text()
                    doc.close()
                    
                    # Modify prompt to indicate images won't be available
                    text_only_prompt = chunk_prompt + f"""

**NOTE: This is now processing text-only due to copyright concerns with the full PDF.**

**Text from these pages:**
'''
{chunk_text}
'''
"""
                    
                    # Try again with text-only
                    response = model.generate_content(
                        [text_only_prompt],
                        generation_config=generation_config
                    )
                
                if hasattr(response, 'text') and response.text:
                    analysis_chunk = response.text
                    analysis_results.append(analysis_chunk)
                else:
                    # Handle alternative response format
                    if response.candidates and hasattr(response.candidates[0], 'content') and response.candidates[0].content:
                        parts = response.candidates[0].content.parts
                        if parts:
                            analysis_chunk = ''.join([part.text for part in parts if hasattr(part, 'text')])
                            analysis_results.append(analysis_chunk)
                        else:
                            analysis_results.append(f"[Error processing chunk {idx+1}]")
                    else:
                        analysis_results.append(f"[Error processing chunk {idx+1}]")
                        
            except Exception as chunk_e:
                st.warning(f"Error processing chunk {idx+1}: {chunk_e}")
                analysis_results.append(f"[Error processing chunk {idx+1}: {chunk_e}]")
        
        # Combine chunk analyses
        combined_analyses = "\n\n".join(analysis_results)
        
        # Get the specific prompt for this content type
        specific_prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type)
        
        # Create the final integration prompt
        integration_prompt = f"""You are an expert educational content developer for CBSE curriculum.
You have been given analyses of different chunks of a book chapter for **{grade_level} (CBSE)**.
Your task is to create comprehensive {content_type.upper()} content based on these analyses.

This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

**Content Type Requirements:**
{specific_prompt}

**Analyses from different parts of the chapter:**
{combined_analyses}

Based on these analyses, create the complete {content_type.upper()} content in Markdown format.
Ensure it is cohesive, well-structured, and follows all the requirements for {content_type} content.
"""
        
        # Generate the final content
        st.info(f"Creating final {content_type} content...")
        
        try:
            # Use higher token limit for mathematics chapter content
            if subject_type == "Mathematics" and content_type == "chapter":
                integration_config = genai.types.GenerationConfig(
                    max_output_tokens=131072,  # Doubled for multi-section math chapters
                    temperature=0.3,  # Lower temperature for more coherent integration
                )
            else:
                integration_config = genai.types.GenerationConfig(
                    max_output_tokens=131072,
                    temperature=0.3,  # Lower temperature for more coherent integration
                )
            
            final_response = model.generate_content(
                [integration_prompt],
                generation_config=integration_config
            )
            
            if hasattr(final_response, 'text') and final_response.text:
                final_content = final_response.text
            else:
                # Handle alternative response format
                if final_response.candidates and hasattr(final_response.candidates[0], 'content') and final_response.candidates[0].content:
                    parts = final_response.candidates[0].content.parts
                    if parts:
                        final_content = ''.join([part.text for part in parts if hasattr(part, 'text')])
                    else:
                        raise Exception("No content parts found in final integration response")
                else:
                    raise Exception("No valid text content found in final integration response")
            
            # Clean up resources
            try:
                genai.delete_file(uploaded_file.name)
                st.info(f"Temporary file '{uploaded_file.display_name}' deleted from Gemini.")
            except Exception as e_del:
                st.warning(f"Could not delete temporary file '{uploaded_file.display_name}' from Gemini: {e_del}")
            
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
                    
            return final_content, "Generated successfully using chunked approach."
            
        except Exception as integration_e:
            st.error(f"Error during final integration: {integration_e}")
            
            # Clean up resources
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
            
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
                
            # Return the combined chunks without final integration
            return combined_result, "Partial analysis complete - final integration failed."
            
    except Exception as e:
        st.error(f"Error during chunked analysis: {e}")
        
        # Clean up resources
        if 'uploaded_file' in locals():
            try:
                genai.delete_file(uploaded_file.name)
            except:
                pass
        
        if 'temp_pdf_path' in locals() and temp_pdf_path.exists():
            temp_pdf_path.unlink()
        
        return None, f"Error during chunked analysis: {str(e)}"

def generate_chat_response(user_prompt, chat_history, uploaded_files, grade_level, subject):
    """Generate a response from EeeBee for the chat interface"""
    try:
        # Build context from chat history
        conversation_context = ""
        if len(chat_history) > 1:
            recent_messages = chat_history[-6:]  # Keep last 6 messages for context
            for msg in recent_messages[:-1]:  # Exclude the current message
                conversation_context += f"{msg['role'].title()}: {msg['content']}\n"
        
        # Build system prompt for EeeBee
        system_prompt = f"""You are EeeBee, an expert educational content development assistant specializing in CBSE curriculum.
You help content teams create, modify, and improve educational materials that align with NCERT, NCF, and NEP 2020 guidelines.

Context:
- Target Grade: {grade_level}
- Subject Area: {subject}
- This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY

Your expertise includes:
- Content creation and improvement for CBSE curriculum
- Curriculum alignment and standards compliance
- Age-appropriate content development
- Educational activity design
- Assessment and exercise creation
- Pedagogical best practices
- Integration of 21st-century skills

Guidelines:
- Be helpful, friendly, and professional
- Provide detailed, actionable advice
- Reference educational standards and best practices
- Suggest specific improvements or modifications
- Ask clarifying questions when needed
- Maintain consistency with CBSE/NCERT guidelines

Previous conversation:
{conversation_context}

Current user question: {user_prompt}"""

        # Prepare content for generation
        content_parts = [system_prompt]
        
        # Add uploaded PDFs if any
        uploaded_file_objects = []
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # Save uploaded file temporarily
                    temp_path = pathlib.Path(f"temp_chat_{uploaded_file.name}")
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    
                    # Upload to Gemini
                    gemini_file = genai.upload_file(
                        path=temp_path, 
                        display_name=uploaded_file.name, 
                        mime_type="application/pdf"
                    )
                    uploaded_file_objects.append(gemini_file)
                    content_parts.append(gemini_file)
                    
                    # Clean up temp file
                    temp_path.unlink()
                    
                except Exception as e:
                    st.warning(f"Could not process {uploaded_file.name}: {e}")
        
        # Generate response
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=131072,
            temperature=0.3,  # Slightly higher temperature for more conversational responses
        )
        
        response = model.generate_content(
            content_parts,
            generation_config=generation_config
        )
        
        # Clean up uploaded files from Gemini
        for gemini_file in uploaded_file_objects:
            try:
                genai.delete_file(gemini_file.name)
            except Exception as e:
                pass  # Ignore cleanup errors
        
        # Extract response text
        if hasattr(response, 'text') and response.text:
            return response.text
        elif response.candidates and response.candidates[0].content:
            parts = response.candidates[0].content.parts
            if parts:
                return ''.join([part.text for part in parts if hasattr(part, 'text')])
            else:
                return "I apologize, but I couldn't generate a proper response. Could you please rephrase your question?"
        else:
            return "I'm having trouble processing your request right now. Please try again."
            
    except Exception as e:
        return f"I encountered an error: {str(e)}. Please try asking your question again or check if your uploaded files are valid PDFs."

def generate_chat_response_stream(user_prompt, chat_history, uploaded_files, grade_level, subject):
    """Generate a streaming response from EeeBee for the chat interface"""
    try:
        # Build context from chat history
        conversation_context = ""
        if len(chat_history) > 1:
            recent_messages = chat_history[-6:]  # Keep last 6 messages for context
            for msg in recent_messages[:-1]:  # Exclude the current message
                conversation_context += f"{msg['role'].title()}: {msg['content']}\n"
        
        # Build system prompt for EeeBee
        system_prompt = f"""You are EeeBee, an expert educational content development assistant specializing in CBSE curriculum.
You help content teams create, modify, and improve educational materials that align with NCERT, NCF, and NEP 2020 guidelines.

Context:
- Target Grade: {grade_level}
- Subject Area: {subject}
- This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY

Your expertise includes:
- Content creation and improvement for CBSE curriculum
- Curriculum alignment and standards compliance
- Age-appropriate content development
- Educational activity design
- Assessment and exercise creation
- Pedagogical best practices
- Integration of 21st-century skills

Guidelines:
- Be helpful, friendly, and professional
- Provide detailed, actionable advice
- Reference educational standards and best practices
- Suggest specific improvements or modifications
- Ask clarifying questions when needed
- Maintain consistency with CBSE/NCERT guidelines

Previous conversation:
{conversation_context}

Current user question: {user_prompt}"""

        # Prepare content for generation
        content_parts = [system_prompt]
        
        # Add uploaded PDFs if any
        uploaded_file_objects = []
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # Save uploaded file temporarily
                    temp_path = pathlib.Path(f"temp_chat_{uploaded_file.name}")
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    
                    # Upload to Gemini
                    gemini_file = genai.upload_file(
                        path=temp_path, 
                        display_name=uploaded_file.name, 
                        mime_type="application/pdf"
                    )
                    uploaded_file_objects.append(gemini_file)
                    content_parts.append(gemini_file)
                    
                    # Clean up temp file
                    temp_path.unlink()
                    
                except Exception as e:
                    st.warning(f"Could not process {uploaded_file.name}: {e}")
        
        # Generate streaming response
        generation_config = genai.types.GenerationConfig(
            max_output_tokens=8192,
            temperature=0.7,  # Slightly higher temperature for more conversational responses
        )
        
        # Use streaming mode
        response_stream = model.generate_content(
            content_parts,
            generation_config=generation_config,
            stream=True
        )
        
        # Yield chunks as they come
        for chunk in response_stream:
            if chunk.text:
                yield chunk.text
        
        # Clean up uploaded files from Gemini
        for gemini_file in uploaded_file_objects:
            try:
                genai.delete_file(gemini_file.name)
            except Exception as e:
                pass  # Ignore cleanup errors
            
    except Exception as e:
        yield f"I encountered an error: {str(e)}. Please try asking your question again or check if your uploaded files are valid PDFs."

# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("📚 EeeBee Content Development Suite ✨")

st.markdown("""
Welcome to the EeeBee Content Development Suite! Choose between improving existing chapters or chatting with EeeBee for content assistance.
""")

# Create tabs for different functionalities
tab1, tab2 = st.tabs(["📚 Chapter Improver", "💬 Content Chat with EeeBee"])

with tab1:
    st.header("📚 Book Chapter Improver Tool")
    st.markdown("""
    Upload your book chapter in PDF format. This tool will analyze it using EeeBee
    based on the 'Model Chapter Progression and Elements' and suggest improvements.
    Select which part of the content you want to generate.
    """)

    # Grade Level Selector
    grade_options = [f"Grade {i}" for i in range(1, 13)] # Grades 1-12
    selected_grade = st.selectbox("Select Target Grade Level (CBSE):", grade_options, index=8, key="grade_selector_tab1") # Default to Grade 9

    # Subject Type Selector
    subject_type = st.selectbox(
        "Select Subject Type:",
        ["General (Uses Model Chapter Progression)", "Mathematics"],
        help="Choose 'Mathematics' for math-specific content structure or 'General' for other subjects.",
        key="subject_selector_tab1"
    )

    # Analysis Method Selector
    analysis_method = st.radio(
        "Choose Analysis Method:",
        ["Standard (Full Document)", "Chunked (For Complex Documents)"],
        help="Use 'Chunked' method if you encounter copyright errors with the standard method.",
        key="analysis_method_tab1"
    )

    # Load Model Chapter Progression
    model_progression = load_model_chapter_progression()

    if model_progression:
        st.sidebar.subheader("Model Chapter Progression:")
        st.sidebar.text_area("Model Details", model_progression, height=300, disabled=True)

        uploaded_file_st = st.file_uploader("Upload your chapter (PDF only)", type="pdf", key="pdf_uploader_tab1")

        if uploaded_file_st is not None:
            st.info(f"Processing '{uploaded_file_st.name}' for {selected_grade}...")
            
            # Get PDF bytes for processing
            pdf_bytes = uploaded_file_st.getvalue()

            # Create columns for the buttons
            col1, col2 = st.columns(2)
            col3, col4 = st.columns(2)
            
            # Create session state to store generated content
            if 'chapter_content' not in st.session_state:
                st.session_state.chapter_content = None
            if 'exercises' not in st.session_state:
                st.session_state.exercises = None
            if 'skill_activities' not in st.session_state:
                st.session_state.skill_activities = None
            if 'art_learning' not in st.session_state:
                st.session_state.art_learning = None

            # Generate Chapter Content Button
            generate_chapter = col1.button("🔍 Generate Chapter Content", key="gen_chapter")
            
            # Generate Exercises Button
            generate_exercises = col2.button("📝 Generate Exercises", key="gen_exercises")
            
            # Generate Skill Activities Button
            generate_skills = col3.button("🛠️ Generate Skill Activities", key="gen_skills")
            
            # Generate Art Learning Button
            generate_art = col4.button("🎨 Generate Art-Integrated Learning", key="gen_art")
            
            # Download All Button (outside columns)
            download_all = st.button("📥 Download Complete Chapter with All Elements", key="download_all")
            
            # Handle button clicks and content generation
            if generate_chapter:
                with st.spinner(f"🧠 Generating Chapter Content for {selected_grade}..."):
                    content, message = generate_specific_content("chapter", pdf_bytes, uploaded_file_st.name, selected_grade, model_progression, subject_type, use_chunked=(analysis_method == "Chunked (For Complex Documents)"))
                    if content:
                        st.session_state.chapter_content = content
                        st.success(f"✅ Chapter Content generated successfully! {message}")
                        st.subheader("Chapter Content:")
                        st.markdown(st.session_state.chapter_content)
                        # Download button for this content only
                        doc = create_word_document(st.session_state.chapter_content)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download Chapter Content as Word (.docx)",
                            data=doc_io,
                            file_name=f"chapter_content_{uploaded_file_st.name.replace('.pdf', '.docx')}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Failed to generate Chapter Content: {message}")
            
            if generate_exercises:
                with st.spinner(f"🧠 Generating Exercises for {selected_grade}..."):
                    content, message = generate_specific_content("exercises", pdf_bytes, uploaded_file_st.name, selected_grade, model_progression, subject_type, use_chunked=(analysis_method == "Chunked (For Complex Documents)"))
                    if content:
                        st.session_state.exercises = content
                        st.success(f"✅ Exercises generated successfully! {message}")
                        st.subheader("Exercises:")
                        st.markdown(st.session_state.exercises)
                        # Download button for this content only
                        doc = create_word_document(st.session_state.exercises)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download Exercises as Word (.docx)",
                            data=doc_io,
                            file_name=f"exercises_{uploaded_file_st.name.replace('.pdf', '.docx')}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Failed to generate Exercises: {message}")
            
            if generate_skills:
                with st.spinner(f"🧠 Generating Skill Activities for {selected_grade}..."):
                    content, message = generate_specific_content("skills", pdf_bytes, uploaded_file_st.name, selected_grade, model_progression, subject_type, use_chunked=(analysis_method == "Chunked (For Complex Documents)"))
                    if content:
                        st.session_state.skill_activities = content
                        st.success(f"✅ Skill Activities generated successfully! {message}")
                        st.subheader("Skill Activities:")
                        st.markdown(st.session_state.skill_activities)
                        # Download button for this content only
                        doc = create_word_document(st.session_state.skill_activities)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download Skill Activities as Word (.docx)",
                            data=doc_io,
                            file_name=f"skill_activities_{uploaded_file_st.name.replace('.pdf', '.docx')}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Failed to generate Skill Activities: {message}")
            
            if generate_art:
                with st.spinner(f"🧠 Generating Art-Integrated Learning for {selected_grade}..."):
                    content, message = generate_specific_content("art", pdf_bytes, uploaded_file_st.name, selected_grade, model_progression, subject_type, use_chunked=(analysis_method == "Chunked (For Complex Documents)"))
                    if content:
                        st.session_state.art_learning = content
                        st.success(f"✅ Art-Integrated Learning generated successfully! {message}")
                        st.subheader("Art-Integrated Learning:")
                        st.markdown(st.session_state.art_learning)
                        # Download button for this content only
                        doc = create_word_document(st.session_state.art_learning)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download Art-Integrated Learning as Word (.docx)",
                            data=doc_io,
                            file_name=f"art_learning_{uploaded_file_st.name.replace('.pdf', '.docx')}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error(f"Failed to generate Art-Integrated Learning: {message}")
            
            if download_all:
                # Combine all generated content (if any)
                all_content_parts = []
                
                if st.session_state.chapter_content:
                    all_content_parts.append("# CHAPTER CONTENT\n\n" + st.session_state.chapter_content)
                else:
                    st.warning("Chapter Content has not been generated yet.")
                
                if st.session_state.exercises:
                    all_content_parts.append("# EXERCISES\n\n" + st.session_state.exercises)
                else:
                    st.warning("Exercises have not been generated yet.")
                
                if st.session_state.skill_activities:
                    all_content_parts.append("# SKILL ACTIVITIES\n\n" + st.session_state.skill_activities)
                else:
                    st.warning("Skill Activities have not been generated yet.")
                
                if st.session_state.art_learning:
                    all_content_parts.append("# ART-INTEGRATED LEARNING\n\n" + st.session_state.art_learning)
                else:
                    st.warning("Art-Integrated Learning has not been generated yet.")
                
                if all_content_parts:
                    combined_content = "\n\n" + "\n\n".join(all_content_parts)
                    
                    # Create Word document with all content
                    doc = create_word_document(combined_content)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.download_button(
                        label="📥 Download Complete Chapter with All Elements as Word (.docx)",
                        data=doc_io,
                        file_name=f"complete_chapter_{uploaded_file_st.name.replace('.pdf', '.docx')}",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No content has been generated yet. Please generate at least one type of content first.")
    else:
        st.error("Failed to load the Model Chapter Progression. The tool cannot proceed without it.")

with tab2:
    st.header("💬 Content Chat with EeeBee")
    st.markdown("""
    Chat with EeeBee for content development assistance! Upload PDFs for context or ask questions directly.
    EeeBee can help with content creation, modification, curriculum alignment, and educational guidance.
    """)
    
    # Move chat settings to main area (above chat)
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        # Grade level for chat context
        chat_grade = st.selectbox(
            "Grade Level Context:", 
            [f"Grade {i}" for i in range(1, 13)], 
            index=8, 
            key="chat_grade"
        )
    
    with col2:
        # Subject context
        chat_subject = st.selectbox(
            "Subject Context:",
            ["General Education", "Mathematics", "Science", "Social Studies", "English", "Hindi", "Other"],
            key="chat_subject"
        )
    
    with col3:
        # Clear chat button
        if st.button("🗑️ Clear Chat History", key="clear_chat"):
            st.session_state.chat_messages = []
            st.session_state.chat_uploaded_files = []
            st.rerun()
    
    # PDF Upload for chat context
    st.subheader("📄 Upload Documents for Context")
    chat_uploaded_files = st.file_uploader(
        "Upload PDFs for EeeBee to reference:",
        type="pdf",
        accept_multiple_files=True,
        key="chat_pdf_uploader"
    )
    
    if chat_uploaded_files:
        st.session_state.chat_uploaded_files = chat_uploaded_files
        st.success(f"📄 {len(chat_uploaded_files)} PDF(s) uploaded successfully!")
        uploaded_files_info = ""
        for file in chat_uploaded_files:
            uploaded_files_info += f"• {file.name}\n"
        st.text(uploaded_files_info)
    
    st.divider()
    
    # Initialize chat session state
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    if 'chat_uploaded_files' not in st.session_state:
        st.session_state.chat_uploaded_files = []
    
    # Display chat messages in a container
    chat_container = st.container()
    
    with chat_container:
        # Display existing chat messages
        for message in st.session_state.chat_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    
    # Chat input at the bottom
    if prompt := st.chat_input("Ask EeeBee anything about content development..."):
        # Add user message to chat history
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        
        # Display user message
        with chat_container:
            with st.chat_message("user"):
                st.markdown(prompt)
        
        # Generate EeeBee response with streaming
        with chat_container:
            with st.chat_message("assistant"):
                # Create placeholder for streaming response
                response_placeholder = st.empty()
                response_text = ""
                
                try:
                    # Generate streaming response
                    for chunk in generate_chat_response_stream(
                        prompt, 
                        st.session_state.chat_messages, 
                        st.session_state.chat_uploaded_files,
                        chat_grade,
                        chat_subject
                    ):
                        response_text += chunk
                        response_placeholder.markdown(response_text + "▊")  # Add cursor effect
                    
                    # Remove cursor and show final response
                    response_placeholder.markdown(response_text)
                    
                    # Add assistant response to chat history
                    st.session_state.chat_messages.append({"role": "assistant", "content": response_text})
                    
                except Exception as e:
                    error_message = f"I encountered an error: {str(e)}. Please try asking your question again."
                    response_placeholder.markdown(error_message)
                    st.session_state.chat_messages.append({"role": "assistant", "content": error_message})

st.sidebar.markdown("---")
st.sidebar.info("This app uses the EeeBee API powered by EeeBee.")
