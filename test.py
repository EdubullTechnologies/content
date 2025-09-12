import streamlit as st
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import io
import base64
import json
import requests
from threading import Event, Thread
import time
import re
from typing import List, Dict, Any
import hashlib
from datetime import datetime

# --- Streamlit Cloud Content Protection System ---
def save_content_safely(content_type, content, grade_level=None):
    """Save content with multiple backup strategies for Streamlit Cloud"""
    if not content:
        return
    
    try:
        timestamp = datetime.now().isoformat()
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        
        # Strategy 1: Primary session state
        st.session_state[content_type] = content
        
        # Strategy 2: Multiple backup session states
        backup_data = {
            'content': content,
            'timestamp': timestamp,
            'hash': content_hash,
            'length': len(content)
        }
        
        st.session_state[f"{content_type}_backup_1"] = backup_data
        st.session_state[f"{content_type}_backup_2"] = backup_data.copy()
        st.session_state[f"{content_type}_backup_3"] = backup_data.copy()
        
        # Strategy 3: Browser localStorage (works on Streamlit Cloud)
        save_to_browser_storage(content_type, content, timestamp)
        
        # Strategy 4: Compressed backup for large content
        if len(content) > 10000:
            # Store truncated version as emergency backup
            emergency_backup = content[:5000] + "\n\n[CONTENT TRUNCATED - PARTIAL BACKUP]"
            st.session_state[f"{content_type}_emergency"] = emergency_backup
        
        # Mark as successfully saved
        st.session_state[f"{content_type}_saved_timestamp"] = timestamp
        
    except Exception as e:
        st.error(f"⚠️ Failed to save {content_type}: {e}")

def save_to_browser_storage(content_type, content, timestamp):
    """Save content to browser localStorage for persistence"""
    try:
        # Create minimal JavaScript to save to localStorage
        content_json = json.dumps(content).replace("'", "\\'")
        js_code = f"""
        <script>
        try {{
            localStorage.setItem('eeebee_{content_type}', '{content_json}');
            localStorage.setItem('eeebee_{content_type}_time', '{timestamp}');
            console.log('Saved {content_type} to localStorage');
        }} catch(e) {{
            console.warn('localStorage save failed:', e);
        }}
        </script>
        """
        st.components.v1.html(js_code, height=0)
    except Exception:
        # Fail silently for browser storage
        pass

def recover_content_safely(content_type):
    """Attempt to recover content from available backups"""
    recovered_content = None
    source = "none"
    
    # Check primary session state
    if st.session_state.get(content_type):
        recovered_content = st.session_state[content_type]
        source = "primary"
    
    # Check backup session states
    elif st.session_state.get(f"{content_type}_backup_1"):
        backup = st.session_state[f"{content_type}_backup_1"]
        recovered_content = backup.get('content')
        source = "backup_1"
    
    elif st.session_state.get(f"{content_type}_backup_2"):
        backup = st.session_state[f"{content_type}_backup_2"]
        recovered_content = backup.get('content')
        source = "backup_2"
    
    elif st.session_state.get(f"{content_type}_backup_3"):
        backup = st.session_state[f"{content_type}_backup_3"]
        recovered_content = backup.get('content')
        source = "backup_3"
    
    # Check emergency backup
    elif st.session_state.get(f"{content_type}_emergency"):
        recovered_content = st.session_state[f"{content_type}_emergency"]
        source = "emergency"
    
    return recovered_content, source

def verify_and_recover_all_content():
    """Check all content and recover if missing"""
    content_types = ['chapter_content', 'exercises', 'skill_activities', 'art_learning']
    recovered_any = False
    
    for content_type in content_types:
        current_content = st.session_state.get(content_type)
        
        if not current_content:
            recovered_content, source = recover_content_safely(content_type)
            if recovered_content:
                st.session_state[content_type] = recovered_content
                st.success(f"🔄 Recovered {content_type.replace('_', ' ')} from {source}")
                recovered_any = True
                
                # Re-save with all protection strategies
                save_content_safely(content_type, recovered_content)
    
    return recovered_any

def auto_save_during_streaming(content_type, partial_content, interval=1000):
    """Auto-save content during streaming at regular intervals"""
    if not partial_content:
        return
        
    # Save every ~1000 characters to prevent loss
    if len(partial_content) % interval < 50:
        save_content_safely(content_type, partial_content)

def display_content_status():
    """Display content status in sidebar for monitoring"""
    content_types = ['chapter_content', 'exercises', 'skill_activities', 'art_learning']
    
    with st.sidebar:
        st.markdown("### 💾 Content Status")
        for content_type in content_types:
            content = st.session_state.get(content_type)
            if content:
                saved_time = st.session_state.get(f"{content_type}_saved_timestamp", "Unknown")
                length = len(content)
                
                # Show status with emoji
                if saved_time != "Unknown":
                    status = "🟢"
                    time_str = saved_time[11:19] if len(saved_time) > 19 else saved_time
                else:
                    status = "🟡"
                    time_str = "Not saved"
                
                name = content_type.replace('_', ' ').title()
                st.markdown(f"{status} **{name}**")
                st.markdown(f"   📏 {length:,} chars | 🕒 {time_str}")
        
        # Recovery button
        if st.button("🔄 Check & Recover All", key="recover_all_content"):
            recovered = verify_and_recover_all_content()
            if not recovered:
                st.info("✅ All content is safe!")
            st.rerun()

# --- Configuration ---
# Get API key from Streamlit secrets
try:
    OPENROUTER_API_KEY = st.secrets["OPENROUTER_API_KEY"]
    YOUR_SITE_URL = st.secrets.get("YOUR_SITE_URL", "https://your-site.com")
    YOUR_SITE_NAME = st.secrets.get("YOUR_SITE_NAME", "EeeBee Content Suite")
except KeyError:
    st.error("""
    🔑 **OpenRouter API Key Not Found!**
    
    Please add your OpenRouter API key to Streamlit secrets:
    
    **For local development:**
    Create a `.streamlit/secrets.toml` file in your project directory with:
    ```toml
    OPENROUTER_API_KEY = "your_openrouter_api_key_here"
    YOUR_SITE_URL = "https://your-site.com"  # Optional
    YOUR_SITE_NAME = "EeeBee Content Suite"  # Optional
    ```
    
    **For Streamlit Cloud deployment:**
    1. Go to your app's settings in Streamlit Cloud
    2. Click on "Secrets" 
    3. Add the following:
    ```toml
    OPENROUTER_API_KEY = "your_openrouter_api_key_here"
    YOUR_SITE_URL = "https://your-site.com"
    YOUR_SITE_NAME = "EeeBee Content Suite"
    ```
    
    **How to get an OpenRouter API Key:**
    1. Go to [OpenRouter](https://openrouter.ai/)
    2. Sign up and create an API key
    3. Copy the key and add it to your secrets
    """)
    st.stop()

# Initialize OpenAI Client with OpenRouter
try:
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=OPENROUTER_API_KEY,
    )
    # Test the client
    models = client.models.list()
except Exception as e:
    st.error(f"Error configuring OpenRouter client: {e}")
    st.stop()

# Model to use
MODEL_NAME = "google/gemini-2.5-pro"
# --- Helper Functions ---

def load_model_chapter_progression(file_path="Model Chapter Progression and Elements.txt"):
    """Loads the model chapter progression text."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        st.error(f"Error: The file {file_path} was not found. Please make sure it's in the same directory as app.py.")
        return None

def extract_text_from_pdf(pdf_file_bytes):
    """Extracts text from PDF bytes."""
    text = ""
    try:
        doc = fitz.open(stream=pdf_file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += f"\n\n--- Page {page_num + 1} ---\n"
            text += page.get_text()
        doc.close()
    except Exception as e:
        st.error(f"Could not extract text from PDF: {e}")
    return text

def extract_images_from_pdf(pdf_file_bytes):
    """Extracts images from PDF and returns them as base64 encoded strings."""
    images = []
    try:
        doc = fitz.open(stream=pdf_file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                # Convert to base64
                base64_image = base64.b64encode(image_bytes).decode('utf-8')
                ext = base_image["ext"]
                images.append({
                    "page": page_num + 1,
                    "base64": f"data:image/{ext};base64,{base64_image}",
                    "description": f"Image from page {page_num + 1}"
                })
        doc.close()
    except Exception as e:
        st.warning(f"Could not extract images from PDF: {e}")
    return images

def create_messages_with_pdf_content(prompt, pdf_text, pdf_images=None):
    """Creates messages array for OpenAI API with PDF content."""
    messages = []
    
    # Add the main prompt with text content
    content_parts = [{"type": "text", "text": prompt}]
    
    # Add PDF text
    if pdf_text:
        content_parts.append({
            "type": "text", 
            "text": f"\n\nPDF Content:\n{pdf_text}"
        })
    
    # Add all images if available
    if pdf_images:  # No limit on number of images
        for img in pdf_images:  # Include all images
            content_parts.append({
                "type": "image_url",
                "image_url": {
                    "url": img["base64"],
                    "detail": "high"  # Use high detail for best quality
                }
            })
    
    messages.append({
        "role": "user",
        "content": content_parts
    })
    
    return messages

# Mistral OCR Functions
def process_pdf_with_mistral_ocr(pdf_bytes, pdf_filename):
    """Process PDF using Mistral OCR API for advanced text extraction with structure preservation."""
    try:
        import requests
        import base64
        import os
        
        # Get Mistral API key from Streamlit secrets
        try:
            mistral_api_key = st.secrets["MISTRAL_API_KEY"]
        except KeyError:
            st.error("❌ **Mistral API Key Required:** Please add MISTRAL_API_KEY to Streamlit secrets")
            st.info("📋 **How to add:** Create `.streamlit/secrets.toml` file with: `MISTRAL_API_KEY = \"your-key-here\"`")
            return None, None
        
        # Encode PDF to base64
        base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
        
        # Prepare the request
        url = "https://api.mistral.ai/v1/ocr"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {mistral_api_key}"
        }
        
        data = {
            "model": "mistral-ocr-latest",
            "document": {
                "type": "document_url",
                "document_url": f"data:application/pdf;base64,{base64_pdf}"
            },
            "include_image_base64": True
        }
        
        st.info("🔬 Processing PDF with Mistral OCR...")
        
        # Make the API request
        response = requests.post(url, headers=headers, json=data, timeout=300)
        
        if response.status_code == 200:
            ocr_result = response.json()
            
            # Extract text and images from OCR result
            extracted_text = ""
            extracted_images = []
            
            pages_processed = len(ocr_result.get('pages', []))
            st.success(f"✅ Successfully processed {pages_processed} pages with Mistral OCR")
            
            for page_data in ocr_result.get('pages', []):
                page_num = page_data.get('index', 1)
                markdown_content = page_data.get('markdown', '')
                
                # Add page header and content
                extracted_text += f"\n\n--- Page {page_num} (Mistral OCR) ---\n"
                extracted_text += markdown_content
                
                # Extract images if available
                if 'images' in page_data:
                    for img_idx, img_data in enumerate(page_data['images']):
                        if 'base64' in img_data:
                            extracted_images.append({
                                "page": page_num,
                                "base64": f"data:image/png;base64,{img_data['base64']}",
                                "description": f"Image {img_idx + 1} from page {page_num} (Mistral OCR)",
                                "bbox": img_data.get('bbox', {})
                            })
            
            # Add processing summary
            usage_info = ocr_result.get('usage_info', {})
            extracted_text += f"\n\n=== MISTRAL OCR SUMMARY ===\n"
            extracted_text += f"Pages processed: {usage_info.get('pages_processed', pages_processed)}\n"
            extracted_text += f"Document size: {usage_info.get('doc_size_bytes', 'Unknown')} bytes\n"
            extracted_text += f"OCR Model: {ocr_result.get('model', 'mistral-ocr-latest')}\n"
            extracted_text += f"Structure preserved: Headers, tables, lists, and formatting maintained in markdown\n"
            
            return extracted_text, extracted_images
            
        else:
            error_msg = f"Mistral OCR API error: {response.status_code} - {response.text}"
            st.error(f"❌ {error_msg}")
            return None, None
            
    except ImportError:
        st.error("❌ **Missing dependency:** Install requests with `pip install requests`")
        return None, None
    except Exception as e:
        st.error(f"❌ **Mistral OCR Error:** {str(e)}")
        return None, None

def create_messages_with_mistral_ocr_content(prompt, ocr_text, ocr_images=None):
    """Creates messages array for OpenAI API with Mistral OCR content."""
    messages = []
    
    # Add the main prompt with OCR text content
    content_parts = [{"type": "text", "text": prompt}]
    
    # Add OCR text (already in markdown format)
    if ocr_text:
        content_parts.append({
            "type": "text", 
            "text": f"\n\nDocument Content (Processed with Mistral OCR):\n{ocr_text}"
        })
    
    # Add images if available
    if ocr_images:
        for img in ocr_images:
            content_parts.append({
                "type": "image_url",
                "image_url": {
                    "url": img["base64"],
                    "detail": "high"  # Use high detail for best quality
                }
            })
    
    messages.append({
        "role": "user",
        "content": content_parts
    })
    
    return messages

def analyze_with_llm(pdf_file_bytes, pdf_filename, model_progression_text, grade_level):
    """Analyzes the uploaded PDF chapter using OpenRouter API."""
    st.info("Extracting content from PDF...")
    
    # Extract text from PDF
    pdf_text = extract_text_from_pdf(pdf_file_bytes)
    if not pdf_text:
        return "Error: Could not extract text from PDF.", "Error"
    
    # Extract images (optional - you can disable this if token usage is too high)
    pdf_images = extract_images_from_pdf(pdf_file_bytes)
    
    prompt_content = f"""You are an expert in educational content development, specifically for CBSE curriculum.
Your task is to analyze the provided PDF document, which is a book chapter intended for **{grade_level} (CBSE)**.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes. Please proceed with the analysis and rewriting.

You need to evaluate it against the 'Model Chapter Progression and Elements' provided below.
The final output should be a completely rewritten and improved chapter that incorporates all necessary corrections and adheres to the model, suitable for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines. The book should be according to the NCERT books for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**REQUIRED CONTENT STRUCTURE:**
The improved chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
10. Exercise (with the following 5 question types, each with appropriate number of questions):
   - MCQ
   - Assertion and Reason
   - Fill in the Blanks
   - True False
   - Define the following terms
   - Match the column
   - Give Reason for the following Statement (Easy Level)
   - Answer in Brief (Moderate Level)
   - Answer in Detail (Hard Level)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Formatting Requirements:**
* Format like a proper textbook with well-structured paragraphs (not too long but sufficiently developed)
* Use clear headings and subheadings to organize content
* Include multiple Bloom's Taxonomy questions at different levels in each relevant section
* Create concept-wise summaries (visuals are optional but helpful)
* Integrate special features within the core concept flow rather than as separate sections
* DO NOT change the original concept names/terminology from the uploaded PDF

**Instructions for Analysis and Rewrite:**
1.  **Read and understand the entire PDF chapter.** Pay attention to text, images, diagrams, and overall structure, keeping the **{grade_level}** student in mind.
2.  **Compare the chapter to each element** in the 'Model Chapter Progression and Elements'.
3.  **Identify deficiencies:** Where does the uploaded chapter fall short for a **{grade_level}** audience and the model? Be specific.
4.  **Analyze Images:** For each image in the PDF, assess its:
    *   Relevance to the surrounding text and learning objectives for **{grade_level}**.
    *   Clarity and quality.
    *   Age-appropriateness for **{grade_level}**.
    *   Presence of any specific characters if expected (e.g., 'EeeBee').
5.  **Suggest Image Actions:**
    *   If an image is good, state that it should be kept.
    *   If an image needs improvement (e.g., better clarity, different focus for **{grade_level}**), describe the improvement.
    *   If an image is irrelevant, unsuitable, or missing, provide a detailed prompt for an AI image generator to create a new, appropriate image suitable for **{grade_level}**. The prompt should be clearly marked, e.g., "[PROMPT FOR NEW IMAGE: A detailed description of the image needed, including style, content, characters like 'EeeBee' if applicable, and educational purpose for **{grade_level}**.]"
6.  **Rewrite the Chapter:** Create a new version of the chapter that:
    *   Follows the 'Model Chapter Progression and Elements' strictly.
    *   MUST include ALL elements from the REQUIRED CONTENT STRUCTURE in the order specified.
    *   Corrects all identified mistakes and deficiencies.
    *   Integrates image feedback directly into the text.
    *   Ensures content is aligned with NCERT, NCF, and NEP 2020 principles (e.g., inquiry-based learning, 21st-century skills, real-world connections) appropriate for **{grade_level}**.
    *   Is written in clear, engaging, and age-appropriate language for **{grade_level} (CBSE)**.
    *   Use Markdown for formatting (headings, bold, italics, lists) to make it easy to convert to a Word document. For example:
        # Chapter Title
        ## Section Title
        **Bold Text**
        *Italic Text*
        - Item 1
        - Item 2
    *   Present content in proper textbook paragraphs like a NCERT textbook
    *   Include multiple Bloom's Taxonomy questions at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create) in each appropriate section
    *   Organize summaries by individual concepts rather than as a single summary
    *   Weave special features (misconceptions, 21st century skills, etc.) into the core concept flow
    *   PRESERVE all original concept names and terminology from the source PDF

IMPORTANT: Do NOT directly copy large portions of text from the PDF. Instead, use your own words to rewrite and improve the content while maintaining the original concepts and terminology.

**Output Format:**
Provide the complete, rewritten chapter text in Markdown format, incorporating all analyses and image prompts as described. Do not include a preamble or a summary of your analysis; just the final chapter content.
"""
    
    st.info(f"Sending request to Claude via OpenRouter for {grade_level}... This may take some time for larger documents.")
    
    try:
        # Create messages with PDF content
        messages = create_messages_with_pdf_content(prompt_content, pdf_text, pdf_images)
        
        # Make API call
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": YOUR_SITE_URL,
                "X-Title": YOUR_SITE_NAME,
            },
            model=MODEL_NAME,
            messages=messages,
            max_tokens=65536,  # Adjust based on your needs and model limits
            temperature=0.3,
        )
        
        improved_text = completion.choices[0].message.content
        
        return improved_text, "LLM analysis and rewrite complete using Claude via OpenRouter."
        
    except Exception as e:
        st.error(f"Error during OpenRouter API call: {e}")
        return f"Error: Could not complete analysis. {e}", "Error"

def analyze_with_chunked_approach(pdf_file_bytes, pdf_filename, model_progression_text, grade_level):
    """Analyzes the PDF in smaller chunks to handle large documents."""
    st.info("Using chunked approach to analyze PDF...")
    
    try:
        # Extract full text
        doc = fitz.open(stream=pdf_file_bytes, filetype="pdf")
        total_pages = len(doc)
        
        # Determine chunk size
        pages_per_chunk = 5
        
        # Process chunks
        analysis_results = []
        
        for start_page in range(0, total_pages, pages_per_chunk):
            end_page = min(start_page + pages_per_chunk - 1, total_pages - 1)
            st.info(f"Analyzing pages {start_page+1}-{end_page+1} of {total_pages}...")
            
            # Extract text for this chunk
            chunk_text = ""
            chunk_images = []
            
            for page_num in range(start_page, end_page + 1):
                page = doc.load_page(page_num)
                chunk_text += f"\n\n--- Page {page_num + 1} ---\n"
                chunk_text += page.get_text()
                
                # Extract images from this page
                image_list = page.get_images(full=True)
                for img in image_list:  # Include all images from page
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        base64_image = base64.b64encode(image_bytes).decode('utf-8')
                        ext = base_image["ext"]
                        chunk_images.append({
                            "page": page_num + 1,
                            "base64": f"data:image/{ext};base64,{base64_image}",
                        })
                    except:
                        pass
            
            # Analyze this chunk
            chunk_prompt = f"""You are an expert in educational content development for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: You are analyzing CHUNK (Pages {start_page+1}-{end_page+1} of {total_pages}) from a book chapter for **{grade_level} (CBSE)**.
This is just a PORTION of the full chapter - focus only on improving this section according to the model.

The book chapter is intended to align with NCERT, NCF, and NEP 2020 guidelines for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Instructions:**
* Rewrite and improve ONLY the content from these pages
* Format like a proper textbook with well-structured paragraphs
* Use clear headings and subheadings where appropriate
* Include Bloom's Taxonomy questions if this section is suitable for them
* Maintain the original concept names/terminology
* Write in your own words - DO NOT copy text directly
* For any images in these pages:
  - Analyze their relevance, clarity, and age-appropriateness
  - Suggest keeping good images or provide image prompts for new/improved ones
  - Image prompts should be marked as: [PROMPT FOR NEW IMAGE: description]

Please provide ONLY the improved version of this specific chunk in Markdown format.
"""
            
            try:
                messages = create_messages_with_pdf_content(chunk_prompt, chunk_text, chunk_images)
                
                completion = client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": YOUR_SITE_URL,
                        "X-Title": YOUR_SITE_NAME,
                    },
                    model=MODEL_NAME,
                    messages=messages,
                    max_tokens=60000,
                    temperature=0.3,
                )
                
                improved_chunk = completion.choices[0].message.content
                analysis_results.append(improved_chunk)
                
            except Exception as chunk_e:
                st.warning(f"Error processing chunk: {chunk_e}")
                analysis_results.append(f"[Error processing pages {start_page+1}-{end_page+1}]")
        
        doc.close()
        
        # Combine results
        combined_result = "\n\n".join(analysis_results)
        
        # Final integration
        integration_prompt = f"""You are an expert educational content developer for CBSE curriculum.
You have been given several chunks of rewritten content for a **{grade_level} (CBSE)** book chapter.
Your task is to integrate these chunks into a coherent, well-structured chapter that flows naturally.

This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

**REQUIRED CONTENT STRUCTURE:**
The final integrated chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
8. Link and Learn based Question
9. Image based question
10. Exercise (with various question types)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Important:**
* Ensure smooth transitions between chunks
* Remove any redundancies
* Create a consistent style and tone
* Format as a proper textbook chapter
* STRICTLY follow the REQUIRED CONTENT STRUCTURE

**Content to integrate:**
{combined_result}

Provide the complete, integrated chapter in Markdown format.
"""
        
        try:
            messages = [{"role": "user", "content": integration_prompt}]
            
            final_completion = client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": YOUR_SITE_URL,
                    "X-Title": YOUR_SITE_NAME,
                },
                model=MODEL_NAME,
                messages=messages,
                max_tokens=65536,
                temperature=0.4,
            )
            
            final_improved_text = final_completion.choices[0].message.content
            
            return final_improved_text, "LLM analysis and rewrite complete using chunked approach."
            
        except Exception as integration_e:
            st.error(f"Error during final integration: {integration_e}")
            return combined_result, "Partial analysis complete - final integration failed."
            
    except Exception as e:
        st.error(f"Error during chunked PDF analysis: {e}")
        return f"Error: Could not complete chunked analysis. {e}", "Error"

def create_word_document(improved_text_markdown):
    """Creates a Word document from the improved text (Markdown formatted)."""
    doc = Document()
    
    lines = improved_text_markdown.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line, style='ListBullet')
        else:
            doc.add_paragraph(line)
            
    return doc

# Helper Functions for Content Generation
def create_specific_prompt(content_type, grade_level, model_progression_text, subject_type="Science", word_limits=None):
    """Creates a prompt focused on a specific content type"""
    
    if subject_type == "Mathematics":
        if content_type == "chapter":
            return create_math_chapter_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "exercises":
            return create_math_exercises_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "skills":
            return create_math_skills_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "art":
            return create_math_art_prompt(grade_level, model_progression_text, word_limits)
    elif subject_type == "Mathematics Primary (Classes 1-5)":
        if content_type == "chapter":
            return create_math_primary_chapter_prompt(grade_level, word_limits)
        elif content_type == "exercises":
            return create_math_primary_exercises_prompt(grade_level, word_limits)
        else:
            # For Mathematics Primary, we only support chapter and exercises generation
            return f"""This content type '{content_type}' is not supported for Mathematics Primary (Classes 1-5).
            
For primary mathematics (Classes 1-5), we support:
- Chapter Content: Complete chapter transformation with all components
- Exercises: MCQs, Match the Column, Puzzles, Mental Math, Thinking Activities, and Math Lab Activities

Please use the appropriate generation option."""
    elif subject_type == "Science & E.V.S. (Classes 1-2)":
        # Extract class number to validate
        class_num = grade_level.split()[-1]
        if class_num not in ['1', '2']:
            return "Science & E.V.S. (Classes 1-2) can only be used with Grade 1 or Grade 2."
        
        if content_type == "chapter":
            return create_science_evs_foundational_chapter_prompt(grade_level, word_limits)
        elif content_type == "exercises":
            return create_science_evs_foundational_exercises_prompt(grade_level, word_limits)
        else:
            # For Science EVS Classes 1-2, we only support chapter and exercises (activities) generation
            return f"""This content type '{content_type}' is not supported for Science & E.V.S. (Classes 1-2).
            
For Science & E.V.S. Classes 1-2, we support:
- Chapter Content: Complete play-based chapter transformation
- Activities: Play-based learning activities (not formal exercises)

Please use the appropriate generation option."""
    elif subject_type == "Science & E.V.S. (Classes 3-5)":
        # Extract class number to validate
        class_num = grade_level.split()[-1]
        if class_num not in ['3', '4', '5']:
            return "Science & E.V.S. (Classes 3-5) can only be used with Grade 3, 4, or 5."
        
        if content_type == "chapter":
            return create_science_evs_preparatory_chapter_prompt(grade_level, word_limits)
        elif content_type == "exercises":
            return create_science_evs_preparatory_exercises_prompt(grade_level, word_limits)
        else:
            # For Science EVS Classes 3-5, we only support chapter and exercises generation
            return f"""This content type '{content_type}' is not supported for Science & E.V.S. (Classes 3-5).
            
For Science & E.V.S. Classes 3-5, we support:
- Chapter Content: Complete inquiry-based chapter transformation
- Exercises: Comprehensive assessment exercises

Note: Skills and Art activities are integrated within the chapter content for this subject type.

Please use the appropriate generation option."""
    elif subject_type == "Computer Science":
        if content_type == "chapter":
            return create_computer_chapter_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "exercises":
            return create_computer_exercises_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "skills":
            return create_computer_skills_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "art":
            return create_computer_art_prompt(grade_level, model_progression_text, word_limits)
    elif subject_type == "English Communication & Grammar (Classes 1-8)":
        # Extract class number to validate
        class_num = int(grade_level.split()[-1])
        if class_num not in range(1, 9):  # Classes 1-8
            return "English Communication & Grammar (Classes 1-8) can only be used with Grades 1 through 8."
        
        if content_type == "chapter":
            return create_english_chapter_prompt(grade_level, word_limits)
        elif content_type == "exercises":
            return create_english_exercises_prompt(grade_level, word_limits)
        elif content_type == "skills":
            return create_english_skills_prompt(grade_level, word_limits)
        elif content_type == "art":
            return create_english_projects_prompt(grade_level, word_limits)
        else:
            return f"""This content type '{content_type}' is not supported for English Communication & Grammar.
            
For English Communication & Grammar (Classes 1-8), we support:
- Chapter Content: Complete language learning chapter with Oxford/Cambridge standards
- Exercises: Grammar, vocabulary, comprehension, and communication exercises
- Skills: Speaking, listening, reading, and writing skill activities
- Projects: Creative language projects and communication activities

Please use the appropriate generation option."""
    elif subject_type == "Artificial Intelligence":
        if content_type == "chapter":
            return create_ai_chapter_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "exercises":
            return create_ai_exercises_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "skills":
            return create_ai_skills_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "art":
            return create_ai_projects_prompt(grade_level, model_progression_text, word_limits)
        else:
            return f"""This content type '{content_type}' is not supported for Artificial Intelligence.
            
For Artificial Intelligence, we support:
- Chapter Content: Complete AI chapter with kit integration and domain applications
- Exercises: AI coding challenges, algorithm exercises, and concept assessments  
- Skills: AI lab activities, machine learning projects, and data analysis
- Projects: Creative AI applications and real-world problem solving

Please use the appropriate generation option."""
    elif subject_type == "Robotics":
        if content_type == "chapter":
            return create_robotics_chapter_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "exercises":
            return create_robotics_exercises_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "skills":
            return create_robotics_skills_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "art":
            return create_robotics_projects_prompt(grade_level, model_progression_text, word_limits)
        else:
            return f"""This content type '{content_type}' is not supported for Robotics.
            
For Robotics, we support:
- Chapter Content: Complete Robotics chapter with kit integration and domain applications
- Exercises: Robotics coding challenges, hardware exercises, and concept assessments
- Skills: Robotics lab activities, automation projects, and sensor integration
- Projects: Creative robotics applications and engineering solutions

Please use the appropriate generation option."""
    else:
        if content_type == "chapter":
            return create_science_chapter_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "exercises":
            return create_science_exercises_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "skills":
            return create_science_skills_prompt(grade_level, model_progression_text, word_limits)
        elif content_type == "art":
            return create_science_art_prompt(grade_level, model_progression_text, word_limits)

# Mathematics Primary Classes (1-5) specific prompt function
def create_math_primary_chapter_prompt(grade_level, word_limits=None):
    """Creates a mathematics-specific chapter content prompt for Classes 1-5 following the lean structure"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'hook': 150,
            'discover': 1500,
            'activity': 300,
            'recap': 200
        }
    
    return f"""You are an Expert Indian Educator and Content Editor specializing in Mathematics textbook transformation for Classes 1-5.

**CRITICAL INSTRUCTIONS**: This prompt is EXCLUSIVELY for Mathematics Classes 1-5. DO NOT apply Model Chapter Progression or any other structure. Follow ONLY the structure specified below.

**Section 0: Prompt Scope and Limitation**
- Specific Application: This is exclusively for transforming Mathematics textbooks for Classes 1 through 5
- Do NOT extend this structure to any other subject or higher grade levels
- This requires a fundamentally different pedagogical approach than secondary education

**Your Role**: Transform the provided PDF chapter into a new, market-leading format using expert pedagogical judgment.

**Target Audience**: {grade_level} (Primary Mathematics - Ages 6-11)

**Three Golden Rules for Transformation**:
1. **PRESERVE THE CORE CONCEPT**: Identify the core mathematical concept(s) and fundamental pedagogical sequence in the PDF
2. **RESTRUCTURE THE LAYOUT**: Completely discard the old layout and use our mandatory four-part lean structure
3. **DETERMINE OPTIMAL LENGTH**: Use expert judgment based on topic complexity (no fixed page limits)

**Guiding Philosophy**:
- **Crystal Clear Language**: Extremely simple, clear language for children aged 6-11
- **Brevity and Impact**: Every page must be purposeful and packed with value
- **Experiential Learning First**: Let children DO or EXPERIENCE the concept
- **CPA Approach**: Concrete-Pictorial-Abstract sequence
- **NEP/NCF Spirit**: Focus on why over how, build thinking skills

**MANDATORY NEW STRUCTURE** (Transform PDF content to fit this EXACT structure):

## 1. Chapter Opener - The Hook (Target: {word_limits.get('hook', 150)} words)
**Action**: Discard any existing opener from PDF. Create completely new content.

**Requirements**:
- Create a captivating full-page illustration description and story
- Set in rich Indian context (Indian names, places, festivals, objects, currency)
- Age-appropriate for 6-11 year olds (positive, encouraging themes)
- Introduce the mathematical concept through story/scenario
- Include detailed image prompt for illustration
- Make it experiential - child should feel connected to the math concept

**Format**:
- Story/scenario that hooks the child
- Visual description for illustration
- Connection to the mathematical concept they'll learn

## 2. Concept & Practice - Let's Discover (Target: {word_limits.get('discover', 1500)} words)
**Action**: Take core lesson from PDF and rewrite completely using CPA approach.

**Requirements**:
- **Variable length**: Just enough pages to cover concept thoroughly without rushing
- **Integrate learning and practice**: Introduce micro-concept, then 2-4 practice questions immediately
- **Replace generic examples**: Use authentic Indian contexts throughout
- **CPA Sequence**: Concrete (manipulatives/real objects) → Pictorial (drawings/images) → Abstract (numbers/symbols)
- **Short sentences**: Break multi-step instructions into numbered points
- **Include "Brain Booster" questions**: Mark with 💡 for Higher Order Thinking Skills
- **Teacher's Notes**: Add "For the Teacher" notes to 70% of content

**Structure for each micro-concept**:
- Simple introduction with concrete example
- Pictorial representation with Indian context
- Abstract mathematical representation
- **2 Concept-wise Examples**: Provide 2 different examples for each concept using Indian context
- **Fun Fact**: Include 1 interesting mathematical fun fact related to each concept
- 2-4 immediate practice questions
- 1 Brain Booster question (💡)

**Content Safety**: Zero bias, show boys and girls equally, diverse representation of India

## 3. Activity Zone - Minds-on, Hands-on (Target: {word_limits.get('activity', 300)} words)
**Action**: Discard any existing activities from PDF. Create completely new content.

**Requirements**:
- **Two fresh activities**: 
  1. **Maths Lab Activity** (tactile, hands-on with real objects)
  2. **Fun with Maths** (game/puzzle format)
- Use easily available materials
- Indian context and examples
- Step-by-step instructions in simple language
- Promote collaboration and positive values
- Age-appropriate for 6-11 year olds

**Format**:
- Activity 1: Maths Lab Activity (hands-on exploration)
- Activity 2: Fun with Maths (engaging game/puzzle)
- Clear materials list
- Step-by-step instructions
- Learning outcome for each activity

## 4. Quick Recap / Revision (Target: {word_limits.get('recap', 200)} words)
**Action**: Create this page from scratch.

**Requirements**:
- **Visual summary**: Mind map or concept map of core concepts from PDF
- **"Let's Revise" section**: 4-5 mixed problems for assessment
- **Simple language**: Child-friendly explanations
- **Indian contexts**: All examples should be culturally relevant
- **Varied question types**: Mix of formats appropriate for the grade level

**Format**:
- Visual concept summary (mind map description)
- "Let's Revise" with 4-5 assessment questions
- Mix of question difficulties
- Answer key or hints

**Content Enhancement Requirements**:
- **Teacher's Notes**: Add practical teaching tips throughout
- **Indian Context**: Replace ALL generic examples with Indian ones (names, places, currency, festivals)
- **Language**: Kind, encouraging, simple voice throughout
- **Safety & Appropriateness**: All content suitable for ages 6-11
- **Inclusivity**: Gender-neutral, culturally diverse, ability-inclusive

**What to Remove from PDF and Replace**:
- Remove: Direct jumps to formulas without conceptual build-up
- Replace: Pages of repetitive drills with varied, contextual problems  
- Replace: Generic or culturally ambiguous examples
- Rewrite: Long, complex paragraphs into short, simple instructions

**Visual Requirements**:
- Include detailed image prompts for illustrations throughout
- Specify Indian cultural elements in visuals
- Ensure age-appropriate, colorful, engaging imagery
- Support the CPA learning progression

Analyze the provided PDF thoroughly and transform it according to this structure. Focus on making mathematics fun, accessible, and meaningful for young Indian children while preserving the core mathematical concepts and learning sequence.

Provide ONLY the transformed mathematics chapter content in Markdown format following the four-part structure above.
"""

def create_math_primary_exercises_prompt(grade_level, word_limits=None):
    """Creates a mathematics-specific exercises prompt for Classes 1-5"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'exercises': 800
        }
    
    return f"""You are an Expert Indian Educator and Content Editor specializing in Mathematics exercises for Classes 1-5.

**CRITICAL INSTRUCTIONS**: This prompt is EXCLUSIVELY for Mathematics Classes 1-5 exercises. Create age-appropriate, engaging mathematical exercises.

**Target Audience**: {grade_level} (Primary Mathematics - Ages 6-11)

Your task is to generate COMPREHENSIVE MATHEMATICS EXERCISES based on the chapter content in the PDF.
**Target Total Word Count for All Exercises**: {word_limits.get('exercises', 800)} words

**REQUIRED EXERCISE TYPES:**

## A. Multiple Choice Questions (MCQs) - 3 Questions
- Create 3 age-appropriate MCQs based on the mathematical concepts
- Use simple language suitable for Classes 1-5
- Include 4 options (A, B, C, D) for each question
- Use Indian context and examples (names, objects, currency)
- Ensure questions test understanding, not just memorization

## B. Match the Column - 5 Pairs
- Create 5 matching pairs related to the mathematical concepts
- Column A: Mathematical concepts, numbers, or problems
- Column B: Corresponding answers, definitions, or examples
- Use visual elements where appropriate (shapes, numbers, objects)
- Include Indian cultural context

## C. Mathematical Puzzle - 1 Puzzle
- Create 1 engaging mathematical puzzle that reinforces the chapter concepts
- Make it fun and challenging but age-appropriate
- Can be number puzzles, pattern puzzles, or logic puzzles
- Include clear instructions and solution approach
- Use colorful, engaging presentation

## D. Mental Math - 3 Questions
- Create 3 mental mathematics questions for quick calculation practice
- Focus on the mathematical concepts from the chapter
- Encourage mental calculation strategies
- Use real-life scenarios familiar to Indian children
- Provide tips for mental calculation where helpful

## E. Thinking Based Activity - 1 Activity
- Create 1 higher-order thinking activity that applies the mathematical concepts
- Encourage problem-solving and critical thinking
- Can involve real-world applications or creative scenarios
- Should promote discussion and multiple solution approaches
- Include guiding questions to support thinking

## F. Math Lab Activity - 1 Activity
- Create 1 hands-on mathematical laboratory activity
- Use easily available materials (household items, simple manipulatives)
- Include step-by-step instructions
- Connect to the mathematical concepts from the chapter
- Promote experiential learning and discovery
- Include observation questions and conclusions

**Content Requirements:**
- **Age-Appropriate Language**: Simple, clear instructions suitable for Classes 1-5
- **Indian Context**: Use Indian names, places, currency, festivals, and cultural elements
- **Visual Elements**: Include descriptions for diagrams, pictures, or visual aids where needed
- **Progressive Difficulty**: Start with easier questions and gradually increase complexity
- **Engaging Content**: Make mathematics fun and interesting for young learners
- **Clear Instructions**: Provide clear, step-by-step instructions for all activities
- **Safety Considerations**: Ensure all activities are safe for young children

**Formatting Requirements:**
- Use clear headings for each exercise type
- Number all questions and activities clearly
- Include answer keys or solution approaches where appropriate
- Use bullet points and numbered lists for clarity
- Mark any materials needed for activities

Ensure that all exercises directly relate to the mathematical concepts covered in the PDF chapter and provide meaningful practice opportunities for young learners.

Provide ONLY the comprehensive mathematics exercises in Markdown format.
"""

# Science & E.V.S. Classes 1-2 (Foundational Stage) specific prompt functions
def create_science_evs_foundational_chapter_prompt(grade_level, word_limits=None):
    """Creates a Science & E.V.S. chapter content prompt for Classes 1-2 following play-based pedagogy"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'opener': 200,
            'activity': 250,
            'concept': 200,
            'closer': 150
        }
    
    # Extract the class number from grade_level (e.g., "Grade 1" -> "1")
    class_num = grade_level.split()[-1]
    if class_num not in ['1', '2']:
        return "This prompt is only for Science & E.V.S. Classes 1-2. Please select Grade 1 or 2."
    
    return f"""You are an expert in Early Childhood Education for Science & E.V.S. Classes 1-2.

**Target Audience**: {grade_level} (Ages 6-8, Foundational Stage)

Create a play-based Science & E.V.S. chapter following this structure:

## Part 1: Chapter Opener (Target: {word_limits.get('opener', 200)} words)
- EeeBee's welcome message
- Simple story related to the topic
- Wonder questions for discussion

## Part 2: Learning Activities (Target: {word_limits.get('activity', 250)} words per activity)
- **Hands-on activities** for each concept
- **Simple explanations**: Basic theoretical understanding in child-friendly language
- **What we learn**: Key facts and concepts children should understand
- **Simple, pictorial instructions**: Visual guidance for activities
- **Use everyday materials**: Accessible resources for exploration

## Part 3: Concept Understanding (Target: {word_limits.get('concept', 200)} words per concept)
- **Theoretical Explanation**: Clear, simple explanations of the concept using age-appropriate language
- **Why This Happens**: Basic scientific reasoning behind the concept
- **Key Facts**: Important information children should understand about the concept
- **Visual infographics**: Supporting diagrams and illustrations
- **EeeBee's key takeaways**: Summary points for reinforcement

## Part 4: Chapter Closer (Target: {word_limits.get('closer', 150)} words)
- Student checklist with achievements
- Teacher observation guide
- Parent home connection activities

**Requirements**:
- **Balanced approach**: Combine hands-on activities with simple theoretical explanations
- **Age-appropriate language** for {grade_level} with basic concept explanations
- **Focus on understanding**: Help children learn both through doing AND knowing
- **Large, colorful illustrations** to support theoretical concepts
- **Safety-conscious activities** with educational value
- **Indian cultural context** in both activities and explanations

Transform the PDF content into this playful format suitable for young learners.

Provide the Science & E.V.S. chapter content in Markdown format.
"""

def create_science_evs_foundational_exercises_prompt(grade_level, word_limits=None):
    """Creates a Science & E.V.S. exercises prompt for Classes 1-2 focusing on activities rather than formal exercises"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'activities': 600
        }
    
    # Extract the class number from grade_level
    class_num = grade_level.split()[-1]
    if class_num not in ['1', '2']:
        return "This prompt is only for Science & E.V.S. Classes 1-2. Please select Grade 1 or 2."
    
    return f"""You are an expert in Early Childhood Education for Science & E.V.S. Classes 1-2.

**Target Audience**: {grade_level} (Ages 6-8, Foundational Stage)

Create play-based learning activities based on the chapter content.
**Target Total Word Count**: {word_limits.get('activities', 600)} words

**Required Activity Types:**
- Observation activities (2 activities)
- Sorting and classifying games (1 activity)
- Story-based activities (1 activity)
- Nature walk activity (1 activity)
- Art and craft (1 activity)
- Circle time games (2 activities)

**Activity Requirements**:
- Picture-based instructions with minimal text
- Safe for young children
- Use familiar Indian materials
- Fun and engaging
- No writing skills required

**Format for Each Activity**:
- Activity name with emoji
- Materials needed
- Step-by-step instructions
- What to observe/discover
- Extension ideas for home

Focus on exploration, discovery, and play rather than formal exercises.

Provide the play-based learning activities in Markdown format.
"""

# Science & E.V.S. Classes 3-5 (Preparatory Stage) specific prompt functions  
def create_science_evs_preparatory_chapter_prompt(grade_level, word_limits=None):
    """Creates a Science & E.V.S. chapter content prompt for Classes 3-5 following inquiry-based approach"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'opener': 300,
            'exploration': 2000,
            'project': 400,
            'summary': 250,
            'exercises': 600
        }
    
    # Extract the class number from grade_level
    class_num = grade_level.split()[-1]
    if class_num not in ['3', '4', '5']:
        return "This prompt is only for Science & E.V.S. Classes 3-5. Please select Grade 3, 4, or 5."
    
    return f"""You are an expert in Science & E.V.S. education for Classes 3-5.

**Target Audience**: {grade_level} (Ages 8-11, Preparatory Stage)

Create an inquiry-based Science & E.V.S. chapter following this structure:

## Part 1: Chapter Opener (Target: {word_limits.get('opener', 300)} words)
- Learning roadmap with EeeBee
- Previous knowledge connections
- Learning outcomes
- Wonder questions to spark curiosity

## Part 2: Exploration - "Let's Uncover the Secrets" (Target: {word_limits.get('exploration', 2000)} words)
For each concept:
- **Quick activity** or intriguing question to introduce the concept
- **Theoretical Foundation**: Detailed explanations of the scientific/environmental concepts
- **How and Why**: Scientific reasoning and cause-effect relationships
- **Concept Breakdown**: Step-by-step explanation of processes and phenomena
- **Real-world connections**: Practical applications and examples
- **Scientific Facts**: Important theoretical knowledge students should learn
- **Fun facts and discussion prompts**: Engaging additional information
- **Summary boxes**: Key theoretical points for each concept

## Part 3: Chapter Project (Target: {word_limits.get('project', 400)} words)
- Hands-on project synthesizing learning
- Materials and step-by-step procedure
- Observation recording format
- Conclusion questions

## Part 4: Summary & Revision (Target: {word_limits.get('summary', 250)} words)
- Key concepts summary
- Glossary of new terms
- Concept map
- Quick revision questions

## Part 5: Exercises (Target: {word_limits.get('exercises', 600)} words)
**Section A**: Fill blanks, True/False, MCQs, Name items, Match columns
**Section B**: Give reasons, Differentiate, Brief answers, Detail answers, HOTS questions

**Requirements**:
- **Theoretical depth**: Include substantial conceptual explanations and scientific reasoning
- **Age-appropriate language** for {grade_level} with proper scientific terminology
- **Indian cultural context** in both theoretical examples and practical applications
- **Scientific accuracy** in all theoretical explanations and concepts
- **Balanced learning**: Combine hands-on activities with strong theoretical foundation
- **Environmental awareness** supported by scientific understanding

Transform the PDF content into this inquiry-based format.

Provide the Science & E.V.S. chapter content in Markdown format.
"""

def create_science_evs_preparatory_exercises_prompt(grade_level, word_limits=None):
    """Creates a Science & E.V.S. exercises prompt for Classes 3-5 with comprehensive assessment"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'exercises': 800
        }
    
    # Extract the class number from grade_level
    class_num = grade_level.split()[-1]
    if class_num not in ['3', '4', '5']:
        return "This prompt is only for Science & E.V.S. Classes 3-5. Please select Grade 3, 4, or 5."
    
    return f"""You are an expert in Science & E.V.S. education for Classes 3-5.

**Target Audience**: {grade_level} (Ages 8-11, Preparatory Stage)

Create comprehensive Science & E.V.S. exercises based on the chapter content.
**Target Total Word Count**: {word_limits.get('exercises', 800)} words

**Exercise Structure:**

## Section A: Remembering & Understanding
- Fill in the Blanks (5 questions)
- True/False with Corrections (5 questions)
- Multiple Choice Questions (5 questions)
- Name the Following (5 items)
- Match the Column (5 pairs)

## Section B: Application, Analysis & Evaluation
- Give Reasons (3 questions)
- Differentiate Between (2 pairs)
- Answer in Brief (3 questions)
- Answer in Detail (2 questions)
- HOTS Questions (2 questions)

## Additional Components:
- Activity-Based Questions (2 questions)
- Diagram-Based Questions (2 questions)
- Value-Based Questions (1 question)

**Requirements**:
- Progressive difficulty levels
- Age-appropriate vocabulary for {grade_level}
- Indian cultural context
- Include diagrams where helpful
- Connect to practical applications
- Build environmental awareness

Ensure all questions relate to the chapter content and are appropriate for {grade_level} students.

Provide the comprehensive Science & E.V.S. exercises in Markdown format.
"""

# Mathematics-specific prompt functions
def create_math_chapter_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a mathematics-specific chapter content prompt following Haese and Harris structure"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'opening': 100,
            'objectives': 80,
            'section_intro': 150,
            'theory': 800,
            'worked_examples': 600,
            'investigation': 300,
            'exercise': 400,
            'review': 500,
            'technology': 200
        }
    
    return f"""You are an expert in mathematical education content development, following the Haese and Harris textbook methodology.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.
The content should follow Haese and Harris structure while aligning with NCERT, NCF, and NEP 2020 guidelines.

**CRITICAL INSTRUCTION**: The PDF may contain MULTIPLE MAJOR SECTIONS. You MUST include ALL sections present in the PDF following the Haese and Harris systematic approach.

**Model Chapter Progression and Elements (Base Structure):**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Mathematics Syllabus)

Your task is to generate COMPREHENSIVE MATHEMATICS CHAPTER CONTENT following the **Haese and Harris methodology**.

## HAESE AND HARRIS CHAPTER STRUCTURE

### CHAPTER OPENER

**Chapter [Number]: [Chapter Title]**

1. **Opening Scenario/Problem** (Target: {word_limits.get('opening', 100)} words)
   - Present a real-world mathematical scenario that motivates the entire chapter
   - Use engaging, practical situations that students can relate to
   - Provide a challenging question that will be answered by the end of the chapter
   - Include a detailed image prompt for the opening scenario

2. **Learning Objectives** (Target: {word_limits.get('objectives', 80)} words)
   - Clear, specific learning outcomes using mathematical language
   - Organized as bullet points with action verbs
   - Aligned with curriculum standards
   - Example format: "By the end of this chapter, you should be able to:"

3. **Chapter Overview** (Target: 50 words)
   - Brief roadmap of sections and key concepts
   - Show logical progression of mathematical ideas

### MAIN CONTENT SECTIONS

**For EACH major section in the PDF, create the following structure:**

## Section [X.Y]: [Section Title]

### A. SECTION OPENING (Target: {word_limits.get('section_intro', 150)} words)
- **Opening Investigation/Question**
  - Pose a specific mathematical question or scenario
  - Connect to real-world applications
  - Motivate the need for the mathematical concepts in this section

### B. THEORY AND DEVELOPMENT (Target: {word_limits.get('theory', 800)} words per section)

**[Concept Name]**

1. **Mathematical Development**
   - Systematic introduction of concepts
   - Clear definitions in highlighted boxes
   - Step-by-step mathematical reasoning
   - Logical progression from simple to complex

2. **Key Concepts and Formulas**
   - Important mathematical results in clearly marked boxes
   - Theorems and properties with brief explanations
   - Formula derivations where appropriate

3. **Mathematical Notation**
   - Introduce and explain new mathematical symbols
   - Consistent use of notation throughout

4. **Connections and Links**
   - Connect to previously learned mathematics
   - Show relationships between concepts
   - Cross-references to other sections/chapters

### C. WORKED EXAMPLES (Target: {word_limits.get('worked_examples', 600)} words per section)

**Example [Number]: [Descriptive Title]**

For each worked example, include:
- **Given/Find/Solution** format
- Clear step-by-step working
- Explanatory notes alongside calculations
- Alternative methods where applicable
- Checking/verification of answers
- Common errors to avoid

Provide 3-4 worked examples per section covering:
- Basic application of concepts
- More complex problems
- Real-world applications
- Technology-enhanced solutions

### D. INVESTIGATION/DISCOVERY ACTIVITY (Target: {word_limits.get('investigation', 300)} words per section)

**Investigation [Number]: [Title]**

- **Aim**: Clear statement of what students will discover
- **Method**: Step-by-step instructions
- **Materials**: List of required materials/technology
- **Questions**: Guided questions leading to discovery
- **Conclusion**: Summary of mathematical findings
- **Extension**: Further exploration opportunities

### E. TECHNOLOGY INTEGRATION (Target: {word_limits.get('technology', 200)} words per section)

**Using Technology: [Tool Name]**

- Specific instructions for using calculators/software
- Screenshots or detailed descriptions of inputs
- Interpretation of technological outputs
- Comparison with manual calculations
- When to use technology vs. manual methods

### F. EXERCISE SETS (Target: {word_limits.get('exercise', 400)} words per section)

**Exercise [Section Number]**

Organize questions in progressive difficulty:

1. **Skills Practice** (Questions 1-8)
   - Basic application of formulas and concepts
   - Drill exercises for fluency
   - Simple substitution problems

2. **Problem Solving** (Questions 9-15)
   - Multi-step problems
   - Application to real-world contexts
   - Synthesis of multiple concepts

3. **Mathematical Reasoning** (Questions 16-20)
   - Proof and justification problems
   - Analysis and evaluation tasks
   - Open-ended investigations

4. **Extension/Challenge** (Questions 21-25)
   - Higher-order thinking problems
   - Competition-style questions
   - Cross-curricular applications

**Question Types to Include:**
- Multiple choice (with detailed explanations)
- Short answer calculations
- Extended response problems
- Graphical/visual problems
- Real-world modeling tasks
- Proof and reasoning questions

### CHAPTER REVIEW AND ASSESSMENT

## Chapter [Number] Review

### A. CHAPTER SUMMARY (Target: {word_limits.get('review', 200)} words)
- **Key Concepts**: Bullet-point summary of main ideas
- **Important Formulas**: Comprehensive formula list
- **Key Skills**: Summary of mathematical skills developed

### B. REVIEW EXERCISES (Target: {word_limits.get('review', 500)} words)

**Mixed Practice**
- 20-25 questions covering all chapter concepts
- Progressive difficulty levels
- Integration of multiple concepts
- Real-world application problems

**Question Categories:**
1. **Concept Review** (Questions 1-8)
2. **Skills Application** (Questions 9-16)  
3. **Problem Solving** (Questions 17-22)
4. **Challenge Problems** (Questions 23-25)

### C. CHAPTER TEST (Target: 300 words)
- Formal assessment covering all learning objectives
- Time allocation guidelines
- Mark allocation clearly shown
- Balanced coverage of all sections

### SPECIAL HAESE AND HARRIS FEATURES

### Mathematical Modeling Projects
- Extended real-world applications
- Cross-curricular connections
- Use of mathematical software/calculators
- Presentation and communication components

### Historical Context Boxes
- Brief historical notes about mathematical development
- Famous mathematicians and their contributions
- Evolution of mathematical concepts

### Career Connections
- Real-world applications in various careers
- Mathematical professions and pathways
- Industry applications of chapter concepts

### Common Errors and Misconceptions
- Typical student mistakes highlighted
- Correct mathematical reasoning emphasized
- Prevention strategies provided

**FORMATTING REQUIREMENTS:**
- Use clear section numbering (X.Y format)
- Include mathematical formatting with LaTeX notation
- Provide detailed image prompts for diagrams and graphs
- Use consistent mathematical terminology
- Include cross-references between sections
- Maintain logical flow and progression

**MATHEMATICAL STANDARDS:**
- Ensure mathematical accuracy and rigor
- Use appropriate mathematical language
- Include proper mathematical notation
- Provide complete solutions and working
- Connect to curriculum standards

**IMPORTANT**: For each major section in the PDF (e.g., Section 1, Section 2, etc.), apply the complete Haese and Harris structure above. DO NOT stop after one section - generate comprehensive content for ALL sections found in the document.

IMPORTANT: Keep the total word count to approximately {word_limits.get('total', 8000)} words per major section. Adjust the detail level accordingly.

For mathematical examples, ALWAYS show each step on a new line with proper spacing, use LaTeX notation ($...$) for all mathematical expressions, and show the work clearly with explanation first, then mathematical expression on the next line.

Generate ONLY the comprehensive mathematics chapter content in Markdown format following the Haese and Harris methodology. Include EVERY section found in the PDF document with complete structure for each.
"""

def create_math_exercises_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a mathematics-specific exercises prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'exercises': 800
        }
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPREHENSIVE MATHEMATICS EXERCISES based on the chapter content in the PDF.
**Target Total Word Count for All Exercises**: {word_limits.get('exercises', 800)} words

**Core Mathematical Exercise Types:**
1. **MCQ (Multiple Choice Questions)** - at least 12 questions with detailed solutions
2. **Short Answer Mathematical Problems** - at least 8 questions  
3. **Long Answer Mathematical Problems** - at least 5 questions
4. **Assertion & Reason (Mathematical)** - at least 5 questions
5. **True/False with Mathematical Justification** - at least 10 statements
6. **Fill in the Blanks (Mathematical)** - at least 10 questions
7. **Match the Following (Mathematical)** - at least 2 sets with 5 matches each
8. **Mathematical Concept Mapping** - at least 1 comprehensive exercise
9. **Mathematical Case Studies** - at least 2 scenarios
10. **Open-ended Mathematical Problems** - at least 3 questions

**Special Mathematical Features:**
- **EeeBee Integration**: Include questions where EeeBee guides mathematical thinking
- **Technology Integration**: Questions involving mathematical software, calculators, or digital tools
- **Differentiation**: Include challenge problems for advanced learners and support questions for revision
- **21st Century Skills**: Collaborative mathematical problems and communication-based questions

Ensure that:
* Questions cover ALL important mathematical concepts from the PDF
* Questions follow Bloom's Taxonomy at various levels
* Mathematical language is clear and appropriate for {grade_level}
* Questions increase in difficulty progressively
* All exercises include detailed mathematical solutions with step-by-step working
* Content is formatted in Markdown with proper mathematical notation

Provide ONLY the comprehensive mathematical exercises in Markdown format.
"""

def create_math_skills_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a mathematics-specific skills and STEM activities prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'skill_activity': 400,
            'stem_activity': 400
        }
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate MATHEMATICAL SKILL-BASED ACTIVITIES and STEM projects.

## Mathematical Skill-Based Activities (Target: {word_limits.get('skill_activity', 400)} words)

Create at least 3 comprehensive mathematical activities:

**Activity Structure for Each:**
1. **Clear Mathematical Objective**
2. **Materials Required**
3. **Step-by-Step Mathematical Procedure**
4. **Inquiry-Based Mathematical Exploration**
5. **Real-Life Mathematical Connections**
6. **EeeBee Integration**
7. **Mathematical Reflection Questions**
8. **Expected Mathematical Outcomes**

**Types of Mathematical Activities:**
- Mathematical Experiments/Investigations
- Mathematical Manipulative Activities
- Mathematical Problem-Solving Challenges
- Mathematical Pattern Recognition Activities
- Mathematical Measurement and Data Collection

## Mathematical STEM Projects (Target: {word_limits.get('stem_activity', 400)} words)

Create at least 2 comprehensive projects that integrate mathematics with Science, Technology, and Engineering:

**Project Structure for Each:**
1. **Mathematical Integration Focus**
2. **Real-World Mathematical Application**
3. **Technology Integration**
4. **Collaborative Mathematical Work**
5. **Mathematical Design Challenge**
6. **Mathematical Analysis and Calculations**
7. **Mathematical Presentation Component**
8. **Assessment Criteria**

**STEM Integration Examples:**
- Mathematical Modeling in Science
- Engineering Design with Mathematical Constraints
- Technology-Enhanced Mathematical Problem Solving
- Mathematical Data Analysis Projects

Format the content in Markdown with proper mathematical notation.

Provide ONLY the Mathematical Skill Activities and STEM Projects in Markdown format.
"""

def create_math_art_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a mathematics-specific art integration prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'art_learning': 400
        }
    return f"""You are an expert in mathematical education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a mathematics book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate MATHEMATICS-INTEGRATED CREATIVE LEARNING projects.
**Target Total Word Count for All Art-Integrated Learning**: {word_limits.get('art_learning', 400)} words

## Mathematical Art Projects

Create at least 3 creative projects that connect mathematical concepts to various art forms:

**Project Structure for Each:**
1. **Mathematical Learning Objective**
2. **Art Form Integration**
3. **Mathematical Concepts Highlighted**
4. **Materials and Tools**
5. **Step-by-Step Mathematical-Artistic Process**
6. **EeeBee's Creative Guidance**
7. **Mathematical Reflection and Analysis**
8. **Showcase and Communication**

**Types of Mathematical Art Integration:**
- Geometric Art and Patterns
- Mathematical Music and Rhythm
- Mathematical Drama and Storytelling
- Mathematical Digital Art
- Mathematical Sculpture and 3D Models

## Mathematical Case Studies

**Case Study – Level 1 (Accessible Mathematical Analysis):**
Create at least 1 simpler case study that:
- Presents a real-world mathematical scenario appropriate for {grade_level}
- Includes guided mathematical analysis questions
- Is accessible to all students with basic mathematical skills

**Case Study – Level 2 (Advanced Mathematical Challenge):**
Create at least 1 more complex case study that:
- Challenges students with multi-step mathematical problems
- Requires deeper application of mathematical concepts
- Encourages higher-order mathematical thinking skills

**Case Study Structure for Each:**
1. **Real-World Mathematical Context**
2. **Mathematical Background Information**
3. **Guided Mathematical Questions**
4. **Mathematical Analysis Requirements**
5. **Creative Mathematical Solutions**
6. **Mathematical Communication Component**
7. **EeeBee's Mathematical Insights**
8. **Extension Mathematical Challenges**

Format the content in Markdown with proper mathematical notation.

Provide ONLY the Mathematics-Integrated Creative Learning content in Markdown format.
"""

# Science subject prompt functions
def create_science_chapter_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a science subject chapter content prompt with dynamic word limits"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'hook': 70,
            'learning_outcome': 70,
            'real_world': 50,
            'previous_class': 100,
            'history': 100,
            'current_concepts': 1200,
            'summary': 700,
            'link_learn': 250,
            'image_based': 250,
            'exercises': 800,
            'skill_activity': 400,
            'stem_activity': 400,
            'art_learning': 400
        }
    
    return f"""You are an expert in science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes.

You are analyzing a science book chapter intended for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines for Science education.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Science Syllabus)

Your task is to generate COMPREHENSIVE CORE SCIENCE CHAPTER CONTENT that should be equivalent to a complete science textbook chapter.

**REQUIRED SECTIONS (Generate ALL with substantial content):**

1. **Current Concepts** (Target: {word_limits.get('current_concepts', 1200)} words)
   
   For each major scientific concept in the chapter, include ALL of the following:
   
   **A. Concept Introduction** (Target: 120 words per concept)
   - Clear introduction to each scientific concept
   - Simple, clear scientific language
   - Use analogies and real-world scientific examples
   - Identify all subconcepts that will be covered under this main concept
   
   **B. Subconcepts (As per NCERT Science Books)** (Target: 500 - 600 words per concept)
   - **IMPORTANT**: Identify and include ALL subconcepts present in NCERT science books for this topic
   - Each main concept typically has 2-5 subconcepts in NCERT science books
   - Subconcepts should be clearly labeled and integrated within the main concept
   - Examples of subconcepts:
     * For "Matter": States of matter, Properties of matter, Changes of state, etc.
     * For "Light": Reflection, Refraction, Dispersion, etc.
     * For "Cell": Plant cells, Animal cells, Cell organelles, Cell division, etc.
   - Each subconcept should include:
     * Definition and explanation 
     * Scientific examples and illustrations
     * Key scientific points to remember
     * Common scientific errors to avoid
   
   **C. Scientific Explanation** (Target: 450 words per concept)
   - Detailed theoretical understanding of the scientific concept and its subconcepts
   - Include step-by-step scientific reasoning and explanations
   - Show different scientific approaches where applicable
   - Ensure each subconcept is thoroughly explained with connections to the main concept
   
   **D. Scientific Examples and Applications** (Target: 350 words per concept)
   - Provide 4-5 different real-world examples for each concept
   - Include examples that cover different subconcepts
   - Include step-by-step explanations with clear scientific reasoning
   - Use varied complexity levels and application formats
   - Ensure examples demonstrate application of all subconcepts
   
   **E. Concept-Based Questions** (Target: 300 words per concept)
   - Give questions for each current concept - 3 MCQs, 2 Short questions, 1 Long question
   - Questions should test different subconcepts
   - Include scientific reasoning and application-based questions
   
   **F. Scientific Activity** (Target: 200 words per concept)
   - Give scientific activity according to the concept
   - Step-by-step scientific experiments or investigations
   - Inquiry-based scientific activities
   
   **G. Scientific Facts and Applications** (Target: 150 words per concept)
   - Give fun scientific fact for each concept
   - Show relevance to daily life, scientific careers, technology
   - Connect to current scientific research and discoveries
   
   **H. Key Scientific Points** (Target: 120 words per concept)
   - Give key scientific points for each concept
   - Highlighted scientific terminology and definitions
   - Essential scientific principles and formulas
   
   **I. Common Scientific Misconceptions** (Target: 150 words per concept)
   - Give common scientific misconceptions for each concept
   - Correct early scientific misunderstandings
   - Explain why these misconceptions occur
   
   **J. Visual Integration** (Throughout each concept)
   - Include detailed image prompts for scientific diagrams, experiments, and illustrations
   
   - Mark concepts which are exactly coming from the PDF and give some extra concepts for higher level understanding
   - Make sure there is no repetition of concepts

2. **Hook (with Image Prompt)** (Target: {word_limits.get('hook', 80)} words)
   - Create an engaging scientific opening that captures student interest
   - Use scientific storytelling, surprising scientific facts, or thought-provoking scientific questions
   - Connect to students' daily scientific experiences or current scientific events
   - Include a detailed image prompt for a compelling scientific visual

3. **Learning Outcome** (Target: {word_limits.get('learning_outcome', 70)} words)
   - List specific, measurable scientific learning objectives
   - Use action verbs (analyze, evaluate, create, experiment, investigate, etc.)
   - Align with Bloom's Taxonomy levels for science education
   - Connect to CBSE science curriculum standards

4. **Real World Connection** (Target: {word_limits.get('real_world', 50)} words)
   - Provide multiple real-world applications of the scientific concepts
   - Include current examples from technology, environment, health, space science, etc.
   - Explain how the scientific concepts impact daily life
   - Connect to scientific careers and future studies

5. **Previous Class Concept** (Target: {word_limits.get('previous_class', 100)} words)
   - Give the scientific concept name and the previous class it was studied in according to NCERT science textbooks
   - Link to prior scientific knowledge and foundations

6. **History** (Target: {word_limits.get('history', 100)} words)
   - Provide comprehensive historical background of scientific discoveries
   - Include key scientists, inventors, or scientific figures and their contributions
   - Explain the timeline of scientific developments and discoveries
   - Connect historical scientific context to modern understanding

7. **Summary** (Target: {word_limits.get('summary', 700)} words)
   - Create detailed concept-wise scientific summaries (not just one overall summary)
   - Include key scientific points, formulas, and important scientific facts
   - Organize by individual scientific concepts covered in the chapter
   - Provide clear, concise explanations that reinforce scientific learning

8. **Link and Learn Based Question** (Target: {word_limits.get('link_learn', 250)} words)
   - Create 3-5 questions that connect different scientific concepts
   - Include questions that link to other science subjects or real-world scientific scenarios
   - Provide detailed explanations for the scientific connections

9. **Image Based Question** (Target: {word_limits.get('image_based', 250)} words)
   - Create 3-5 questions based on scientific images/diagrams from the chapter
   - Include detailed scientific image descriptions if creating new image prompts
   - Ensure questions test scientific understanding, not just observation

**CONTENT REQUIREMENTS:**
* **Scientific Accuracy**: Ensure all scientific content is accurate and up-to-date
* **Detailed Scientific Explanations**: Each scientific concept should be explained thoroughly with multiple paragraphs
* **Scientific Examples and Illustrations**: Include numerous scientific examples, case studies, and practical applications
* **Age-Appropriate Scientific Language**: Use scientific vocabulary suitable for {grade_level} but don't oversimplify
* **Engaging Scientific Tone**: Write in an engaging, conversational style that maintains student interest in science
* **Clear Scientific Structure**: Use proper headings, subheadings, and formatting for science content
* **Visual Integration**: Include detailed image prompts for scientific diagrams, experiments, and illustrations

**FORMATTING REQUIREMENTS:**
* Use Markdown formatting with clear headings (# ## ###)
* Include bullet points and numbered lists where appropriate for scientific content
* Use **bold** for key scientific terms and *italics* for emphasis
* Create well-structured paragraphs (3-5 sentences each) explaining scientific concepts
* Include image prompts marked as: [PROMPT FOR NEW IMAGE: detailed scientific description]

**SCIENTIFIC QUALITY STANDARDS:**
* Each section should be scientifically substantial and comprehensive
* Avoid superficial coverage - go deep into each scientific topic
* Include multiple scientific perspectives and approaches to concepts
* Ensure content flows logically from one scientific concept to the next
* Maintain consistency in scientific terminology and explanations
* Connect scientific concepts to real-world applications and current scientific research
* Include age-appropriate scientific investigations and inquiry-based learning

Analyze the PDF document thoroughly and create improved scientific content that expands significantly on what's provided while maintaining all original concept names and terminology.

Provide ONLY the comprehensive science chapter content in Markdown format. Do not include exercises, activities, or art projects.
"""

def create_science_exercises_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a science subject exercises prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'exercises': 800
        }
    return f"""You are an expert in science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPREHENSIVE SCIENCE EXERCISES based on the chapter content in the PDF.
**Target Total Word Count for All Exercises**: {word_limits.get('exercises', 800)} words

**Core Science Exercise Types:**

Create the following science exercise types:
1. **MCQ (Multiple Choice Questions)** - at least 10 questions with scientific reasoning
2. **Assertion and Reason (Scientific)** - at least 5 questions testing scientific logic
3. **Fill in the Blanks (Scientific)** - at least 10 questions focusing on scientific terms
4. **True False (Scientific)** - at least 10 statements with scientific justification
5. **Define the following scientific terms** - at least 10 terms from the chapter
6. **Match the column (Scientific)** - at least 2 sets with 5 scientific matches each
7. **Give Reason for the following Scientific Statement (Easy Level)** - at least 5 questions
8. **Answer in Brief (Moderate Level)** - at least 5 questions on scientific concepts
9. **Answer in Detail (Hard Level)** - at least 5 questions requiring scientific explanation

**Special Science Features:**
- **Scientific Investigation Questions**: Include questions that require scientific thinking
- **Experimental Design**: Questions about planning and conducting scientific experiments
- **Data Analysis**: Questions involving interpretation of scientific data and graphs
- **Real-world Applications**: Connect scientific concepts to everyday phenomena
- **Cross-curricular Connections**: Link science to mathematics, technology, and environment

Ensure that:
* Questions cover ALL important scientific concepts from the PDF
* Questions follow Bloom's Taxonomy at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create)
* Scientific language is clear and appropriate for {grade_level}
* Questions increase in difficulty from basic scientific recall to higher-order scientific thinking
* All exercises include correct answers or model scientific solutions
* Content is formatted in Markdown with proper scientific notation
* Questions encourage scientific inquiry and critical thinking
* Include questions that test understanding of scientific processes and methods

Do NOT directly copy questions from the PDF. Create new, original scientific questions based on the content.

Provide ONLY the comprehensive science exercises in Markdown format.
"""

def create_science_skills_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a science subject skills and STEM activities prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'skill_activity': 400,
            'stem_activity': 400
        }
    return f"""You are an expert in science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate SCIENCE SKILL-BASED ACTIVITIES and STEM projects based on the chapter content in the PDF.

## Science Skill-Based Activities (Target: {word_limits.get('skill_activity', 400)} words)

Create at least 3 hands-on science activities that:
* Reinforce the key scientific concepts from the chapter
* Develop practical scientific skills relevant to the subject
* Can be completed with easily available scientific materials
* Include clear step-by-step scientific procedures
* Encourage scientific observation and data collection
* Connect to scientific method and inquiry-based learning

**Activity Structure for Each:**
1. **Clear Scientific Objective**
2. **Scientific Materials Required**
3. **Step-by-Step Scientific Procedure**
4. **Scientific Observations to Record**
5. **Data Collection and Analysis**
6. **Safety Precautions for Science Activities**
7. **Scientific Reflection Questions**
8. **Expected Scientific Outcomes and Learning**

## Science STEM Projects (Target: {word_limits.get('stem_activity', 400)} words)

Create at least 2 comprehensive STEM projects that:
* Integrate Science, Technology, Engineering and Mathematics
* Connect to real-world scientific applications of the chapter concepts
* Encourage scientific problem-solving and critical thinking
* Include scientific materials needed, procedure, and expected outcomes
* Promote scientific investigation and evidence-based conclusions
* Link to current scientific research and technological developments

**STEM Project Structure for Each:**
1. **Scientific Integration Focus**
2. **Real-World Scientific Application**
3. **Technology Integration in Science**
4. **Engineering Design Challenge**
5. **Mathematical Analysis of Scientific Data**
6. **Scientific Methodology and Process**
7. **Scientific Communication and Presentation**
8. **Assessment Criteria for Scientific Understanding**

**Special Science Features:**
- **Scientific Investigation**: Include projects that require hypothesis formation and testing
- **Environmental Connections**: Link activities to environmental science and sustainability
- **Technology Integration**: Use digital tools for data collection and analysis
- **Career Connections**: Connect to science, technology, engineering careers
- **Safety First**: Emphasize scientific safety protocols in all activities

For each activity/project, include:
* A clear scientific title and learning objective
* Complete list of scientific materials required
* Detailed scientific procedure with numbered steps
* Safety precautions for scientific activities
* Scientific observation charts and data collection sheets
* Reflection questions that promote scientific thinking
* Expected scientific outcomes and learning points
* Connections to scientific careers and real-world applications

Format the content in Markdown with proper scientific headings, lists, and organization.

Provide ONLY the Science Skill Activities and STEM Projects in Markdown format.
"""

def create_science_art_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a science subject art integration prompt with dynamic word limits"""
    if word_limits is None:
        word_limits = {
            'art_learning': 400
        }
    return f"""You are an expert in science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate SCIENCE-INTEGRATED CREATIVE LEARNING projects based on the chapter content in the PDF.
**Target Total Word Count for All Art-Integrated Learning**: {word_limits.get('art_learning', 400)} words

## Science Art Projects

Create at least 3 creative projects that connect scientific concepts to various art forms:

**Project Structure for Each:**
1. **Scientific Learning Objective**
2. **Art Form Integration with Science**
3. **Scientific Concepts Highlighted**
4. **Materials and Scientific Tools**
5. **Step-by-Step Scientific-Artistic Process**
6. **Scientific Observation and Documentation**
7. **Scientific Reflection and Analysis**
8. **Showcase and Scientific Communication**

**Types of Science Art Integration:**
- **Scientific Visualization**: Create artistic representations of scientific processes
- **Science Through Music**: Use rhythm and sound to represent scientific patterns
- **Scientific Drama**: Act out scientific processes and phenomena
- **Digital Science Art**: Use technology to create scientific visualizations
- **Scientific Models and Sculptures**: 3D representations of scientific concepts

## Scientific Case Studies

**Case Study – Level 1 (Accessible Scientific Analysis):**
Create at least 1 simpler case study that:
- Presents a real-world scientific scenario appropriate for {grade_level}
- Includes guided scientific analysis questions
- Is accessible to all students with basic scientific knowledge
- Connects to current scientific events or discoveries

**Case Study – Level 2 (Advanced Scientific Challenge):**
Create at least 1 more complex case study that:
- Challenges students with multi-step scientific problems
- Requires deeper application of scientific concepts
- Encourages higher-order scientific thinking skills
- Includes data analysis and scientific reasoning

**Scientific Case Study Structure for Each:**
1. **Real-World Scientific Context**
2. **Scientific Background Information**
3. **Guided Scientific Questions**
4. **Scientific Data and Evidence**
5. **Analysis Requirements**
6. **Creative Scientific Solutions**
7. **Scientific Communication Component**
8. **Extension Scientific Challenges**

**Special Science Integration Features:**
- **Environmental Connections**: Link art projects to environmental science
- **Scientific Method Integration**: Include hypothesis, observation, and conclusion
- **Technology Tools**: Use digital tools for scientific art creation
- **Scientific Communication**: Present findings through artistic expression
- **Cross-Curricular Connections**: Connect science to mathematics, geography, and history

For each project/case study, include:
* A clear scientific title and learning objective
* Complete list of materials needed (scientific and artistic supplies)
* Detailed scientific-artistic instructions or scenario description
* Scientific observation sheets and reflection questions
* Assessment criteria for both scientific understanding and creative expression
* Connections to scientific careers and real-world applications
* Safety considerations for scientific art activities

Format the content in Markdown with proper scientific and artistic headings, lists, and organization.

Provide ONLY the Science-Integrated Creative Learning content in Markdown format.
"""

# Computer Science specific prompt functions
def create_computer_chapter_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a computer science-specific chapter content prompt"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'hook': 80,
            'learning_outcome': 70,
            'real_world': 70,
            'previous_class': 100,
            'current_concepts': 4500,
            'summary': 100,
            'exercises': 1000,
            'skill_activity': 100
        }
    
    return f"""You are an expert in computer science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes.

You are analyzing a computer science book chapter intended for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines for Computer Science education.

**CRITICAL INSTRUCTION**: The PDF may contain MULTIPLE MAJOR SECTIONS (e.g., Section 1, Section 2, Section 3, etc.). You MUST include ALL sections present in the PDF. Do NOT stop after completing just one or two sections. Generate comprehensive content for EVERY section found in the document.

**Model Chapter Progression and Elements (Base Structure):**
---
{model_progression_text}
---

**Target Audience:** {grade_level} (CBSE Computer Science Syllabus)

Your task is to generate COMPREHENSIVE COMPUTER SCIENCE CHAPTER CONTENT following the Model Chapter Progression structure enhanced with computer science-specific elements.

**IMPORTANT**: If the PDF contains multiple sections, you MUST generate complete content for EACH section. The structure below should be applied to EACH major section in the PDF.

**REQUIRED SECTIONS (Generate ALL with substantial content for EACH major section in the PDF):**

## I. Chapter Opener

1. **Chapter Title** - Engaging and technology-focused

2. **Hook (with Image Prompt)** (Target: {word_limits.get('hook', 80)} words)
   - Create an engaging technological opening that captures student interest
   - Use real-life technology scenarios, digital innovations, or computational challenges
   - Connect to students' digital experiences and interests
   - Include a detailed image prompt for a compelling tech visual

3. **Real-World Connection** (Target: {word_limits.get('real_world', 70)} words)
   - Provide multiple real-world applications of the computer science concepts
   - Show how technology is transforming various industries
   - Include examples from gaming, apps, AI, robotics, web development, etc.
   - Connect to tech careers and future opportunities

4. **Learning Outcomes** (Target: {word_limits.get('learning_outcome', 70)} words)
   - List specific, measurable computer science learning objectives
   - Use action verbs (code, debug, design, implement, analyze, create, etc.)
   - Align with Bloom's Taxonomy for computational thinking
   - Connect to CBSE computer science curriculum standards

5. **Previous Class Link** (Target: {word_limits.get('previous_class', 100)} words)
   - Link to prior computer science knowledge from previous classes
   - Explain how previous concepts build into current learning
   - Provide a brief review of essential prerequisites

6. **Chapter Map/Overview** (100 words)
   - Visual layout of computer science concepts (flowchart or mind map description)
   - Show progression from basic to advanced concepts

7. **Meet the Character (EeeBee)** (30-50 words)
   - Introduce EeeBee as a digital guide/coding buddy throughout the chapter

## II. Core Content Sections (REPEAT FOR EACH MAJOR SECTION IN THE PDF)

8. **Introduction of Section** (150 words per section)
   - Give a related section introduction that sets the technological context
   - Explain the importance of the computer science concepts to be learned
   - Connect to the broader digital literacy curriculum

9. **Warm-up Activities** (100 words per section)
   - Create 5-7 engaging warm-up activities that connect to prior tech knowledge
   - Include unplugged activities, quick challenges, or digital puzzles

10. **Current Concepts** (Target: {word_limits.get('current_concepts', 4500)} words minimum PER SECTION)
    
    For each major concept in EACH section, include ALL of the following:
    
    **A. Concept Introduction** (100 words per concept)
    - Clear introduction to each computer science concept
    - Simple, accessible language with technical terms explained
    - Use analogies from everyday life
    
    **B. Technical Explanation** (400-500 words per concept)
    - Detailed theoretical understanding of the concept
    - Include technical specifications, algorithms, or processes
    - Show different approaches and methodologies
    
    **C. Hands-on Examples** (500-600 words per concept)
    - Provide 4-5 different practical examples
    - Include step-by-step implementations
    - Show code snippets (if applicable) with clear explanations
    - Use varied difficulty levels
    
    **D. Practice Exercises** (400-500 words per concept)
    Create the following specific computer science question types for each concept:
    
    1. **Code Completion** - 3-4 exercises
       - Fill in missing parts of code or algorithms
    
    2. **Debug the Code** - 2 exercises
       - Find and fix errors in given code snippets
    
    3. **Algorithm Design** - 3 exercises
       - Design solutions for given problems
    
    4. **Output Prediction** - 3 exercises
       - Predict the output of given code
    
    5. **Tech Application Questions** - 5 questions
       - Real-world scenarios using the concept
    
    6. **Computational Thinking Puzzles** - 3 questions
       - Logic puzzles and problem-solving challenges
    
    **E. Lab Activities** (300-400 words per concept)
    - Step-by-step computer lab exercises
    - Include both online and offline activities
    - Specify software/tools required
    
    **F. Visual Representations** (Throughout each concept)
    - Include detailed image prompts for flowcharts, diagrams, screenshots
    - User interface mockups where applicable
    
    **G. Try It Yourself** (200-250 words per concept)
    - Mini-projects or coding challenges
    - Encourage experimentation and creativity
    
    **H. Key Technical Terms** (100-150 words per concept)
    - Highlighted technical vocabulary
    - Clear definitions with examples
    
    **I. Tech Facts** (100-150 words per concept)
    - Interesting facts about the technology
    - Current trends and future possibilities
    
    **J. Quick Tips** (150-200 words per concept)
    - Best practices and shortcuts
    - Common mistakes to avoid
    
    **K. Fun With Colors** (50-80 words per concept)
    - Color-coded programming concepts or visual elements
    - Creative activities using colors to understand technical concepts
    
    **L. Unscramble the Letters** (30-50 words per concept)
    - Word puzzles related to computer science terminology
    - Technical vocabulary building exercises
    
    **M. Health, Safety & Ethics Corner/Computer Etiquettes** (100-120 words per concept)
    - Digital wellness and ergonomics
    - Ethical considerations in technology use
    - Proper computer usage and digital citizenship
    
    **N. Shortcut Keys** (50-80 words per concept)
    - Relevant keyboard shortcuts for the concept
    - Efficiency tips for computer operations

## IV. Chapter Wrap-Up (For the ENTIRE chapter)

11. **Common Programming Errors** (80-100 words)
    - 2-3 common mistakes per concept
    - How to identify and fix them

12. **Summary/Points to Remember** (Target: {word_limits.get('summary', 100)} words)
    - Key takeaways from the chapter
    - Essential concepts and skills learned
    - Important technical points to remember

13. **Self-Assessment Checklist** (80-100 words)
    - Programming skills checklist
    - Concept understanding verification

14. **Chapter-wise Lab Exercise** (Target: {word_limits.get('exercises', 1000)} words)
    - Comprehensive Lab Project
    - Multiple Choice Questions (5)
    - Fill in the Blanks (5)
    - Abbreviations (5)
    - Debugging Exercises (3)
    - Match The Following (5)
    - Code Writing Tasks (3)
    - Technical Concept Mapping
    - True/False with Explanation (5)
    - Case Study Analysis (1)
    - Innovation Challenge (1)
    - Answer the Following Questions - Short (5)
    - Answer the Following Questions - Long (5)

15. **Apply Your Digital Skills** (Target: {word_limits.get('skill_activity', 100)} words)
    - Real-world technology project
    - Cross-curricular integration

**CONTENT REQUIREMENTS:**
* **CRITICAL**: Include ALL major sections from the PDF
* **Technical Accuracy**: Ensure all code and technical content is accurate
* **Age-Appropriate Complexity**: Match technical depth to grade level
* **Practical Focus**: Balance theory with hands-on practice
* **Progressive Learning**: Structure from basic to advanced
* Include actual code examples appropriate for the grade level
* DO NOT copy code directly from the PDF - create new examples

Provide ONLY the comprehensive computer science chapter content in Markdown format. Remember to include EVERY section found in the PDF document.
"""

def create_computer_exercises_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a computer science-specific exercises prompt"""
    if word_limits is None:
        word_limits = {
            'exercises': 800
        }
    return f"""You are an expert in computer science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a computer science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPREHENSIVE COMPUTER SCIENCE EXERCISES based on the chapter content in the PDF.
**Target Total Word Count for All Exercises**: {word_limits.get('exercises', 800)} words

**Core Computer Science Exercise Types:**
1. **MCQ (Multiple Choice Questions)** - at least 12 questions covering theory and concepts
2. **Code Completion Exercises** - at least 8 exercises
3. **Debug the Code** - at least 5 exercises with errors to fix
4. **Write the Code** - at least 5 programming tasks
5. **Algorithm Design** - at least 5 problems requiring algorithmic solutions
6. **Output Prediction** - at least 8 code snippets to analyze
7. **True/False with Technical Justification** - at least 10 statements
8. **Fill in the Blanks (Technical)** - at least 10 questions
9. **Match the Following (Tech Terms)** - at least 2 sets with 5 matches each
10. **Flowchart/Diagram Exercises** - at least 3 exercises
11. **Case Studies (Tech Scenarios)** - at least 2 real-world scenarios
12. **Project-Based Questions** - at least 3 mini-project ideas

**Special Computer Science Features:**
- **Lab Exercises**: Step-by-step practical implementations
- **EeeBee Integration**: Include exercises where EeeBee guides problem-solving
- **Digital Tools**: Questions involving specific software/platforms
- **Cyber Ethics**: Include digital citizenship and ethics questions
- **Differentiation**: Basic to advanced level questions

**Programming Languages** (adjust based on grade):
- For lower grades: Block-based coding, Scratch concepts
- For middle grades: Basic Python, HTML/CSS
- For higher grades: Python, Java, SQL, Web technologies

Ensure that:
* Questions cover ALL important computer science concepts from the PDF
* Include practical coding exercises, not just theory
* Questions progress from basic recall to complex problem-solving
* Language and complexity are appropriate for {grade_level}
* All exercises include detailed solutions with explanations
* Code examples use proper syntax and best practices

Provide ONLY the comprehensive computer science exercises in Markdown format.
"""

def create_computer_skills_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a computer science-specific skills and lab activities prompt"""
    if word_limits is None:
        word_limits = {
            'skill_activity': 400,
            'stem_activity': 400
        }
    return f"""You are an expert in computer science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a computer science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPUTER SCIENCE SKILL-BASED ACTIVITIES and LAB PROJECTS.

## Computer Science Skill-Based Activities (Target: {word_limits.get('skill_activity', 400)} words)

Create at least 3 comprehensive hands-on activities:

**Activity Structure for Each:**
1. **Clear Technical Objective**
2. **Software/Hardware Requirements**
3. **Step-by-Step Implementation**
4. **Code Templates (if applicable)**
5. **Testing and Debugging Steps**
6. **Real-World Applications**
7. **EeeBee's Coding Tips**
8. **Reflection and Documentation**
9. **Expected Output/Results**

**Types of Computer Science Activities:**
- Programming Challenges
- App/Game Development Mini-Projects
- Web Development Tasks
- Database Design Exercises
- Algorithm Implementation
- Digital Art/Animation Creation
- Robotics/IoT Projects (conceptual if hardware unavailable)

## Computer Lab Projects (STEM Integration) (Target: {word_limits.get('stem_activity', 400)} words)

Create at least 2 comprehensive lab projects that integrate computer science with other STEM fields:

**Project Structure for Each:**
1. **Project Title and Overview**
2. **Technical Requirements**
3. **Learning Objectives**
4. **Detailed Implementation Steps**
5. **Code Structure and Templates**
6. **Testing Procedures**
7. **Troubleshooting Guide**
8. **Extensions and Modifications**
9. **Presentation Guidelines**
10. **Assessment Rubric**

**STEM Integration Examples:**
- Data Science Projects (CS + Math)
- Simulation Projects (CS + Science)
- Engineering Design with Code
- Digital Solutions for Real Problems
- AI/ML Demonstrations (age-appropriate)

**Digital Citizenship Component:**
Include for each activity:
- Ethical considerations
- Copyright and licensing awareness
- Online safety practices
- Responsible technology use

Format the content in Markdown with proper code formatting using triple backticks for code blocks.

Provide ONLY the Computer Science Skill Activities and Lab Projects in Markdown format.
"""

def create_computer_art_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a computer science-specific creative and case study prompt"""
    if word_limits is None:
        word_limits = {
            'art_learning': 400
        }
    return f"""You are an expert in computer science education content development, specifically for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

You are analyzing a computer science book chapter intended for **{grade_level} (CBSE)**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

Your task is to generate COMPUTER SCIENCE CREATIVE PROJECTS and CASE STUDIES.
**Target Total Word Count for All Art-Integrated Learning**: {word_limits.get('art_learning', 400)} words

## Digital Creativity Projects

Create at least 3 creative projects that blend computer science with artistic expression:

**Project Structure for Each:**
1. **Creative Tech Objective**
2. **Tools and Technologies Required**
3. **Technical Concepts Applied**
4. **Step-by-Step Creative Process**
5. **Code Components (if applicable)**
6. **EeeBee's Creative Coding Tips**
7. **Showcase and Portfolio Building**
8. **Peer Review Guidelines**

**Types of Digital Creative Projects:**
- Digital Art and Graphics Programming
- Interactive Storytelling with Code
- Game Design and Development
- Animation and Multimedia Projects
- Music and Sound with Code
- AR/VR Concepts (age-appropriate)
- Creative Web Design
- Data Visualization Art

## Technology Case Studies

**Case Study – Level 1 (Foundational Analysis):**
Create at least 1 accessible case study that:
- Presents a real-world technology scenario appropriate for {grade_level}
- Includes guided analysis of technical solutions
- Focuses on problem-solving with technology
- Is accessible to students with basic computer skills

**Case Study – Level 2 (Advanced Challenge):**
Create at least 1 complex case study that:
- Presents multi-faceted technical challenges
- Requires system design thinking
- Involves multiple technologies/concepts
- Encourages innovative solutions

**Case Study Structure for Each:**
1. **Real-World Tech Scenario**
2. **Problem Statement**
3. **Technical Background**
4. **Current Technology Analysis**
5. **Proposed Digital Solutions**
6. **Implementation Considerations**
7. **Code Snippets/Pseudocode**
8. **Ethical and Social Impact**
9. **EeeBee's Tech Insights**
10. **Future Enhancements**

**Innovation Challenges:**
Include for each case study:
- "What if?" scenarios
- Brainstorming prompts
- Prototype development ideas
- Presentation formats

**Digital Portfolio Component:**
- How to document projects
- Building a digital portfolio
- Showcasing work online safely

Format the content in Markdown with proper formatting for code examples.

Provide ONLY the Computer Science Creative Projects and Case Studies in Markdown format.
"""

# English Communication & Grammar (Classes 1-8) specific prompt functions
def create_english_chapter_prompt(grade_level, word_limits=None):
    """Creates an English Communication & Grammar chapter prompt with Oxford/Cambridge standards"""
    # Extract class number for appropriate structure
    class_num = int(grade_level.split()[-1])
    
    # Default word limits based on class level
    if word_limits is None:
        if class_num <= 3:
            word_limits = {
                'warm_up': 200,
                'vocabulary': 300,
                'grammar_intro': 250,
                'practice': 400,
                'communication': 300,
                'exercises': 600
            }
        elif class_num <= 5:
            word_limits = {
                'introduction': 250,
                'reading': 600,
                'grammar': 500,
                'vocabulary': 350,
                'writing': 450,
                'speaking': 350,
                'exercises': 700
            }
        else:
            word_limits = {
                'introduction': 300,
                'reading': 800,
                'grammar': 600,
                'vocabulary': 400,
                'writing': 600,
                'speaking': 400,
                'literature': 500,
                'exercises': 900,
                'projects': 300
            }
    
    if class_num <= 3:
        return f"""You are an Expert English Language Educator specializing in foundational English Communication & Grammar for Classes 1-3, following the best practices of Oxford, Cambridge, and Wren & Martin publications.

**CRITICAL INSTRUCTIONS**: This prompt is EXCLUSIVELY for English Communication & Grammar Classes 1-3. Create content that builds strong language foundations using play-based and interactive methodologies.

**Target Audience**: {grade_level} (Primary English Learners - Ages 6-9)

**Educational Philosophy**:
- **Oxford Method**: Systematic phonics and vocabulary building
- **Cambridge Approach**: Communicative language teaching with real-world contexts
- **Wren & Martin Foundation**: Clear grammar rules with simple explanations
- **Indian Context**: Use familiar Indian names, places, festivals, and cultural elements

**MANDATORY CHAPTER STRUCTURE** (Transform PDF content to fit this EXACT structure):

## 1. Chapter Introduction with EeeBee (Target: {word_limits.get('warm_up', 200)} words)
- Chapter number and colorful title
- EeeBee's friendly greeting and welcome message
- Learning Outcomes (3-4 simple, clear objectives)
- Visual chapter overview with illustrations

## 2. Lead In - Let's Begin! (Target: {word_limits.get('warm_up', 200)} words)
- Picture-based warm-up activity
- Interactive observation exercises
- Simple questions to activate prior knowledge
- Fun, engaging entry point to the lesson

## 3. Vocabulary Building - New Words, New Worlds (Target: {word_limits.get('vocabulary', 300)} words)
- 8-12 new words with:
  * Clear pronunciations
  * Simple definitions
  * Picture associations
  * Usage in simple sentences
- Indian context examples (mango, rickshaw, diwali, etc.)
- Word games and activities

## 4. Grammar Introduction - Language Rules Made Easy (Target: {word_limits.get('grammar_intro', 250)} words)
- One core grammar concept per chapter
- Visual representation of rules
- "Grammar Byte" boxes with key points
- Color-coded examples
- Step-by-step explanation with pictures

## 5. Practice Activities - Let's Try Together (Target: {word_limits.get('practice', 400)} words)
- Exercise A: Basic identification/matching with pictures
- Exercise B: Fill in the blanks with visual clues
- Exercise C: Circle/underline activities
- Progressive difficulty from guided to independent
- All exercises with Indian cultural context

## 6. Communication Skills - Let's Talk and Share (Target: {word_limits.get('communication', 300)} words)
- Speaking activities with picture prompts
- Simple listening exercises
- Show and tell opportunities
- Pair work and group activities
- "My Activity Corner" for creative expression

## 7. Assessment & Review - How Well Did I Learn? (Target: {word_limits.get('exercises', 600)} words)
- "Points to Remember" summary box
- Self-assessment checklist with smiley faces
- Fun review activities (puzzles, matching, coloring)
- "EeeBee Interactive Activities" section
- QR codes for digital practice
- Skills tracking (Listening, Speaking, Reading, Writing)

**CONTENT REQUIREMENTS**:
- **Simple, Clear Language**: Age-appropriate vocabulary
- **Visual Learning**: Heavy emphasis on pictures and colors
- **Indian Integration**: Names like Raj, Priya, Amit; places like Delhi, Mumbai
- **Progressive Learning**: Build from known to unknown
- **Interactive Elements**: Games, songs, activities
- **Positive Reinforcement**: Encouraging tone throughout

Transform the PDF content following Oxford precision, Cambridge engagement, and Wren & Martin clarity.

Provide ONLY the comprehensive English chapter content in Markdown format.
"""
    
    elif class_num <= 5:
        return f"""You are an Expert English Language Educator for Classes 4-5, following Oxford, Cambridge, and Wren & Martin excellence standards.

**Target Audience**: {grade_level} (Elementary English Learners - Ages 9-11)

**Educational Philosophy**:
- **Oxford Approach**: Structured grammar progression with clear explanations
- **Cambridge Method**: Integrated skills development with communicative focus
- **Wren & Martin Tradition**: Systematic grammar rules with practical application
- **CEFR Alignment**: A2 level competencies

**MANDATORY CHAPTER STRUCTURE**:

## 1. Chapter Introduction & Learning Outcomes (Target: {word_limits.get('introduction', 250)} words)
- Clear chapter number and engaging title
- EeeBee's introduction (mature but friendly tone)
- 3-4 specific, measurable learning outcomes
- Chapter roadmap showing skill integration
- Connection to previous learning

## 2. Lead In Activity - Discover & Explore (Target: {word_limits.get('reading', 600)} words)
- Thought-provoking opening (puzzles, jumbled text, incomplete ideas)
- Critical thinking questions ("Why doesn't this make sense?")
- Discovery-based learning approach
- Picture-supported but not picture-dependent activities
- Introduction of key concepts through exploration

## 3. Grammar Focus - Understanding Language (Target: {word_limits.get('grammar', 500)} words)
- Clear concept definitions with examples
- Step-by-step rule explanation
- "Grammar Byte" boxes highlighting key rules
- Progressive examples (simple → complex)
- Common errors and corrections
- Comparison with mother tongue where relevant
- I.Q. Test sections for advanced learners

## 4. Vocabulary Expansion - Words in Context (Target: {word_limits.get('vocabulary', 350)} words)
- 15-20 new words with:
  * Pronunciation guides
  * Contextual meanings
  * Word families and derivatives
  * Synonyms and antonyms
- Usage in varied sentence structures
- Indian and global contexts balanced
- Word formation exercises

## 5. Integrated Skills Practice (Target: {word_limits.get('writing', 450)} words)
**Reading Comprehension**:
- Age-appropriate passages (150-200 words)
- Mix of fiction and non-fiction
- Comprehension questions at different levels

**Writing Skills**:
- Guided writing (sentence → paragraph)
- Creative writing prompts
- Functional writing (letters, messages)
- Process writing approach

## 6. Speaking & Listening Activities (Target: {word_limits.get('speaking', 350)} words)
- Structured oral presentations ("Tell about yourself")
- Listening comprehension exercises
- Role-play and dialogues
- Group discussion topics
- Pronunciation practice
- "My Activity Corner" for interactive tasks

## 7. Comprehensive Assessment (Target: {word_limits.get('exercises', 700)} words)
- Varied exercise types:
  * Tick/cross for correct sentences
  * Error identification and correction
  * Sentence transformation
  * Gap filling with grammar focus
  * Matching and categorization
- "Points to Remember" comprehensive summary
- Self-assessment opportunities
- EeeBee Interactive Activities with QR codes
- LSRW skills integration chart
- Progressive difficulty levels

**CONTENT SPECIFICATIONS**:
- **Skill Level Indicators**: Mark exercises as A2/B1 level
- **Balanced Approach**: Equal emphasis on all four skills
- **Cultural Sensitivity**: Indian contexts with global awareness
- **Digital Integration**: References to Communication Lab
- **Critical Thinking**: "Why" questions and analytical tasks
- **Age-Appropriate Complexity**: Challenging but achievable
- **Grammar Depth**: Sentence structure, tenses, parts of speech

Transform the PDF content maintaining academic rigor while ensuring engagement.

Provide ONLY the comprehensive English chapter content in Markdown format.
"""
    
    else:  # Classes 6-8
        return f"""You are an Expert English Language Educator for Classes 6-8, following Oxford, Cambridge, and Wren & Martin excellence standards.

**Target Audience**: {grade_level} (Middle School English Learners - Ages 11-14)

**Educational Philosophy**:
- **Oxford Excellence**: Academic rigor with systematic progression
- **Cambridge Innovation**: Critical thinking and real-world application
- **Wren & Martin Mastery**: Comprehensive grammar with advanced usage
- **CEFR Alignment**: B1-B2 level competencies

**MANDATORY CHAPTER STRUCTURE**:

## 1. Chapter Introduction & Context Setting (Target: {word_limits.get('introduction', 300)} words)
- Sophisticated chapter title with thematic focus
- Learning objectives aligned with board examinations
- Chapter overview with interdisciplinary connections
- EeeBee as a study companion (minimal, mature presence)
- Pre-reading questions to activate critical thinking
- Skills to be developed clearly outlined

## 2. Reading & Critical Analysis (Target: {word_limits.get('reading', 800)} words)
**Main Reading Text**:
- Complex passages (400-500 words)
- Literary and non-literary texts
- Multiple text types (narrative, expository, argumentative)
- Global themes with Indian perspectives

**Comprehension & Analysis**:
- Literal, inferential, and critical questions
- Text analysis and interpretation
- Author's purpose and tone identification
- Literary device recognition

## 3. Advanced Grammar & Language Usage (Target: {word_limits.get('grammar', 600)} words)
- Complex grammar concepts:
  * Advanced tenses and aspects
  * Clauses and sentence types
  * Voice and narration
  * Reported speech
  * Modals and conditionals
- "Grammar in Context" showing real usage
- Error analysis and correction
- Grammar for effective communication
- Comparison with other languages
- Advanced punctuation rules

## 4. Vocabulary & Etymology (Target: {word_limits.get('vocabulary', 400)} words)
- 20-30 advanced vocabulary items
- Word etymology and roots (Greek/Latin origins)
- Academic word list integration
- Collocations and phrasal verbs
- Idioms and expressions
- Register and appropriacy
- Vocabulary for specific purposes
- Word transformation exercises

## 5. Advanced Writing Skills (Target: {word_limits.get('writing', 600)} words)
**Composition Types**:
- Essay writing (descriptive, narrative, argumentative)
- Report writing and data interpretation
- Creative writing with literary techniques
- Email and formal letter writing
- Note-making and summarization

**Writing Process**:
- Planning and outlining
- Drafting and revising
- Peer review integration
- Style and tone adaptation

## 6. Speaking & Presentation Excellence (Target: {word_limits.get('speaking', 400)} words)
- Formal presentation skills
- Debate and argumentation
- Group discussion strategies
- Interview techniques
- Multimedia presentation integration
- Public speaking confidence
- Accent and intonation practice
- Extempore speaking

## 7. Literature Appreciation & Analysis (Target: {word_limits.get('literature', 500)} words)
- Poetry analysis and appreciation
- Prose interpretation
- Drama elements and performance
- Literary movements introduction
- Indian English literature
- Comparative literature concepts
- Creative response to literature
- Critical reviews and analysis

## 8. Comprehensive Assessment & Evaluation (Target: {word_limits.get('exercises', 900)} words)
**Question Types**:
- Multiple choice with justification
- Short answer (2-3 sentences)
- Long answer (150-200 words)
- Reference to context questions
- Grammar transformation exercises
- Integrated grammar and vocabulary
- Error correction passages
- Cloze tests and gap filling

**Additional Components**:
- Board exam pattern questions
- Previous years' question analysis
- Self-evaluation rubrics
- Peer assessment guidelines
- Digital assessment integration

## 9. Extended Learning & Projects (Target: {word_limits.get('projects', 300)} words)
- Research-based projects
- Collaborative assignments
- Technology integration tasks
- Real-world application projects
- Cross-curricular connections
- Portfolio development
- Presentation projects
- Community engagement tasks

**ADVANCED SPECIFICATIONS**:
- **Academic Language**: Formal register with subject-specific terminology
- **Exam Preparation**: Alignment with board examination patterns
- **21st Century Skills**: Critical thinking, collaboration, creativity
- **Digital Literacy**: Online research, digital presentations
- **Global Competence**: International perspectives with Indian roots
- **Differentiation**: Challenge activities for gifted learners
- **Study Skills**: Note-taking, summarizing, mind-mapping
- **Assessment Focus**: Continuous and comprehensive evaluation

Transform the PDF content to meet the rigorous academic standards expected at the middle school level.

Provide ONLY the comprehensive English chapter content in Markdown format.
"""

def create_english_exercises_prompt(grade_level, word_limits=None):
    """Creates comprehensive English exercises following Oxford/Cambridge/Wren & Martin standards"""
    class_num = int(grade_level.split()[-1])
    
    if word_limits is None:
        word_limits = {'exercises': 800 if class_num <= 5 else 1000}
    
    if class_num <= 3:
        return f"""You are an Expert English Language Assessment Specialist for {grade_level}, following Oxford, Cambridge, and Wren & Martin assessment standards for primary learners.

**Target Word Count**: {word_limits.get('exercises', 800)} words

Create COMPREHENSIVE ENGLISH EXERCISES based on the chapter content in the PDF.

**PRIMARY LEVEL EXERCISE STRUCTURE** (Classes 1-3):

## A. Picture-Based Vocabulary (25% of content)
- Match words to pictures with Indian contexts
- Circle the correct word for given pictures
- Fill in missing letters with visual clues
- Word-picture association games

## B. Simple Grammar Practice (25% of content)
- Fill in the blanks with picture support
- Circle/underline correct options
- Match sentence parts with visual aids
- Grammar games and activities

## C. Reading & Comprehension (25% of content)
- Simple passages (50-75 words) with pictures
- Tick/cross questions about the story
- Choose the right answer from pictures
- Complete sentences based on reading

## D. Fun Writing Activities (25% of content)
- Complete sentences with given words
- Write 2-3 sentences about pictures
- Simple creative writing with prompts
- Copy and complete activities

**Special Features**:
- EeeBee Interactive sections
- Smiley face self-assessment
- Coloring and drawing activities
- Indian cultural contexts (festivals, food, places)

Provide ONLY the comprehensive English exercises in Markdown format.
"""
    
    elif class_num <= 5:
        return f"""You are an Expert English Language Assessment Specialist for {grade_level}, following Oxford, Cambridge, and Wren & Martin assessment standards for elementary learners.

**Target Word Count**: {word_limits.get('exercises', 800)} words

Create COMPREHENSIVE ENGLISH EXERCISES based on the chapter content in the PDF.

**ELEMENTARY LEVEL EXERCISE STRUCTURE** (Classes 4-5):

## A. Advanced Vocabulary (20% of content)
- Word meanings, synonyms, antonyms
- Word families and derivatives
- Usage in varied sentence structures
- Contextual vocabulary with Indian examples

## B. Grammar Mastery (25% of content)
- Error identification and correction
- Sentence transformation exercises
- Grammar rules application
- Progressive difficulty levels

## C. Reading Comprehension (25% of content)
- Passages (150-200 words) with mixed themes
- Multiple choice and short answer questions
- Inference and critical thinking questions
- Text analysis and interpretation

## D. Integrated Writing Skills (20% of content)
- Guided paragraph writing
- Creative writing prompts
- Functional writing (letters, messages)
- Process writing activities

## E. Communication Assessment (10% of content)
- Speaking activities and presentation topics
- Listening comprehension exercises
- Role-play and dialogue practice
- Group discussion topics

**Assessment Features**:
- A2/B1 level indicators
- Self-evaluation opportunities
- EeeBee Interactive Activities
- LSRW skills integration

Provide ONLY the comprehensive English exercises in Markdown format.
"""
    
    else:  # Classes 6-8
        return f"""You are an Expert English Language Assessment Specialist for {grade_level}, following Oxford, Cambridge, and Wren & Martin assessment standards for middle school learners.

**Target Word Count**: {word_limits.get('exercises', 1000)} words

Create COMPREHENSIVE ENGLISH EXERCISES based on the chapter content in the PDF.

**MIDDLE SCHOOL EXERCISE STRUCTURE** (Classes 6-8):

## A. Advanced Vocabulary & Etymology (15% of content)
- Word roots, prefixes, suffixes analysis
- Academic vocabulary and collocations
- Phrasal verbs and idiomatic expressions
- Register and appropriacy exercises

## B. Complex Grammar & Usage (20% of content)
- Advanced grammar transformation
- Error analysis and correction
- Sentence structure and syntax
- Grammar in context applications

## C. Critical Reading & Analysis (25% of content)
- Complex passages (400-500 words)
- Literary and non-literary text analysis
- Critical thinking and inference questions
- Author's purpose and tone identification

## D. Advanced Writing Skills (20% of content)
- Essay writing (descriptive, narrative, argumentative)
- Report writing and data interpretation
- Creative writing with literary techniques
- Formal and informal writing styles

## E. Literature Appreciation (10% of content)
- Poetry analysis and interpretation
- Prose comprehension and criticism
- Literary device identification
- Creative response to literature

## F. Communication Excellence (10% of content)
- Formal presentation and debate skills
- Interview and group discussion techniques
- Public speaking and confidence building
- Multimedia presentation integration

**Advanced Assessment Features**:
- Board exam pattern alignment
- B1-B2 level competencies
- Self and peer evaluation rubrics
- Digital assessment integration
- Previous years' question analysis

Provide ONLY the comprehensive English exercises in Markdown format.
"""

def create_english_skills_prompt(grade_level, word_limits=None):
    """Creates English skills activities following international best practices"""
    class_num = int(grade_level.split()[-1])
    
    if word_limits is None:
        word_limits = {'skills': 600 if class_num <= 5 else 800}
    
    if class_num <= 3:
        return f"""You are an Expert English Language Skills Developer for {grade_level}, specializing in foundational language skills for primary learners.

**Target Word Count**: {word_limits.get('skills', 600)} words

Create COMPREHENSIVE ENGLISH SKILLS ACTIVITIES based on the chapter content.

**PRIMARY SKILLS DEVELOPMENT** (Classes 1-3):

## 1. Listening Skills - Hear and Learn (25% of content)
- Simple audio descriptions with picture support
- Listen and circle/match activities
- Sound recognition and phonics practice
- Indian songs, rhymes, and stories
- EeeBee's audio adventures

## 2. Speaking Skills - Talk and Share (25% of content)
- Show and tell with picture prompts
- Simple role-play and dialogues
- Pronunciation practice with visual cues
- Confidence-building through games
- "My Activity Corner" speaking tasks

## 3. Reading Skills - Read and Discover (25% of content)
- Picture reading and story sequencing
- Simple text with visual support
- Reading aloud with expression
- Comprehension through pictures
- Indian cultural context stories

## 4. Writing Skills - Write and Create (25% of content)
- Tracing and copying activities
- Complete sentences with pictures
- Simple creative writing prompts
- Drawing and writing combinations
- Pattern writing and letter formation

**Special Features**:
- Visual learning emphasis
- Indian cultural integration
- Interactive games and activities
- Self-assessment with smiley faces
- EeeBee skill-building sections

Provide ONLY the comprehensive English skills activities in Markdown format.
"""
    
    elif class_num <= 5:
        return f"""You are an Expert English Language Skills Developer for {grade_level}, specializing in integrated language skills for elementary learners.

**Target Word Count**: {word_limits.get('skills', 600)} words

Create COMPREHENSIVE ENGLISH SKILLS ACTIVITIES based on the chapter content.

**ELEMENTARY SKILLS DEVELOPMENT** (Classes 4-5):

## 1. Listening Skills - Active Listening (25% of content)
- Structured listening comprehension
- Audio passages with questions
- Listening for specific information
- Indian accents and global contexts
- Communication Lab integration

## 2. Speaking Skills - Confident Communication (25% of content)
- Structured presentations ("Tell about yourself")
- Role-play and interactive dialogues
- Pronunciation and intonation practice
- Group discussion participation
- Interview and conversation skills

## 3. Reading Skills - Critical Reading (25% of content)
- Reading strategies and techniques
- Multiple text types and genres
- Comprehension at different levels
- Text analysis and interpretation
- Reading for pleasure and information

## 4. Writing Skills - Process Writing (25% of content)
- Guided writing from sentence to paragraph
- Creative and functional writing
- Planning, drafting, and revising
- Different writing purposes and audiences
- Grammar integration in writing

**Assessment Integration**:
- A2/B1 level skill indicators
- Self and peer evaluation
- LSRW skills tracking chart
- Digital skill development
- Portfolio-based assessment

Provide ONLY the comprehensive English skills activities in Markdown format.
"""
    
    else:  # Classes 6-8
        return f"""You are an Expert English Language Skills Developer for {grade_level}, specializing in advanced language skills for middle school learners.

**Target Word Count**: {word_limits.get('skills', 800)} words

Create COMPREHENSIVE ENGLISH SKILLS ACTIVITIES based on the chapter content.

**ADVANCED SKILLS DEVELOPMENT** (Classes 6-8):

## 1. Advanced Listening Skills (25% of content)
- Complex audio passages and lectures
- Note-taking during listening
- Listening for implied meaning and tone
- Academic listening strategies
- Multimedia listening integration

## 2. Presentation & Speaking Excellence (25% of content)
- Formal presentation skills development
- Debate and argumentation techniques
- Interview and group discussion mastery
- Public speaking confidence building
- Accent, intonation, and fluency practice

## 3. Critical Reading & Analysis (25% of content)
- Advanced reading strategies
- Literary and academic text analysis
- Critical thinking and evaluation skills
- Research and information literacy
- Comparative reading and synthesis

## 4. Advanced Writing Mastery (25% of content)
- Essay writing (descriptive, narrative, argumentative)
- Academic writing and research skills
- Creative writing with literary techniques
- Digital writing and presentation
- Peer review and collaborative writing

**21st Century Skills Integration**:
- Digital literacy and online research
- Collaborative learning projects
- Critical thinking and problem-solving
- Cross-curricular connections
- Global competence development

**Assessment Excellence**:
- B1-B2 level competency markers
- Board exam preparation alignment
- Self-directed learning strategies
- Portfolio and project assessment
- Technology-enhanced evaluation

Provide ONLY the comprehensive English skills activities in Markdown format.
"""

def create_english_projects_prompt(grade_level, word_limits=None):
    """Creates creative English language projects following international standards"""
    class_num = int(grade_level.split()[-1])
    
    if word_limits is None:
        word_limits = {'projects': 400 if class_num <= 5 else 600}
    
    if class_num <= 3:
        return f"""You are an Expert English Language Project Designer for {grade_level}, creating fun and engaging language projects for primary learners.

**Target Word Count**: {word_limits.get('projects', 400)} words

Create INNOVATIVE ENGLISH LANGUAGE PROJECTS based on the chapter content.

**PRIMARY PROJECT CATEGORIES** (Classes 1-3):

## 1. Creative Art & Language Project (35% of content)
- Picture story creation with simple sentences
- Drawing and coloring with English labels
- Craft activities with English instructions
- Indian festival celebration projects
- EeeBee adventure story creation

## 2. Fun Communication Project (35% of content)
- Show and tell with family photos
- Simple puppet shows and role-play
- Singing English songs and rhymes
- Family interview with basic questions
- Classroom presentation with pictures

## 3. Hands-On Learning Project (30% of content)
- Make and label classroom objects
- Create picture dictionaries
- Simple cooking with English instructions
- Garden project with plant names
- Community helpers identification

**Project Features**:
- Heavy visual and hands-on components
- Family involvement opportunities
- Indian cultural integration
- Simple language objectives
- Fun, play-based learning approach

Provide ONLY the comprehensive English language projects in Markdown format.
"""
    
    elif class_num <= 5:
        return f"""You are an Expert English Language Project Designer for {grade_level}, creating engaging language projects for elementary learners.

**Target Word Count**: {word_limits.get('projects', 400)} words

Create INNOVATIVE ENGLISH LANGUAGE PROJECTS based on the chapter content.

**ELEMENTARY PROJECT CATEGORIES** (Classes 4-5):

## 1. Creative Expression Project (30% of content)
- Story writing and illustration
- Drama performance and script writing
- Poetry creation and recitation
- Indian folktale adaptation
- Creative magazine or newsletter

## 2. Communication & Research Project (30% of content)
- Interview family members or community
- Create presentations on Indian culture
- Pen pal correspondence project
- School newspaper article writing
- Survey and data presentation

## 3. Digital & Multimedia Project (25% of content)
- Simple video creation with narration
- Digital storytelling with images
- Audio recording of stories/poems
- Online research with presentation
- Communication Lab integration

## 4. Collaborative Learning Project (15% of content)
- Group drama or skit performance
- Collaborative story writing
- Classroom debate preparation
- Team presentation on chosen topics
- Peer teaching and learning

**Assessment Integration**:
- Clear project guidelines and rubrics
- Self and peer evaluation components
- LSRW skills development focus
- Cultural sensitivity and awareness
- Process and product evaluation

Provide ONLY the comprehensive English language projects in Markdown format.
"""
    
    else:  # Classes 6-8
        return f"""You are an Expert English Language Project Designer for {grade_level}, creating sophisticated language projects for middle school learners.

**Target Word Count**: {word_limits.get('projects', 600)} words

Create INNOVATIVE ENGLISH LANGUAGE PROJECTS based on the chapter content.

**ADVANCED PROJECT CATEGORIES** (Classes 6-8):

## 1. Literary & Creative Projects (25% of content)
- Original short story or novella writing
- Poetry anthology with analysis
- Drama script writing and performance
- Literary magazine creation
- Creative response to literature

## 2. Research & Academic Projects (25% of content)
- In-depth research report on chosen topics
- Comparative literature analysis
- Historical or cultural investigation
- Academic presentation with citations
- Data analysis and interpretation

## 3. Digital & Technology Projects (25% of content)
- Multimedia presentation creation
- Website or blog development
- Video documentary production
- Podcast or audio program creation
- Digital storytelling with advanced tools

## 4. Community & Global Projects (25% of content)
- Community service with documentation
- Cultural exchange programs
- Social issue awareness campaigns
- Interview and oral history projects
- International collaboration initiatives

**Advanced Project Features**:
- Board exam preparation alignment
- 21st century skills development
- Global competence and awareness
- Technology integration and digital literacy
- Real-world application and relevance
- Cross-curricular connections
- Independent and collaborative options

**Assessment Excellence**:
- Comprehensive rubrics with criteria
- Self-reflection and peer evaluation
- Portfolio development and maintenance
- Presentation and defense components
- Process documentation and analysis

Provide ONLY the comprehensive English language projects in Markdown format.
"""

# Artificial Intelligence specific prompt functions
# CBSE Class 10 AI Curriculum Structure with Sub-units
CBSE_AI_CLASS_10_CURRICULUM = {
    "Unit 1: Introduction to Artificial Intelligence": {
        "sub_units": [
            {
                "name": "Foundational concepts of AI",
                "learning_outcomes": "Understand the concept of human intelligence and its various components such as reasoning, problem-solving, and creativity",
                "activities": [
                    "Session: What is Intelligence?",
                    "Session: Decision Making - How do you make decisions?",
                    "Session: What is Artificial Intelligence and what is not?"
                ]
            },
            {
                "name": "Basics of AI: Let's Get Started",
                "learning_outcomes": "Understand the concept of Artificial Intelligence (AI) and its domains",
                "activities": [
                    "Introduction to AI, ML & DL",
                    "Introduction to AI Domains (Data Sciences, CV & NLP)",
                    "Data Sciences- Impact Filter: https://artsexperiments.withgoogle.com/impactfilter/",
                    "CV- Autodraw: https://www.autodraw.com/",
                    "NLP- Wordtune: https://www.wordtune.com/"
                ]
            },
            {
                "name": "AI in Real Life",
                "learning_outcomes": "Explore the use of AI in real Life",
                "activities": ["Session: Applications of AI – A look at Real-life AI implementations"]
            },
            {
                "name": "AI Ethics",
                "learning_outcomes": "Learn about the ethical concerns involved in AI development, such as AI bias, data privacy and how they can be addressed",
                "activities": [
                    "Session: AI Ethics",
                    "Moral Machine Activity: http://moralmachine.mit.edu/"
                ]
            }
        ]
    },
    "Unit 2: AI Project Cycle": {
        "sub_units": [
            {
                "name": "Introduction",
                "learning_outcomes": "Understand the stages involved in the AI project cycle, such as problem scoping, data collection, data exploration, modeling, evaluation",
                "activities": ["Session: Introduction to AI Project Cycle"]
            },
            {
                "name": "Problem Scoping",
                "learning_outcomes": "Learn about the importance of project planning in AI development and how to define project goals and objectives",
                "activities": ["Session: Understanding Problem Scoping & Sustainable Development Goals"]
            },
            {
                "name": "Data Acquisition",
                "learning_outcomes": "Develop an understanding of the importance of data collection in AI and how to choose the right data sources",
                "activities": ["Session: Simplifying Data Acquisition"]
            },
            {
                "name": "Data Exploration",
                "learning_outcomes": "Know various data exploration techniques and its importance",
                "activities": ["Session: Visualising Data"]
            },
            {
                "name": "Modelling",
                "learning_outcomes": "Know about the different machine learning algorithms used to train AI models",
                "activities": [
                    "Introduction to Rule Based & Learning Based AI Approaches",
                    "Activity: Teachable machine (Supervised Learning): https://teachablemachine.withgoogle.com/",
                    "Activity: Infinite Drum Machine (Unsupervised learning): https://experiments.withgoogle.com/ai/drum-machine/view/",
                    "Introduction to Supervised, Unsupervised & Reinforcement Learning Models",
                    "Neural Networks"
                ]
            },
            {
                "name": "Evaluation",
                "learning_outcomes": "Know the importance of evaluation and various metrics available for evaluation",
                "activities": ["Session: Evaluating the idea!"]
            }
        ]
    },
    "Unit 3: Advance Python": {
        "sub_units": [
            {
                "name": "Recap",
                "learning_outcomes": "Understand to work with Jupyter Notebook, creating virtual environment, installing Python Packages",
                "activities": [
                    "Session: Jupyter Notebook",
                    "Session: Introduction to Python",
                    "Session: Python Basics"
                ]
            }
        ],
        "note": "To be assessed through Practicals only"
    },
    "Unit 4: Data Sciences": {
        "theory_sub_units": [
            {
                "name": "Introduction",
                "learning_outcomes": "Define the concept of Data Science and understand its applications in various fields",
                "activities": [
                    "Session: Introduction to Data Science",
                    "Session: Applications of Data Science"
                ]
            },
            {
                "name": "Getting Started",
                "learning_outcomes": "Understand the basic concepts of data acquisition, visualization, and exploration",
                "activities": [
                    "Session: Revisiting AI Project Cycle, Data Collection, Data Access",
                    "Game: Rock, Paper & Scissors: https://next.rockpaperscissors.ai/"
                ]
            }
        ],
        "practical_sub_units": [
            {
                "name": "Python Packages",
                "learning_outcomes": "Use Python libraries such as NumPy, Pandas, and Matplotlib for data analysis and visualization",
                "activities": [
                    "Session: Python for Data Sciences - Numpy, Pandas, Matplotlib"
                ]
            },
            {
                "name": "Concepts of Data Sciences",
                "learning_outcomes": "Understand the basic concepts of statistics, such as mean, median, mode, and standard deviation",
                "activities": ["Session: Statistical Learning & Data Visualisation"]
            },
            {
                "name": "K-nearest neighbour model (Optional)",
                "learning_outcomes": "Understand the basic concepts of the KNN algorithm and its applications in supervised learning",
                "activities": [
                    "Activity: Personality Prediction (Optional)",
                    "Session: Understanding K-nearest neighbour model (Optional)"
                ]
            }
        ]
    },
    "Unit 5: Computer Vision": {
        "theory_sub_units": [
            {
                "name": "Introduction",
                "learning_outcomes": "Define the concept of Computer Vision and understand its applications in various fields",
                "activities": [
                    "Session: Introduction to Computer Vision",
                    "Session: Applications of CV"
                ]
            },
            {
                "name": "Concepts of Computer Vision",
                "learning_outcomes": "Understand the basic concepts of image representation, feature extraction, object detection, and segmentation",
                "activities": [
                    "Session: Understanding CV Concepts",
                    "Computer Vision Tasks",
                    "Basics of Images-Pixel, Resolution, Pixel value",
                    "Grayscale and RGB images",
                    "Game- Emoji Scavenger Hunt: https://emojiscavengerhunt.withgoogle.com/",
                    "RGB Calculator: https://www.w3schools.com/colors/colors_rgb.asp",
                    "Create your own pixel art: www.piskelapp.com",
                    "Create your own convolutions: http://setosa.io/ev/image-kernels/"
                ]
            }
        ],
        "practical_sub_units": [
            {
                "name": "OpenCV",
                "learning_outcomes": "Use Python libraries such as OpenCV for basic image processing and computer vision tasks",
                "activities": [
                    "Session: Introduction to OpenCV",
                    "Hands-on: Image Processing"
                ]
            },
            {
                "name": "Convolution Operator (Optional)",
                "learning_outcomes": "Apply the convolution operator to process images and extract useful features",
                "activities": [
                    "Session: Understanding Convolution operator (Optional)",
                    "Activity: Convolution Operator (Optional)"
                ]
            },
            {
                "name": "Convolution Neural Network (Optional)",
                "learning_outcomes": "Understand the basic architecture of a CNN and its applications in computer vision and image recognition",
                "activities": [
                    "Session: Introduction to CNN (Optional)",
                    "Session: Understanding CNN - Kernel, Layers of CNN (Optional)",
                    "Activity: Testing CNN (Optional)"
                ]
            }
        ]
    },
    "Unit 6: Natural Language Processing": {
        "sub_units": [
            {
                "name": "Introduction",
                "learning_outcomes": "Understand the concept of Natural Language Processing (NLP) and its importance in the field of AI",
                "activities": [
                    "Session: Introduction to Natural Language Processing",
                    "Activity: Use of Google Translate for same spelling words",
                    "Session: NLP Applications",
                    "Session: Revisiting AI Project Cycle"
                ]
            },
            {
                "name": "Chatbots",
                "learning_outcomes": "Explore the various applications of NLP in everyday life, such as chatbots, sentiment analysis, and automatic summarization",
                "activities": ["Activity: Introduction to Chatbots"]
            },
            {
                "name": "Language Differences",
                "learning_outcomes": "Gain an understanding of the challenges involved in understanding human language by machine",
                "activities": ["Session: Human Language VS Computer Language"]
            },
            {
                "name": "Concepts of Natural Language Processing",
                "learning_outcomes": "Learn about the Text Normalization technique used in NLP and popular NLP model - Bag-of-Words",
                "activities": [
                    "Session: Data Processing - Text Normalisation, Bag of Words",
                    "Hands-on: Text processing - Data Processing, Bag of Words",
                    "TFIDF (Optional)",
                    "NLTK (Optional)"
                ]
            }
        ]
    },
    "Unit 7: Evaluation": {
        "sub_units": [
            {
                "name": "Introduction",
                "learning_outcomes": "Understand the role of evaluation in the development and implementation of AI systems",
                "activities": [
                    "Session: Introduction to Model Evaluation",
                    "What is Evaluation?",
                    "Different types of Evaluation techniques - Underfit, Perfect Fit, OverFit"
                ]
            },
            {
                "name": "Model Evaluation Terminology",
                "learning_outcomes": "Learn various Model Evaluation Terminologies",
                "activities": [
                    "Session: Model Evaluation Terminologies",
                    "The Scenario - Prediction, Reality, True Positive, True Negative, False Positive, False Negative",
                    "Confusion Matrix",
                    "Activity: Make a confusion matrix for Containment Zone Prediction Model"
                ]
            },
            {
                "name": "Confusion Matrix",
                "learning_outcomes": "Learn to make a confusion matrix for given Scenario",
                "activities": ["Session & Activity: Confusion Matrix"]
            },
            {
                "name": "Evaluation Methods",
                "learning_outcomes": "Learn about the different types of evaluation techniques in AI, such as Accuracy, Precision, Recall and F1 Score, and their significance",
                "activities": [
                    "Session: Evaluation Methods - Accuracy, Precision, Recall",
                    "Which Metric is Important? - Precision or Recall",
                    "F1 Score",
                    "Activity: Practice Evaluation"
                ]
            }
        ]
    }
}

def create_ai_chapter_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates an Artificial Intelligence chapter content prompt aligned with CBSE curriculum and publisher best practices"""
    
    # Extract class number from grade level
    import re
    match = re.search(r'(\d+)', grade_level)
    class_num = int(match.group(1)) if match else 9
    
    # Get CBSE curriculum focus for the class
    if class_num == 9:
        curriculum_focus = """**CBSE Class 9 AI (Code 417) Curriculum Focus**:
- Introduction to AI and foundational concepts
- AI Project Cycle methodology
- Basic domains: Data Science, Computer Vision, NLP
- AI Ethics and responsible development
- Practical applications in daily life

**Include for each topic**:
- SUB-UNITS with specific learning outcomes
- Hands-on ACTIVITIES and gamified tools
- PRACTICAL sessions with real tools
- Links to online resources when relevant"""
    elif class_num == 10:
        curriculum_focus = """**CBSE Class 10 AI (Code 417) Curriculum Focus**:
- Unit 1: Introduction to AI (Foundational concepts, AI domains)
- Unit 2: AI Project Cycle (Problem scoping, Data acquisition, Modeling, Evaluation)
- Unit 3: Advanced Python (Jupyter, Python basics, Libraries)
- Unit 4: Data Science (Statistics, Visualization, Python packages)
- Unit 5: Computer Vision (Image processing, OpenCV, CNN basics)
- Unit 6: Natural Language Processing (Chatbots, Text processing, Bag of Words)
- Unit 7: Evaluation (Confusion Matrix, Accuracy, Precision, Recall, F1 Score)

**Include for each topic**:
- SUB-UNITS with specific LEARNING OUTCOMES
- SESSION plans and step-by-step ACTIVITIES
- PRACTICAL exercises with Python code
- Gamified tools and online resources (Teachable Machine, AutoDraw, etc.)
- Real-world applications and case studies"""
    elif class_num == 11:
        curriculum_focus = """**CBSE Class 11 AI (Code 843) Curriculum Focus**:
- Python programming for AI development
- Statistics and probability foundations
- Data science and visualization techniques
- Machine Learning algorithms (supervised/unsupervised)
- Introduction to Neural Networks"""
    else:  # Class 12
        curriculum_focus = """**CBSE Class 12 AI (Code 843) Curriculum Focus**:
- Deep Learning and advanced architectures
- Convolutional and Recurrent Neural Networks
- Advanced NLP with transformers
- Computer Vision with Deep Learning
- Capstone AI project development"""
    
    # Default word limits based on class level
    if word_limits is None:
        if class_num <= 10:
            word_limits = {
                'introduction': 400,
                'concepts': 2500,
                'hands_on': 1500,
                'case_studies': 800,
                'exercises': 1000,
                'projects': 600
            }
        else:
            word_limits = {
                'introduction': 500,
                'concepts': 3500,
                'hands_on': 2000,
                'case_studies': 1000,
                'exercises': 1200,
                'projects': 800
            }
    
    # Get selected chapter from session state if available
    import streamlit as st
    selected_chapter = ""
    chapter_directive = ""
    curriculum_details = ""
    
    if 'ai_chapter_selected' in st.session_state:
        selected_chapter = st.session_state['ai_chapter_selected']
        # Extract just the chapter topic without the number
        chapter_topic = selected_chapter.split(': ', 1)[1] if ': ' in selected_chapter else selected_chapter
        chapter_num = selected_chapter.split(':')[0] if ':' in selected_chapter else "Chapter"
        
        # Map chapter to curriculum unit and get detailed structure
        if class_num == 10:
            # Map chapters to units based on topic
            chapter_to_unit_map = {
                "Introduction to AI": "Unit 1: Introduction to Artificial Intelligence",
                "AI Basics": "Unit 1: Introduction to Artificial Intelligence",
                "AI Project Cycle": "Unit 2: AI Project Cycle",
                "Data Acquisition": "Unit 2: AI Project Cycle",
                "AI Modeling": "Unit 2: AI Project Cycle",
                "Advanced Python": "Unit 3: Advance Python",
                "Data Science": "Unit 4: Data Sciences",
                "Computer Vision": "Unit 5: Computer Vision",
                "Natural Language Processing": "Unit 6: Natural Language Processing",
                "Model Evaluation": "Unit 7: Evaluation",
                "AI Ethics": "Unit 1: Introduction to Artificial Intelligence"
            }
            
            # Find matching unit
            unit_key = None
            for key_phrase, unit in chapter_to_unit_map.items():
                if key_phrase.lower() in chapter_topic.lower():
                    unit_key = unit
                    break
            
            if unit_key and unit_key in CBSE_AI_CLASS_10_CURRICULUM:
                unit_data = CBSE_AI_CLASS_10_CURRICULUM[unit_key]
                sub_units_text = "\n\n**CBSE CURRICULUM SUB-UNITS TO COVER:**\n\n"
                
                # Process regular sub_units
                if "sub_units" in unit_data:
                    for sub_unit in unit_data["sub_units"]:
                        sub_units_text += f"\n**{sub_unit['name']}**\n"
                        sub_units_text += f"- Learning Outcome: {sub_unit['learning_outcomes']}\n"
                        sub_units_text += "- Activities:\n"
                        for activity in sub_unit['activities']:
                            sub_units_text += f"  • {activity}\n"
                
                # Process theory and practical sub_units separately if they exist
                if "theory_sub_units" in unit_data:
                    sub_units_text += "\n**THEORY COMPONENTS:**\n"
                    for sub_unit in unit_data["theory_sub_units"]:
                        sub_units_text += f"\n**{sub_unit['name']}**\n"
                        sub_units_text += f"- Learning Outcome: {sub_unit['learning_outcomes']}\n"
                        sub_units_text += "- Activities:\n"
                        for activity in sub_unit['activities']:
                            sub_units_text += f"  • {activity}\n"
                
                if "practical_sub_units" in unit_data:
                    sub_units_text += "\n**PRACTICAL COMPONENTS:**\n"
                    for sub_unit in unit_data["practical_sub_units"]:
                        sub_units_text += f"\n**{sub_unit['name']}**\n"
                        sub_units_text += f"- Learning Outcome: {sub_unit['learning_outcomes']}\n"
                        sub_units_text += "- Activities:\n"
                        for activity in sub_unit['activities']:
                            sub_units_text += f"  • {activity}\n"
                
                curriculum_details = sub_units_text
        
        chapter_directive = f"""

**CHAPTER TO GENERATE**: 
{selected_chapter}

**IMPORTANT**: 
- Focus ONLY on this specific chapter topic: {chapter_topic}
- This is a standalone chapter - do not reference other chapters
- All content must be directly relevant to {chapter_topic}
- Provide comprehensive, in-depth coverage of this single topic
{curriculum_details}

**INSTRUCTIONS**:
1. Cover ALL the sub-units and learning outcomes listed above
2. Include ALL the activities and sessions mentioned
3. Add clickable links for online tools where provided
4. THEN go BEYOND the curriculum by adding:
   - Additional real-world examples
   - Extra hands-on projects
   - Industry case studies
   - Advanced optional topics
   - Career connections
   - Latest AI trends related to the topic"""
    else:
        # If no chapter selected, provide a default
        chapter_directive = """

**NOTE**: No specific chapter selected. Please generate a foundational AI chapter appropriate for the grade level."""
    
    return f"""You are an Expert AI Textbook Author specializing in creating high-quality educational content aligned with CBSE curriculum requirements.

**Your Mission**: Create a single, comprehensive AI textbook chapter that meets the highest educational standards.

**Target Audience**: {grade_level} (CBSE AI Curriculum)

{curriculum_focus}{chapter_directive}

**Publisher Best Practices to Incorporate**:

1. **O'Reilly Style Elements**:
   - "In a Nutshell" boxes for key concepts
   - "Try This" hands-on exercises throughout
   - Real-world case studies from tech companies
   - Code examples with detailed annotations
   - "Common Pitfalls" and "Best Practices" sidebars
   - Progressive difficulty with clear learning paths

2. **Manning's Hands-On Approach**:
   - "Think Like an AI Engineer" problem-solving sections
   - Project-based learning with incremental builds
   - Interactive coding challenges
   - Industry-standard tools and frameworks

3. **Packt's Practical Focus**:
   - Step-by-step tutorials with screenshots
   - "Quick Start" sections for immediate application
   - GitHub integration and version control practices
   - Cloud deployment considerations

**Model Chapter Progression (Enhanced for AI)**:
---
{model_progression_text}
---

**COMPREHENSIVE CHAPTER STRUCTURE FOR AI**:

## 1. Chapter Opening - The Big Picture (Target: {word_limits.get('introduction', 400)} words)

### A. Chapter Title & Learning Quest
- Engaging title with tech-appeal (e.g., "Building Your First AI Brain", "Teaching Machines to See")
- Chapter number and position in curriculum
- "What You'll Build" showcase box
- Prerequisites checklist

### B. Real-World Hook
- Start with a current AI breakthrough or application
- Connect to students' daily tech experiences (smartphones, games, social media)
- Include data/statistics on AI growth and opportunities
- Feature spotlight on an AI pioneer or company

### C. Learning Objectives (SMART Goals)
- 5-6 specific, measurable objectives
- Mapped to CBSE curriculum competencies
- Skills badges to be earned
- Connection to AI career paths

### D. Chapter Roadmap
- Visual journey map of concepts
- Estimated time for each section
- Difficulty indicators (Beginner/Intermediate/Advanced)
- Optional deep-dive sections marked

## 2. Core AI Concepts - Building Blocks (Target: {word_limits.get('concepts', 2500)} words)

### A. Conceptual Foundation (30% of content)
- **"AI in a Nutshell"** boxes for each major concept
- Visual diagrams and infographics (specify detailed descriptions)
- Analogies to familiar concepts
- Historical context and evolution
- Mathematical foundations (age-appropriate)

### B. Technical Deep Dive (40% of content)
- Detailed explanations with progressive complexity
- Algorithm walkthroughs with pseudocode
- Data structures and representations
- Model architectures and components
- Performance metrics and evaluation

### C. Code Implementation (30% of content)
- **Language-appropriate examples**:
  - Classes 9-10: Scratch blocks, Python basics, visual programming
  - Classes 11-12: Python (NumPy, Pandas, Scikit-learn, TensorFlow/Keras basics)
- Code snippets with line-by-line explanations
- Common debugging scenarios
- Performance optimization tips
- Version control best practices

## 3. Hands-On Labs - Learn by Doing (Target: {word_limits.get('hands_on', 1500)} words)

### A. Guided Project Build
- **Project Title**: Something exciting and relevant
- **Tools Required**: Specific software, libraries, datasets
- **Step-by-Step Implementation**:
  1. Environment setup
  2. Data preparation
  3. Model building
  4. Training and testing
  5. Evaluation and improvement
  6. Deployment considerations

### B. Experimentation Sandbox
- "What happens if..." scenarios
- Parameter tuning exercises
- A/B testing concepts
- Performance comparison activities

### C. Troubleshooting Guide
- Common errors and solutions
- Debugging strategies
- Performance bottlenecks
- Resource optimization

## 4. Real-World Applications & Case Studies (Target: {word_limits.get('case_studies', 800)} words)

### A. Industry Case Studies
- 2-3 detailed examples from different sectors:
  - Healthcare (disease prediction, drug discovery)
  - Finance (fraud detection, algorithmic trading)
  - Entertainment (recommendation systems, game AI)
  - Transportation (autonomous vehicles, route optimization)
  - Education (personalized learning, automated grading)

### B. Indian Context Applications
- AI in Indian agriculture (crop prediction, pest detection)
- Smart cities initiatives
- Language processing for Indian languages
- AI in Indian healthcare and telemedicine

### C. Ethical Considerations
- Bias in AI systems
- Privacy and data protection
- Responsible AI development
- Future implications and careers

## 5. Practice & Assessment (Target: {word_limits.get('exercises', 1000)} words)

### A. Knowledge Check (Quick Review)
- 10 MCQs with explanations
- 5 True/False with justification
- 5 Fill in the blanks (technical terms)

### B. Conceptual Questions
- 5 short answer questions (2-3 sentences)
- 3 long answer questions (paragraph length)
- 2 comparison/contrast questions

### C. Practical Challenges
- 3 coding challenges (increasing difficulty)
- 2 debugging exercises
- 1 optimization problem
- 1 design challenge

### D. Project Extensions
- 3 ways to extend the chapter project
- Cross-curricular connections
- Competition preparation tips

## 6. Beyond the Chapter (Target: {word_limits.get('projects', 600)} words)

### A. Capstone Project Ideas
- 3-4 comprehensive project suggestions
- Resource requirements
- Expected outcomes
- Assessment rubrics

### B. Further Learning Resources
- Recommended books on this topic
- Online courses and MOOCs
- Open datasets for practice
- AI communities and forums
- Competitions and hackathons

### C. Career Connections
- AI career paths and roles
- Required skills and certifications
- Indian and global opportunities
- Building a portfolio

### D. Next Chapter Preview
- What's coming next
- How current knowledge will be used
- Preparation suggestions

**QUALITY STANDARDS**:

1. **Technical Accuracy**: All code must be tested and working
2. **Progressive Difficulty**: Content should build logically
3. **Engagement**: Use current examples and relatable scenarios
4. **Inclusivity**: Examples from diverse contexts and applications
5. **Future-Ready**: Include emerging trends and technologies
6. **Assessment-Aligned**: Include CBSE board exam pattern questions

**IMPORTANT NOTES**:
- Do NOT create the entire textbook at once - focus on ONE chapter
- Specify clear chapter number and title
- Include "Chapter Summary" and "Key Takeaways" boxes
- Add "Further Reading" suggestions relevant to the topic
- Use Indian names and contexts in examples where appropriate
- Balance theory with practical application (40:60 ratio)
- Include QR codes placeholders for additional resources

Generate a single, complete AI chapter following this enhanced structure, incorporating best pedagogical practices while meeting CBSE curriculum requirements.

Provide the complete AI chapter content in Markdown format.
"""

def create_ai_exercises_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Artificial Intelligence specific exercises and assessments"""
    if word_limits is None:
        word_limits = {
            'exercises': 1000
        }
    
    return f"""You are an Expert Artificial Intelligence Education Assessment Developer.

**Subject Focus**: ARTIFICIAL INTELLIGENCE

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

Create comprehensive AI exercises based on the chapter content.
**Target Total Word Count**: {word_limits.get('exercises', 1000)} words

**AI Exercise Categories**:

## 1. AI Conceptual Understanding (30%)
- Multiple Choice Questions on AI principles, machine learning, pattern recognition (8 questions)
- True/False with AI reasoning justification (10 statements)
- Fill in the blanks with AI terminology (decision trees, neural networks, algorithms, etc.) (10 questions)
- Match AI components to their functions (data, algorithms, models, predictions) (2 sets of 5 matches)

## 2. Algorithm Analysis & Implementation (25%)
- AI algorithm analysis and prediction (5 exercises)
- Debug AI code challenges (3 code blocks with logic errors)
- Complete AI algorithm implementations (5 exercises)
- AI decision-making problems (3 challenges)

## 3. Data & Pattern Recognition (20%)
- Data analysis and interpretation (3 exercises)
- Pattern identification challenges (3 scenarios)
- AI training data evaluation (3 exercises)
- Model accuracy assessment (2 exercises)

## 4. Real-World AI Applications (15%)
- AI application scenarios (3 case studies in healthcare, transportation, smart homes)
- AI career pathway connections (2 detailed scenarios)
- AI ethics and societal impact (3 discussion questions)
- Industry AI problem-solving (2 challenges)

## 5. Creative AI Problem Solving (10%)
- AI system design challenges (2 open-ended problems)
- Innovation with AI proposals (2 creative extensions)
- Cross-domain AI applications (2 scenarios)

**Exercise Requirements**:
- Include detailed solutions with AI reasoning
- Connect to specific kit sensors and data collection when applicable
- Reference real-world AI applications and case studies
- Provide multiple difficulty levels within each category
- Include data visualization and interpretation exercises
- Emphasize ethical AI considerations and responsible development

Ensure all exercises relate directly to the chapter content and demonstrate practical AI concepts.

Provide comprehensive Artificial Intelligence exercises in Markdown format.
"""

def create_ai_skills_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Artificial Intelligence hands-on lab activities and skill development exercises"""
    if word_limits is None:
        word_limits = {
            'skill_activity': 600,
            'lab_project': 600
        }
    
    return f"""You are an Expert Artificial Intelligence Lab Instructor and Skills Developer.

**Subject Focus**: ARTIFICIAL INTELLIGENCE

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

## AI Skill-Based Activities (Target: {word_limits.get('skill_activity', 600)} words)

Create at least 4 comprehensive hands-on activities that develop practical AI skills:

**Activity Structure for Each**:
1. **AI Skill Development Objective** (what specific AI capability)
2. **Kit Components Required** (sensors for data collection, processing units, etc.)
3. **Data Collection Setup** (how to gather training/test data)
4. **Step-by-Step AI Implementation** (algorithm development and testing)
5. **Model Training & Validation** (how to train and test AI models)
6. **Real-World AI Connection** (how this skill applies in AI industry)
7. **Performance Evaluation** (metrics for measuring AI system success)
8. **AI Skill Validation** (how to verify learning objectives met)

**Types of AI Skills Activities**:
- **Data Collection & Processing**: Learn to gather, clean, and prepare data for AI systems
- **Pattern Recognition Implementation**: Build systems that identify patterns in sensor data
- **Decision-Making Algorithms**: Create AI that makes intelligent choices based on inputs
- **Learning System Development**: Implement basic machine learning with kit components
- **AI Model Testing & Validation**: Learn to evaluate and improve AI system performance

## AI Lab Projects (Target: {word_limits.get('lab_project', 600)} words)

Create at least 3 comprehensive lab projects that integrate multiple AI concepts:

**Project Structure for Each**:
1. **AI Learning Objectives** (comprehensive intelligence goals)
2. **Domain Application Focus** (Smart Home AI, Healthcare AI, Educational AI, etc.)
3. **Data Collection Strategy** (using kit sensors for real-world data)
4. **AI Algorithm Development** (incremental intelligence building)
5. **Training & Testing Protocol** (systematic AI model development)
6. **Intelligence Optimization** (improving AI performance and accuracy)
7. **AI Documentation Standards** (technical AI reporting requirements)
8. **Advanced AI Extensions** (sophisticated AI modifications)

**AI Project Categories**:
- **Intelligent Sensing System**: Implement AI that learns from environmental data
- **Predictive Analytics Project**: Build AI that predicts future states from current data
- **Adaptive Behavior System**: Create AI that adjusts behavior based on learning
- **Human-AI Collaboration Interface**: Develop AI assistant that works with users

**AI Lab Safety & Ethics**:
- Data privacy and security protocols
- Responsible AI development practices
- Bias detection and mitigation
- Ethical decision-making in AI systems

**AI Assessment Criteria**:
- Algorithm implementation quality
- Data handling and processing skills
- AI model accuracy and reliability
- Ethical considerations in AI development
- Innovation in AI applications

Each activity should build progressively toward the lab projects, with clear connections to real-world AI applications and career paths.

Provide comprehensive Artificial Intelligence skills activities in Markdown format.
"""

def create_ai_projects_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Artificial Intelligence creative projects and applications"""
    if word_limits is None:
        word_limits = {
            'creative_projects': 800
        }
    
    return f"""You are an Expert Artificial Intelligence Innovation Mentor and Project Designer.

**Subject Focus**: ARTIFICIAL INTELLIGENCE

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

## AI Creative Innovation Projects (Target: {word_limits.get('creative_projects', 800)} words)

Create at least 4 innovative projects that combine AI with creative applications:

**Project Structure for Each**:
1. **AI Innovation Objective** (intelligent + creative goals)
2. **Cross-Domain AI Integration** (how AI enhances other fields)
3. **Kit Components & AI Extensions** (sensors, processors, AI algorithms)
4. **AI Design Thinking Process** (data-driven empathy, intelligent definition, AI ideation, smart prototyping, intelligent testing)
5. **AI Technical Implementation Plan** (algorithms + data + hardware integration)
6. **Intelligent User Experience Design** (AI-powered interface and interaction)
7. **AI Social Impact Assessment** (how AI project benefits society)
8. **AI Portfolio & Demonstration** (showcasing AI capabilities)

**Creative AI Project Categories**:

### 🎨 AI Art & Creative Intelligence
- **Generative AI Art**: AI systems that create visual art based on learned patterns
- **AI Music Composition**: Machine learning for music generation and performance
- **Intelligent Storytelling**: AI-powered narrative systems that adapt to user preferences
- **Creative AI Collaboration**: Human-AI partnership in creative expression

### 🏠 Smart Living AI Solutions
- **Adaptive Home Intelligence**: AI systems that learn family patterns and optimize living
- **Personal AI Health Assistant**: Intelligent health monitoring and recommendation systems
- **Sustainable AI Living**: AI for energy optimization and environmental management
- **Accessibility AI Enhancement**: Intelligent tools for inclusive design and assistance

### 🌍 AI for Social Good
- **Community Intelligence Projects**: Apply AI to solve local problems with data-driven solutions
- **Educational AI Tutors**: Personalized learning systems that adapt to individual needs
- **Environmental AI Monitoring**: Intelligent systems for ecological data analysis
- **Cultural AI Preservation**: Using AI to document, analyze, and share cultural heritage

### 🔬 AI Research & Discovery Tools
- **Scientific AI Assistants**: Intelligent systems for research data analysis and hypothesis generation
- **AI Exploration Systems**: Smart systems for investigating and mapping complex environments
- **Pattern Discovery AI**: Intelligent tools for finding hidden patterns in large datasets
- **Collaborative AI Research**: Multi-user intelligent platforms for team-based discovery

**AI Innovation Methodologies**:
- Data-driven design thinking workshops
- AI prototyping with rapid iteration
- Human-centered AI development principles
- Ethical AI development frameworks
- Open source AI collaboration methods

**AI Portfolio Development**:
- AI technical documentation standards
- AI demonstration and explanation techniques
- Presentation skills for AI audiences
- AI peer review and feedback protocols
- AI industry mentorship connections

**AI Impact Metrics**:
- Problem-solving intelligence effectiveness
- AI system adoption and usability
- Scalability of AI solutions
- Social benefit of AI applications
- Technical innovation in AI implementation

**AI Extension Opportunities**:
- AI competition participation (machine learning contests, AI challenges)
- Community AI demonstration events
- AI industry collaboration projects
- Academic AI research partnerships
- AI entrepreneurship development paths

Each project should demonstrate how AI can be applied creatively to solve real problems while building technical AI expertise and fostering innovation mindsets.

Provide comprehensive Artificial Intelligence creative projects in Markdown format.
"""

# Robotics specific prompt functions
def create_robotics_chapter_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates a Robotics chapter content prompt using the enhanced template"""
    # Default word limits if none provided
    if word_limits is None:
        word_limits = {
            'mission_briefing': 300,
            'domain_analysis': 800,
            'core_concepts': 3000,
            'hands_on_project': 1500,
            'assessment': 600
        }
    
    # Read the enhanced AI/Robotics template
    try:
        with open('ai_robo.txt', 'r', encoding='utf-8') as f:
            robotics_template = f.read()
    except FileNotFoundError:
        robotics_template = """Enhanced Robotics template not found. Using fallback template."""
    
    return f"""You are an Expert Robotics Education Content Developer specializing in hands-on learning with IIT-developed kits.

**Subject Focus**: ROBOTICS - Focus on automation, sensors, actuators, mechanical systems, control systems, human-robot interaction, and physical world manipulation.

**Target Audience**: {grade_level}

**Enhanced Robotics Education Template**:
---
{robotics_template}
---

**Model Chapter Progression Context**:
---
{model_progression_text}
---

**Word Count Targets**:
- Mission Briefing: {word_limits.get('mission_briefing', 300)} words
- Domain Analysis & Applications: {word_limits.get('domain_analysis', 800)} words
- Core Concepts (all concepts combined): {word_limits.get('core_concepts', 3000)} words
- Hands-On Project: {word_limits.get('hands_on_project', 1500)} words
- Assessment & Debrief: {word_limits.get('assessment', 600)} words

**ROBOTICS-SPECIFIC FOCUS AREAS**:
1. **Sensors & Perception**: How robots sense and understand their environment
2. **Actuators & Movement**: How robots move and manipulate objects in the physical world
3. **Control Systems**: How robots are programmed and controlled to perform tasks
4. **Automation & Efficiency**: How robots automate repetitive tasks and improve efficiency
5. **Human-Robot Interaction**: How robots work safely and effectively with humans

**CRITICAL INSTRUCTIONS**:
1. Follow the enhanced template structure EXACTLY
2. Preserve all original IIT code and technical content
3. Analyze kit components to determine robotics applications (motors for movement, sensors for feedback, controllers for automation)
4. Connect every concept to real-world robotics applications (manufacturing robots, service robots, exploration robots, etc.)
5. Handle multiple programming languages (Scratch for beginners, Python for advanced, Arduino for hardware control)
6. Include comprehensive troubleshooting for robotics system integration
7. Provide multi-level extension challenges focusing on robotics capabilities

Transform the provided PDF content into a comprehensive Robotics chapter following the enhanced template structure.

Provide the complete Robotics chapter content in Markdown format.
"""

def create_robotics_exercises_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Robotics specific exercises and assessments"""
    if word_limits is None:
        word_limits = {
            'exercises': 1000
        }
    
    return f"""You are an Expert Robotics Education Assessment Developer.

**Subject Focus**: ROBOTICS

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

Create comprehensive Robotics exercises based on the chapter content.
**Target Total Word Count**: {word_limits.get('exercises', 1000)} words

**Robotics Exercise Categories**:

## 1. Robotics Conceptual Understanding (30%)
- Multiple Choice Questions on robotics principles, automation, sensors, actuators (8 questions)
- True/False with robotics engineering justification (10 statements)
- Fill in the blanks with robotics terminology (servo motors, sensors, controllers, feedback loops) (10 questions)
- Match robotics components to their functions (sensors to perception, motors to movement, etc.) (2 sets of 5 matches)

## 2. Control Systems & Programming (25%)
- Robotics control algorithm analysis (5 exercises)
- Debug robotics code challenges (3 code blocks with control errors)
- Complete robotics control implementations (5 exercises)
- Robotics automation problems (3 challenges)

## 3. Hardware Integration & Troubleshooting (20%)
- Circuit analysis and wiring problems (3 exercises)
- Sensor calibration and testing (3 scenarios)
- Motor control and feedback systems (3 exercises)
- System integration troubleshooting (2 exercises)

## 4. Real-World Robotics Applications (15%)
- Robotics application scenarios (3 case studies in manufacturing, healthcare, exploration)
- Robotics career pathway connections (2 detailed scenarios)
- Robotics safety and ethics (3 discussion questions)
- Industrial robotics problem-solving (2 challenges)

## 5. Creative Robotics Design (10%)
- Robotics system design challenges (2 open-ended problems)
- Innovation in robotics proposals (2 creative extensions)
- Cross-domain robotics applications (2 scenarios)

**Exercise Requirements**:
- Include detailed solutions with engineering reasoning
- Connect to specific kit motors, sensors, and controllers
- Reference real-world robotics applications and case studies
- Provide multiple difficulty levels within each category
- Include mechanical design and motion planning exercises
- Emphasize safety protocols and responsible robotics development

Ensure all exercises relate directly to the chapter content and demonstrate practical robotics engineering concepts.

Provide comprehensive Robotics exercises in Markdown format.
"""

def create_robotics_skills_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Robotics hands-on lab activities and skill development exercises"""
    if word_limits is None:
        word_limits = {
            'skill_activity': 600,
            'lab_project': 600
        }
    
    return f"""You are an Expert Robotics Lab Instructor and Skills Developer.

**Subject Focus**: ROBOTICS

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

## Robotics Skill-Based Activities (Target: {word_limits.get('skill_activity', 600)} words)

Create at least 4 comprehensive hands-on activities that develop practical robotics skills:

**Activity Structure for Each**:
1. **Robotics Skill Development Objective** (what specific robotics capability)
2. **Kit Components Required** (motors, sensors, controllers, mechanical parts)
3. **Mechanical Assembly Setup** (physical construction and connections)
4. **Step-by-Step Control Implementation** (programming robot behavior)
5. **Testing & Calibration Protocol** (ensuring accurate robot performance)
6. **Real-World Robotics Connection** (how this skill applies in robotics industry)
7. **Performance Optimization** (improving robot efficiency and accuracy)
8. **Robotics Skill Validation** (how to verify engineering objectives met)

**Types of Robotics Skills Activities**:
- **Sensor Integration & Feedback**: Learn to connect sensors and use feedback for control
- **Motor Control & Precision Movement**: Master precise robot movement and positioning
- **Automation Sequence Programming**: Create robots that perform complex task sequences
- **Human-Robot Interface Development**: Build systems for safe human-robot interaction
- **Robotic System Integration**: Combine multiple subsystems into working robots

## Robotics Lab Projects (Target: {word_limits.get('lab_project', 600)} words)

Create at least 3 comprehensive lab projects that integrate multiple robotics concepts:

**Project Structure for Each**:
1. **Robotics Engineering Objectives** (comprehensive automation goals)
2. **Domain Application Focus** (Manufacturing Robotics, Service Robotics, Exploration Robotics, etc.)
3. **Mechanical Design Strategy** (using kit components for robot construction)
4. **Control System Development** (incremental robot behavior programming)
5. **Testing & Validation Protocol** (systematic robot performance evaluation)
6. **System Optimization** (improving robot speed, accuracy, and reliability)
7. **Robotics Documentation Standards** (technical engineering reporting requirements)
8. **Advanced Robotics Extensions** (sophisticated robot modifications)

**Robotics Project Categories**:
- **Autonomous Navigation Robot**: Implement robot that navigates and avoids obstacles
- **Automated Task Execution System**: Build robot that performs specific work tasks
- **Adaptive Robotics Behavior**: Create robot that adjusts to changing conditions
- **Collaborative Robotics Interface**: Develop robot that works safely with humans

**Robotics Lab Safety & Ethics**:
- Mechanical safety and proper tool use
- Electrical safety with motors and controllers
- Human-robot interaction safety protocols
- Responsible robotics development practices

**Robotics Assessment Criteria**:
- Mechanical design and construction quality
- Control system implementation effectiveness
- Robot performance accuracy and reliability
- Safety compliance and risk management
- Innovation in robotics applications

Each activity should build progressively toward the lab projects, with clear connections to real-world robotics applications and engineering careers.

Provide comprehensive Robotics skills activities in Markdown format.
"""

def create_robotics_projects_prompt(grade_level, model_progression_text, word_limits=None):
    """Creates Robotics creative projects and engineering applications"""
    if word_limits is None:
        word_limits = {
            'creative_projects': 800
        }
    
    return f"""You are an Expert Robotics Innovation Mentor and Project Designer.

**Subject Focus**: ROBOTICS

**Target Audience**: {grade_level}

**Model Chapter Progression Context**:
---
{model_progression_text}
---

## Robotics Creative Innovation Projects (Target: {word_limits.get('creative_projects', 800)} words)

Create at least 4 innovative projects that combine robotics with creative engineering applications:

**Project Structure for Each**:
1. **Robotics Innovation Objective** (engineering + creative goals)
2. **Cross-Domain Robotics Integration** (how robotics enhances other fields)
3. **Kit Components & Mechanical Extensions** (motors, sensors, structural elements, custom additions)
4. **Robotics Design Thinking Process** (empathize with users, define robot requirements, ideate solutions, prototype robots, test performance)
5. **Robotics Technical Implementation Plan** (mechanical design + control systems + hardware integration)
6. **Human-Robot Interaction Design** (safe and effective robot-human collaboration)
7. **Robotics Social Impact Assessment** (how robot project benefits society)
8. **Robotics Portfolio & Demonstration** (showcasing robot capabilities and engineering)

**Creative Robotics Project Categories**:

### 🎨 Artistic & Creative Robotics
- **Robotic Art Creation**: Robots that create physical art through programmed movements
- **Performance Robotics**: Robots designed for entertainment and artistic expression
- **Interactive Robotic Installations**: Robots that respond to human presence and create experiences
- **Collaborative Art Robotics**: Human-robot teams working together on creative projects

### 🏠 Service & Assistance Robotics
- **Home Assistant Robots**: Automated systems that help with household tasks
- **Personal Care Robotics**: Robots designed to assist with daily living activities
- **Accessibility Robotics**: Robots that help people with disabilities navigate and interact
- **Educational Robotics Tutors**: Robots that assist with learning and skill development

### 🌍 Social Impact Robotics
- **Community Service Robots**: Automated systems that address local community needs
- **Environmental Monitoring Robots**: Robots that collect environmental data and maintain ecosystems
- **Emergency Response Robotics**: Robots designed for search, rescue, and disaster response
- **Agricultural Robotics**: Automated systems for sustainable farming and food production

### 🔬 Exploration & Research Robotics
- **Scientific Research Robots**: Automated systems for data collection and experimentation
- **Exploration Robotics**: Robots designed to investigate hard-to-reach or dangerous environments
- **Surveillance & Monitoring Robotics**: Automated systems for security and observation
- **Collaborative Research Platforms**: Multi-robot systems for team-based investigation

**Robotics Innovation Methodologies**:
- Human-centered robotics design workshops
- Rapid robotics prototyping with iterative testing
- User-centered robotics development principles
- Ethical robotics development frameworks
- Open source robotics collaboration methods

**Robotics Portfolio Development**:
- Robotics technical documentation standards
- Robot demonstration and operation techniques
- Presentation skills for engineering audiences
- Robotics peer review and feedback protocols
- Robotics industry mentorship connections

**Robotics Impact Metrics**:
- Task automation effectiveness
- Robot system reliability and safety
- Scalability of robotics solutions
- Social benefit of robotics applications
- Technical innovation in robotics engineering

**Robotics Extension Opportunities**:
- Robotics competition participation (robotics contests, engineering challenges)
- Community robotics demonstration events
- Robotics industry collaboration projects
- Academic robotics research partnerships
- Robotics entrepreneurship development paths

Each project should demonstrate how robotics can be applied creatively to solve real problems while building technical engineering expertise and fostering innovation mindsets.

Provide comprehensive Robotics creative projects in Markdown format.
"""

def generate_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, model_progression_text, subject_type="Science", word_limits=None, use_chunked=False, use_openrouter_method=False, pdf_method="Text Extraction (Original)"):
    """Generates specific content based on content type"""
    
    # Special handling for AI subject without PDF
    if subject_type == "Artificial Intelligence" and pdf_bytes is None:
        # Direct generation without PDF for AI textbook creation
        try:
            prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type, word_limits)
            
            # Show detailed info for AI chapter generation
            if 'ai_chapter_selected' in st.session_state:
                chapter_info = st.session_state['ai_chapter_selected']
                st.info(f"📚 Generating {content_type} content for {grade_level}\n\n📖 Chapter: {chapter_info}")
            else:
                st.info(f"Generating {content_type} content for {grade_level}...")
            
            # Direct API call without PDF content
            messages = [{"role": "user", "content": prompt}]
            
            completion = client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": YOUR_SITE_URL,
                    "X-Title": YOUR_SITE_NAME,
                },
                model=MODEL_NAME,
                messages=messages,
                max_tokens=70000,  # Increased for comprehensive AI textbook chapters
                temperature=0.4,
            )
            
            content = completion.choices[0].message.content
            return content, "Success"
            
        except Exception as e:
            return None, f"Error generating AI content: {str(e)}"
    
    # Regular flow for other subjects with PDF
    if not use_chunked:
        # Standard approach
        try:
            # Get the specific prompt (Mathematics Primary doesn't use model_progression_text)
            if subject_type == "Mathematics Primary (Classes 1-5)":
                prompt = create_specific_prompt(content_type, grade_level, None, subject_type, word_limits)
            else:
                prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type, word_limits)
            
            # Show detailed info for AI chapter generation
            if subject_type == "Artificial Intelligence" and 'ai_chapter_selected' in st.session_state:
                chapter_info = st.session_state['ai_chapter_selected']
                st.info(f"📚 Generating {content_type} content for {grade_level}\n\n📖 Chapter: {chapter_info}")
            else:
                st.info(f"Generating {content_type} content for {grade_level}...")
            
            if use_openrouter_method:
                # Use OpenRouter's recommended direct PDF upload
                messages = create_messages_with_pdf_openrouter(prompt, pdf_bytes, pdf_filename)
                
                # Configure PDF processing engine
                plugins = [
                    {
                        "id": "file-parser",
                        "pdf": {
                            "engine": "pdf-text"  # or "mistral-ocr" for OCR
                        }
                    }
                ]
                
                # Determine max tokens based on content type
                max_tokens = 131072 if (subject_type == "Mathematics" and content_type == "chapter") else 65536
                
                # Make API call with plugins
                completion = client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": YOUR_SITE_URL,
                        "X-Title": YOUR_SITE_NAME,
                    },
                    model=MODEL_NAME,
                    messages=messages,
                    plugins=plugins,
                    max_tokens=max_tokens,
                    temperature=0.3,
                )
            else:
                # Use text extraction method (original or Mistral OCR)
                if pdf_method == "Mistral OCR (Advanced)":
                    st.info("🔍 **DEBUG:** Attempting Mistral OCR processing...")
                    # Use Mistral OCR for advanced text extraction
                    pdf_text, pdf_images = process_pdf_with_mistral_ocr(pdf_bytes, pdf_filename)
                    if pdf_text is None:
                        st.error("🔍 **DEBUG:** Mistral OCR failed - returning error")
                        return None, "Failed to process PDF with Mistral OCR"
                    st.success("🔍 **DEBUG:** Mistral OCR succeeded - proceeding with OCR content")
                    # Create messages with Mistral OCR content
                    messages = create_messages_with_mistral_ocr_content(prompt, pdf_text, pdf_images)
                else:
                    st.info(f"🔍 **DEBUG:** Using method: {pdf_method}")
                # Use original text extraction method
                pdf_text = extract_text_from_pdf(pdf_bytes)
                pdf_images = extract_images_from_pdf(pdf_bytes)
                # Create messages with PDF content
                messages = create_messages_with_pdf_content(prompt, pdf_text, pdf_images)
                
                # Determine max tokens based on content type
                max_tokens = 131072 if (subject_type == "Mathematics" and content_type == "chapter") else 65536
                
                # Make API call
                completion = client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": YOUR_SITE_URL,
                        "X-Title": YOUR_SITE_NAME,
                    },
                    model=MODEL_NAME,
                    messages=messages,
                    max_tokens=max_tokens,
                    temperature=0.3,
                )
            
            result_text = completion.choices[0].message.content
            
            return result_text, "Generated successfully using standard approach."
            
        except Exception as e:
            st.error(f"Error during content generation: {e}")
            if "token" in str(e).lower() or "limit" in str(e).lower():
                st.info("Document might be too large. Trying chunked approach...")
                return generate_specific_content(content_type, pdf_bytes, pdf_filename, grade_level, 
                                               model_progression_text, subject_type, word_limits, use_chunked=True, use_openrouter_method=use_openrouter_method)
            return None, f"Error: {str(e)}"
    else:
        # Chunked approach
        return analyze_with_chunked_approach_for_specific_content(content_type, pdf_bytes, pdf_filename, 
                                                                 grade_level, model_progression_text, subject_type, word_limits, use_openrouter_method, pdf_method)

def analyze_with_chunked_approach_for_specific_content(content_type, pdf_bytes, pdf_filename, 
                                                       grade_level, model_progression_text, subject_type, word_limits, use_openrouter_method, pdf_method="Text Extraction (Original)"):
    """Specialized chunked approach for specific content types"""
    st.info(f"Using chunked approach to generate {content_type} content...")
    
    try:
        # Extract text from PDF for determining page chunks
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(doc)
        
        # Determine chunk size
        pages_per_chunk = 5
        
        # Process chunks
        analysis_results = []
        
        for start_page in range(0, total_pages, pages_per_chunk):
            end_page = min(start_page + pages_per_chunk - 1, total_pages - 1)
            st.info(f"Analyzing chunk: Pages {start_page+1}-{end_page+1} of {total_pages}...")
            
            # Extract text and images for this chunk
            chunk_text = ""
            chunk_images = []
            
            for page_num in range(start_page, end_page + 1):
                page = doc.load_page(page_num)
                chunk_text += f"\n\n--- Page {page_num + 1} ---\n"
                chunk_text += page.get_text()
                
                # Extract all images
                image_list = page.get_images(full=True)
                for img in image_list:  # Include all images from page
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        base64_image = base64.b64encode(image_bytes).decode('utf-8')
                        ext = base_image["ext"]
                        chunk_images.append({
                            "page": page_num + 1,
                            "base64": f"data:image/{ext};base64,{base64_image}",
                        })
                    except:
                        pass
            
            # Analyze this chunk
            chunk_prompt = f"""You are an expert in educational content development for CBSE curriculum.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: You are analyzing CHUNK (Pages {start_page+1}-{end_page+1} of {total_pages}) from a book chapter for **{grade_level} (CBSE)**.
This is just a PORTION of the full chapter - focus only on this section.

You are extracting information relevant for: {content_type.upper()} CONTENT

Please analyze these specific pages and extract:
* Key concepts, topics, and terminology relevant to {content_type}
* Important facts, examples, and explanations
* For any images, their content and relevance
* Any specific information that would be useful for creating {content_type} content

Format your analysis in Markdown. This is just an intermediate step - don't create final content yet.
"""
            
            try:
                if use_openrouter_method:
                    # For chunked OpenRouter approach, we'll need to create a temporary PDF for this chunk
                    # This is a limitation - OpenRouter expects a full PDF, not partial text
                    # So we'll fall back to text-based approach for chunks
                    messages = create_messages_with_pdf_content(chunk_prompt, chunk_text, chunk_images)
                else:
                    messages = create_messages_with_pdf_content(chunk_prompt, chunk_text, chunk_images)
                
                completion = client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": YOUR_SITE_URL,
                        "X-Title": YOUR_SITE_NAME,
                    },
                    model=MODEL_NAME,
                    messages=messages,
                    max_tokens=32768,
                    temperature=0.3,
                )
                
                analysis_chunk = completion.choices[0].message.content
                analysis_results.append(analysis_chunk)
                
            except Exception as chunk_e:
                st.warning(f"Error processing chunk: {chunk_e}")
                analysis_results.append(f"[Error processing pages {start_page+1}-{end_page+1}]")
        
        doc.close()
        
        # Combine chunk analyses
        combined_analyses = "\n\n".join(analysis_results)
        
        # Get the specific prompt for this content type
        specific_prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type, word_limits)
        
        # Create the final integration prompt
        integration_prompt = f"""You are an expert educational content developer for CBSE curriculum.
You have been given analyses of different chunks of a book chapter for **{grade_level} (CBSE)**.
Your task is to create comprehensive {content_type.upper()} content based on these analyses.

This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

**Content Type Requirements:**
{specific_prompt}

**Analyses from different parts of the chapter:**
{combined_analyses}

Based on these analyses, create the complete {content_type.upper()} content in Markdown format.
Ensure it is cohesive, well-structured, and follows all the requirements for {content_type} content.
"""
        
        # Generate the final content
        st.info(f"Creating final {content_type} content...")
        
        try:
            # Determine max tokens
            max_tokens = 131072 if (subject_type == "Mathematics" and content_type == "chapter") else 65536
            
            messages = [{"role": "user", "content": integration_prompt}]
            
            final_completion = client.chat.completions.create(
                extra_headers={
                    "HTTP-Referer": YOUR_SITE_URL,
                    "X-Title": YOUR_SITE_NAME,
                },
                model=MODEL_NAME,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.3,
            )
            
            final_content = final_completion.choices[0].message.content
            
            return final_content, "Generated successfully using chunked approach."
            
        except Exception as integration_e:
            st.error(f"Error during final integration: {integration_e}")
            return combined_analyses, "Partial analysis complete - final integration failed."
            
    except Exception as e:
        st.error(f"Error during chunked analysis: {e}")
        return None, f"Error during chunked analysis: {str(e)}"

def encode_pdf_to_base64(pdf_bytes):
    """Encodes PDF bytes to base64 string for OpenRouter API."""
    return base64.b64encode(pdf_bytes).decode('utf-8')

def create_messages_with_pdf_openrouter(prompt, pdf_bytes, pdf_filename):
    """Creates messages array for OpenRouter API with direct PDF upload."""
    # Encode PDF to base64
    base64_pdf = encode_pdf_to_base64(pdf_bytes)
    data_url = f"data:application/pdf;base64,{base64_pdf}"
    
    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": prompt
                },
                {
                    "type": "file",
                    "file": {
                        "filename": pdf_filename,
                        "file_data": data_url
                    }
                }
            ]
        }
    ]
    
    return messages

def analyze_with_llm_openrouter(pdf_file_bytes, pdf_filename, model_progression_text, grade_level):
    """Analyzes the uploaded PDF chapter using OpenRouter's direct PDF upload approach."""
    st.info("Processing PDF for analysis...")
    
    prompt_content = f"""You are an expert in educational content development, specifically for CBSE curriculum.
Your task is to analyze the provided PDF document, which is a book chapter intended for **{grade_level} (CBSE)**.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes. Please proceed with the analysis and rewriting.

You need to evaluate it against the 'Model Chapter Progression and Elements' provided below.
The final output should be a completely rewritten and improved chapter that incorporates all necessary corrections and adheres to the model, suitable for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines. The book should be according to the NCERT books for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**REQUIRED CONTENT STRUCTURE:**
The improved chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
10. Exercise (with the following 5 question types, each with appropriate number of questions):
   - MCQ
   - Assertion and Reason
   - Fill in the Blanks
   - True False
   - Define the following terms
   - Match the column
   - Give Reason for the following Statement (Easy Level)
   - Answer in Brief (Moderate Level)
   - Answer in Detail (Hard Level)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Formatting Requirements:**
* Format like a proper textbook with well-structured paragraphs (not too long but sufficiently developed)
* Use clear headings and subheadings to organize content
* Include multiple Bloom's Taxonomy questions at different levels in each relevant section
* Create concept-wise summaries (visuals are optional but helpful)
* Integrate special features within the core concept flow rather than as separate sections
* DO NOT change the original concept names/terminology from the uploaded PDF

**Instructions for Analysis and Rewrite:**
1.  **Read and understand the entire PDF chapter.** Pay attention to text, images, diagrams, and overall structure, keeping the **{grade_level}** student in mind.
2.  **Compare the chapter to each element** in the 'Model Chapter Progression and Elements'.
3.  **Identify deficiencies:** Where does the uploaded chapter fall short for a **{grade_level}** audience and the model? Be specific.
4.  **Analyze Images:** For each image in the PDF, assess its:
    *   Relevance to the surrounding text and learning objectives for **{grade_level}**.
    *   Clarity and quality.
    *   Age-appropriateness for **{grade_level}**.
    *   Presence of any specific characters if expected (e.g., 'EeeBee').
5.  **Suggest Image Actions:**
    *   If an image is good, state that it should be kept.
    *   If an image needs improvement (e.g., better clarity, different focus for **{grade_level}**), describe the improvement.
    *   If an image is irrelevant, unsuitable, or missing, provide a detailed prompt for an AI image generator to create a new, appropriate image suitable for **{grade_level}**. The prompt should be clearly marked, e.g., "[PROMPT FOR NEW IMAGE: A detailed description of the image needed, including style, content, characters like 'EeeBee' if applicable, and educational purpose for **{grade_level}**.]"
6.  **Rewrite the Chapter:** Create a new version of the chapter that:
    *   Follows the 'Model Chapter Progression and Elements' strictly.
    *   MUST include ALL elements from the REQUIRED CONTENT STRUCTURE in the order specified.
    *   Corrects all identified mistakes and deficiencies.
    *   Integrates image feedback directly into the text.
    *   Ensures content is aligned with NCERT, NCF, and NEP 2020 principles (e.g., inquiry-based learning, 21st-century skills, real-world connections) appropriate for **{grade_level}**.
    *   Is written in clear, engaging, and age-appropriate language for **{grade_level} (CBSE)**.
    *   Use Markdown for formatting (headings, bold, italics, lists) to make it easy to convert to a Word document. For example:
        # Chapter Title
        ## Section Title
        **Bold Text**
        *Italic Text*
        - Item 1
        - Item 2
    *   Present content in proper textbook paragraphs like a NCERT textbook
    *   Include multiple Bloom's Taxonomy questions at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create) in each appropriate section
    *   Organize summaries by individual concepts rather than as a single summary
    *   Weave special features (misconceptions, 21st century skills, etc.) into the core concept flow
    *   PRESERVE all original concept names and terminology from the source PDF

IMPORTANT: Do NOT directly copy large portions of text from the PDF. Instead, use your own words to rewrite and improve the content while maintaining the original concepts and terminology.

**Output Format:**
Provide the complete, rewritten chapter text in Markdown format, incorporating all analyses and image prompts as described. Do not include a preamble or a summary of your analysis; just the final chapter content.
"""
    
    st.info(f"Sending request to Claude via OpenRouter for {grade_level}... This may take some time for larger documents.")
    
    try:
        # Create messages with direct PDF upload
        messages = create_messages_with_pdf_openrouter(prompt_content, pdf_file_bytes, pdf_filename)
        
        # Optional: Configure PDF processing engine
        plugins = [
            {
                "id": "file-parser",
                "pdf": {
                    "engine": "pdf-text"  # or "mistral-ocr" for OCR
                }
            }
        ]
        
        # Make API call with plugins
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": YOUR_SITE_URL,
                "X-Title": YOUR_SITE_NAME,
            },
            model=MODEL_NAME,
            messages=messages,
            plugins=plugins,  # Add plugins for PDF processing
            max_tokens=65536,
            temperature=0.3,
        )
        
        improved_text = completion.choices[0].message.content
        
        return improved_text, "LLM analysis and rewrite complete using Claude via OpenRouter with direct PDF upload."
        
    except Exception as e:
        st.error(f"Error during OpenRouter API call: {e}")
        # Fallback to the original text extraction method
        st.info("Trying alternative method with text extraction...")
        return analyze_with_llm(pdf_file_bytes, pdf_filename, model_progression_text, grade_level)

def generate_chat_response(user_prompt, chat_history, uploaded_files, grade_level, subject):
    """Generate a response from EeeBee for the chat interface"""
    try:
        # Build context from chat history
        conversation_context = ""
        if len(chat_history) > 1:
            recent_messages = chat_history[-6:]  # Keep last 6 messages for context
            for msg in recent_messages[:-1]:  # Exclude the current message
                conversation_context += f"{msg['role'].title()}: {msg['content']}\n"
        
        # Build system prompt for EeeBee
        system_prompt = f"""You are EeeBee, an expert educational content development assistant specializing in CBSE curriculum.
You help content teams create, modify, and improve educational materials that align with NCERT, NCF, and NEP 2020 guidelines.

Context:
- Target Grade: {grade_level}
- Subject Area: {subject}
- This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY

Your expertise includes:
- Content creation and improvement for CBSE curriculum
- Curriculum alignment and standards compliance
- Age-appropriate content development
- Educational activity design
- Assessment and exercise creation
- Pedagogical best practices
- Integration of 21st-century skills

Guidelines:
- Be helpful, friendly, and professional
- Provide detailed, actionable advice
- Reference educational standards and best practices
- Suggest specific improvements or modifications
- Ask clarifying questions when needed
- Maintain consistency with CBSE/NCERT guidelines

Previous conversation:
{conversation_context}

Current user question: {user_prompt}"""

        # Prepare content parts
        content_parts = [{"type": "text", "text": system_prompt}]
        
        # Add uploaded PDFs if any
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # Extract text from PDF
                    pdf_bytes = uploaded_file.getvalue()
                    pdf_text = extract_text_from_pdf(pdf_bytes)
                    
                    if pdf_text:
                        content_parts.append({
                            "type": "text",
                            "text": f"\n\nContent from {uploaded_file.name}:\n{pdf_text[:10000]}"  # Limit text length
                        })
                    
                except Exception as e:
                    st.warning(f"Could not process {uploaded_file.name}: {e}")
        
        # Create messages
        messages = [{"role": "user", "content": content_parts}]
        
        # Generate response
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": YOUR_SITE_URL,
                "X-Title": YOUR_SITE_NAME,
            },
            model=MODEL_NAME,
            messages=messages,
            max_tokens=8192,
            temperature=0.7,
        )
        
        return completion.choices[0].message.content
            
    except Exception as e:
        return f"I encountered an error: {str(e)}. Please try asking your question again or check if your uploaded files are valid PDFs."

def generate_chat_response_stream(user_prompt, chat_history, uploaded_files, grade_level, subject):
    """Generate a streaming response from EeeBee for the chat interface using requests"""
    try:
        # Build context from chat history
        conversation_context = ""
        if len(chat_history) > 1:
            recent_messages = chat_history[-6:]  # Keep last 6 messages for context
            for msg in recent_messages[:-1]:  # Exclude the current message
                conversation_context += f"{msg['role'].title()}: {msg['content']}\n"
        
        # Build system prompt for EeeBee
        system_prompt = f"""You are EeeBee, an expert educational content development assistant specializing in CBSE curriculum.
You help content teams create, modify, and improve educational materials that align with NCERT, NCF, and NEP 2020 guidelines.

Context:
- Target Grade: {grade_level}
- Subject Area: {subject}
- This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY

Your expertise includes:
- Content creation and improvement for CBSE curriculum
- Curriculum alignment and standards compliance
- Age-appropriate content development
- Educational activity design
- Assessment and exercise creation
- Pedagogical best practices
- Integration of 21st-century skills

Guidelines:
- Be helpful, friendly, and professional
- Provide detailed, actionable advice
- Reference educational standards and best practices
- Suggest specific improvements or modifications
- Ask clarifying questions when needed
- Maintain consistency with CBSE/NCERT guidelines

Previous conversation:
{conversation_context}

Current user question: {user_prompt}"""

        # Prepare content parts
        content_parts = [{"type": "text", "text": system_prompt}]
        
        # Add uploaded PDFs if any
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # For chat, we'll use the Direct PDF Upload method
                    pdf_bytes = uploaded_file.getvalue()
                    pdf_filename = uploaded_file.name
                    
                    # Encode PDF to base64
                    base64_pdf = encode_pdf_to_base64(pdf_bytes)
                    data_url = f"data:application/pdf;base64,{base64_pdf}"
                    
                    # Add as file in content
                    content_parts.append({
                        "type": "file",
                        "file": {
                            "filename": pdf_filename,
                            "file_data": data_url
                        }
                    })
                    
                except Exception as e:
                    st.warning(f"Could not process {uploaded_file.name}: {e}")
        
        # Create messages
        messages = [{"role": "user", "content": content_parts}]
        
        # Use the streaming function
        cancel_event = st.session_state.get('chat_cancel_event', Event())
        
        for chunk in stream_openrouter_response(
            messages=messages,
            model_name=MODEL_NAME,
            max_tokens=8192,
            temperature=0.7,
            plugins=[{"id": "file-parser", "pdf": {"engine": "pdf-text"}}] if uploaded_files else None,
            cancel_event=cancel_event
        ):
            yield chunk
            
    except Exception as e:
        yield f"I encountered an error: {str(e)}. Please try asking your question again or check if your uploaded files are valid PDFs."

# --- Streaming Support Functions ---
def stream_openrouter_response(messages, model_name, max_tokens, temperature, plugins=None, cancel_event=None):
    """
    Stream responses from OpenRouter API with cancellation support.
    Returns a generator that yields response chunks.
    """
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": YOUR_SITE_URL,
        "X-Title": YOUR_SITE_NAME,
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": model_name,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": temperature,
        "stream": True
    }
    
    if plugins:
        payload["plugins"] = plugins
    
    try:
        with requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=payload,
            stream=True,
            timeout=300  # 5 minute timeout
        ) as response:
            response.raise_for_status()
            
            for line in response.iter_lines():
                if cancel_event and cancel_event.is_set():
                    response.close()
                    yield "[CANCELLED] Generation cancelled by user."
                    return
                    
                if line:
                    line_text = line.decode('utf-8')
                    if line_text.startswith('data: '):
                        data_str = line_text[6:]  # Remove 'data: ' prefix
                        if data_str == '[DONE]':
                            return
                        try:
                            data = json.loads(data_str)
                            if 'choices' in data and len(data['choices']) > 0:
                                delta = data['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    yield delta['content']
                        except json.JSONDecodeError:
                            continue
                            
    except requests.exceptions.RequestException as e:
        yield f"\n\n[ERROR] Request failed: {str(e)}"
    except Exception as e:
        yield f"\n\n[ERROR] Unexpected error: {str(e)}"

def analyze_with_llm_streaming(pdf_file_bytes, pdf_filename, model_progression_text, grade_level, use_openrouter_method=False):
    """
    Analyzes the uploaded PDF chapter with streaming support.
    Returns a generator that yields response chunks.
    """
    # Create the prompt (same as before)
    prompt_content = f"""You are an expert in educational content development, specifically for CBSE curriculum.
Your task is to analyze the provided PDF document, which is a book chapter intended for **{grade_level} (CBSE)**.
This is the user's OWN CONTENT being used for EDUCATIONAL PURPOSES ONLY.

IMPORTANT: This is the user's own copyright material, and they have explicitly authorized its analysis and transformation for educational purposes. Please proceed with the analysis and rewriting.

You need to evaluate it against the 'Model Chapter Progression and Elements' provided below.
The final output should be a completely rewritten and improved chapter that incorporates all necessary corrections and adheres to the model, suitable for **{grade_level} (CBSE)**.
The book is intended to align with NCERT, NCF, and NEP 2020 guidelines. The book should be according to the NCERT books for **{grade_level}**.

**Model Chapter Progression and Elements:**
---
{model_progression_text}
---

**REQUIRED CONTENT STRUCTURE:**
The improved chapter MUST follow this specific structure:
1. Current Concepts
2. Hook (with Image Prompt)
3. Learning Outcome
4. Real world connection
5. Previous class concept
6. History
7. Summary
10. Exercise (with the following 5 question types, each with appropriate number of questions):
   - MCQ
   - Assertion and Reason
   - Fill in the Blanks
   - True False
   - Define the following terms
   - Match the column
   - Give Reason for the following Statement (Easy Level)
   - Answer in Brief (Moderate Level)
   - Answer in Detail (Hard Level)
11. Skill based Activity
12. Activity Time – STEM
13. Creativity – Art
14. Case Study – Level 1
15. Case Study – Level 2

**Target Audience:** {grade_level} (CBSE Syllabus)

**Important Formatting Requirements:**
* Format like a proper textbook with well-structured paragraphs (not too long but sufficiently developed)
* Use clear headings and subheadings to organize content
* Include multiple Bloom's Taxonomy questions at different levels in each relevant section
* Create concept-wise summaries (visuals are optional but helpful)
* Integrate special features within the core concept flow rather than as separate sections
* DO NOT change the original concept names/terminology from the uploaded PDF

**Instructions for Analysis and Rewrite:**
1.  **Read and understand the entire PDF chapter.** Pay attention to text, images, diagrams, and overall structure, keeping the **{grade_level}** student in mind.
2.  **Compare the chapter to each element** in the 'Model Chapter Progression and Elements'.
3.  **Identify deficiencies:** Where does the uploaded chapter fall short for a **{grade_level}** audience and the model? Be specific.
4.  **Analyze Images:** For each image in the PDF, assess its:
    *   Relevance to the surrounding text and learning objectives for **{grade_level}**.
    *   Clarity and quality.
    *   Age-appropriateness for **{grade_level}**.
    *   Presence of any specific characters if expected (e.g., 'EeeBee').
5.  **Suggest Image Actions:**
    *   If an image is good, state that it should be kept.
    *   If an image needs improvement (e.g., better clarity, different focus for **{grade_level}**), describe the improvement.
    *   If an image is irrelevant, unsuitable, or missing, provide a detailed prompt for an AI image generator to create a new, appropriate image suitable for **{grade_level}**. The prompt should be clearly marked, e.g., "[PROMPT FOR NEW IMAGE: A detailed description of the image needed, including style, content, characters like 'EeeBee' if applicable, and educational purpose for **{grade_level}**.]"
6.  **Rewrite the Chapter:** Create a new version of the chapter that:
    *   Follows the 'Model Chapter Progression and Elements' strictly.
    *   MUST include ALL elements from the REQUIRED CONTENT STRUCTURE in the order specified.
    *   Corrects all identified mistakes and deficiencies.
    *   Integrates image feedback directly into the text.
    *   Ensures content is aligned with NCERT, NCF, and NEP 2020 principles (e.g., inquiry-based learning, 21st-century skills, real-world connections) appropriate for **{grade_level}**.
    *   Is written in clear, engaging, and age-appropriate language for **{grade_level} (CBSE)**.
    *   Use Markdown for formatting (headings, bold, italics, lists) to make it easy to convert to a Word document.
    *   Present content in proper textbook paragraphs like a NCERT textbook
    *   Include multiple Bloom's Taxonomy questions at various levels (Remember, Understand, Apply, Analyze, Evaluate, Create) in each appropriate section
    *   Organize summaries by individual concepts rather than as a single summary
    *   Weave special features (misconceptions, 21st century skills, etc.) into the core concept flow
    *   PRESERVE all original concept names and terminology from the source PDF

IMPORTANT: Do NOT directly copy large portions of text from the PDF. Instead, use your own words to rewrite and improve the content while maintaining the original concepts and terminology.

**Output Format:**
Provide the complete, rewritten chapter text in Markdown format, incorporating all analyses and image prompts as described. Do not include a preamble or a summary of your analysis; just the final chapter content.
"""
    
    if use_openrouter_method:
        # Create messages with direct PDF upload
        messages = create_messages_with_pdf_openrouter(prompt_content, pdf_file_bytes, pdf_filename)
        plugins = [{"id": "file-parser", "pdf": {"engine": "pdf-text"}}]
    else:
        # Extract text and images from PDF
        pdf_text = extract_text_from_pdf(pdf_file_bytes)
        pdf_images = extract_images_from_pdf(pdf_file_bytes)
        messages = create_messages_with_pdf_content(prompt_content, pdf_text, pdf_images)
        plugins = None
    
    # Create cancel event (can be controlled from UI)
    cancel_event = st.session_state.get('cancel_event', Event())
    
    # Stream the response
    for chunk in stream_openrouter_response(
        messages=messages,
        model_name=MODEL_NAME,
        max_tokens=65536,
        temperature=0.3,
        plugins=plugins,
        cancel_event=cancel_event
    ):
        yield chunk

def generate_specific_content_streaming(content_type, pdf_bytes, pdf_filename, grade_level, 
                                       model_progression_text, subject_type, word_limits, use_openrouter_method, pdf_method="Text Extraction (Original)"):
    """
    Generates specific content with streaming support.
    Returns a generator that yields response chunks.
    """
    # Get the specific prompt (Mathematics Primary doesn't use model_progression_text)
    if subject_type == "Mathematics Primary (Classes 1-5)":
        prompt = create_specific_prompt(content_type, grade_level, None, subject_type, word_limits)
    else:
        prompt = create_specific_prompt(content_type, grade_level, model_progression_text, subject_type, word_limits)
    
    # Special handling for AI subject without PDF
    if subject_type == "Artificial Intelligence" and pdf_bytes is None:
        # Direct streaming without PDF for AI textbook creation
        messages = [{"role": "user", "content": prompt}]
        
        # Stream the response
        stream = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": YOUR_SITE_URL,
                "X-Title": YOUR_SITE_NAME,
            },
            model=MODEL_NAME,
            messages=messages,
            max_tokens=70000,  # Increased for comprehensive AI textbook chapters
            temperature=0.4,
            stream=True
        )
        
        for chunk in stream:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content
        return
    
    # Regular flow with PDF
    if use_openrouter_method:
        # Create messages with direct PDF upload
        messages = create_messages_with_pdf_openrouter(prompt, pdf_bytes, pdf_filename)
        plugins = [{"id": "file-parser", "pdf": {"engine": "pdf-text"}}]
    else:
        # Use text extraction method (original, Mistral OCR, or other)
        if pdf_method == "Mistral OCR (Advanced)":
            st.info("🔍 **DEBUG (Streaming):** Attempting Mistral OCR processing...")
            # Use Mistral OCR for advanced text extraction
            pdf_text, pdf_images = process_pdf_with_mistral_ocr(pdf_bytes, pdf_filename)
            if pdf_text is None:
                st.error("🔍 **DEBUG (Streaming):** Mistral OCR failed - cannot proceed with streaming")
                return  # Stop the generator
            st.success("🔍 **DEBUG (Streaming):** Mistral OCR succeeded - proceeding with streaming")
            # Create messages with Mistral OCR content
            messages = create_messages_with_mistral_ocr_content(prompt, pdf_text, pdf_images)
        else:
            st.info(f"🔍 **DEBUG (Streaming):** Using method: {pdf_method}")
            # Extract text and images from PDF using original method
        pdf_text = extract_text_from_pdf(pdf_bytes)
        pdf_images = extract_images_from_pdf(pdf_bytes)
        messages = create_messages_with_pdf_content(prompt, pdf_text, pdf_images)
        plugins = None
    
    # Determine max tokens based on content type and subject
    if subject_type == "Artificial Intelligence":
        max_tokens = 70000  # Maximum tokens for comprehensive AI textbook chapters
    elif subject_type == "Mathematics" and content_type == "chapter":
        max_tokens = 32768
    else:
        max_tokens = 16384
    
    # Create cancel event (can be controlled from UI)
    cancel_event = st.session_state.get('cancel_event', Event())
    
    # Stream the response
    for chunk in stream_openrouter_response(
        messages=messages,
        model_name=MODEL_NAME,
        max_tokens=max_tokens,
        temperature=0.3,
        plugins=plugins,
        cancel_event=cancel_event
    ):
        yield chunk

def handle_streaming_generation(content_type, pdf_bytes, pdf_filename, selected_grade, 
                               model_progression, subject_type, word_limits, button_key, pdf_method="Text Extraction (Original)"):
    """Helper function to handle streaming content generation with UI"""
    # Initialize cancel event for streaming
    st.session_state.cancel_event = Event()
    cancel_col1, cancel_col2 = st.columns([5, 1])
    with cancel_col2:
        cancel_button = st.button("🛑 Cancel", key=f"cancel_{button_key}")
        if cancel_button:
            st.session_state.cancel_event.set()
    
    # Create a container for streaming content
    content_container = st.container()
    with content_container:
        content_placeholder = st.empty()
        accumulated_content = ""
        content_saved = False
        
        try:
            # Stream the content with the selected PDF method
            for chunk in generate_specific_content_streaming(
                content_type,
                pdf_bytes, 
                pdf_filename, 
                selected_grade,
                model_progression, 
                subject_type,
                word_limits,
                use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                pdf_method=pdf_method
            ):
                if st.session_state.cancel_event.is_set():
                    st.warning("⚠️ Generation cancelled by user.")
                    break
                
                accumulated_content += chunk
                
                # CRITICAL: Auto-save during streaming to prevent loss
                auto_save_during_streaming(content_type, accumulated_content)
                
                # Update the placeholder with accumulated content
                with content_placeholder.container():
                    st.markdown(f"### Generated {content_type.title()} Content:")
                    st.markdown(accumulated_content + " ⏳")
            
            # Always try to save content if we have any, regardless of completion status
            if accumulated_content:
                # CRITICAL: Final save with all protection strategies
                save_content_safely(content_type, accumulated_content)
                
                with content_placeholder.container():
                    if st.session_state.cancel_event.is_set():
                        st.markdown(f"### ⚠️ {content_type.title()} Content (Cancelled - Partial):")
                    else:
                        st.markdown(f"### ✅ {content_type.title()} Content (Complete):")
                    st.markdown(accumulated_content)
                
                # Determine success message
                if st.session_state.cancel_event.is_set():
                    return accumulated_content, "Partial content saved (generation cancelled by user)"
                else:
                    return accumulated_content, "Generated successfully using streaming!"
            else:
                return None, "No content generated"
                
        except Exception as e:
            st.error(f"❌ Error during streaming: {e}")
            # Always try to save any content we managed to generate
            if accumulated_content:
                st.info("💾 Saving partial content despite error...")
                
                # CRITICAL: Save partial content with protection
                save_content_safely(content_type, accumulated_content)
                
                with content_placeholder.container():
                    st.markdown(f"### ⚠️ {content_type.title()} Content (Partial - Error Occurred):")
                    st.markdown(accumulated_content)
                return accumulated_content, f"Partial content saved due to error: {str(e)}"
            return None, f"Error: {str(e)}"

# --- Hybrid Content Expansion System ---
def parse_content_sections(content: str) -> List[Dict[str, Any]]:
    """Automatically detect expandable sections in generated content"""
    sections = []
    
    if not content:
        return sections
    
    # Split content into lines for processing
    lines = content.split('\n')
    current_section = ""
    section_type = "paragraph"
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
            
        # Detect different section types
        if re.match(r'^#{1,6}\s+(.+)$', line):  # Headings
            if current_section:
                sections.append({
                    'type': section_type,
                    'text': current_section.strip(),
                    'expandable': True,
                    'line_start': max(0, i-10),
                    'line_end': min(len(lines), i+10)
                })
            current_section = line
            section_type = 'heading'
            
        elif re.match(r'^\*\*([^*]+)\*\*', line):  # Bold concepts
            if current_section:
                sections.append({
                    'type': section_type,
                    'text': current_section.strip(),
                    'expandable': True,
                    'line_start': max(0, i-5),
                    'line_end': min(len(lines), i+15)
                })
            current_section = line
            section_type = 'concept'
            
        elif re.match(r'^[-*+]\s+(.+)$', line):  # List items
            if section_type != 'list':
                if current_section:
                    sections.append({
                        'type': section_type,
                        'text': current_section.strip(),
                        'expandable': True,
                        'line_start': max(0, i-5),
                        'line_end': min(len(lines), i+10)
                    })
                current_section = line
                section_type = 'list'
            else:
                current_section += "\n" + line
                
        elif len(line) > 50:  # Regular paragraphs
            if section_type in ['heading', 'concept'] and current_section:
                current_section += "\n" + line
            else:
                if current_section and section_type != 'paragraph':
                    sections.append({
                        'type': section_type,
                        'text': current_section.strip(),
                        'expandable': True,
                        'line_start': max(0, i-5),
                        'line_end': min(len(lines), i+10)
                    })
                if section_type != 'paragraph':
                    current_section = line
                    section_type = 'paragraph'
                else:
                    current_section += "\n" + line
    
    # Add the last section
    if current_section:
        sections.append({
            'type': section_type,
            'text': current_section.strip(),
            'expandable': True,
            'line_start': 0,
            'line_end': len(lines)
        })
    
    return sections

def expand_text_with_ai(selected_text: str, expansion_type: str, context: str, content_type: str, grade_level: str, subject_type: str) -> str:
    """Generate expanded content for selected text using AI"""
    
    expansion_prompts = {
        'detail': "Provide more detailed explanation and depth",
        'examples': "Add relevant examples, case studies, and practical applications",
        'activities': "Add hands-on activities, experiments, and interactive exercises",
        'simplify': "Make the explanation simpler and clearer for students",
        'questions': "Add thought-provoking questions and assessments",
        'connections': "Add connections to other concepts and real-world applications",
        'shorten': "Make the content more concise and brief while keeping key information",
        'summarize': "Create a condensed summary of the main points",
        'bullet_points': "Convert the content into clear, concise bullet points"
    }
    
    prompt = f"""You are an expert educational content developer for CBSE {grade_level} curriculum.

TASK: Expand the selected educational content with {expansion_prompts.get(expansion_type, 'more detail')}.

CONTEXT:
- Content Type: {content_type}
- Grade Level: {grade_level} (CBSE)
- Subject: {subject_type}
- Selected Text: "{selected_text}"
- Surrounding Context: "{context[:500]}..."

EXPANSION TYPE: {expansion_type.upper()}

REQUIREMENTS:
1. Keep the same educational tone and CBSE alignment
2. Make content appropriate for {grade_level} students
3. Maintain consistency with existing content
4. Target length: {"50-70% of original length (make it more concise)" if expansion_type in ['shorten', 'summarize', 'bullet_points'] else "2-3x the original length (expand it)"}
5. Use clear, engaging language
6. Include proper formatting (markdown)

EXPANDED CONTENT:"""
    
    try:
        # Create messages for AI
        messages = [{"role": "user", "content": prompt}]
        
        # Generate expansion
        completion = client.chat.completions.create(
            extra_headers={
                "HTTP-Referer": YOUR_SITE_URL,
                "X-Title": YOUR_SITE_NAME,
            },
            model=MODEL_NAME,
            messages=messages,
            max_tokens=8192,
            temperature=0.4,
        )
        
        return completion.choices[0].message.content
        
    except Exception as e:
        return f"Error generating expansion: {str(e)}"

def display_section_expander(sections: List[Dict[str, Any]], content_type: str, grade_level: str, subject_type: str):
    """Display expandable sections with expansion buttons"""
    
    st.subheader("🔍 Section-by-Section Expansion")
    st.markdown("*Click the expand button next to any section you want to develop further.*")
    
    # Get fresh content from session state to ensure we have the latest version
    def get_current_content():
        if content_type == "chapter":
            return st.session_state.get("chapter_content", "")
        elif content_type == "exercises":
            return st.session_state.get("exercises", "")
        elif content_type == "skills":
            return st.session_state.get("skill_activities", "")
        elif content_type == "art":
            return st.session_state.get("art_learning", "")
        return ""
    
    for i, section in enumerate(sections):
        # Create container for each section
        with st.container():
            col1, col2 = st.columns([5, 1])
            
            with col1:
                # Display section based on type
                if section['type'] == 'heading':
                    st.markdown(section['text'])
                elif section['type'] == 'concept':
                    st.markdown(section['text'])
                else:
                    # Truncate long paragraphs for display
                    display_text = section['text'][:200] + "..." if len(section['text']) > 200 else section['text']
                    st.markdown(display_text)
            
            with col2:
                # Expansion button
                if st.button("🔍 Expand", key=f"expand_section_{i}", help=f"Expand this {section['type']}"):
                    # Show expansion options
                    st.session_state[f"show_expansion_options_{i}"] = True
                    st.rerun()
            
            # Show expansion options if button was clicked
            if st.session_state.get(f"show_expansion_options_{i}", False):
                with st.container():
                    st.markdown("**Expand Content:**")
                    exp_col1, exp_col2, exp_col3, exp_col4 = st.columns(4)
                    
                    expansion_type = None
                    if exp_col1.button("📝 More Detail", key=f"detail_{i}"):
                        expansion_type = "detail"
                    if exp_col2.button("💡 Add Examples", key=f"examples_{i}"):
                        expansion_type = "examples"
                    if exp_col3.button("🎯 Add Activities", key=f"activities_{i}"):
                        expansion_type = "activities"
                    if exp_col4.button("📖 Simplify", key=f"simplify_{i}"):
                        expansion_type = "simplify"
                    
                    st.markdown("**Shorten Content:**")
                    short_col1, short_col2, short_col3 = st.columns(3)
                    
                    if short_col1.button("✂️ Make Shorter", key=f"shorten_{i}"):
                        expansion_type = "shorten"
                    if short_col2.button("📋 Summarize", key=f"summarize_{i}"):
                        expansion_type = "summarize"
                    if short_col3.button("• Bullet Points", key=f"bullets_{i}"):
                        expansion_type = "bullet_points"
                    
                    if expansion_type:
                        with st.spinner(f"🧠 Expanding section with {expansion_type}..."):
                            expanded_content = expand_text_with_ai(
                                selected_text=section['text'],
                                expansion_type=expansion_type,
                                context=section['text'],  # For now, use section itself as context
                                content_type=content_type,
                                grade_level=grade_level,
                                subject_type=subject_type
                            )
                            
                            st.session_state[f"expanded_content_{i}"] = expanded_content
                            st.session_state[f"show_expansion_options_{i}"] = False
                            st.rerun()
            
            # Show expanded content if available
            if st.session_state.get(f"expanded_content_{i}"):
                with st.expander("✨ Expanded Content", expanded=True):
                    st.markdown(st.session_state[f"expanded_content_{i}"])
                    
                    # Option to replace original content
                    col1, col2 = st.columns(2)
                    if col1.button("🔄 Replace Section", key=f"replace_{i}"):
                        # Get the current content (fresh from session state)
                        original_content = get_current_content()
                        
                        if original_content and section['text'] in original_content:
                            # Replace the section in the original content
                            updated_content = original_content.replace(section['text'], st.session_state[f"expanded_content_{i}"])
                            
                            # Update the session state with the modified content
                            if content_type == "chapter":
                                st.session_state.chapter_content = updated_content
                            elif content_type == "exercises":
                                st.session_state.exercises = updated_content
                            elif content_type == "skills":
                                st.session_state.skill_activities = updated_content
                            elif content_type == "art":
                                st.session_state.art_learning = updated_content
                            
                            # Save the updated content
                            save_content_safely(content_type, updated_content)
                            
                            # Clear the expanded content and expansion options to prevent stale data
                            st.session_state.pop(f"expanded_content_{i}", None)
                            st.session_state.pop(f"show_expansion_options_{i}", None)
                            
                            st.success("✅ Section replaced in original content!")
                            st.rerun()
                        else:
                            # Try to find a partial match if exact match fails
                            section_words = section['text'].split()[:10]  # First 10 words
                            section_start = " ".join(section_words)
                            
                            if section_start in original_content:
                                # Find the full section in original content
                                start_idx = original_content.find(section_start)
                                # Try to find the end of the section (next paragraph or section)
                                remaining_content = original_content[start_idx:]
                                lines = remaining_content.split('\n')
                                
                                # Reconstruct the full section
                                full_section = ""
                                for line in lines:
                                    if line.strip():
                                        full_section += line + "\n"
                                    else:
                                        break
                                
                                full_section = full_section.strip()
                                
                                if full_section:
                                    # Replace with the reconstructed section
                                    updated_content = original_content.replace(full_section, st.session_state[f"expanded_content_{i}"])
                                    
                                    # Update the session state
                                    if content_type == "chapter":
                                        st.session_state.chapter_content = updated_content
                                    elif content_type == "exercises":
                                        st.session_state.exercises = updated_content
                                    elif content_type == "skills":
                                        st.session_state.skill_activities = updated_content
                                    elif content_type == "art":
                                        st.session_state.art_learning = updated_content
                                    
                                    # Save the updated content
                                    save_content_safely(content_type, updated_content)
                                    
                                    # Clear the expanded content and expansion options
                                    st.session_state.pop(f"expanded_content_{i}", None)
                                    st.session_state.pop(f"show_expansion_options_{i}", None)
                                    
                                    st.success("✅ Section replaced in original content!")
                                    st.rerun()
                                else:
                                    st.error("❌ Could not reconstruct the original section to replace.")
                            else:
                                st.error("❌ Could not find the original section to replace. Try using the Manual Text Expansion instead.")
                    
                    if col2.button("💾 Keep Both", key=f"keep_{i}"):
                        st.success("✅ Expanded content saved separately! Original content unchanged.")
            
            st.divider()

def display_manual_text_expander(original_content: str, content_type: str, grade_level: str, subject_type: str):
    """Allow users to manually specify text for expansion"""
    
    st.subheader("🎯 Manual Text Expansion")
    st.markdown("*Copy and paste any specific text from your content that you want to expand.*")
    
    # User input for text selection
    selected_text = st.text_area(
        "Paste the specific text you want to expand:",
        height=100,
        placeholder="Example: 'Current Concepts' or 'Mathematical explanation' or any specific paragraph...",
        key="manual_text_input"
    )
    
    if selected_text.strip():
        # Expansion type selection
        st.markdown("**Expand Content:**")
        exp_col1, exp_col2, exp_col3 = st.columns(3)
        exp_col4, exp_col5, exp_col6 = st.columns(3)
        
        expansion_type = None
        if exp_col1.button("📝 More Detail", key="manual_detail"):
            expansion_type = "detail"
        if exp_col2.button("💡 Add Examples", key="manual_examples"):
            expansion_type = "examples"
        if exp_col3.button("🎯 Add Activities", key="manual_activities"):
            expansion_type = "activities"
        if exp_col4.button("📖 Simplify", key="manual_simplify"):
            expansion_type = "simplify"
        if exp_col5.button("❓ Add Questions", key="manual_questions"):
            expansion_type = "questions"
        if exp_col6.button("🔗 Add Connections", key="manual_connections"):
            expansion_type = "connections"
        
        st.markdown("**Shorten Content:**")
        short_col1, short_col2, short_col3 = st.columns(3)
        
        if short_col1.button("✂️ Make Shorter", key="manual_shorten"):
            expansion_type = "shorten"
        if short_col2.button("📋 Summarize", key="manual_summarize"):
            expansion_type = "summarize"
        if short_col3.button("• Bullet Points", key="manual_bullets"):
            expansion_type = "bullet_points"
        
        if expansion_type:
            with st.spinner(f"🧠 Expanding your text with {expansion_type}..."):
                # Find context around the selected text
                context = ""
                if selected_text in original_content:
                    start_idx = original_content.find(selected_text)
                    context_start = max(0, start_idx - 200)
                    context_end = min(len(original_content), start_idx + len(selected_text) + 200)
                    context = original_content[context_start:context_end]
                
                expanded_content = expand_text_with_ai(
                    selected_text=selected_text,
                    expansion_type=expansion_type,
                    context=context,
                    content_type=content_type,
                    grade_level=grade_level,
                    subject_type=subject_type
                )
                
                st.markdown("### ✨ Expanded Content:")
                st.markdown(expanded_content)
                
                # Options for the expanded content
                col1, col2, col3 = st.columns(3)
                
                if col1.button("🔄 Replace Original", key="replace_manual_expansion"):
                    # Replace the selected text in the original content
                    original_content_state = ""
                    if content_type == "chapter":
                        original_content_state = st.session_state.get("chapter_content", "")
                    elif content_type == "exercises":
                        original_content_state = st.session_state.get("exercises", "")
                    elif content_type == "skills":
                        original_content_state = st.session_state.get("skill_activities", "")
                    elif content_type == "art":
                        original_content_state = st.session_state.get("art_learning", "")
                    
                    if original_content_state and selected_text in original_content_state:
                        # Replace the selected text with expanded content
                        updated_content = original_content_state.replace(selected_text, expanded_content)
                        
                        # Update the session state
                        if content_type == "chapter":
                            st.session_state.chapter_content = updated_content
                        elif content_type == "exercises":
                            st.session_state.exercises = updated_content
                        elif content_type == "skills":
                            st.session_state.skill_activities = updated_content
                        elif content_type == "art":
                            st.session_state.art_learning = updated_content
                        
                        # Save the updated content
                        save_content_safely(content_type, updated_content)
                        st.success("✅ Original content updated with expansion!")
                        st.rerun()
                    else:
                        st.error("❌ Could not find the selected text in original content to replace.")
                
                if col2.button("💾 Save Expansion", key="save_manual_expansion"):
                    # Store in session state for later use
                    if 'saved_expansions' not in st.session_state:
                        st.session_state.saved_expansions = []
                    
                    st.session_state.saved_expansions.append({
                        'original': selected_text,
                        'expanded': expanded_content,
                        'type': expansion_type,
                        'timestamp': time.time()
                    })
                    st.success("✅ Expansion saved! You can view all saved expansions below.")
                
                if col3.button("📋 Copy to Clipboard", key="copy_manual_expansion"):
                    # This would copy to clipboard if we had JavaScript support
                    st.info("💡 You can manually copy the expanded content above.")

def display_global_content_expander(original_content: str, content_type: str, grade_level: str, subject_type: str):
    """Global content expansion options"""
    
    st.subheader("🚀 Global Content Enhancement")
    st.markdown("*Enhance the entire content with specific improvements.*")
    
    st.markdown("**Expand Content:**")
    col1, col2, col3 = st.columns(3)
    
    global_action = None
    if col1.button("📏 Make Content Longer", key="global_longer"):
        global_action = "longer"
    if col2.button("🎯 Add More Activities", key="global_activities"):
        global_action = "activities"
    if col3.button("💡 Add More Examples", key="global_examples"):
        global_action = "examples"
    
    st.markdown("**Shorten Content:**")
    short_col1, short_col2, short_col3 = st.columns(3)
    
    if short_col1.button("✂️ Make Shorter", key="global_shorten"):
        global_action = "shorten"
    if short_col2.button("📋 Summarize All", key="global_summarize"):
        global_action = "summarize"
    if short_col3.button("• Key Points Only", key="global_bullets"):
        global_action = "bullet_points"
    
    if global_action:
        with st.spinner(f"🧠 Enhancing entire content..."):
            if global_action == "longer":
                prompt = f"""Make this {content_type} content significantly longer and more detailed for {grade_level} CBSE students.
                
Original Content:
{original_content}

Add more depth, explanations, and comprehensive coverage while maintaining the same structure and educational quality."""
                
            elif global_action == "activities":
                prompt = f"""Add more hands-on activities, exercises, and interactive elements throughout this {content_type} content for {grade_level} CBSE students.

Original Content:
{original_content}

Integrate practical activities that reinforce learning and engage students actively."""
                
            elif global_action == "examples":
                prompt = f"""Add more real-world examples, case studies, and practical applications throughout this {content_type} content for {grade_level} CBSE students.

Original Content:
{original_content}

Include diverse examples that help students understand concepts better."""
                
            elif global_action == "shorten":
                prompt = f"""Make this {content_type} content more concise and brief for {grade_level} CBSE students while keeping all essential information and concepts.

Original Content:
{original_content}

Remove unnecessary details, redundant explanations, and verbose sections while maintaining educational quality and completeness."""
                
            elif global_action == "summarize":
                prompt = f"""Create a condensed summary of this {content_type} content for {grade_level} CBSE students, focusing on key concepts and main points.

Original Content:
{original_content}

Extract and present the most important information in a clear, organized summary format."""
                
            elif global_action == "bullet_points":
                prompt = f"""Convert this {content_type} content into clear, concise bullet points for {grade_level} CBSE students.

Original Content:
{original_content}

Organize the content into well-structured bullet points that capture all key concepts and information."""
            
            try:
                messages = [{"role": "user", "content": prompt}]
                
                completion = client.chat.completions.create(
                    extra_headers={
                        "HTTP-Referer": YOUR_SITE_URL,
                        "X-Title": YOUR_SITE_NAME,
                    },
                    model=MODEL_NAME,
                    messages=messages,
                    max_tokens=16384,
                    temperature=0.4,
                )
                
                enhanced_content = completion.choices[0].message.content
                
                st.markdown("### ✨ Enhanced Content:")
                st.markdown(enhanced_content)
                
                # Option to replace original content
                col1, col2 = st.columns(2)
                if col1.button("🔄 Replace Original Content", key=f"replace_global_{global_action}"):
                    # Update the session state with enhanced content
                    if content_type == "chapter":
                        st.session_state.chapter_content = enhanced_content
                    elif content_type == "exercises":
                        st.session_state.exercises = enhanced_content
                    elif content_type == "skills":
                        st.session_state.skill_activities = enhanced_content
                    elif content_type == "art":
                        st.session_state.art_learning = enhanced_content
                    
                    st.success("✅ Original content replaced with enhanced version!")
                    st.rerun()
                
                if col2.button("💾 Save as New Version", key=f"save_global_{global_action}"):
                    # Save as a new version
                    st.session_state[f"{content_type}_enhanced_{global_action}"] = enhanced_content
                    st.success("✅ Enhanced content saved as new version!")
                
            except Exception as e:
                st.error(f"Error enhancing content: {str(e)}")

def hybrid_content_expander(content: str, content_type: str, grade_level: str, subject_type: str):
    """Complete hybrid content expansion system"""
    
    if not content:
        st.warning("No content available for expansion.")
        return
    
    st.markdown("---")
    st.header("✨ Content Expansion Studio")
    st.markdown("*Enhance your generated content with AI-powered expansions*")
    
    # Get fresh content from session state
    def get_fresh_content():
        if content_type == "chapter":
            return st.session_state.get("chapter_content", "")
        elif content_type == "exercises":
            return st.session_state.get("exercises", "")
        elif content_type == "skills":
            return st.session_state.get("skill_activities", "")
        elif content_type == "art":
            return st.session_state.get("art_learning", "")
        return content  # fallback to original content
    
    # Tab interface for different expansion methods
    tab1, tab2, tab3, tab4 = st.tabs(["🔍 Auto-Sections", "🎯 Manual Select", "🚀 Global Enhance", "💾 Saved Expansions"])
    
    with tab1:
        # Get fresh content and re-parse sections each time the tab is accessed
        fresh_content = get_fresh_content()
        sections = parse_content_sections(fresh_content)
        
        if sections:
            display_section_expander(sections, content_type, grade_level, subject_type)
        else:
            st.info("No expandable sections detected. Try the Manual Select tab to expand specific text.")
    
    with tab2:
        fresh_content = get_fresh_content()
        display_manual_text_expander(fresh_content, content_type, grade_level, subject_type)
    
    with tab3:
        fresh_content = get_fresh_content()
        display_global_content_expander(fresh_content, content_type, grade_level, subject_type)
    
    with tab4:
        st.subheader("💾 Your Saved Expansions")
        if 'saved_expansions' in st.session_state and st.session_state.saved_expansions:
            for i, expansion in enumerate(st.session_state.saved_expansions):
                with st.expander(f"Expansion {i+1}: {expansion['type'].title()}", expanded=False):
                    st.markdown("**Original Text:**")
                    st.markdown(expansion['original'][:200] + "..." if len(expansion['original']) > 200 else expansion['original'])
                    st.markdown("**Expanded Version:**")
                    st.markdown(expansion['expanded'])
                    
                    if st.button(f"🗑️ Delete", key=f"delete_expansion_{i}"):
                        st.session_state.saved_expansions.pop(i)
                        st.rerun()
        else:
            st.info("No saved expansions yet. Use the other tabs to create expansions.")

# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("📚 EeeBee Content Development Suite ✨ (OpenRouter Edition)")

st.markdown("""
Welcome to the EeeBee Content Development Suite powered by Claude via OpenRouter! 
Choose between improving existing chapters or chatting with EeeBee for content assistance.
""")

# Create tabs for different functionalities
tab1, tab2, tab3, tab4 = st.tabs(["📚 Chapter Improver", "💬 Content Chat with EeeBee", "🔍 PDF Checker", "🩹 Remedial Content"])

with tab1:
    st.header("📚 Book Chapter Improver Tool")
    st.markdown("""
    Upload your book chapter in PDF format. This tool will analyze it using EeeBee (Claude)
    based on the 'Model Chapter Progression and Elements' and suggest improvements.
    Select which part of the content you want to generate.
    
    ✨ **New Features**: 
    - **Content Protection**: Your content is automatically saved with multiple backups to prevent loss
    - **Expand Content**: Use the **"Expand Content"** buttons to enhance any generated content with:
      - **Auto-Section Detection**: Click specific sections to expand
      - **Manual Text Selection**: Paste any text you want to develop further  
      - **Global Enhancement**: Make entire content longer or add more examples/activities
    
    💾 **Content Status**: Check the sidebar to monitor your content safety and recovery options.
    """)

    # Subject Type Selector
    subject_type = st.selectbox(
        "Select Subject Type:",
        ["Science (Uses Model Chapter Progression)", "Mathematics", "Mathematics Primary (Classes 1-5)", "Science & E.V.S. (Classes 1-2)", "Science & E.V.S. (Classes 3-5)", "Computer Science", "English Communication & Grammar (Classes 1-8)", "Artificial Intelligence", "Robotics"],
        help="Choose the appropriate subject type based on your needs. Science EVS options are specialized for primary grades. English option uses best practices from Oxford, Cambridge, and Wren & Martin. AI and Robotics options are designed for hands-on learning with IIT kits.",
        key="subject_selector_tab1"
    )

    # Grade/Level Selector - Show appropriate options based on subject
    if subject_type == "Artificial Intelligence":
        # CBSE Grade Selector for AI (Classes 9-12)
        col1, col2 = st.columns(2)
        
        with col1:
            grade_options = [
                "Class 9 (CBSE Code 417)",
                "Class 10 (CBSE Code 417)",
                "Class 11 (CBSE Code 843)",
                "Class 12 (CBSE Code 843)"
            ]
            selected_grade = st.selectbox(
                "Select CBSE Grade:", 
                grade_options, 
                index=0,
                key="ai_grade_selector",
                help="Select the CBSE grade level for AI content generation"
            )
        
        with col2:
            # Chapter selection based on grade
            chapter_options = {
                "Class 9 (CBSE Code 417)": [
                    "Chapter 1: Introduction to AI - Foundational Concepts",
                    "Chapter 2: AI Project Cycle - Problem Scoping",
                    "Chapter 3: Data Acquisition and Exploration",
                    "Chapter 4: Modelling Basics",
                    "Chapter 5: Model Evaluation",
                    "Chapter 6: Introduction to Python for AI",
                    "Chapter 7: Data Science Domain",
                    "Chapter 8: Computer Vision Basics",
                    "Chapter 9: Natural Language Processing Basics",
                    "Chapter 10: AI Ethics and Bias"
                ],
                "Class 10 (CBSE Code 417)": [
                    "Chapter 1: Introduction to AI - Foundational Concepts, Intelligence & Decision Making",
                    "Chapter 2: AI Basics - AI vs ML vs DL, AI Domains & Applications",
                    "Chapter 3: AI Project Cycle - Problem Scoping & SDGs",
                    "Chapter 4: Data Acquisition & Data Exploration - Visualizing Data",
                    "Chapter 5: AI Modeling - Rule-Based vs Learning-Based, Supervised & Unsupervised Learning",
                    "Chapter 6: Advanced Python & Data Science - NumPy, Pandas, Matplotlib",
                    "Chapter 7: Computer Vision - Image Basics, Pixels, RGB, OpenCV Introduction",
                    "Chapter 8: Natural Language Processing - Chatbots, Text Processing, Bag of Words",
                    "Chapter 9: Model Evaluation - Confusion Matrix, Accuracy, Precision, Recall, F1 Score",
                    "Chapter 10: AI Ethics, Bias & Real-World Applications"
                ],
                "Class 11 (CBSE Code 843)": [
                    "Chapter 1: Python Programming for AI",
                    "Chapter 2: Statistics and Probability for AI",
                    "Chapter 3: Data Visualization with Python",
                    "Chapter 4: Introduction to NumPy and Pandas",
                    "Chapter 5: Supervised Learning - Classification",
                    "Chapter 6: Supervised Learning - Regression",
                    "Chapter 7: Unsupervised Learning",
                    "Chapter 8: Introduction to Neural Networks",
                    "Chapter 9: Natural Language Processing with Python",
                    "Chapter 10: Computer Vision with OpenCV"
                ],
                "Class 12 (CBSE Code 843)": [
                    "Chapter 1: Deep Learning Fundamentals",
                    "Chapter 2: Convolutional Neural Networks (CNNs)",
                    "Chapter 3: Recurrent Neural Networks (RNNs)",
                    "Chapter 4: Generative Adversarial Networks (GANs)",
                    "Chapter 5: Transfer Learning and Fine-tuning",
                    "Chapter 6: Advanced NLP - Transformers and BERT",
                    "Chapter 7: Computer Vision - Object Detection and Segmentation",
                    "Chapter 8: Reinforcement Learning Basics",
                    "Chapter 9: AI Ethics, Bias, and Fairness",
                    "Chapter 10: AI Capstone Project"
                ]
            }
            
            selected_chapter = st.selectbox(
                "Select Chapter to Generate:",
                chapter_options[selected_grade],
                key="ai_chapter_selector",
                help="Choose a specific chapter aligned with CBSE curriculum"
            )
            
            # Store chapter info for use in prompt
            st.session_state['ai_chapter_selected'] = selected_chapter
            
    elif subject_type == "Robotics":
        # Level Selector for Robotics
        level_options = [
            "JL1 (Classes 1-3)",
            "JL2 (Classes 4-5)", 
            "SL1 (Classes 6-7)",
            "SL2 (Classes 8-10)",
            "SL3 (Class 11)"
        ]
        selected_level = st.selectbox("Select Target Level:", level_options, index=2, key="level_selector_tab1") # Default to SL1
        
        # For Robotics, use the level as grade_level for prompts
        selected_grade = selected_level
    else:
        # Regular Grade Selector for other subjects
        grade_options = [f"Grade {i}" for i in range(1, 13)] # Grades 1-12
        selected_grade = st.selectbox("Select Target Grade Level (CBSE):", grade_options, index=8, key="grade_selector_tab1") # Default to Grade 9
        selected_level = None  # No level for other subjects

    # Word Limit Controls
    st.subheader("📝 Content Length Settings")
    with st.expander("Configure Word Limits for Each Section", expanded=False):
        if subject_type == "Mathematics Primary (Classes 1-5)":
            # Special word limits for Primary Mathematics
            st.markdown("**Mathematics Primary (Classes 1-5) Structure**")
            col1, col2 = st.columns(2)
            
            with col1:
                hook_words = st.number_input("Chapter Opener - The Hook (words)", min_value=100, value=150, step=10, key="primary_hook_words")
                discover_words = st.number_input("Let's Discover - Concept & Practice (words)", min_value=1000, value=1500, step=100, key="primary_discover_words")
                activity_words = st.number_input("Activity Zone - Hands-on (words)", min_value=200, value=300, step=25, key="primary_activity_words")
            
            with col2:
                recap_words = st.number_input("Quick Recap - Revision (words)", min_value=150, value=200, step=25, key="primary_recap_words")
                exercises_words = st.number_input("Exercises (words)", min_value=500, value=800, step=50, key="primary_exercises_words")
            
            # Store word limits for primary mathematics
            st.session_state.word_limits = {
                'hook': hook_words,
                'discover': discover_words,
                'activity': activity_words,
                'recap': recap_words,
                'exercises': exercises_words
            }
        elif subject_type == "Science & E.V.S. (Classes 1-2)":
            # Special word limits for Science EVS Classes 1-2
            st.markdown("**Science & E.V.S. (Classes 1-2) Play-Based Structure**")
            col1, col2 = st.columns(2)
            
            with col1:
                opener_words = st.number_input("Chapter Opener - Story and Wonder (words)", min_value=100, value=200, step=25, key="evs12_opener_words")
                activity_words = st.number_input("Let's Play & Do - Per Activity (words)", min_value=150, value=250, step=25, key="evs12_activity_words")
                concept_words = st.number_input("Concept Connect - Per Concept (words)", min_value=100, value=200, step=25, key="evs12_concept_words")
            
            with col2:
                closer_words = st.number_input("Chapter Closer - Learning Together (words)", min_value=100, value=150, step=25, key="evs12_closer_words")
                activities_words = st.number_input("Play-Based Activities Total (words)", min_value=400, value=600, step=50, key="evs12_activities_words")
            
            # Store word limits for Science EVS Classes 1-2
            st.session_state.word_limits = {
                'opener': opener_words,
                'activity': activity_words,
                'concept': concept_words,
                'closer': closer_words,
                'activities': activities_words
            }
        elif subject_type == "Science & E.V.S. (Classes 3-5)":
            # Special word limits for Science EVS Classes 3-5
            st.markdown("**Science & E.V.S. (Classes 3-5) Inquiry-Based Structure**")
            col1, col2 = st.columns(2)
            
            with col1:
                opener_words = st.number_input("Chapter Opener Page (words)", min_value=200, value=300, step=25, key="evs35_opener_words")
                exploration_words = st.number_input("Let's Uncover the Secrets (words)", min_value=1500, value=2000, step=100, key="evs35_exploration_words")
                project_words = st.number_input("Capstone Project (words)", min_value=300, value=400, step=50, key="evs35_project_words")
            
            with col2:
                summary_words = st.number_input("Summary & Revision Tools (words)", min_value=200, value=250, step=25, key="evs35_summary_words")
                exercises_words = st.number_input("End-of-Chapter Exercises (words)", min_value=500, value=600, step=50, key="evs35_exercises_words")
            
            # Store word limits for Science EVS Classes 3-5
            st.session_state.word_limits = {
                'opener': opener_words,
                'exploration': exploration_words,
                'project': project_words,
                'summary': summary_words,
                'exercises': exercises_words
            }
        elif subject_type == "English Communication & Grammar (Classes 1-8)":
            # Special word limits for English Communication & Grammar
            st.markdown("**English Communication & Grammar (Oxford/Cambridge Style)**")
            
            # Determine class level for appropriate structure
            class_num = int(selected_grade.split()[-1])
            
            if class_num <= 3:
                st.markdown("*Primary Level (Classes 1-3): Foundation Building*")
                col1, col2 = st.columns(2)
                
                with col1:
                    warm_up_words = st.number_input("Warm-Up Activity (words)", min_value=100, value=200, step=25, key="eng_warmup_words")
                    vocabulary_words = st.number_input("Vocabulary Building (words)", min_value=200, value=300, step=25, key="eng_vocab_words")
                    grammar_intro_words = st.number_input("Grammar Introduction (words)", min_value=150, value=250, step=25, key="eng_grammar_words")
                
                with col2:
                    practice_words = st.number_input("Practice Activities (words)", min_value=300, value=400, step=50, key="eng_practice_words")
                    communication_words = st.number_input("Communication Skills (words)", min_value=200, value=300, step=25, key="eng_comm_words")
                    exercises_words = st.number_input("Exercises & Assessment (words)", min_value=400, value=600, step=50, key="eng_exercises_words")
                
                st.session_state.word_limits = {
                    'warm_up': warm_up_words,
                    'vocabulary': vocabulary_words,
                    'grammar_intro': grammar_intro_words,
                    'practice': practice_words,
                    'communication': communication_words,
                    'exercises': exercises_words
                }
            
            elif class_num <= 5:
                st.markdown("*Elementary Level (Classes 4-5): Skill Development*")
                col1, col2 = st.columns(2)
                
                with col1:
                    introduction_words = st.number_input("Chapter Introduction (words)", min_value=150, value=250, step=25, key="eng_intro_words")
                    reading_words = st.number_input("Reading Comprehension (words)", min_value=400, value=600, step=50, key="eng_reading_words")
                    grammar_words = st.number_input("Grammar Focus (words)", min_value=300, value=500, step=50, key="eng_grammar_focus_words")
                    vocabulary_words = st.number_input("Vocabulary Expansion (words)", min_value=250, value=350, step=25, key="eng_vocab_exp_words")
                
                with col2:
                    writing_words = st.number_input("Writing Skills (words)", min_value=300, value=450, step=50, key="eng_writing_words")
                    speaking_words = st.number_input("Speaking & Listening (words)", min_value=250, value=350, step=25, key="eng_speaking_words")
                    exercises_words = st.number_input("Practice Exercises (words)", min_value=500, value=700, step=50, key="eng_exercises_elem_words")
                
                st.session_state.word_limits = {
                    'introduction': introduction_words,
                    'reading': reading_words,
                    'grammar': grammar_words,
                    'vocabulary': vocabulary_words,
                    'writing': writing_words,
                    'speaking': speaking_words,
                    'exercises': exercises_words
                }
            
            else:  # Classes 6-8
                st.markdown("*Middle School Level (Classes 6-8): Advanced Communication*")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("**Core Skills**")
                    introduction_words = st.number_input("Chapter Introduction (words)", min_value=200, value=300, step=25, key="eng_intro_mid_words")
                    reading_words = st.number_input("Reading & Comprehension (words)", min_value=500, value=800, step=50, key="eng_reading_mid_words")
                    grammar_words = st.number_input("Grammar & Usage (words)", min_value=400, value=600, step=50, key="eng_grammar_mid_words")
                
                with col2:
                    st.markdown("**Communication**")
                    vocabulary_words = st.number_input("Vocabulary & Etymology (words)", min_value=300, value=400, step=25, key="eng_vocab_mid_words")
                    writing_words = st.number_input("Composition & Writing (words)", min_value=400, value=600, step=50, key="eng_writing_mid_words")
                    speaking_words = st.number_input("Speaking & Presentation (words)", min_value=300, value=400, step=25, key="eng_speaking_mid_words")
                
                with col3:
                    st.markdown("**Assessment**")
                    literature_words = st.number_input("Literature Appreciation (words)", min_value=300, value=500, step=50, key="eng_lit_words")
                    exercises_words = st.number_input("Comprehensive Exercises (words)", min_value=600, value=900, step=50, key="eng_exercises_mid_words")
                    projects_words = st.number_input("Language Projects (words)", min_value=200, value=300, step=25, key="eng_projects_words")
                
                st.session_state.word_limits = {
                    'introduction': introduction_words,
                    'reading': reading_words,
                    'grammar': grammar_words,
                    'vocabulary': vocabulary_words,
                    'writing': writing_words,
                    'speaking': speaking_words,
                    'literature': literature_words,
                    'exercises': exercises_words,
                    'projects': projects_words
                }
        elif subject_type == "Artificial Intelligence":
            # Enhanced word limits for comprehensive AI textbook chapters
            st.markdown("**AI Textbook Chapter Structure (O'Reilly/CBSE Standards)**")
            st.info("📚 Creating publisher-quality AI textbook chapters with comprehensive content")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Core Content**")
                introduction_words = st.number_input(
                    "Chapter Introduction & Hook (words)", 
                    min_value=300, value=400, step=50, 
                    key="ai_intro_words",
                    help="Real-world hook, objectives, roadmap"
                )
                concepts_words = st.number_input(
                    "Core AI Concepts & Theory (words)", 
                    min_value=2000, value=2500, step=100, 
                    key="ai_concepts_words",
                    help="Detailed explanations, algorithms, code examples"
                )
                hands_on_words = st.number_input(
                    "Hands-On Labs & Projects (words)", 
                    min_value=1200, value=1500, step=100, 
                    key="ai_handson_words",
                    help="Step-by-step project builds, experiments"
                )
            
            with col2:
                st.markdown("**Applications & Assessment**")
                case_studies_words = st.number_input(
                    "Case Studies & Applications (words)", 
                    min_value=600, value=800, step=50, 
                    key="ai_cases_words",
                    help="Industry examples, Indian context"
                )
                exercises_words = st.number_input(
                    "Exercises & Practice (words)", 
                    min_value=800, value=1000, step=50, 
                    key="ai_exercises_words",
                    help="MCQs, coding challenges, problems"
                )
                projects_words = st.number_input(
                    "Extended Projects & Resources (words)", 
                    min_value=500, value=600, step=50, 
                    key="ai_projects_words",
                    help="Capstone ideas, career connections"
                )
            
            # Display selected chapter info
            if 'ai_chapter_selected' in st.session_state:
                st.success(f"📖 Generating: {st.session_state['ai_chapter_selected']}")
            
            # Store word limits for AI
            st.session_state.word_limits = {
                'introduction': introduction_words,
                'concepts': concepts_words,
                'hands_on': hands_on_words,
                'case_studies': case_studies_words,
                'exercises': exercises_words,
                'projects': projects_words
            }
        elif subject_type == "Robotics":
            # Special word limits for Robotics
            st.markdown("**Robotics Structure**")
            col1, col2 = st.columns(2)
            
            with col1:
                mission_briefing_words = st.number_input("Mission Briefing (words)", min_value=200, value=300, step=25, key="robotics_mission_words")
                domain_analysis_words = st.number_input("Domain Analysis & Applications (words)", min_value=600, value=800, step=50, key="robotics_domain_words")
                core_concepts_words = st.number_input("Core Concepts (words)", min_value=2000, value=3000, step=100, key="robotics_concepts_words")
            
            with col2:
                hands_on_project_words = st.number_input("Hands-On Project (words)", min_value=1000, value=1500, step=100, key="robotics_project_words")
                assessment_words = st.number_input("Assessment & Debrief (words)", min_value=400, value=600, step=50, key="robotics_assessment_words")
                exercises_words = st.number_input("Exercises (words)", min_value=800, value=1000, step=50, key="robotics_exercises_words")
            
            # Store word limits for Robotics
            st.session_state.word_limits = {
                'mission_briefing': mission_briefing_words,
                'domain_analysis': domain_analysis_words,
                'core_concepts': core_concepts_words,
                'hands_on_project': hands_on_project_words,
                'assessment': assessment_words,
                'exercises': exercises_words
            }
        else:
            # Standard word limits for other subjects
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown("**Core Sections**")
                hook_words = st.number_input("Hook (words)", min_value=20, value=80, step=10, key="hook_words")
                learning_outcome_words = st.number_input("Learning Outcome (words)", min_value=30, value=70, step=10, key="learning_outcome_words")
                real_world_words = st.number_input("Real World Connection (words)", min_value=30, value=50, step=10, key="real_world_words")
                previous_class_words = st.number_input("Previous Class Concept (words)", min_value=30, value=100, step=10, key="previous_class_words")
                history_words = st.number_input("History (words)", min_value=50, value=100, step=10, key="history_words")
            
            with col2:
                st.markdown("**Main Content**")
                current_concepts_words = st.number_input("Current Concepts (words)", min_value=500, value=1200, step=100, key="current_concepts_words")
                summary_words = st.number_input("Summary (words)", min_value=300, value=700, step=50, key="summary_words")
                link_learn_words = st.number_input("Link and Learn Questions (words)", min_value=100, value=250, step=25, key="link_learn_words")
                image_based_words = st.number_input("Image Based Questions (words)", min_value=100, value=250, step=25, key="image_based_words")
            
            with col3:
                st.markdown("**Activities & Exercises**")
                exercises_words = st.number_input("Exercise Questions (words)", min_value=300, value=800, step=50, key="exercises_words")
                skill_activity_words = st.number_input("Skill Activities (words)", min_value=200, value=400, step=50, key="skill_activity_words")
                stem_activity_words = st.number_input("STEM Activities (words)", min_value=200, value=400, step=50, key="stem_activity_words")
                art_learning_words = st.number_input("Art Learning (words)", min_value=200, value=400, step=50, key="art_learning_words")

            # Store word limits in session state
            st.session_state.word_limits = {
                'hook': hook_words,
                'learning_outcome': learning_outcome_words,
                'real_world': real_world_words,
                'previous_class': previous_class_words,
                'history': history_words,
                'current_concepts': current_concepts_words,
                'summary': summary_words,
                'link_learn': link_learn_words,
                'image_based': image_based_words,
                'exercises': exercises_words,
                'skill_activity': skill_activity_words,
                'stem_activity': stem_activity_words,
                'art_learning': art_learning_words
            }

    # Only show analysis and PDF processing options for non-AI subjects
    if subject_type != "Artificial Intelligence":
        # Analysis Method Selector
        analysis_method = st.radio(
            "Choose Analysis Method:",
            ["Standard (Full Document)", "Chunked (For Complex Documents)"],
            help="Use 'Chunked' method for very large documents or if you encounter errors with the standard method.",
            key="analysis_method_tab1"
        )

        # PDF Processing Method Selector
        pdf_method = st.radio(
            "Choose PDF Processing Method:",
            ["Text Extraction (Original)", "Mistral OCR (Advanced)", "Direct PDF Upload (OpenRouter Recommended)"],
            help="Original: Basic text extraction | Mistral OCR: Advanced OCR with structure preservation | Direct Upload: Sends PDF to OpenRouter",
            key="pdf_method_tab1"
        )
        
        # Show info for Mistral OCR
        if pdf_method == "Mistral OCR (Advanced)":
            st.info("🔬 **Mistral OCR Features:** Preserves document structure, handles complex layouts, extracts text from images, maintains tables and formulas in markdown format")
            st.info("📋 **Requirements:** Mistral API key (add MISTRAL_API_KEY to Streamlit secrets)")
    else:
        # Set defaults for AI subject
        analysis_method = "Standard (Full Document)"
        pdf_method = "Text Extraction (Original)"

    # Streaming Mode Toggle
    use_streaming = st.checkbox(
        "Enable Streaming Mode",
        value=True,
        help="Stream responses in real-time as they are generated. You can cancel generation at any time.",
        key="streaming_mode_tab1"
    )

    # Load Model Chapter Progression
    model_progression = load_model_chapter_progression()

    # Initialize session state for content persistence (MOVED OUTSIDE FILE UPLOAD CONDITION)
    if 'chapter_content' not in st.session_state:
        st.session_state.chapter_content = None
    if 'exercises' not in st.session_state:
        st.session_state.exercises = None
    if 'skill_activities' not in st.session_state:
        st.session_state.skill_activities = None
    if 'art_learning' not in st.session_state:
        st.session_state.art_learning = None
    
    # Auto-recover any lost content on startup
    verify_and_recover_all_content()
    
    # Display content protection status
    display_content_status()

    if model_progression:
        st.sidebar.subheader("Model Chapter Progression:")
        st.sidebar.text_area("Model Details", model_progression, height=300, disabled=True)

        # Display previously generated content (MOVED OUTSIDE FILE UPLOAD CONDITION)
        col_header1, col_header2 = st.columns([4, 1])
        with col_header1:
            st.subheader("📋 Previously Generated Content")
        with col_header2:
            if st.button("🗑️ Clear All Content", key="clear_all_content", help="Clear all generated content from session"):
                st.session_state.chapter_content = None
                st.session_state.exercises = None
                st.session_state.skill_activities = None
                st.session_state.art_learning = None
                st.success("✅ All content cleared!")
                st.rerun()
        
        prev_col1, prev_col2, prev_col3, prev_col4 = st.columns(4)
        
        with prev_col1:
            if st.session_state.chapter_content:
                with st.expander("📖 Chapter Content Available", expanded=False):
                    st.markdown(st.session_state.chapter_content[:500] + "..." if len(st.session_state.chapter_content) > 500 else st.session_state.chapter_content)
                    
                    col_dl, col_expand = st.columns(2)
                    with col_dl:
                        doc = create_word_document(st.session_state.chapter_content)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download",
                            data=doc_io,
                            file_name="chapter_content.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="prev_download_chapter"
                        )
                    with col_expand:
                        if st.button("✨ Expand Content", key="expand_chapter_btn"):
                            st.session_state.show_chapter_expander = True
                            st.rerun()
            else:
                st.info("📖 No chapter content generated yet")
        
        with prev_col2:
            if st.session_state.exercises:
                with st.expander("📝 Exercises Available", expanded=False):
                    st.markdown(st.session_state.exercises[:500] + "..." if len(st.session_state.exercises) > 500 else st.session_state.exercises)
                    
                    col_dl, col_expand = st.columns(2)
                    with col_dl:
                        doc = create_word_document(st.session_state.exercises)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download",
                            data=doc_io,
                            file_name="exercises.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="prev_download_exercises"
                        )
                    with col_expand:
                        if st.button("✨ Expand Content", key="expand_exercises_btn"):
                            st.session_state.show_exercises_expander = True
                            st.rerun()
            else:
                st.info("📝 No exercises generated yet")
        
        with prev_col3:
            if st.session_state.skill_activities:
                with st.expander("🛠️ Skills Available", expanded=False):
                    st.markdown(st.session_state.skill_activities[:500] + "..." if len(st.session_state.skill_activities) > 500 else st.session_state.skill_activities)
                    
                    col_dl, col_expand = st.columns(2)
                    with col_dl:
                        doc = create_word_document(st.session_state.skill_activities)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download",
                            data=doc_io,
                            file_name="skill_activities.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="prev_download_skills"
                        )
                    with col_expand:
                        if st.button("✨ Expand Content", key="expand_skills_btn"):
                            st.session_state.show_skills_expander = True
                            st.rerun()
            else:
                st.info("🛠️ No skill activities generated yet")
        
        with prev_col4:
            if st.session_state.art_learning:
                with st.expander("🎨 Art Learning Available", expanded=False):
                    st.markdown(st.session_state.art_learning[:500] + "..." if len(st.session_state.art_learning) > 500 else st.session_state.art_learning)
                    
                    col_dl, col_expand = st.columns(2)
                    with col_dl:
                        doc = create_word_document(st.session_state.art_learning)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        st.download_button(
                            label="📥 Download",
                            data=doc_io,
                            file_name="art_learning.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="prev_download_art"
                        )
                    with col_expand:
                        if st.button("✨ Expand Content", key="expand_art_btn"):
                            st.session_state.show_art_expander = True
                            st.rerun()
            else:
                st.info("🎨 No art learning generated yet")

        # Content Expansion Displays (when expand buttons are clicked)
        if st.session_state.get('show_chapter_expander', False):
            hybrid_content_expander(
                st.session_state.chapter_content, 
                "chapter", 
                selected_grade, 
                subject_type.replace(" (Uses Model Chapter Progression)", "")
            )
            if st.button("❌ Close Expander", key="close_chapter_expander"):
                st.session_state.show_chapter_expander = False
                st.rerun()
        
        if st.session_state.get('show_exercises_expander', False):
            hybrid_content_expander(
                st.session_state.exercises, 
                "exercises", 
                selected_grade, 
                subject_type.replace(" (Uses Model Chapter Progression)", "")
            )
            if st.button("❌ Close Expander", key="close_exercises_expander"):
                st.session_state.show_exercises_expander = False
                st.rerun()
        
        if st.session_state.get('show_skills_expander', False):
            hybrid_content_expander(
                st.session_state.skill_activities, 
                "skills", 
                selected_grade, 
                subject_type.replace(" (Uses Model Chapter Progression)", "")
            )
            if st.button("❌ Close Expander", key="close_skills_expander"):
                st.session_state.show_skills_expander = False
                st.rerun()
        
        if st.session_state.get('show_art_expander', False):
            hybrid_content_expander(
                st.session_state.art_learning, 
                "art", 
                selected_grade, 
                subject_type.replace(" (Uses Model Chapter Progression)", "")
            )
            if st.button("❌ Close Expander", key="close_art_expander"):
                st.session_state.show_art_expander = False
                st.rerun()

        # Check if AI subject is selected - no PDF needed for AI
        if subject_type == "Artificial Intelligence":
            st.success("🎯 AI Textbook Generation Mode - No PDF Required!")
            st.markdown("""
            ### 📚 Generate Original AI Content
            You're creating a fresh AI textbook chapter from scratch based on CBSE curriculum.
            Select the content type below to generate:
            """)
            
            # Skip PDF upload and go directly to generation buttons
            uploaded_file_st = None  # No file needed
            pdf_bytes = None  # No PDF bytes needed
        else:
            # Regular PDF upload for other subjects
            uploaded_file_st = st.file_uploader("Upload your chapter (PDF only)", type="pdf", key="pdf_uploader_tab1")

            if uploaded_file_st is not None:
                st.info(f"Processing '{uploaded_file_st.name}' for {selected_grade}...")
                
                # Get PDF bytes for processing
                pdf_bytes = uploaded_file_st.getvalue()
            else:
                pdf_bytes = None

        st.divider()
        
        st.subheader("🚀 Generate New Content")

        # Show buttons for all subjects including AI
        if subject_type == "Artificial Intelligence":
            # AI-specific buttons
            col1, col2 = st.columns(2)
            col3, col4 = st.columns(2)
            
            # Generate Chapter Content Button
            generate_chapter = col1.button("📚 Generate AI Chapter", key="gen_ai_chapter")
            
            # Generate Exercises Button
            generate_exercises = col2.button("📝 Generate AI Exercises", key="gen_ai_exercises")
            
            # Generate Skill Activities Button
            generate_skills = col3.button("🤖 Generate AI Labs", key="gen_ai_skills")
            
            # Generate Art Learning Button
            generate_art = col4.button("🎨 Generate AI Projects", key="gen_ai_art")
            
        elif subject_type == "Mathematics Primary (Classes 1-5)":
            # For primary mathematics, show chapter and exercises generation
            st.info("📘 **Mathematics Primary Mode**: For Classes 1-5, we provide specialized content generation designed for young learners.")
            
            # Show chapter and exercises buttons for primary mathematics
            prim_col1, prim_col2 = st.columns(2)
            
            with prim_col1:
                generate_chapter = st.button("🔍 Generate Complete Mathematics Chapter", key="gen_primary_chapter")
            
            with prim_col2:
                generate_exercises = st.button("📝 Generate Mathematics Exercises", key="gen_primary_exercises")
            
            generate_skills = False  
            generate_art = False
        elif subject_type == "Science & E.V.S. (Classes 1-2)":
            # For Science EVS Classes 1-2, show chapter and activities generation
            st.info("🌱 **Science & E.V.S. Classes 1-2**: Play-based learning approach for foundational stage. Activities are integrated within the chapter.")
            
            # Show chapter and activities buttons for Science EVS Classes 1-2
            evs12_col1, evs12_col2 = st.columns(2)
            
            with evs12_col1:
                generate_chapter = st.button("🔍 Generate Play-Based Chapter", key="gen_evs12_chapter")
            
            with evs12_col2:
                generate_exercises = st.button("🎮 Generate Play Activities", key="gen_evs12_activities")
            
            generate_skills = False  
            generate_art = False
        elif subject_type == "Science & E.V.S. (Classes 3-5)":
            # For Science EVS Classes 3-5, show chapter and exercises generation
            st.info("🔬 **Science & E.V.S. Classes 3-5**: Inquiry-based learning approach for preparatory stage. Skills and art activities are integrated within the chapter.")
            
            # Show chapter and exercises buttons for Science EVS Classes 3-5
            evs35_col1, evs35_col2 = st.columns(2)
            
            with evs35_col1:
                generate_chapter = st.button("🔍 Generate Inquiry-Based Chapter", key="gen_evs35_chapter")
            
            with evs35_col2:
                generate_exercises = st.button("📝 Generate Assessment Exercises", key="gen_evs35_exercises")
            
            generate_skills = False  
            generate_art = False
        else:
                # Standard buttons for other subjects - only show if PDF is uploaded
                if uploaded_file_st is not None:
                    col1, col2 = st.columns(2)
                    col3, col4 = st.columns(2)
                    
                    # Generate Chapter Content Button
                    generate_chapter = col1.button("🔍 Generate Chapter Content", key="gen_chapter")
                    
                    # Generate Exercises Button
                    generate_exercises = col2.button("📝 Generate Exercises", key="gen_exercises")
                    
                    # Generate Skill Activities Button
                    generate_skills = col3.button("🛠️ Generate Skill Activities", key="gen_skills")
                    
                    # Generate Art Learning Button
                    generate_art = col4.button("🎨 Generate Art-Integrated Learning", key="gen_art")
                else:
                    st.info("📤 Please upload a PDF file to generate content.")
                    generate_chapter = False
                    generate_exercises = False
                    generate_skills = False
                    generate_art = False
            
        # Download All Button (outside columns)
        download_all = st.button("📥 Download Complete Chapter with All Elements", key="download_all")
        
        # Handle button clicks and content generation
        if generate_chapter:
                # Get word limits from session state
                word_limits = st.session_state.get('word_limits', {})
                
                with st.spinner(f"🧠 Generating Chapter Content for {selected_grade}..."):
                    if use_streaming:
                        # Initialize cancel event for streaming
                        st.session_state.cancel_event = Event()
                        
                        # Create cancel button and content container
                        cancel_col1, cancel_col2 = st.columns([5, 1])
                        with cancel_col2:
                            cancel_button = st.button("🛑 Cancel", key="cancel_chapter")
                            if cancel_button:
                                st.session_state.cancel_event.set()
                        
                        # Create container for streaming content
                        content_container = st.container()
                        with content_container:
                            content_placeholder = st.empty()
                            accumulated_content = ""
                            
                            try:
                                # Stream the content
                                # For AI subject without PDF, pass None for pdf_bytes and a placeholder name
                                pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Chapter_Generation"
                                
                                for chunk in generate_specific_content_streaming(
                                    "chapter",
                                    pdf_bytes,  # Will be None for AI
                                    pdf_filename, 
                                    selected_grade,
                                    model_progression, 
                                    subject_type,
                                    word_limits,
                                    use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                                    pdf_method=pdf_method
                                ):
                                    if st.session_state.cancel_event.is_set():
                                        st.warning("⚠️ Generation cancelled by user.")
                                        break
                                    accumulated_content += chunk
                                    
                                    # CRITICAL: Auto-save during streaming to prevent loss
                                    auto_save_during_streaming("chapter_content", accumulated_content)
                                    
                                    # Update the placeholder with accumulated content (in a code block to preserve formatting)
                                    with content_placeholder.container():
                                        st.markdown("### Generated Chapter Content:")
                                        st.markdown(accumulated_content + " ⏳")
                                
                                # Always save content if we have any, regardless of completion status
                                if accumulated_content:
                                    st.session_state.chapter_content = accumulated_content
                                    
                                    # CRITICAL: Final save with all protection strategies
                                    save_content_safely("chapter_content", accumulated_content, selected_grade)
                                    
                                    with content_placeholder.container():
                                        if st.session_state.cancel_event.is_set():
                                            st.markdown("### ⚠️ Chapter Content (Cancelled - Partial):")
                                        else:
                                            st.markdown("### ✅ Chapter Content (Complete):")
                                        st.markdown(st.session_state.chapter_content)
                                    
                                    if st.session_state.cancel_event.is_set():
                                        st.warning("⚠️ Chapter Content partially generated (cancelled by user) but saved!")
                                    else:
                                        st.success(f"✅ Chapter Content generated successfully!")
                                    
                                    # Download button (available even for partial content)
                                    doc = create_word_document(st.session_state.chapter_content)
                                    doc_io = io.BytesIO()
                                    doc.save(doc_io)
                                    doc_io.seek(0)
                                    download_filename = f"chapter_content_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"chapter_content_{selected_grade.replace(' ', '_')}.docx"
                                    st.download_button(
                                        label="📥 Download Chapter Content as Word (.docx)",
                                        data=doc_io,
                                        file_name=download_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_chapter_streaming"
                                    )
                                else:
                                    st.error("❌ No content was generated.")
                                    
                            except Exception as e:
                                st.error(f"❌ Error during streaming: {e}")
                                # Always save any partial content
                                if accumulated_content:
                                    st.info("💾 Saving partial content despite error...")
                                    st.session_state.chapter_content = accumulated_content
                                    
                                    # CRITICAL: Save partial content with protection
                                    save_content_safely("chapter_content", accumulated_content, selected_grade)
                                    
                                    with content_placeholder.container():
                                        st.markdown("### ⚠️ Chapter Content (Partial - Error Occurred):")
                                        st.markdown(st.session_state.chapter_content)
                                    
                                    # Download button for partial content
                                    doc = create_word_document(st.session_state.chapter_content)
                                    doc_io = io.BytesIO()
                                    doc.save(doc_io)
                                    doc_io.seek(0)
                                    download_filename = f"partial_chapter_content_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"partial_chapter_content_{selected_grade.replace(' ', '_')}.docx"
                                    st.download_button(
                                        label="📥 Download Partial Chapter Content as Word (.docx)",
                                        data=doc_io,
                                        file_name=download_filename,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_chapter_partial"
                                    )
                    else:
                        # Use non-streaming approach
                        # For AI subject without PDF, pass None for pdf_bytes and a placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Chapter_Generation"
                        
                        content, message = generate_specific_content(
                            "chapter", 
                            pdf_bytes, 
                            pdf_filename, 
                            selected_grade, 
                            model_progression, 
                            subject_type, 
                            word_limits,
                            use_chunked=(analysis_method == "Chunked (For Complex Documents)"), 
                            use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                            pdf_method=pdf_method
                        )
                        if content:
                            st.session_state.chapter_content = content
                            
                            # CRITICAL: Save content with protection
                            save_content_safely("chapter_content", content, selected_grade)
                            
                            st.success(f"✅ Chapter Content generated successfully! {message}")
                            st.subheader("📖 Chapter Content:")
                            with st.expander("View Chapter Content", expanded=True):
                                st.markdown(st.session_state.chapter_content)
                            
                            # Download and Expand buttons
                            dl_col, expand_col = st.columns(2)
                            with dl_col:
                                doc = create_word_document(st.session_state.chapter_content)
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                download_filename = f"chapter_content_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"chapter_content_{selected_grade.replace(' ', '_')}.docx"
                                st.download_button(
                                    label="📥 Download Chapter Content as Word (.docx)",
                                    data=doc_io,
                                    file_name=download_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_chapter_standard"
                                )
                            with expand_col:
                                if st.button("✨ Expand This Content", key="expand_new_chapter"):
                                    st.session_state.show_new_chapter_expander = True
                                    st.rerun()
                        
                        # Show expander for newly generated content
                        if st.session_state.get('show_new_chapter_expander', False):
                            hybrid_content_expander(
                                st.session_state.chapter_content, 
                                "chapter", 
                                selected_grade, 
                                subject_type.replace(" (Uses Model Chapter Progression)", "")
                            )
                            if st.button("❌ Close Expander", key="close_new_chapter_expander"):
                                st.session_state.show_new_chapter_expander = False
                                st.rerun()
                        else:
                            st.error(f"❌ Failed to generate Chapter Content: {message}")
            
        if generate_exercises:
                # Get word limits from session state
                word_limits = st.session_state.get('word_limits', {})
                
                with st.spinner(f"🧠 Generating Exercises for {selected_grade}..."):
                    if use_streaming:
                        # For AI subject without PDF, pass placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Exercises_Generation"
                        content, message = handle_streaming_generation(
                            "exercises", pdf_bytes, pdf_filename, selected_grade, 
                            model_progression, subject_type, word_limits, "exercises", pdf_method
                        )
                        if content:
                            st.session_state.exercises = content
                            st.success(f"✅ Exercises generated successfully! {message}")
                            
                            # Download button
                            doc = create_word_document(st.session_state.exercises)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            # Generate filename based on whether PDF was uploaded
                            if uploaded_file_st:
                                download_filename = f"exercises_{uploaded_file_st.name.replace('.pdf', '.docx')}"
                            else:
                                download_filename = f"exercises_{selected_grade.replace(' ', '_')}.docx"
                            
                            st.download_button(
                                label="📥 Download Exercises as Word (.docx)",
                                data=doc_io,
                                file_name=download_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_exercises_streaming"
                            )
                        else:
                            st.error(f"❌ Failed to generate Exercises: {message}")
                    else:
                        # For AI subject without PDF, pass None for pdf_bytes and a placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Exercises_Generation"
                        
                        content, message = generate_specific_content(
                            "exercises", 
                            pdf_bytes, 
                            pdf_filename, 
                            selected_grade, 
                            model_progression, 
                            subject_type, 
                            word_limits,
                            use_chunked=(analysis_method == "Chunked (For Complex Documents)"), 
                            use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                            pdf_method=pdf_method
                        )
                        if content:
                            st.session_state.exercises = content
                            
                            # CRITICAL: Save content with protection
                            save_content_safely("exercises", content, selected_grade)
                            
                            st.success(f"✅ Exercises generated successfully! {message}")
                            st.subheader("📝 Exercises:")
                            with st.expander("View Exercises", expanded=True):
                                st.markdown(st.session_state.exercises)
                            
                            # Download and Expand buttons
                            dl_col, expand_col = st.columns(2)
                            with dl_col:
                                doc = create_word_document(st.session_state.exercises)
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                download_filename = f"exercises_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"exercises_{selected_grade.replace(' ', '_')}.docx"
                                st.download_button(
                                    label="📥 Download Exercises as Word (.docx)",
                                    data=doc_io,
                                    file_name=download_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_exercises_standard"
                                )
                            with expand_col:
                                if st.button("✨ Expand This Content", key="expand_new_exercises"):
                                    st.session_state.show_new_exercises_expander = True
                                    st.rerun()
                        
                        # Show expander for newly generated exercises
                        if st.session_state.get('show_new_exercises_expander', False):
                            hybrid_content_expander(
                                st.session_state.exercises, 
                                "exercises", 
                                selected_grade, 
                                subject_type.replace(" (Uses Model Chapter Progression)", "")
                            )
                            if st.button("❌ Close Expander", key="close_new_exercises_expander"):
                                st.session_state.show_new_exercises_expander = False
                                st.rerun()
                        else:
                            st.error(f"❌ Failed to generate Exercises: {message}")
            
        if generate_skills:
                # Get word limits from session state
                word_limits = st.session_state.get('word_limits', {})
                
                with st.spinner(f"🧠 Generating Skill Activities for {selected_grade}..."):
                    if use_streaming:
                        # For AI subject without PDF, pass placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Skills_Generation"
                        content, message = handle_streaming_generation(
                            "skills", pdf_bytes, pdf_filename, selected_grade, 
                            model_progression, subject_type, word_limits, "skills", pdf_method
                        )
                        if content:
                            st.session_state.skill_activities = content
                            st.success(f"✅ Skill Activities generated successfully! {message}")
                            
                            # Download button
                            doc = create_word_document(st.session_state.skill_activities)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            # Generate filename based on whether PDF was uploaded
                            if uploaded_file_st:
                                download_filename = f"skill_activities_{uploaded_file_st.name.replace('.pdf', '.docx')}"
                            else:
                                download_filename = f"skill_activities_{selected_grade.replace(' ', '_')}.docx"
                            
                            st.download_button(
                                label="📥 Download Skill Activities as Word (.docx)",
                                data=doc_io,
                                file_name=download_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_skills_streaming"
                            )
                        else:
                            st.error(f"❌ Failed to generate Skill Activities: {message}")
                    else:
                        # For AI subject without PDF, pass None for pdf_bytes and a placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Skills_Generation"
                        
                        content, message = generate_specific_content(
                            "skills", 
                            pdf_bytes, 
                            pdf_filename, 
                            selected_grade, 
                            model_progression, 
                            subject_type, 
                            word_limits,
                            use_chunked=(analysis_method == "Chunked (For Complex Documents)"), 
                            use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                            pdf_method=pdf_method
                        )
                        if content:
                            st.session_state.skill_activities = content
                            
                            # CRITICAL: Save content with protection
                            save_content_safely("skill_activities", content, selected_grade)
                            
                            st.success(f"✅ Skill Activities generated successfully! {message}")
                            st.subheader("🛠️ Skill Activities:")
                            with st.expander("View Skill Activities", expanded=True):
                                st.markdown(st.session_state.skill_activities)
                            
                            # Download and Expand buttons
                            dl_col, expand_col = st.columns(2)
                            with dl_col:
                                doc = create_word_document(st.session_state.skill_activities)
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                download_filename = f"skill_activities_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"skill_activities_{selected_grade.replace(' ', '_')}.docx"
                                st.download_button(
                                    label="📥 Download Skill Activities as Word (.docx)",
                                    data=doc_io,
                                    file_name=download_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_skills_standard"
                                )
                            with expand_col:
                                if st.button("✨ Expand This Content", key="expand_new_skills"):
                                    st.session_state.show_new_skills_expander = True
                                    st.rerun()
                        
                        # Show expander for newly generated skills
                        if st.session_state.get('show_new_skills_expander', False):
                            hybrid_content_expander(
                                st.session_state.skill_activities, 
                                "skills", 
                                selected_grade, 
                                subject_type.replace(" (Uses Model Chapter Progression)", "")
                            )
                            if st.button("❌ Close Expander", key="close_new_skills_expander"):
                                st.session_state.show_new_skills_expander = False
                                st.rerun()
                        else:
                            st.error(f"❌ Failed to generate Skill Activities: {message}")
            
        if generate_art:
                # Get word limits from session state
                word_limits = st.session_state.get('word_limits', {})
                
                with st.spinner(f"🧠 Generating Art-Integrated Learning for {selected_grade}..."):
                    if use_streaming:
                        # For AI subject without PDF, pass placeholder name
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Projects_Generation"
                        content, message = handle_streaming_generation(
                            "art", pdf_bytes, pdf_filename, selected_grade, 
                            model_progression, subject_type, word_limits, "art", pdf_method
                        )
                        if content:
                            st.session_state.art_learning = content
                            st.success(f"✅ Art-Integrated Learning generated successfully! {message}")
                            
                            # Download button
                            doc = create_word_document(st.session_state.art_learning)
                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            # Generate filename based on whether PDF was uploaded
                            if uploaded_file_st:
                                download_filename = f"art_learning_{uploaded_file_st.name.replace('.pdf', '.docx')}"
                            else:
                                download_filename = f"art_learning_{selected_grade.replace(' ', '_')}.docx"
                            
                            st.download_button(
                                label="📥 Download Art-Integrated Learning as Word (.docx)",
                                data=doc_io,
                                file_name=download_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_art_streaming"
                            )
                        else:
                            st.error(f"❌ Failed to generate Art-Integrated Learning: {message}")
                    else:
                        # For AI subject without PDF, pass None for pdf_bytes and a placeholder name  
                        pdf_filename = uploaded_file_st.name if uploaded_file_st else "AI_Projects_Generation"
                        
                        content, message = generate_specific_content(
                            "art", 
                            pdf_bytes, 
                            pdf_filename, 
                            selected_grade, 
                            model_progression, 
                            subject_type, 
                            word_limits,
                            use_chunked=(analysis_method == "Chunked (For Complex Documents)"), 
                            use_openrouter_method=(pdf_method == "Direct PDF Upload (OpenRouter Recommended)"),
                            pdf_method=pdf_method
                        )
                        if content:
                            st.session_state.art_learning = content
                            
                            # CRITICAL: Save content with protection
                            save_content_safely("art_learning", content, selected_grade)
                            
                            st.success(f"✅ Art-Integrated Learning generated successfully! {message}")
                            st.subheader("🎨 Art-Integrated Learning:")
                            with st.expander("View Art-Integrated Learning", expanded=True):
                                st.markdown(st.session_state.art_learning)
                            
                            # Download and Expand buttons
                            dl_col, expand_col = st.columns(2)
                            with dl_col:
                                doc = create_word_document(st.session_state.art_learning)
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)
                                download_filename = f"art_learning_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"art_learning_{selected_grade.replace(' ', '_')}.docx"
                                st.download_button(
                                    label="📥 Download Art-Integrated Learning as Word (.docx)",
                                    data=doc_io,
                                    file_name=download_filename,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_art_standard"
                                )
                            with expand_col:
                                if st.button("✨ Expand This Content", key="expand_new_art"):
                                    st.session_state.show_new_art_expander = True
                                    st.rerun()
                        
                        # Show expander for newly generated art learning
                        if st.session_state.get('show_new_art_expander', False):
                            hybrid_content_expander(
                                st.session_state.art_learning, 
                                "art", 
                                selected_grade, 
                                subject_type.replace(" (Uses Model Chapter Progression)", "")
                            )
                            if st.button("❌ Close Expander", key="close_new_art_expander"):
                                st.session_state.show_new_art_expander = False
                                st.rerun()
                        else:
                            st.error(f"❌ Failed to generate Art-Integrated Learning: {message}")
            
        if download_all:
                # Combine all generated content (if any)
                all_content_parts = []
                
                if st.session_state.chapter_content:
                    all_content_parts.append("# CHAPTER CONTENT\n\n" + st.session_state.chapter_content)
                else:
                    st.warning("Chapter Content has not been generated yet.")
                
                if st.session_state.exercises:
                    all_content_parts.append("# EXERCISES\n\n" + st.session_state.exercises)
                else:
                    st.warning("Exercises have not been generated yet.")
                
                if st.session_state.skill_activities:
                    all_content_parts.append("# SKILL ACTIVITIES\n\n" + st.session_state.skill_activities)
                else:
                    st.warning("Skill Activities have not been generated yet.")
                
                if st.session_state.art_learning:
                    all_content_parts.append("# ART-INTEGRATED LEARNING\n\n" + st.session_state.art_learning)
                else:
                    st.warning("Art-Integrated Learning has not been generated yet.")
                
                if all_content_parts:
                    combined_content = "\n\n" + "\n\n".join(all_content_parts)
                    
                    # Create Word document with all content
                    doc = create_word_document(combined_content)
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    download_filename = f"complete_chapter_{uploaded_file_st.name.replace('.pdf', '.docx')}" if uploaded_file_st else f"complete_chapter_{selected_grade.replace(' ', '_')}.docx"
                    st.download_button(
                        label="📥 Download Complete Chapter with All Elements as Word (.docx)",
                        data=doc_io,
                        file_name=download_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("No content has been generated yet. Please generate at least one type of content first.")
    else:
        st.error("Failed to load the Model Chapter Progression. The tool cannot proceed without it.")

with tab2:
    st.header("💬 Content Chat with EeeBee")
    st.markdown("""
    Chat with EeeBee (powered by Claude) for content development assistance! Upload PDFs for context or ask questions directly.
    EeeBee can help with content creation, modification, curriculum alignment, and educational guidance.
    """)
    
    # Move chat settings to main area (above chat)
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        # Subject context first to determine grade/level options
        chat_subject = st.selectbox(
            "Subject Context:",
            ["Science Education", "Mathematics", "Mathematics Primary (Classes 1-5)", "Science & E.V.S. (Classes 1-2)", "Science & E.V.S. (Classes 3-5)", "Computer Science", "Artificial Intelligence", "Robotics", "Social Studies", "English", "Hindi", "General Education", "Other"],
            key="chat_subject"
        )
    
    with col2:
        # Conditional Grade/Level Selection based on subject
        if chat_subject in ["Artificial Intelligence", "Robotics"]:
            # Level Selector for AI/Robotics
            level_options = [
                "JL1 (Classes 1-3)",
                "JL2 (Classes 4-5)", 
                "SL1 (Classes 6-7)",
                "SL2 (Classes 8-10)",
                "SL3 (Class 11)"
            ]
            chat_grade = st.selectbox("Level Context:", level_options, index=2, key="chat_level")
        else:
            # Regular Grade Selector for other subjects
            chat_grade = st.selectbox(
                "Grade Level Context:", 
                [f"Grade {i}" for i in range(1, 13)], 
                index=8, 
                key="chat_grade"
            )
    
    with col3:
        # Clear chat button
        if st.button("🗑️ Clear Chat History", key="clear_chat"):
            st.session_state.chat_messages = []
            st.session_state.chat_uploaded_files = []
            st.rerun()
    
    # PDF Upload for chat context
    st.subheader("📄 Upload Documents for Context")
    chat_uploaded_files = st.file_uploader(
        "Upload PDFs for EeeBee to reference:",
        type="pdf",
        accept_multiple_files=True,
        key="chat_pdf_uploader"
    )
    
    if chat_uploaded_files:
        st.session_state.chat_uploaded_files = chat_uploaded_files
        st.success(f"📄 {len(chat_uploaded_files)} PDF(s) uploaded successfully!")
        uploaded_files_info = ""
        for file in chat_uploaded_files:
            uploaded_files_info += f"• {file.name}\n"
        st.text(uploaded_files_info)
    
    st.divider()
    
    # Initialize chat session state
    if 'chat_messages' not in st.session_state:
        st.session_state.chat_messages = []
    if 'chat_uploaded_files' not in st.session_state:
        st.session_state.chat_uploaded_files = []
    
    # Display chat messages in a container
    chat_container = st.container()
    
    with chat_container:
        # Display existing chat messages
        for message in st.session_state.chat_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    
    # Chat input at the bottom
    if prompt := st.chat_input("Ask EeeBee anything about content development..."):
        # Add user message to chat history
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        
        # Display user message
        with chat_container:
            with st.chat_message("user"):
                st.markdown(prompt)
        
        # Generate EeeBee response with streaming
        with chat_container:
            with st.chat_message("assistant"):
                # Create placeholder for streaming response
                response_placeholder = st.empty()
                response_text = ""
                
                try:
                    # Generate streaming response
                    for chunk in generate_chat_response_stream(
                        prompt, 
                        st.session_state.chat_messages, 
                        st.session_state.chat_uploaded_files,
                        chat_grade,
                        chat_subject
                    ):
                        response_text += chunk
                        response_placeholder.markdown(response_text + "▊")  # Add cursor effect
                    
                    # Remove cursor and show final response
                    response_placeholder.markdown(response_text)
                    
                    # Add assistant response to chat history
                    st.session_state.chat_messages.append({"role": "assistant", "content": response_text})
                    
                except Exception as e:
                    error_message = f"I encountered an error: {str(e)}. Please try asking your question again."
                    response_placeholder.markdown(error_message)
                    st.session_state.chat_messages.append({"role": "assistant", "content": error_message})

with tab3:
    st.header("🔍 PDF Checker")
    st.markdown("""
    Check your PDF documents for spelling errors, grammar issues, content coherence, and formatting.
    This tool provides comprehensive analysis to ensure your educational content is polished and professional.
    """)
    
    # PDF Checker section
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # Subject Selection first to determine grade/level options
        subject_checker = st.selectbox(
            "Select Subject",
            options=["Science", "Mathematics", "English", "Social Studies", "Computer Science", "Artificial Intelligence", "Robotics"],
            key="subject_checker"
        )
        
        # Conditional Grade/Level Selection
        if subject_checker in ["Artificial Intelligence", "Robotics"]:
            # Level Selector for AI/Robotics
            level_options = [
                "JL1 (Classes 1-3)",
                "JL2 (Classes 4-5)", 
                "SL1 (Classes 6-7)",
                "SL2 (Classes 8-10)",
                "SL3 (Class 11)"
            ]
            selected_level_checker = st.selectbox("Select Target Level:", level_options, index=2, key="level_selector_checker")
            grade_checker = selected_level_checker  # Use level as grade for AI/Robotics
        else:
            # Regular Grade Selector for other subjects
            grade_checker = st.selectbox(
                "Select Grade Level",
                options=[f"Grade {i}" for i in range(1, 13)],
                key="grade_checker"
            )
    
    with col2:
        # Check options
        st.markdown("**Check Options:**")
        check_spelling = st.checkbox("Check Spelling", value=True, key="check_spelling")
        check_grammar = st.checkbox("Check Grammar", value=True, key="check_grammar")
        check_coherence = st.checkbox("Check Content Coherence", value=True, key="check_coherence")
        check_formatting = st.checkbox("Check Formatting", value=True, key="check_formatting")
    
    # File uploader
    uploaded_pdf_checker = st.file_uploader(
        "Upload PDF for checking",
        type=['pdf'],
        key="pdf_checker_upload"
    )
    
    if uploaded_pdf_checker is not None:
        if st.button("🔍 Check PDF", type="primary", key="check_pdf_button"):
            with st.spinner("Analyzing your PDF..."):
                try:
                    # Extract text from PDF
                    pdf_text = extract_text_from_pdf(uploaded_pdf_checker)
                    
                    if not pdf_text.strip():
                        st.warning("No text found in PDF. Attempting OCR...")
                        pdf_text = process_pdf_with_mistral_ocr(uploaded_pdf_checker)
                    
                    if pdf_text.strip():
                        # Prepare the checking prompt
                        check_types = []
                        if check_spelling:
                            check_types.append("spelling errors")
                        if check_grammar:
                            check_types.append("grammar issues")
                        if check_coherence:
                            check_types.append("content coherence and logical flow")
                        if check_formatting:
                            check_types.append("formatting consistency")
                        
                        checking_prompt = f"""You are an expert educational content reviewer. Analyze the following {grade_checker} {subject_checker} content for {', '.join(check_types)}.

Content to analyze:
{pdf_text}

Please provide a detailed analysis with:
1. **Overall Assessment**: Brief overview of the document quality
2. **Issues Found**: List specific issues with page/location references where possible
3. **Suggestions**: Actionable recommendations for improvement
4. **Summary**: Key points and priority fixes

Format your response in a clear, structured manner using markdown."""

                        # Call AI for analysis
                        headers = {
                            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                            "HTTP-Referer": YOUR_SITE_URL,
                            "X-Title": YOUR_SITE_NAME,
                            "Content-Type": "application/json"
                        }
                        
                        response = requests.post(
                            "https://openrouter.ai/api/v1/chat/completions",
                            headers=headers,
                            json={
                                "model": MODEL_NAME,
                                "messages": [{"role": "user", "content": checking_prompt}],
                                "max_tokens": 60000,
                                "temperature": 0.3
                            }
                        )
                        
                        if response.status_code == 200:
                            result = response.json()
                            analysis = result['choices'][0]['message']['content']
                            
                            # Display results
                            st.success("✅ PDF Analysis Complete!")
                            
                            # Save results
                            save_content_safely("pdf_analysis", analysis, grade_checker)
                            
                            # Create expandable sections for results
                            with st.expander("📊 Full Analysis Report", expanded=True):
                                st.markdown(analysis)
                            
                            # Download button for analysis
                            st.download_button(
                                label="📥 Download Analysis Report",
                                data=analysis,
                                file_name=f"pdf_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                mime="text/plain",
                                key="download_pdf_analysis"
                            )
                            
                        else:
                            st.error(f"API Error: {response.status_code}")
                    else:
                        st.error("Could not extract text from PDF")
                        
                except Exception as e:
                    st.error(f"Error analyzing PDF: {str(e)}")
# Tab 4: Remedial Content Generator
with tab4:
    st.header("🩹 Remedial Content Generator")
    st.markdown("""
    Generate comprehensive remedial content for Science and Mathematics based on specific concepts.
    Choose your grade level, subject, and concept to create targeted educational materials.
    """)
    
    # Grade and Subject Selection
    col1, col2 = st.columns([1, 1])
    
    with col1:
        grade_remedial = st.selectbox(
            "Select Grade Level",
            options=[f"Grade {i}" for i in range(1, 13)],
            key="grade_remedial"
        )
        
    with col2:
        # Determine available subjects based on grade
        grade_num = int(grade_remedial.split()[-1])
        if grade_num <= 2:
            subjects = ["Mathematics"]
        else:
            subjects = ["Science", "Mathematics"]
            
        subject_remedial = st.selectbox(
            "Select Subject",
            options=subjects,
            key="subject_remedial"
        )
    
    # Concept input
    concept_name = st.text_input(
        "Enter the concept name",
        placeholder=f"e.g., {'Photosynthesis' if subject_remedial == 'Science' else 'Fractions'}",
        key="concept_input"
    )
    
    # Content type selection based on grade and subject
    st.markdown("### Select Content Types to Generate:")
    
    if subject_remedial == "Science":
        # Science content options
        col1, col2 = st.columns(2)
        with col1:
            gen_video_script = st.checkbox("Video Script with Narration", value=True, key="gen_video_science")
            gen_notes = st.checkbox("Detailed Notes", value=True, key="gen_notes_science")
            gen_worksheet = st.checkbox("Worksheet", value=True, key="gen_worksheet_science")
        with col2:
            gen_mcq_bank = st.checkbox("MCQ Question Bank (40 Questions)", value=True, key="gen_mcq_science")
            gen_summary = st.checkbox("Summary Points", value=True, key="gen_summary_science")
            
    else:  # Mathematics
        if grade_num <= 2:
            # Class 1-2 Mathematics options
            col1, col2 = st.columns(2)
            with col1:
                gen_video_script = st.checkbox("Video Script with Narration", value=True, key="gen_video_math12")
                gen_notes = st.checkbox("Notes with Examples", value=True, key="gen_notes_math12")
                gen_worksheet = st.checkbox("Worksheet", value=True, key="gen_worksheet_math12")
            with col2:
                gen_mcq_bank = st.checkbox("MCQ Question Bank (40 Questions)", value=True, key="gen_mcq_math12")
                gen_fun_activity = st.checkbox("Fun Activity", value=True, key="gen_fun_math12")
        else:
            # Class 3-8 Mathematics options
            col1, col2 = st.columns(2)
            with col1:
                gen_video_script = st.checkbox("Video Script with Narration", value=True, key="gen_video_math38")
                gen_notes = st.checkbox("Notes with Examples", value=True, key="gen_notes_math38")
                gen_worksheet = st.checkbox("Worksheet", value=True, key="gen_worksheet_math38")
            with col2:
                gen_mcq_bank = st.checkbox("MCQ Question Bank (40 Questions)", value=True, key="gen_mcq_math38")
                gen_fun_activity = st.checkbox("Fun Activity", value=True, key="gen_fun_math38")
    
    # Generate button
    if concept_name and st.button("🚀 Generate Remedial Content", type="primary", key="generate_remedial"):
        with st.spinner("Generating remedial content..."):
            try:
                # Prepare content types to generate
                content_to_generate = []
                
                if subject_remedial == "Science":
                    if st.session_state.get("gen_video_science", False):
                        content_to_generate.append("video_script")
                    if st.session_state.get("gen_notes_science", False):
                        content_to_generate.append("notes")
                    if st.session_state.get("gen_worksheet_science", False):
                        content_to_generate.append("worksheet")
                    if st.session_state.get("gen_mcq_science", False):
                        content_to_generate.append("mcq_bank")
                    if st.session_state.get("gen_summary_science", False):
                        content_to_generate.append("summary")
                else:  # Mathematics
                    if grade_num <= 2:
                        if st.session_state.get("gen_video_math12", False):
                            content_to_generate.append("video_script")
                        if st.session_state.get("gen_notes_math12", False):
                            content_to_generate.append("notes")
                        if st.session_state.get("gen_worksheet_math12", False):
                            content_to_generate.append("worksheet")
                        if st.session_state.get("gen_mcq_math12", False):
                            content_to_generate.append("mcq_bank")
                        if st.session_state.get("gen_fun_math12", False):
                            content_to_generate.append("fun_activity")
                    else:
                        if st.session_state.get("gen_video_math38", False):
                            content_to_generate.append("video_script")
                        if st.session_state.get("gen_notes_math38", False):
                            content_to_generate.append("notes")
                        if st.session_state.get("gen_worksheet_math38", False):
                            content_to_generate.append("worksheet")
                        if st.session_state.get("gen_mcq_math38", False):
                            content_to_generate.append("mcq_bank")
                        if st.session_state.get("gen_fun_math38", False):
                            content_to_generate.append("fun_activity")
                
                if not content_to_generate:
                    st.warning("Please select at least one content type to generate.")
                else:
                    # Generate remedial content
                    generated_content = {}
                    
                    for content_type in content_to_generate:
                        st.info(f"Generating {content_type.replace('_', ' ')}...")
                        
                        # Prepare the prompt based on content type
                        if content_type == "video_script":
                            prompt = f"""Create a detailed video script with narration for teaching the concept '{concept_name}' for {grade_remedial} {subject_remedial}.

The script should include:
1. Opening hook to grab attention
2. Clear explanation of the concept with visual descriptions
3. Step-by-step breakdown for complex topics
4. Engaging examples and analogies
5. Interactive elements or questions
6. Summary and closing

Format the script with:
- Scene descriptions in [brackets]
- Narration in regular text
- Timing estimates for each section

Make it engaging and age-appropriate for {grade_remedial} students."""
                        
                        elif content_type == "notes":
                            prompt = f"""Create comprehensive notes for the concept '{concept_name}' for {grade_remedial} {subject_remedial}.

Include:
1. Clear definition and explanation
2. Key points and important terms
3. Detailed examples with solutions
4. Common misconceptions and clarifications
5. Visual aids descriptions (diagrams, charts)
6. Practice problems with step-by-step solutions
7. Summary of main concepts

Format with clear headings, bullet points, and organized sections."""
                        
                        elif content_type == "worksheet":
                            prompt = f"""Create a comprehensive worksheet for practicing '{concept_name}' for {grade_remedial} {subject_remedial}.

Include:
1. Warm-up questions (5 easy questions)
2. Practice problems (10 medium difficulty)
3. Challenge questions (5 difficult)
4. Word problems or application questions (5)
5. Fill in the blanks (5)
6. True/False with corrections (5)
7. Matching exercises if applicable

Provide answer key at the end with detailed solutions."""
                        
                        elif content_type == "mcq_bank":
                            prompt = f"""Create a bank of 40 Multiple Choice Questions (MCQs) on '{concept_name}' for {grade_remedial} {subject_remedial}.

For each question include:
- Clear question stem
- 4 options (A, B, C, D)
- Correct answer marked
- Brief explanation of why the answer is correct
- Difficulty level (Easy/Medium/Hard)

Distribute questions as:
- 15 Easy questions (basic understanding)
- 20 Medium questions (application)
- 5 Hard questions (analysis/problem-solving)

Vary question types: factual, conceptual, application-based."""
                        
                        elif content_type == "summary":
                            prompt = f"""Create a concise summary of key points for '{concept_name}' for {grade_remedial} {subject_remedial}.

Include:
1. Main concept in 2-3 sentences
2. 5-7 key points to remember
3. Important formulas or definitions
4. Quick tips for solving problems
5. Common mistakes to avoid
6. Real-world applications

Format as bullet points for easy revision."""
                        
                        elif content_type == "fun_activity":
                            prompt = f"""Create an engaging fun activity for teaching '{concept_name}' for {grade_remedial} Mathematics.

Include:
1. Activity name and objective
2. Materials needed (simple, easily available)
3. Step-by-step instructions
4. Learning outcomes
5. Variations for different skill levels
6. Extension activities
7. Assessment ideas

Make it hands-on, interactive, and age-appropriate."""
                        
                        # Call API to generate content
                        messages = [{
                            "role": "user",
                            "content": prompt
                        }]
                        
                        headers = {
                            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                            "HTTP-Referer": YOUR_SITE_URL,
                            "X-Title": YOUR_SITE_NAME,
                            "Content-Type": "application/json"
                        }
                        
                        response = requests.post(
                            "https://openrouter.ai/api/v1/chat/completions",
                            headers=headers,
                            json={
                                "model": MODEL_NAME,
                                "messages": messages,
                                "max_tokens": 60000,
                                "temperature": 0.3
                            }
                        )
                        
                        if response.status_code == 200:
                            result = response.json()
                            content = result['choices'][0]['message']['content']
                            generated_content[content_type] = content
                        else:
                            st.error(f"Error generating {content_type}: {response.status_code}")
                    
                    # Display generated content
                    if generated_content:
                        st.success("✅ Remedial content generated successfully!")
                        
                        # Create tabs for different content types
                        tab_names = []
                        for ct in generated_content.keys():
                            readable_name = ct.replace('_', ' ').title()
                            tab_names.append(readable_name)
                        
                        tabs = st.tabs(tab_names)
                        
                        for idx, (content_type, content) in enumerate(generated_content.items()):
                            with tabs[idx]:
                                st.markdown(content)
                                
                                # Download button for each content type
                                st.download_button(
                                    label=f"📥 Download {content_type.replace('_', ' ').title()}",
                                    data=content,
                                    file_name=f"{concept_name}_{content_type}_{grade_remedial}.txt",
                                    mime="text/plain",
                                    key=f"download_{content_type}"
                                )
                        
                        # Combined download button
                        st.markdown("---")
                        combined_content = f"# Remedial Content for {concept_name}\n{grade_remedial} {subject_remedial}\n\n"
                        
                        for content_type, content in generated_content.items():
                            readable_name = content_type.replace('_', ' ').title()
                            combined_content += f"\n\n## {readable_name}\n\n{content}\n\n"
                            combined_content += "="*50 + "\n"
                        
                        # Create Word document
                        doc = create_word_document(combined_content)
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                label="📥 Download All Content (Text)",
                                data=combined_content,
                                file_name=f"{concept_name}_remedial_content_{grade_remedial}.txt",
                                mime="text/plain",
                                key="download_all_text"
                            )
                        with col2:
                            st.download_button(
                                label="📥 Download All Content (Word)",
                                data=doc_io,
                                file_name=f"{concept_name}_remedial_content_{grade_remedial}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_all_word"
                            )
                            
            except Exception as e:
                st.error(f"Error generating content: {str(e)}")

st.sidebar.markdown("---")
st.sidebar.info("This app uses the Claude API via OpenRouter for AI-powered content analysis and generation.")
